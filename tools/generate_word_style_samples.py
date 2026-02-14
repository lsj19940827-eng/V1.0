# -*- coding: utf-8 -*-
"""
生成3种计算书排版风格样例文档，供用户预览选择
"""
import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from copy import deepcopy

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'data')

# ============================================================
# 公共模拟数据
# ============================================================
SAMPLE_PARAMS = [
    ("设计流量 Q", "12.500 m³/s"),
    ("渠底比降 i", "1/5000"),
    ("糙率 n", "0.015"),
    ("边坡系数 m", "1.50"),
    ("不冲流速", "0.80 m/s"),
    ("不淤流速", "0.30 m/s"),
]
SAMPLE_STEPS = [
    ("1. 确定断面类型", "采用梯形断面，边坡系数 m = 1.50"),
    ("2. 水力最佳断面计算", None),
    ("  公式", r"B = 2h(√(1+m²) - m)"),
    ("  代入", r"B = 2×2.165×(√(1+1.5²) - 1.5) = 2×2.165×0.303 = 1.312 m"),
    ("3. 经济断面调整（α=1.20）", None),
    ("  公式", r"B' = α × B_opt = 1.20 × 1.312 = 1.574 m"),
    ("  取整", r"B = 1.60 m"),
    ("4. 计算正常水深", None),
    ("  迭代", r"h = 2.098 m"),
    ("5. 复核流速", None),
    ("  公式", r"V = Q/A = 12.500 / 25.83 = 0.484 m/s"),
    ("  校核", "0.30 m/s < 0.484 m/s < 0.80 m/s  ✓ 满足要求"),
    ("6. 计算超高与堤顶宽", None),
    ("  安全超高", r"Δ = 0.50 m"),
    ("  渠深", r"H = h + Δ = 2.098 + 0.50 = 2.60 m"),
]
SAMPLE_TABLE_HEADERS = ['α值', '方案类型', '底宽B(m)', '水深h(m)', '宽深比β', '流速V(m/s)', '面积增加', '状态']
SAMPLE_TABLE_DATA = [
    ['1.00', '水力最佳', '1.312', '2.165', '0.606', '0.512', '+0%', ''],
    ['1.10', '经济断面', '1.443', '2.131', '0.677', '0.498', '+2%', ''],
    ['1.20', '推荐方案', '1.574', '2.098', '0.750', '0.484', '+4%', '★选中'],
    ['1.30', '偏宽断面', '1.706', '2.066', '0.826', '0.471', '+7%', ''],
    ['1.50', '宽浅断面', '1.969', '2.005', '0.982', '0.447', '+13%', ''],
]

SAMPLE_RESULT = [
    ("底宽 B", "1.60 m"),
    ("正常水深 h", "2.098 m"),
    ("渠深 H", "2.60 m"),
    ("过水面积 A", "25.83 m²"),
    ("湿周 χ", "9.16 m"),
    ("水力半径 R", "2.820 m"),
    ("流速 V", "0.484 m/s"),
]


def set_cell_shading(cell, color_hex):
    """设置单元格底色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    """设置单元格边框 kwargs: top/bottom/left/right = (size, color, style)"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
    for edge, (sz, color, style) in kwargs.items():
        el = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{style}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        )
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_paragraph_border_bottom(paragraph, color="000000", size=6):
    """段落底部边框线"""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:space="1" w:color="{color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def set_paragraph_border_left(paragraph, color="4472C4", size=18):
    """段落左侧边框（竖条装饰）"""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="{size}" w:space="6" w:color="{color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_page_number(section, font_name='Times New Roman', font_size=9):
    """添加页脚页码"""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run1 = p.add_run("— ")
    run1.font.name = font_name
    run1.font.size = Pt(font_size)
    
    # PAGE field
    fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run_field = p.add_run()
    run_field._r.append(fld_char_begin)
    
    instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run_instr = p.add_run()
    run_instr._r.append(instr)
    
    fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run_end = p.add_run()
    run_end._r.append(fld_char_end)
    
    run2 = p.add_run(" —")
    run2.font.name = font_name
    run2.font.size = Pt(font_size)


# ============================================================
# 方案一：经典工程报告风格
# ============================================================
def build_scheme1():
    """经典工程报告 —— 仿国标/设计院风格
    
    特点：
    - 宋体/黑体正式字体组合
    - 三线表
    - 蓝色标题装饰线
    - 封面信息框
    - 页眉工程名称 + 页脚页码
    """
    doc = Document()
    
    # --- 页面设置 ---
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    
    # --- 默认样式 ---
    style_normal = doc.styles['Normal']
    style_normal.font.name = '宋体'
    style_normal.font.size = Pt(12)
    style_normal._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style_normal.paragraph_format.line_spacing = Pt(22)
    
    # ==================== 封面 ====================
    # 上方留白
    for _ in range(4):
        doc.add_paragraph('')
    
    # 标题
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run('明渠水力计算书')
    title_run.font.name = '黑体'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    title_run.font.size = Pt(36)
    title_run.bold = True
    
    # 副标题
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run('（梯形断面）')
    sub_run.font.name = '黑体'
    sub_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    sub_run.font.size = Pt(22)
    sub_run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
    
    doc.add_paragraph('')
    doc.add_paragraph('')
    
    # 封面信息表
    info_table = doc.add_table(rows=5, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_items = [
        ("工程名称", "XX灌区续建配套工程"),
        ("渠道名称", "总干渠 GK0+000~GK12+500"),
        ("设计单位", "XX水利水电勘测设计院"),
        ("计算人员", ""),
        ("计算日期", "2025年02月"),
    ]
    for i, (label, value) in enumerate(info_items):
        c0 = info_table.cell(i, 0)
        c1 = info_table.cell(i, 1)
        c0.text = label
        c1.text = value
        c0.width = Cm(3.5)
        c1.width = Cm(7)
        for c in (c0, c1):
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.name = '宋体'
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    r.font.size = Pt(14)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # 设置边框：只有底线
        for c in (c0, c1):
            set_cell_border(c, 
                top=("0", "FFFFFF", "none"),
                left=("0", "FFFFFF", "none"),
                right=("0", "FFFFFF", "none"),
                bottom=("4", "333333", "single"))
    
    # 分页
    doc.add_page_break()
    
    # ==================== 页眉页脚 ====================
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run('XX灌区续建配套工程 — 明渠水力计算书')
    hr.font.name = '宋体'
    hr._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    hr.font.size = Pt(9)
    hr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    set_paragraph_border_bottom(hp, color="AAAAAA", size=4)
    
    add_page_number(section, '宋体', 9)
    
    # ==================== 正文 ====================
    # 一级标题：基本资料
    h1 = doc.add_paragraph()
    h1_run = h1.add_run('一、基本资料')
    h1_run.font.name = '黑体'
    h1_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    h1_run.font.size = Pt(16)
    h1_run.bold = True
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    set_paragraph_border_bottom(h1, color="4472C4", size=8)
    
    # 参数表 —— 三线表
    param_table = doc.add_table(rows=len(SAMPLE_PARAMS)+1, cols=2)
    param_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for j, hdr in enumerate(["参  数", "取  值"]):
        c = param_table.cell(0, j)
        c.text = hdr
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = '黑体'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                r.font.size = Pt(11)
                r.bold = True
        set_cell_shading(c, "4472C4")
        for p in c.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    
    for i, (name, val) in enumerate(SAMPLE_PARAMS):
        for j, txt in enumerate([name, val]):
            c = param_table.cell(i+1, j)
            c.text = txt
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.name = '宋体'
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    r.font.size = Pt(11)
        if i % 2 == 0:
            for j in range(2):
                set_cell_shading(param_table.cell(i+1, j), "EDF2FA")
    
    # 三线表边框
    tbl = param_table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="12" w:space="0" w:color="4472C4"/>'
        f'  <w:bottom w:val="single" w:sz="12" w:space="0" w:color="4472C4"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="B4C6E7"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="FFFFFF"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="FFFFFF"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="FFFFFF"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)
    
    doc.add_paragraph('')
    
    # 二级标题
    h1_2 = doc.add_paragraph()
    h1_2_run = h1_2.add_run('二、计算过程')
    h1_2_run.font.name = '黑体'
    h1_2_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    h1_2_run.font.size = Pt(16)
    h1_2_run.bold = True
    h1_2.paragraph_format.space_before = Pt(18)
    h1_2.paragraph_format.space_after = Pt(10)
    set_paragraph_border_bottom(h1_2, color="4472C4", size=8)
    
    # 基础公式
    h2 = doc.add_paragraph()
    h2_run = h2.add_run('2.1  基础公式')
    h2_run.font.name = '黑体'
    h2_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    h2_run.font.size = Pt(14)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    
    formula_lines = [
        "曼宁公式：Q = (1/n) · A · R^(2/3) · i^(1/2)",
        "过水面积：A = (B + m·h) · h",
        "湿    周：χ = B + 2h√(1+m²)",
        "水力半径：R = A / χ",
    ]
    for fl in formula_lines:
        fp = doc.add_paragraph()
        fp.paragraph_format.left_indent = Cm(1.5)
        fp.paragraph_format.space_before = Pt(2)
        fp.paragraph_format.space_after = Pt(2)
        fr = fp.add_run(fl)
        fr.font.name = 'Cambria Math'
        fr._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        fr.font.size = Pt(12)
    
    doc.add_paragraph('')
    
    # 计算步骤
    h2_2 = doc.add_paragraph()
    h2_2_run = h2_2.add_run('2.2  计算步骤')
    h2_2_run.font.name = '黑体'
    h2_2_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    h2_2_run.font.size = Pt(14)
    h2_2.paragraph_format.space_before = Pt(12)
    h2_2.paragraph_format.space_after = Pt(6)
    
    for label, content in SAMPLE_STEPS:
        if label.startswith("  "):
            # 公式/子步骤 —— 缩进
            fp = doc.add_paragraph()
            fp.paragraph_format.left_indent = Cm(1.5)
            fp.paragraph_format.space_before = Pt(1)
            fp.paragraph_format.space_after = Pt(1)
            text = f"{label.strip()}：{content}" if content else label.strip()
            fr = fp.add_run(text)
            fr.font.name = 'Cambria Math'
            fr._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            fr.font.size = Pt(11)
        else:
            # 步骤标题
            sp = doc.add_paragraph()
            sp.paragraph_format.space_before = Pt(8)
            sp.paragraph_format.space_after = Pt(4)
            sr = sp.add_run(label)
            sr.font.name = '黑体'
            sr._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            sr.font.size = Pt(12)
            sr.bold = True
            if content:
                cr = sp.add_run(f'\n{content}')
                cr.font.name = '宋体'
                cr._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                cr.font.size = Pt(11)
    
    doc.add_paragraph('')
    
    # 方案对比表
    h1_3 = doc.add_paragraph()
    h1_3_run = h1_3.add_run('三、断面方案对比')
    h1_3_run.font.name = '黑体'
    h1_3_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    h1_3_run.font.size = Pt(16)
    h1_3_run.bold = True
    h1_3.paragraph_format.space_before = Pt(18)
    h1_3.paragraph_format.space_after = Pt(10)
    set_paragraph_border_bottom(h1_3, color="4472C4", size=8)
    
    # 表标题
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_r = cap.add_run('表 1  附录E断面方案对比表')
    cap_r.font.name = '黑体'
    cap_r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    cap_r.font.size = Pt(11)
    cap_r.bold = True
    
    comp_table = doc.add_table(rows=len(SAMPLE_TABLE_DATA)+1, cols=len(SAMPLE_TABLE_HEADERS))
    comp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, hdr in enumerate(SAMPLE_TABLE_HEADERS):
        c = comp_table.cell(0, j)
        c.text = hdr
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = '黑体'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                r.font.size = Pt(9)
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(c, "4472C4")
    
    for i, row_data in enumerate(SAMPLE_TABLE_DATA):
        for j, val in enumerate(row_data):
            c = comp_table.cell(i+1, j)
            c.text = val
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.name = '宋体'
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    r.font.size = Pt(9)
                    if val == '★选中':
                        r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
                        r.bold = True
        if i % 2 == 0:
            for j in range(len(SAMPLE_TABLE_HEADERS)):
                set_cell_shading(comp_table.cell(i+1, j), "EDF2FA")
    
    # 三线表样式
    tbl2 = comp_table._tbl
    tblPr2 = tbl2.tblPr if tbl2.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders2 = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="12" w:space="0" w:color="4472C4"/>'
        f'  <w:bottom w:val="single" w:sz="12" w:space="0" w:color="4472C4"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="D6E0F0"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="D6E0F0"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="FFFFFF"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="FFFFFF"/>'
        f'</w:tblBorders>'
    )
    tblPr2.append(borders2)
    
    doc.add_paragraph('')
    
    # 计算结果汇总
    h1_4 = doc.add_paragraph()
    h1_4_run = h1_4.add_run('四、计算结果')
    h1_4_run.font.name = '黑体'
    h1_4_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    h1_4_run.font.size = Pt(16)
    h1_4_run.bold = True
    h1_4.paragraph_format.space_before = Pt(18)
    h1_4.paragraph_format.space_after = Pt(10)
    set_paragraph_border_bottom(h1_4, color="4472C4", size=8)
    
    res_table = doc.add_table(rows=len(SAMPLE_RESULT)+1, cols=2)
    res_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, hdr in enumerate(["项  目", "计算结果"]):
        c = res_table.cell(0, j)
        c.text = hdr
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = '黑体'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                r.font.size = Pt(11)
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(c, "4472C4")
    
    for i, (name, val) in enumerate(SAMPLE_RESULT):
        for j, txt in enumerate([name, val]):
            c = res_table.cell(i+1, j)
            c.text = txt
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.name = '宋体'
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    r.font.size = Pt(11)
        if i % 2 == 0:
            for j in range(2):
                set_cell_shading(res_table.cell(i+1, j), "EDF2FA")
    
    tbl3 = res_table._tbl
    tblPr3 = tbl3.tblPr if tbl3.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders3 = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="12" w:space="0" w:color="4472C4"/>'
        f'  <w:bottom w:val="single" w:sz="12" w:space="0" w:color="4472C4"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="B4C6E7"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="FFFFFF"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="FFFFFF"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="FFFFFF"/>'
        f'</w:tblBorders>'
    )
    tblPr3.append(borders3)
    
    filepath = os.path.join(OUTPUT_DIR, '排版方案1_经典工程报告风格.docx')
    doc.save(filepath)
    print(f"✓ 方案1已生成: {filepath}")
    return filepath


# ============================================================
# 方案二：现代简约风格
# ============================================================
def build_scheme2():
    """现代简约 —— 科技感、清爽
    
    特点：
    - 微软雅黑无衬线体
    - 深灰+青蓝配色
    - 左侧竖条装饰标题
    - 浅灰交替行表格
    - 参数卡片式布局
    """
    doc = Document()
    
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    
    ACCENT = "2196F3"  # Material Blue
    DARK = "263238"    # Blue Grey 900
    LIGHT_BG = "F5F7FA"
    
    style_normal = doc.styles['Normal']
    style_normal.font.name = '微软雅黑'
    style_normal.font.size = Pt(11)
    style_normal._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style_normal.paragraph_format.line_spacing = Pt(20)
    
    # ==================== 封面 ====================
    for _ in range(3):
        doc.add_paragraph('')
    
    # 顶部装饰线
    deco = doc.add_paragraph()
    deco.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_border_bottom(deco, color=ACCENT, size=24)
    
    doc.add_paragraph('')
    
    # 标题
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    t_run = title_p.add_run('明渠水力计算书')
    t_run.font.name = '微软雅黑'
    t_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    t_run.font.size = Pt(32)
    t_run.font.color.rgb = RGBColor(0x26, 0x32, 0x38)
    t_run.bold = True
    
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    s_run = sub.add_run('HYDRAULIC CALCULATION REPORT — TRAPEZOIDAL SECTION')
    s_run.font.name = 'Segoe UI'
    s_run.font.size = Pt(11)
    s_run.font.color.rgb = RGBColor(0x90, 0xA4, 0xAE)
    
    doc.add_paragraph('')
    doc.add_paragraph('')
    
    # 封面信息 —— 竖排式
    cover_items = [
        ("工 程", "XX灌区续建配套工程"),
        ("渠 道", "总干渠 GK0+000~GK12+500"),
        ("单 位", "XX水利水电勘测设计院"),
        ("日 期", "2025年02月"),
    ]
    for label, value in cover_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        lr = p.add_run(f'{label}    ')
        lr.font.name = '微软雅黑'
        lr._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        lr.font.size = Pt(13)
        lr.font.color.rgb = RGBColor(0x90, 0xA4, 0xAE)
        vr = p.add_run(value)
        vr.font.name = '微软雅黑'
        vr._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        vr.font.size = Pt(13)
        vr.font.color.rgb = RGBColor(0x26, 0x32, 0x38)
    
    doc.add_page_break()
    
    # ==================== 页眉页脚 ====================
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run('明渠水力计算书')
    hr.font.name = '微软雅黑'
    hr._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    hr.font.size = Pt(8)
    hr.font.color.rgb = RGBColor(0xB0, 0xBE, 0xC5)
    
    add_page_number(section, '微软雅黑', 9)
    
    # ==================== 正文 ====================
    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(8)
        set_paragraph_border_left(p, color=ACCENT, size=24)
        r = p.add_run(text)
        r.font.name = '微软雅黑'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        r.font.size = Pt(18)
        r.font.color.rgb = RGBColor(0x26, 0x32, 0x38)
        r.bold = True
        return p
    
    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        r.font.name = '微软雅黑'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        r.font.size = Pt(14)
        r.font.color.rgb = RGBColor(0x21, 0x96, 0xF3)
        r.bold = True
        return p
    
    def add_body(text, indent=False):
        p = doc.add_paragraph()
        if indent:
            p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text)
        r.font.name = '微软雅黑'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x37, 0x47, 0x4F)
        return p
    
    # --- 基本资料 ---
    add_h1('01  基本资料')
    
    # 参数卡片（2列表格无边框带底色）
    card_table = doc.add_table(rows=3, cols=4)
    card_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    flat_params = []
    for name, val in SAMPLE_PARAMS:
        flat_params.append((name, val))
    
    row_idx = 0
    for i in range(0, len(flat_params), 2):
        for offset in range(2):
            if i + offset < len(flat_params):
                name, val = flat_params[i + offset]
                c_name = card_table.cell(row_idx, offset * 2)
                c_val = card_table.cell(row_idx, offset * 2 + 1)
                c_name.text = name
                c_val.text = val
                for p in c_name.paragraphs:
                    for r in p.runs:
                        r.font.name = '微软雅黑'
                        r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                        r.font.size = Pt(10)
                        r.font.color.rgb = RGBColor(0x90, 0xA4, 0xAE)
                for p in c_val.paragraphs:
                    for r in p.runs:
                        r.font.name = '微软雅黑'
                        r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                        r.font.size = Pt(11)
                        r.font.color.rgb = RGBColor(0x26, 0x32, 0x38)
                        r.bold = True
                set_cell_shading(c_name, LIGHT_BG)
                set_cell_shading(c_val, LIGHT_BG)
        row_idx += 1
    
    # 卡片无边框
    for row in card_table.rows:
        for cell in row.cells:
            set_cell_border(cell,
                top=("0", "FFFFFF", "none"),
                bottom=("0", "FFFFFF", "none"),
                left=("0", "FFFFFF", "none"),
                right=("0", "FFFFFF", "none"))
    
    doc.add_paragraph('')
    
    # --- 计算过程 ---
    add_h1('02  计算过程')
    add_h2('2.1  基础公式')
    
    formula_lines = [
        "曼宁公式：Q = (1/n) · A · R^(2/3) · i^(1/2)",
        "过水面积：A = (B + m·h) · h",
        "湿    周：χ = B + 2h√(1+m²)",
        "水力半径：R = A / χ",
    ]
    for fl in formula_lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(fl)
        r.font.name = 'Cambria Math'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x37, 0x47, 0x4F)
    
    add_h2('2.2  计算步骤')
    
    for label, content in SAMPLE_STEPS:
        if label.startswith("  "):
            text = f"{label.strip()}：{content}" if content else label.strip()
            add_body(text, indent=True)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(label)
            r.font.name = '微软雅黑'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(0x26, 0x32, 0x38)
            r.bold = True
            if content:
                add_body(content, indent=True)
    
    doc.add_paragraph('')
    
    # --- 方案对比 ---
    add_h1('03  断面方案对比')
    
    comp_table = doc.add_table(rows=len(SAMPLE_TABLE_DATA)+1, cols=len(SAMPLE_TABLE_HEADERS))
    comp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, hdr in enumerate(SAMPLE_TABLE_HEADERS):
        c = comp_table.cell(0, j)
        c.text = hdr
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = '微软雅黑'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                r.font.size = Pt(9)
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(c, ACCENT)
    
    for i, row_data in enumerate(SAMPLE_TABLE_DATA):
        for j, val in enumerate(row_data):
            c = comp_table.cell(i+1, j)
            c.text = val
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.name = '微软雅黑'
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    r.font.size = Pt(9)
                    if val == '★选中':
                        r.font.color.rgb = RGBColor(0xE5, 0x39, 0x35)
                        r.bold = True
        if i % 2 == 1:
            for j in range(len(SAMPLE_TABLE_HEADERS)):
                set_cell_shading(comp_table.cell(i+1, j), LIGHT_BG)
    
    # 现代表格边框
    tbl = comp_table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="6" w:space="0" w:color="{ACCENT}"/>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="{ACCENT}"/>'
        f'  <w:insideH w:val="single" w:sz="2" w:space="0" w:color="E0E0E0"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="FFFFFF"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="FFFFFF"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="FFFFFF"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)
    
    doc.add_paragraph('')
    
    # --- 结果 ---
    add_h1('04  计算结果')
    
    # 结果用醒目卡片
    res_table = doc.add_table(rows=len(SAMPLE_RESULT), cols=2)
    res_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (name, val) in enumerate(SAMPLE_RESULT):
        c0 = res_table.cell(i, 0)
        c1 = res_table.cell(i, 1)
        c0.text = name
        c1.text = val
        for p in c0.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.name = '微软雅黑'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                r.font.size = Pt(11)
                r.font.color.rgb = RGBColor(0x60, 0x7D, 0x8B)
        for p in c1.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for r in p.runs:
                r.font.name = '微软雅黑'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                r.font.size = Pt(12)
                r.font.color.rgb = RGBColor(0x21, 0x96, 0xF3)
                r.bold = True
        set_cell_shading(c0, LIGHT_BG)
        set_cell_shading(c1, LIGHT_BG)
        set_cell_border(c0, bottom=("2", "E0E0E0", "single"), top=("0","FFFFFF","none"), left=("0","FFFFFF","none"), right=("0","FFFFFF","none"))
        set_cell_border(c1, bottom=("2", "E0E0E0", "single"), top=("0","FFFFFF","none"), left=("0","FFFFFF","none"), right=("0","FFFFFF","none"))
    
    filepath = os.path.join(OUTPUT_DIR, '排版方案2_现代简约风格.docx')
    doc.save(filepath)
    print(f"✓ 方案2已生成: {filepath}")
    return filepath


# ============================================================
# 方案三：高端咨询报告风格
# ============================================================
def build_scheme3():
    """高端咨询报告 —— 深色封面、金色点缀、精致排版
    
    特点：
    - 深藏青+金色封面色块
    - 衬线标题+无衬线正文
    - 精致表格双色交替
    - 结果突出展示卡片
    - 页眉细线 + 页脚居中
    """
    doc = Document()
    
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)
    
    NAVY = "1B2A4A"
    GOLD = "C9A96E"
    WARM_GRAY = "F8F6F3"
    TEXT_DARK = "2C3E50"
    TEXT_MID = "5D6D7E"
    
    style_normal = doc.styles['Normal']
    style_normal.font.name = '微软雅黑'
    style_normal.font.size = Pt(11)
    style_normal._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style_normal.paragraph_format.line_spacing = Pt(21)
    
    # ==================== 封面 ====================
    # 顶部深色色块（用全宽表格模拟）
    cover_bar = doc.add_table(rows=1, cols=1)
    cover_bar.alignment = WD_TABLE_ALIGNMENT.CENTER
    bar_cell = cover_bar.cell(0, 0)
    # 6行空白模拟高度
    bar_cell.text = "\n\n\n\n\n"
    set_cell_shading(bar_cell, NAVY)
    set_cell_border(bar_cell, top=("0","FFFFFF","none"), bottom=("0","FFFFFF","none"), left=("0","FFFFFF","none"), right=("0","FFFFFF","none"))
    for p in bar_cell.paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
    
    doc.add_paragraph('')
    doc.add_paragraph('')
    
    # 标题
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_run = title_p.add_run('明渠水力计算书')
    t_run.font.name = '华文中宋'
    t_run._element.rPr.rFonts.set(qn('w:eastAsia'), '华文中宋')
    t_run.font.size = Pt(34)
    t_run.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
    
    # 金色分割线
    gold_line = doc.add_paragraph()
    gold_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    gl_run = gold_line.add_run('━━━━━━━━━━━━━━━')
    gl_run.font.size = Pt(14)
    gl_run.font.color.rgb = RGBColor(0xC9, 0xA9, 0x6E)
    
    # 副标题
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub_p.add_run('梯形断面  ·  经济断面法')
    sr.font.name = '微软雅黑'
    sr._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    sr.font.size = Pt(16)
    sr.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)
    
    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('')
    
    # 封面信息
    cover_items = [
        ("工程名称", "XX灌区续建配套工程"),
        ("渠道名称", "总干渠 GK0+000~GK12+500"),
        ("设计单位", "XX水利水电勘测设计院"),
        ("计算日期", "2025年02月"),
    ]
    info_tbl = doc.add_table(rows=len(cover_items), cols=2)
    info_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(cover_items):
        c0 = info_tbl.cell(i, 0)
        c1 = info_tbl.cell(i, 1)
        c0.text = label
        c1.text = value
        c0.width = Cm(3)
        c1.width = Cm(8)
        for p in c0.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for r in p.runs:
                r.font.name = '微软雅黑'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                r.font.size = Pt(12)
                r.font.color.rgb = RGBColor(0xC9, 0xA9, 0x6E)
        for p in c1.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent = Cm(0.5)
            for r in p.runs:
                r.font.name = '微软雅黑'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                r.font.size = Pt(12)
                r.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
        for c in (c0, c1):
            set_cell_border(c,
                top=("0","FFFFFF","none"),
                left=("0","FFFFFF","none"),
                right=("0","FFFFFF","none"),
                bottom=("2", GOLD, "single"))
    
    doc.add_page_break()
    
    # ==================== 页眉页脚 ====================
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run('XX灌区续建配套工程')
    hr.font.name = '微软雅黑'
    hr._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    hr.font.size = Pt(8)
    hr.font.color.rgb = RGBColor(0xC9, 0xA9, 0x6E)
    set_paragraph_border_bottom(hp, color=GOLD, size=3)
    
    add_page_number(section, '微软雅黑', 9)
    
    # ==================== 正文 ====================
    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(10)
        r = p.add_run(text)
        r.font.name = '华文中宋'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '华文中宋')
        r.font.size = Pt(18)
        r.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
        set_paragraph_border_bottom(p, color=GOLD, size=6)
        return p
    
    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        r.font.name = '微软雅黑'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
        r.bold = True
        return p
    
    def add_body(text, indent=False):
        p = doc.add_paragraph()
        if indent:
            p.paragraph_format.left_indent = Cm(1.2)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text)
        r.font.name = '微软雅黑'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
        return p
    
    # --- 基本资料 ---
    add_h1('一、基本资料')
    
    # 参数表
    param_table = doc.add_table(rows=len(SAMPLE_PARAMS)+1, cols=2)
    param_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, hdr in enumerate(["参  数", "取  值"]):
        c = param_table.cell(0, j)
        c.text = hdr
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = '微软雅黑'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                r.font.size = Pt(10)
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(c, NAVY)
    
    for i, (name, val) in enumerate(SAMPLE_PARAMS):
        for j, txt in enumerate([name, val]):
            c = param_table.cell(i+1, j)
            c.text = txt
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.name = '微软雅黑'
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    r.font.size = Pt(10)
                    r.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
        if i % 2 == 0:
            for j in range(2):
                set_cell_shading(param_table.cell(i+1, j), WARM_GRAY)
    
    tbl = param_table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="8" w:space="0" w:color="{NAVY}"/>'
        f'  <w:bottom w:val="single" w:sz="8" w:space="0" w:color="{NAVY}"/>'
        f'  <w:insideH w:val="single" w:sz="2" w:space="0" w:color="E8E4DF"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="FFFFFF"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="FFFFFF"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="FFFFFF"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)
    
    doc.add_paragraph('')
    
    # --- 计算过程 ---
    add_h1('二、计算过程')
    add_h2('2.1  基础公式')
    
    formula_lines = [
        "曼宁公式：Q = (1/n) · A · R^(2/3) · i^(1/2)",
        "过水面积：A = (B + m·h) · h",
        "湿    周：χ = B + 2h√(1+m²)",
        "水力半径：R = A / χ",
    ]
    for fl in formula_lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.2)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(fl)
        r.font.name = 'Cambria Math'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    
    add_h2('2.2  计算步骤')
    
    for label, content in SAMPLE_STEPS:
        if label.startswith("  "):
            text = f"{label.strip()}：{content}" if content else label.strip()
            add_body(text, indent=True)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(label)
            r.font.name = '微软雅黑'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
            r.bold = True
            if content:
                add_body(content, indent=True)
    
    doc.add_paragraph('')
    
    # --- 方案对比 ---
    add_h1('三、断面方案对比')
    
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(6)
    cap_r = cap.add_run('表 1  附录E断面方案对比表')
    cap_r.font.name = '微软雅黑'
    cap_r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    cap_r.font.size = Pt(10)
    cap_r.italic = True
    cap_r.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)
    
    comp_table = doc.add_table(rows=len(SAMPLE_TABLE_DATA)+1, cols=len(SAMPLE_TABLE_HEADERS))
    comp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, hdr in enumerate(SAMPLE_TABLE_HEADERS):
        c = comp_table.cell(0, j)
        c.text = hdr
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = '微软雅黑'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                r.font.size = Pt(9)
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(c, NAVY)
    
    for i, row_data in enumerate(SAMPLE_TABLE_DATA):
        for j, val in enumerate(row_data):
            c = comp_table.cell(i+1, j)
            c.text = val
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.name = '微软雅黑'
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    r.font.size = Pt(9)
                    r.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
                    if val == '★选中':
                        r.font.color.rgb = RGBColor(0xC9, 0xA9, 0x6E)
                        r.bold = True
        if i % 2 == 0:
            for j in range(len(SAMPLE_TABLE_HEADERS)):
                set_cell_shading(comp_table.cell(i+1, j), WARM_GRAY)
    
    tbl2 = comp_table._tbl
    tblPr2 = tbl2.tblPr if tbl2.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders2 = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="8" w:space="0" w:color="{NAVY}"/>'
        f'  <w:bottom w:val="single" w:sz="8" w:space="0" w:color="{NAVY}"/>'
        f'  <w:insideH w:val="single" w:sz="2" w:space="0" w:color="E8E4DF"/>'
        f'  <w:insideV w:val="single" w:sz="2" w:space="0" w:color="E8E4DF"/>'
        f'  <w:left w:val="single" w:sz="2" w:space="0" w:color="E8E4DF"/>'
        f'  <w:right w:val="single" w:sz="2" w:space="0" w:color="E8E4DF"/>'
        f'</w:tblBorders>'
    )
    tblPr2.append(borders2)
    
    doc.add_paragraph('')
    
    # --- 计算结果 ---
    add_h1('四、计算结果')
    
    # 结果高亮卡片
    res_table = doc.add_table(rows=len(SAMPLE_RESULT)+1, cols=2)
    res_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for j, hdr in enumerate(["项  目", "计算结果"]):
        c = res_table.cell(0, j)
        c.text = hdr
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = '微软雅黑'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                r.font.size = Pt(10)
                r.bold = True
                r.font.color.rgb = RGBColor(0xC9, 0xA9, 0x6E)
        set_cell_shading(c, NAVY)
    
    for i, (name, val) in enumerate(SAMPLE_RESULT):
        c0 = res_table.cell(i+1, 0)
        c1 = res_table.cell(i+1, 1)
        c0.text = name
        c1.text = val
        for p in c0.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent = Cm(0.3)
            for r in p.runs:
                r.font.name = '微软雅黑'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                r.font.size = Pt(11)
                r.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)
        for p in c1.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = '微软雅黑'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                r.font.size = Pt(11)
                r.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
                r.bold = True
        if i % 2 == 0:
            for j in range(2):
                set_cell_shading(res_table.cell(i+1, j), WARM_GRAY)
    
    tbl3 = res_table._tbl
    tblPr3 = tbl3.tblPr if tbl3.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders3 = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="8" w:space="0" w:color="{NAVY}"/>'
        f'  <w:bottom w:val="single" w:sz="8" w:space="0" w:color="{NAVY}"/>'
        f'  <w:insideH w:val="single" w:sz="2" w:space="0" w:color="E8E4DF"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="FFFFFF"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="FFFFFF"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="FFFFFF"/>'
        f'</w:tblBorders>'
    )
    tblPr3.append(borders3)
    
    filepath = os.path.join(OUTPUT_DIR, '排版方案3_高端咨询报告风格.docx')
    doc.save(filepath)
    print(f"✓ 方案3已生成: {filepath}")
    return filepath


# ============================================================
# 主程序
# ============================================================
if __name__ == '__main__':
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    print("正在生成3种排版方案样例文档...\n")
    f1 = build_scheme1()
    f2 = build_scheme2()
    f3 = build_scheme3()
    print(f"\n全部完成！请打开以下文件预览：")
    print(f"  方案1（经典工程报告）: {f1}")
    print(f"  方案2（现代简约风格）: {f2}")
    print(f"  方案3（高端咨询报告）: {f3}")
