# -*- coding: utf-8 -*-
"""
导出工具 —— Word（含可编辑LaTeX公式）/ TXT 导出辅助函数
"""

import os
import re
import subprocess

# Word导出依赖（可选）
try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    import latex2mathml.converter
    from lxml import etree
    WORD_EXPORT_AVAILABLE = True
except ImportError:
    WORD_EXPORT_AVAILABLE = False


# ============================================================
# OMML转换（LaTeX → Word可编辑公式）
# ============================================================

def _find_mml2omml_xsl():
    """查找Office的MML2OMML.XSL文件"""
    candidates = [
        r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL",
        r"C:\Program Files (x86)\Microsoft Office\root\Office16\MML2OMML.XSL",
        r"C:\Program Files\Microsoft Office\Office16\MML2OMML.XSL",
        r"C:\Program Files\Microsoft Office\root\Office15\MML2OMML.XSL",
    ]
    for p in candidates:
        if os.path.exists(p):
            return p
    return None


_MML2OMML_XSL_PATH = _find_mml2omml_xsl() if WORD_EXPORT_AVAILABLE else None
_MML2OMML_TRANSFORM = None


def _get_omml_transform():
    """延迟加载XSLT转换器"""
    global _MML2OMML_TRANSFORM
    if _MML2OMML_TRANSFORM is None and _MML2OMML_XSL_PATH:
        xslt_tree = etree.parse(_MML2OMML_XSL_PATH)
        _MML2OMML_TRANSFORM = etree.XSLT(xslt_tree)
    return _MML2OMML_TRANSFORM


def latex_to_omml(latex_str):
    """将LaTeX公式转换为Word可编辑的OMML元素"""
    transform = _get_omml_transform()
    if transform is None:
        return None
    try:
        mathml_str = latex2mathml.converter.convert(latex_str)
        mathml_tree = etree.fromstring(mathml_str.encode('utf-8'))
        omml_result = transform(mathml_tree)
        return omml_result.getroot()
    except Exception:
        return None


def add_formula_to_doc(doc, latex_str, prefix=""):
    """向Word文档添加一个可编辑公式段落"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if prefix:
        run = p.add_run(prefix)
        run.font.name = 'Microsoft YaHei'
        run.font.size = Pt(11)
    omml = latex_to_omml(latex_str)
    if omml is not None:
        p._element.append(omml)
    else:
        run = p.add_run(f" {latex_str}")
        run.font.name = 'Cambria Math'
        run.font.size = Pt(11)
    return p


# ============================================================
# 公式行转换辅助
# ============================================================

def try_convert_formula_line(line):
    """尝试将计算行转换为LaTeX公式字符串，不适合则返回None"""
    if not any(c in line for c in ['=', '×', '√', '^']):
        return None
    cn_count = sum(1 for c in line if '\u4e00' <= c <= '\u9fff')
    if cn_count > 4:
        return None
    if '✓' in line or '✗' in line or '通过' in line or '未通过' in line:
        return None
    m = re.match(r'^([A-Za-zα-ωΑ-Ωβχ\u4e00-\u9fff]+\s*=\s*.+)', line)
    if not m:
        m = re.match(r'^(=\s*.+)', line)
    if not m:
        return None
    expr = line.strip()
    latex = expr
    latex = latex.replace('×', r' \times ')
    latex = latex.replace('√', r'\sqrt')
    latex = latex.replace('π', r'\pi')
    latex = latex.replace('χ', r'\chi')
    latex = latex.replace('β', r'\beta')
    latex = latex.replace('θ', r'\theta')
    latex = re.sub(r'\^[\(（]([^)）]+)[\)）]', r'^{\1}', latex)
    latex = latex.replace('²', '^{2}')
    latex = latex.replace('³', '^{3}')
    latex = re.sub(r'm\^?\{?3\}?/s', r'\\text{ m}^3\\text{/s}', latex)
    latex = re.sub(r'\s+m$', r' \\text{ m}', latex)
    latex = re.sub(r'\s+m²$', r' \\text{ m}^2', latex)
    return latex


# ============================================================
# 导出后询问并打开文件
# ============================================================

def ask_open_file(filepath, parent=None):
    """导出成功后询问用户是否用默认程序打开文件"""
    from app_渠系计算前端.styles import fluent_question
    if fluent_question(
        parent, "导出成功",
        f"文件已保存到:\n{filepath}\n\n是否立即打开该文件？",
        yes_text="打开", no_text="关闭"
    ):
        try:
            os.startfile(filepath)
        except Exception:
            try:
                subprocess.Popen(['start', '', filepath], shell=True)
            except Exception:
                pass


# ============================================================
# 高端咨询报告风格排版工具（方案3）
# ============================================================

# --- 配色常量 ---
WORD_NAVY = "1B2A4A"
WORD_GOLD = "C9A96E"
WORD_WARM_GRAY = "F8F6F3"
WORD_TEXT_DARK = "2C3E50"
WORD_TEXT_MID = "5D6D7E"

_NAVY_RGB = RGBColor(0x1B, 0x2A, 0x4A) if WORD_EXPORT_AVAILABLE else None
_GOLD_RGB = RGBColor(0xC9, 0xA9, 0x6E) if WORD_EXPORT_AVAILABLE else None
_TEXT_DARK_RGB = RGBColor(0x2C, 0x3E, 0x50) if WORD_EXPORT_AVAILABLE else None
_TEXT_MID_RGB = RGBColor(0x5D, 0x6D, 0x7E) if WORD_EXPORT_AVAILABLE else None
_WHITE_RGB = RGBColor(0xFF, 0xFF, 0xFF) if WORD_EXPORT_AVAILABLE else None


def _set_cell_shading(cell, color_hex):
    """设置单元格底色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_border(cell, **kwargs):
    """设置单元格边框 kwargs: top/bottom/left/right = (size, color, style)"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
    for edge, (sz, color, style) in kwargs.items():
        el = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{style}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        )
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _set_paragraph_border_bottom(paragraph, color="000000", size=6):
    """段落底部边框线"""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:space="1" w:color="{color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def _set_table_borders(table, top_color, bottom_color, inside_h_color,
                       inside_v_color=None, left_color=None, right_color=None,
                       top_sz="8", bottom_sz="8", inside_h_sz="2"):
    """设置表格边框样式"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    iv = f'<w:insideV w:val="single" w:sz="2" w:space="0" w:color="{inside_v_color}"/>' if inside_v_color else \
         f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="FFFFFF"/>'
    lv = f'<w:left w:val="single" w:sz="2" w:space="0" w:color="{left_color}"/>' if left_color else \
         f'<w:left w:val="none" w:sz="0" w:space="0" w:color="FFFFFF"/>'
    rv = f'<w:right w:val="single" w:sz="2" w:space="0" w:color="{right_color}"/>' if right_color else \
         f'<w:right w:val="none" w:sz="0" w:space="0" w:color="FFFFFF"/>'
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{top_sz}" w:space="0" w:color="{top_color}"/>'
        f'  <w:bottom w:val="single" w:sz="{bottom_sz}" w:space="0" w:color="{bottom_color}"/>'
        f'  <w:insideH w:val="single" w:sz="{inside_h_sz}" w:space="0" w:color="{inside_h_color}"/>'
        f'  {iv}{lv}{rv}'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def _add_page_number(section, font_name='微软雅黑', font_size=9):
    """添加页脚页码"""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p.add_run("— ")
    run1.font.name = font_name
    run1.font.size = Pt(font_size)
    fld_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run_f1 = p.add_run()
    run_f1._r.append(fld_begin)
    instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run_f2 = p.add_run()
    run_f2._r.append(instr)
    fld_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run_f3 = p.add_run()
    run_f3._r.append(fld_end)
    run2 = p.add_run(" —")
    run2.font.name = font_name
    run2.font.size = Pt(font_size)


def _run_styled(paragraph, text, font_name='微软雅黑', font_size=11,
                color=None, bold=False, italic=False):
    """向段落添加一个带样式的run"""
    r = paragraph.add_run(text)
    r.font.name = font_name
    r._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    r.font.size = Pt(font_size)
    if color:
        r.font.color.rgb = color
    if bold:
        r.bold = True
    if italic:
        r.italic = True
    return r


_LOGO_PNG_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'resources', 'logo.png')
_APP_NAME = "渠系建筑物水力计算系统"
_APP_VERSION = "V1.0"
_DEFAULT_DESIGN_UNIT = "四川水发勘测设计研究有限公司"


def create_styled_doc(title, subtitle="", header_text="",
                      channel_name="", channel_level="",
                      start_station="", end_station=""):
    """创建一个方案3风格的Word文档（含封面、页眉页脚），返回doc对象
    
    Args:
        title: 封面大标题，如 "明渠水力计算书"
        subtitle: 封面副标题，如 "梯形断面 · 经济断面法"
        header_text: 页眉文字
        channel_name: 渠道名称（从用户输入读取）
        channel_level: 渠道级别（从用户输入读取）
        start_station: 起始桩号
        end_station: 终止桩号
    """
    import datetime
    doc = DocxDocument()
    
    # --- 页面设置 ---
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)
    
    # --- 默认样式 ---
    style_normal = doc.styles['Normal']
    style_normal.font.name = '微软雅黑'
    style_normal.font.size = Pt(11)
    style_normal._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    # 必须使用显式比例行距（非固定磅值），否则大字号标题/Logo/图片会被裁切
    style_normal.paragraph_format.line_spacing = 1.15
    
    # ==================== 封面 ====================
    # 上方留白
    for _ in range(2):
        doc.add_paragraph('')
    
    # 软件Logo
    if os.path.exists(_LOGO_PNG_PATH):
        logo_p = doc.add_paragraph()
        logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_p.paragraph_format.line_spacing = 1.0
        logo_p.add_run().add_picture(_LOGO_PNG_PATH, width=Cm(2.2))
    
    # 软件名称
    app_p = doc.add_paragraph()
    app_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    app_p.paragraph_format.space_before = Pt(6)
    app_p.paragraph_format.space_after = Pt(2)
    _run_styled(app_p, _APP_NAME, '微软雅黑', 11, _TEXT_MID_RGB)
    
    # 版本号
    ver_p = doc.add_paragraph()
    ver_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ver_p.paragraph_format.space_before = Pt(0)
    ver_p.paragraph_format.space_after = Pt(16)
    _run_styled(ver_p, _APP_VERSION, '微软雅黑', 9, _GOLD_RGB)
    
    # 金色分割线（上）
    line1 = doc.add_paragraph()
    line1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line1.paragraph_format.space_after = Pt(16)
    _set_paragraph_border_bottom(line1, color=WORD_GOLD, size=4)
    
    # 大标题
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(8)
    title_p.paragraph_format.space_after = Pt(8)
    title_p.paragraph_format.line_spacing = 1.0
    _run_styled(title_p, title, '黑体', 32, _NAVY_RGB, bold=True)
    
    # 副标题
    if subtitle:
        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_p.paragraph_format.space_before = Pt(4)
        sub_p.paragraph_format.space_after = Pt(4)
        _run_styled(sub_p, subtitle, '微软雅黑', 14, _TEXT_MID_RGB)
    
    # 金色分割线（下）
    line2 = doc.add_paragraph()
    line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line2.paragraph_format.space_before = Pt(16)
    _set_paragraph_border_bottom(line2, color=WORD_GOLD, size=4)
    
    # 留白
    for _ in range(3):
        doc.add_paragraph('')
    
    # 封面信息表
    cover_items = []
    cover_items.append(("设计单位", _DEFAULT_DESIGN_UNIT))
    if channel_name:
        cover_items.append(("渠道名称", channel_name))
    if channel_level:
        cover_items.append(("渠道级别", channel_level))
    if start_station or end_station:
        station_text = ""
        if start_station and end_station:
            station_text = f"{start_station} ~ {end_station}"
        elif start_station:
            station_text = f"{start_station} ~"
        else:
            station_text = f"~ {end_station}"
        cover_items.append(("桩  号", station_text))
    cover_items.append(("计算日期", datetime.datetime.now().strftime("%Y年%m月")))
    
    _add_cover_info_table(doc, cover_items)
    
    # ==================== 页眉页脚 ====================
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if header_text:
        _run_styled(hp, header_text, '微软雅黑', 8, _GOLD_RGB)
    _set_paragraph_border_bottom(hp, color=WORD_GOLD, size=3)
    
    _add_page_number(section)
    
    return doc


def _add_cover_info_table(doc, items):
    """在封面添加信息表（内部方法）"""
    info_tbl = doc.add_table(rows=len(items), cols=2)
    info_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(items):
        c0 = info_tbl.cell(i, 0)
        c1 = info_tbl.cell(i, 1)
        c0.text = label
        c1.text = value
        c0.width = Cm(3)
        c1.width = Cm(8.5)
        for p in c0.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for r in p.runs:
                r.font.name = '微软雅黑'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                r.font.size = Pt(12)
                r.font.color.rgb = _GOLD_RGB
        for p in c1.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent = Cm(0.5)
            for r in p.runs:
                r.font.name = '微软雅黑'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                r.font.size = Pt(12)
                r.font.color.rgb = _TEXT_DARK_RGB
        for c in (c0, c1):
            _set_cell_border(c,
                top=("0","FFFFFF","none"), left=("0","FFFFFF","none"),
                right=("0","FFFFFF","none"), bottom=("2", WORD_GOLD, "single"))


def doc_add_h1(doc, text):
    """添加一级标题（黑体 + 金色底线）"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.0
    _run_styled(p, text, '黑体', 18, _NAVY_RGB, bold=True)
    _set_paragraph_border_bottom(p, color=WORD_GOLD, size=6)
    return p


def doc_add_h2(doc, text):
    """添加二级标题（微软雅黑加粗 深藏青）"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    _run_styled(p, text, '微软雅黑', 13, _NAVY_RGB, bold=True)
    return p


def doc_add_h3(doc, text):
    """添加三级标题/步骤标题（微软雅黑加粗 12pt）"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    _run_styled(p, text, '微软雅黑', 12, _NAVY_RGB, bold=True)
    return p


def doc_add_body(doc, text, indent=False):
    """添加正文段落"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(1.2)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    _run_styled(p, text, '微软雅黑', 11, _TEXT_DARK_RGB)
    return p


def doc_add_formula(doc, latex_str, prefix=""):
    """添加公式段落（方案3风格缩进 + OMML可编辑公式）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Cm(1.2)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    if prefix:
        _run_styled(p, prefix, '微软雅黑', 11, _TEXT_DARK_RGB)
    omml = latex_to_omml(latex_str)
    if omml is not None:
        p._element.append(omml)
    else:
        _run_styled(p, f" {latex_str}", 'Cambria Math', 11, _TEXT_DARK_RGB)
    return p


def doc_add_styled_table(doc, headers, data, highlight_col=None, highlight_val=None,
                         with_full_border=False):
    """添加方案3风格的数据表格
    
    Args:
        headers: 列标题列表
        data: 二维数据列表
        highlight_col: 高亮匹配的列索引
        highlight_val: 高亮匹配的值
        with_full_border: 是否显示完整网格线
    """
    table = doc.add_table(rows=len(data)+1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 表头
    for j, hdr in enumerate(headers):
        c = table.cell(0, j)
        c.text = hdr
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = '微软雅黑'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                r.font.size = Pt(9)
                r.bold = True
                r.font.color.rgb = _WHITE_RGB
        _set_cell_shading(c, WORD_NAVY)
    
    # 数据行
    for i, row_data in enumerate(data):
        for j, val in enumerate(row_data):
            c = table.cell(i+1, j)
            c.text = str(val)
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.name = '微软雅黑'
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    r.font.size = Pt(9)
                    r.font.color.rgb = _TEXT_DARK_RGB
                    if highlight_col is not None and j == highlight_col and highlight_val and val == highlight_val:
                        r.font.color.rgb = _GOLD_RGB
                        r.bold = True
        if i % 2 == 0:
            for j in range(len(headers)):
                _set_cell_shading(table.cell(i+1, j), WORD_WARM_GRAY)
    
    # 边框
    iv_color = "E8E4DF" if with_full_border else None
    lv_color = "E8E4DF" if with_full_border else None
    rv_color = "E8E4DF" if with_full_border else None
    _set_table_borders(table, WORD_NAVY, WORD_NAVY, "E8E4DF",
                       inside_v_color=iv_color, left_color=lv_color, right_color=rv_color)
    return table


def doc_add_result_table(doc, items):
    """添加计算结果汇总表（项目+结果 两列，金色表头）
    
    Args:
        items: list of (name, value) 元组
    """
    table = doc.add_table(rows=len(items)+1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, hdr in enumerate(["项  目", "计算结果"]):
        c = table.cell(0, j)
        c.text = hdr
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = '微软雅黑'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                r.font.size = Pt(10)
                r.bold = True
                r.font.color.rgb = _GOLD_RGB
        _set_cell_shading(c, WORD_NAVY)
    
    for i, (name, val) in enumerate(items):
        c0 = table.cell(i+1, 0)
        c1 = table.cell(i+1, 1)
        c0.text = name
        c1.text = val
        for p in c0.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent = Cm(0.3)
            for r in p.runs:
                r.font.name = '微软雅黑'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                r.font.size = Pt(11)
                r.font.color.rgb = _TEXT_MID_RGB
        for p in c1.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = '微软雅黑'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                r.font.size = Pt(11)
                r.font.color.rgb = _NAVY_RGB
                r.bold = True
        if i % 2 == 0:
            for j in range(2):
                _set_cell_shading(table.cell(i+1, j), WORD_WARM_GRAY)
    
    _set_table_borders(table, WORD_NAVY, WORD_NAVY, "E8E4DF")
    return table


def doc_add_param_table(doc, params):
    """添加基本参数表（参数+取值 两列）
    
    Args:
        params: list of (name, value) 元组
    """
    table = doc.add_table(rows=len(params)+1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, hdr in enumerate(["参  数", "取  值"]):
        c = table.cell(0, j)
        c.text = hdr
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = '微软雅黑'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                r.font.size = Pt(10)
                r.bold = True
                r.font.color.rgb = _WHITE_RGB
        _set_cell_shading(c, WORD_NAVY)
    
    for i, (name, val) in enumerate(params):
        for j, txt in enumerate([name, val]):
            c = table.cell(i+1, j)
            c.text = txt
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.name = '微软雅黑'
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    r.font.size = Pt(10)
                    r.font.color.rgb = _TEXT_DARK_RGB
        if i % 2 == 0:
            for j in range(2):
                _set_cell_shading(table.cell(i+1, j), WORD_WARM_GRAY)
    
    _set_table_borders(table, WORD_NAVY, WORD_NAVY, "E8E4DF")
    return table


def doc_add_table_caption(doc, text):
    """添加表标题（居中、斜体、灰色）"""
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(6)
    _run_styled(cap, text, '微软雅黑', 10, _TEXT_MID_RGB, italic=True)
    return cap


def doc_add_figure(doc, image_path, width_cm=14):
    """添加居中图片（独立段落，显式比例行距避免裁切）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run()
    run.add_picture(image_path, width=Cm(width_cm))
    return p


def doc_render_calc_text(doc, text, skip_title_keyword=""):
    """将纯文本计算过程渲染为方案3风格的Word内容
    
    自动识别：
    - 【xxx】 → 二级标题
    - 数字. xxx → 三级标题
    - 公式行 → OMML公式
    - 普通行 → 正文段落
    
    Args:
        text: 纯文本计算过程
        skip_title_keyword: 跳过包含此关键词的行（如"水力计算结果"）
    """
    if not text:
        return
    for line in text.split('\n'):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph('')
            continue
        if set(stripped) <= {'=', '-', ' '} and len(stripped) > 5:
            continue
        if skip_title_keyword and skip_title_keyword in stripped:
            continue
        if stripped.startswith('{{') and stripped.endswith('}}'):
            continue
        if stripped.startswith('【') and '】' in stripped:
            doc_add_h2(doc, stripped.lstrip('【').rstrip('】'))
            continue
        step_m = re.match(r'^(\d+)\.\s+(.+)', stripped)
        if step_m:
            doc_add_h3(doc, f'{step_m.group(1)}. {step_m.group(2)}')
            continue
        fl = try_convert_formula_line(stripped)
        if fl:
            doc_add_formula(doc, fl, '    ')
            continue
        doc_add_body(doc, stripped)


# ============================================================
# 工程产品运行卡格式 —— 辅助函数
# ============================================================

def _set_section_margins(section, top=2.5, bottom=2.5, left=3.2, right=3.2):
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)


def _run_song(para, text, size_pt=12, bold=False, color=None):
    """向段落添加宋体 run"""
    r = para.add_run(text)
    r.font.name = '宋体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    r._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    r.font.size = Pt(size_pt)
    if bold:
        r.bold = True
    if color:
        r.font.color.rgb = color
    return r


def _cell_text_song(cell, text, size_pt=10, bold=False,
                    align=None, v_align=None):
    """清空单元格并写入宋体文字"""
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    _run_song(p, text, size_pt, bold)
    if v_align is not None:
        cell.vertical_alignment = v_align


def _add_page1_running_card(doc, meta, calc_title, calc_content_desc):
    """Page 1: 工程阶段产品运行卡（表023格式，L=R=1.9cm）"""
    section = doc.sections[0]
    _set_section_margins(section, top=2.5, bottom=2.5, left=1.9, right=1.9)

    proj = meta.project_name.strip() or "工程名称"

    p0 = doc.add_paragraph()
    p0.paragraph_format.space_before = Pt(0)
    p0.paragraph_format.space_after = Pt(2)
    _run_song(p0, '表023  工程阶段产品运行卡', 10)

    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(6)
    p1.paragraph_format.space_after = Pt(6)
    _run_song(p1, f'{proj}  {meta.design_stage}  阶段  计算  产品运行卡', 14)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(2)
    _run_song(p2, f'产品级别：  {meta.product_level}                                              记录编号：{meta.record_number}', 10)

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after = Pt(2)
    _run_song(p3, ('注：1．产品质量评定结论以校审最高一级的评定结果为准；质量评定等级分优良、合格、'
                   '不合格（评定标准详见《不合格品控制程序》附录：产品质量合格评定准则）。'
                   '2．一个产品由多人计算时，应分别填写在同一张运行卡上。'), 9)

    vol_cur = meta.volume_current.strip() or ' '
    vol_tot = meta.volume_total.strip() or ' '
    p_vol = doc.add_paragraph()
    p_vol.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_vol.paragraph_format.space_before = Pt(2)
    p_vol.paragraph_format.space_after = Pt(4)
    _run_song(p_vol, f'第  {vol_cur}  册    共  {vol_tot}  册', 13)

    # 14行×5列大表
    tbl = doc.add_table(rows=14, cols=5)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    col_widths = [Cm(2.0), Cm(7.5), Cm(2.0), Cm(3.2), Cm(2.5)]
    for row in tbl.rows:
        for ci, w in enumerate(col_widths):
            row.cells[ci].width = w

    # 垂直合并（列0）
    tbl.cell(0, 0).merge(tbl.cell(1, 0))
    tbl.cell(2, 0).merge(tbl.cell(3, 0))
    tbl.cell(4, 0).merge(tbl.cell(6, 0))
    tbl.cell(7, 0).merge(tbl.cell(9, 0))
    tbl.cell(10, 0).merge(tbl.cell(12, 0))
    # 列1 行0-1 垂直合并
    tbl.cell(0, 1).merge(tbl.cell(1, 1))
    # 各行 列1-4 水平合并
    for ri in [2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12, 13]:
        tbl.cell(ri, 1).merge(tbl.cell(ri, 4))

    designer = meta.calculator or ' '
    _cell_text_song(tbl.cell(0, 0), '产品\n名称', 10, bold=True, v_align=WD_ALIGN_VERTICAL.CENTER)
    _cell_text_song(tbl.cell(0, 1), calc_title, 10, align=WD_ALIGN_PARAGRAPH.LEFT, v_align=WD_ALIGN_VERTICAL.CENTER)
    _cell_text_song(tbl.cell(0, 2), '图号或\n页数', 9, v_align=WD_ALIGN_VERTICAL.CENTER)
    _cell_text_song(tbl.cell(0, 3), '电子文件档案号', 9, v_align=WD_ALIGN_VERTICAL.CENTER)
    _cell_text_song(tbl.cell(0, 4), '设计（作业）人', 9, v_align=WD_ALIGN_VERTICAL.CENTER)
    _cell_text_song(tbl.cell(1, 2), ' ', 10)
    _cell_text_song(tbl.cell(1, 3), ' ', 10)
    _cell_text_song(tbl.cell(1, 4), designer, 10, v_align=WD_ALIGN_VERTICAL.CENTER)
    _cell_text_song(tbl.cell(2, 0), '产品\n内容', 10, bold=True, v_align=WD_ALIGN_VERTICAL.CENTER)
    _cell_text_song(tbl.cell(2, 1), calc_content_desc, 10, align=WD_ALIGN_PARAGRAPH.LEFT, v_align=WD_ALIGN_VERTICAL.CENTER)
    clause = meta.mandatory_clause.strip() or '无'
    _cell_text_song(tbl.cell(3, 1), f'产品采用的强标条款：{clause}', 10, align=WD_ALIGN_PARAGRAPH.LEFT)

    def _fill_opinion(r0, label):
        _cell_text_song(tbl.cell(r0, 0), label, 10, bold=True, v_align=WD_ALIGN_VERTICAL.CENTER)
        short = label.replace('\n', '')
        _cell_text_song(tbl.cell(r0, 1),
                        f'已{short}：□所有数据和公式；□图、表与计算一致；□产品格式符合院要求。\n其它较大意见：',
                        9, align=WD_ALIGN_PARAGRAPH.LEFT)
        _cell_text_song(tbl.cell(r0 + 1, 1), f'产品强标{short}意见：', 9, align=WD_ALIGN_PARAGRAPH.LEFT)
        _cell_text_song(tbl.cell(r0 + 2, 1),
                        '收件时间：            产品质量等级：            签名：            出手日期：',
                        9, align=WD_ALIGN_PARAGRAPH.LEFT)

    _fill_opinion(4, '校\n核\n意\n见')
    _fill_opinion(7, '审\n查\n意\n见')
    _fill_opinion(10, '审\n定\n意\n见')
    _cell_text_song(tbl.cell(13, 0), '批准\n意见', 10, bold=True, v_align=WD_ALIGN_VERTICAL.CENTER)
    _cell_text_song(tbl.cell(13, 1), '\n签名：                                    出手日期：', 10, align=WD_ALIGN_PARAGRAPH.LEFT)

    # 行高设置
    rh_list = [Cm(0.9), Cm(0.9), Cm(0.8), Cm(0.8),
               Cm(1.1), Cm(0.6), Cm(0.7),
               Cm(1.1), Cm(0.6), Cm(0.7),
               Cm(1.1), Cm(0.6), Cm(0.7),
               Cm(1.2)]
    for ri, rh in enumerate(rh_list):
        if ri < len(tbl.rows):
            tr = tbl.rows[ri]._tr
            trPr = tr.get_or_add_trPr()
            trHeight = parse_xml(
                f'<w:trHeight {nsdecls("w")} w:val="{int(rh.emu / 12700)}" w:hRule="atLeast"/>'
            )
            trPr.append(trHeight)


def _add_page2_cover(doc, meta, calc_title):
    """Page 2: 封面（新 section，L=R=3.2cm）"""
    import datetime
    new_sec = doc.add_section()
    _set_section_margins(new_sec, top=2.5, bottom=2.5, left=3.2, right=3.2)

    proj = meta.project_name.strip() or "工程名称"
    for _ in range(2):
        doc.add_paragraph('')

    p_proj = doc.add_paragraph()
    p_proj.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_proj.paragraph_format.space_after = Pt(8)
    _run_song(p_proj, f'{proj}  {meta.design_stage}  设计阶段', 18)

    doc.add_paragraph('')
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(6)
    p_title.paragraph_format.space_after = Pt(4)
    _run_song(p_title, calc_title, 22, bold=True)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_before = Pt(2)
    p_sub.paragraph_format.space_after = Pt(2)
    _run_song(p_sub, '计  算  稿', 16)

    for _ in range(4):
        doc.add_paragraph('')

    def _person_row(label, name):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        _run_song(p, f'           {label:<10}  {name}', 14)

    _person_row('专业名称', f'    {meta.specialty or "水工"}')
    _person_row('审    定', f'    {meta.approver}')
    _person_row('审    查', f'    {meta.reviewer}')
    _person_row('校    核', f'    {meta.checker}')
    _person_row('计    算', f'    {meta.calculator}')

    for _ in range(4):
        doc.add_paragraph('')

    from app_渠系计算前端.report_meta import DESIGN_UNIT
    p_unit = doc.add_paragraph()
    p_unit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_unit.paragraph_format.space_before = Pt(4)
    _run_song(p_unit, DESIGN_UNIT, 16, bold=True)

    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run_song(p_date, datetime.datetime.now().strftime('%Y年%m月'), 14)


def _add_page3_mandatory_standards(doc, meta, calc_title, calc_content_desc):
    """Page 3: 强制性标准条文执行情况校审检查表"""
    doc.add_page_break()
    p_ttl = doc.add_paragraph()
    p_ttl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ttl.paragraph_format.space_before = Pt(0)
    p_ttl.paragraph_format.space_after = Pt(6)
    _run_song(p_ttl, '强制性标准条文执行情况校审检查表', 14, bold=True)

    tbl = doc.add_table(rows=6, cols=6)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    col_w3 = [Cm(2.0), Cm(5.0), Cm(1.5), Cm(2.5), Cm(1.5), Cm(1.5)]
    for row in tbl.rows:
        for ci, w in enumerate(col_w3):
            row.cells[ci].width = w

    tbl.cell(0, 0).merge(tbl.cell(0, 5))
    _cell_text_song(tbl.cell(0, 0), '强制性标准条文执行情况校审检查表', 12, bold=True)

    proj = meta.project_name.strip() or ' '
    _cell_text_song(tbl.cell(1, 0), '工程名称', 10, bold=True)
    _cell_text_song(tbl.cell(1, 1), proj, 10, align=WD_ALIGN_PARAGRAPH.LEFT)
    _cell_text_song(tbl.cell(1, 2), '设计阶段', 10, bold=True)
    tbl.cell(1, 3).merge(tbl.cell(1, 5))
    _cell_text_song(tbl.cell(1, 3), meta.design_stage, 10)

    _cell_text_song(tbl.cell(2, 0), '产品名称', 10, bold=True)
    _cell_text_song(tbl.cell(2, 1), calc_title, 10, align=WD_ALIGN_PARAGRAPH.LEFT)
    _cell_text_song(tbl.cell(2, 2), '产品内容', 10, bold=True)
    tbl.cell(2, 3).merge(tbl.cell(2, 5))
    _cell_text_song(tbl.cell(2, 3), calc_content_desc, 10, align=WD_ALIGN_PARAGRAPH.LEFT)

    _cell_text_song(tbl.cell(3, 0), '标准名称\n及编号', 9, bold=True)
    tbl.cell(3, 1).merge(tbl.cell(4, 1))
    _cell_text_song(tbl.cell(3, 1), '强制性标准条款号、条文内容及执行情况', 9, bold=True)
    tbl.cell(3, 2).merge(tbl.cell(3, 5))
    _cell_text_song(tbl.cell(3, 2), '校审意见（符合的画"√"、不符合的画"×"）', 9, bold=True)

    _cell_text_song(tbl.cell(4, 0), '标准名称\n及编号', 9, bold=True)
    _cell_text_song(tbl.cell(4, 2), '校核', 9, bold=True)
    _cell_text_song(tbl.cell(4, 3), '审查', 9, bold=True)
    tbl.cell(4, 4).merge(tbl.cell(4, 5))
    _cell_text_song(tbl.cell(4, 4), '审定', 9, bold=True)

    clause = meta.mandatory_clause.strip() or '无'
    _cell_text_song(tbl.cell(5, 0), clause, 10, align=WD_ALIGN_PARAGRAPH.LEFT)
    _cell_text_song(tbl.cell(5, 1), clause, 10, align=WD_ALIGN_PARAGRAPH.LEFT)
    _cell_text_song(tbl.cell(5, 2), '', 10)
    _cell_text_song(tbl.cell(5, 3), '', 10)
    tbl.cell(5, 4).merge(tbl.cell(5, 5))
    _cell_text_song(tbl.cell(5, 4), '', 10)


def _add_page4_toc(doc):
    """Page 4: 目录"""
    doc.add_page_break()
    p_ttl = doc.add_paragraph()
    p_ttl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ttl.paragraph_format.space_before = Pt(12)
    p_ttl.paragraph_format.space_after = Pt(20)
    _run_song(p_ttl, '目  录', 18, bold=True)
    for item in ['1、计算目的', '2、计算依据', '3、基本资料', '4、计算程序', '5、计算内容']:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(2)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        _run_song(p, item, 14)


def _add_page5_calc_intro(doc, meta, calc_purpose, references, calc_program_text):
    """Page 5: 计算目的/依据/基本资料/计算程序"""
    doc.add_page_break()

    def _h4(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        _run_song(p, text, 14, bold=True)

    def _body(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = Pt(20)
        _run_song(p, text, 12)

    _h4('1、计算目的')
    _body(calc_purpose or '（计算目的未填写）')

    _h4('2、计算依据')
    for i, ref in enumerate(references, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = Pt(20)
        _run_song(p, f'{i}、{ref}', 12)

    _h4('3、基本资料')
    basic = meta.basic_info.strip() if meta.basic_info.strip() else '（基本资料未填写，请在"项目设置"中补充）'
    for seg in basic.split('\n'):
        seg = seg.strip()
        if seg:
            _body(seg)

    _h4('4、计算程序')
    _body(calc_program_text or '渠系建筑物水力计算系统 V1.0')


def _add_eng_header_footer(doc, meta, calc_title):
    """为 Section 1（封面及之后）添加简洁页眉页脚"""
    # 找到 Section 1（index 1）
    if len(doc.sections) < 2:
        return
    sec1 = doc.sections[1]
    sec1.header.is_linked_to_previous = False
    hp = sec1.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    proj = meta.project_name.strip() if meta.project_name.strip() else ""
    header_str = f'{proj}  {calc_title}' if proj else calc_title
    _run_song(hp, header_str, 9)
    _set_paragraph_border_bottom(hp, color="000000", size=3)
    _add_page_number(sec1, font_name='宋体', font_size=9)


_TEMPLATE_PATH = os.path.normpath(
    os.path.join(os.path.dirname(__file__), '..', 'data', 'xxxxx计算稿（计算产品运行卡）.docx')
)


def _fill_template_cell(cell, text, size_pt=10, left=False):
    """在模板单元格中清空内容并写入新文本，保留单元格原有格式结构。"""
    p = cell.paragraphs[0]
    for r in p.runs:
        r.text = ''
    if p.runs:
        p.runs[0].text = text
        p.runs[0].font.size = Pt(size_pt)
    else:
        r = p.add_run(text)
        r.font.size = Pt(size_pt)
    if left:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def create_engineering_report_doc_from_template(meta, calc_title, calc_content_desc,
                                                calc_purpose, references,
                                                calc_program_text=None):
    """直接使用 xxxxx计算稿 模板，替换占位文本后追加计算内容，原汁原味保留排版。"""
    import io
    import datetime

    if not os.path.isfile(_TEMPLATE_PATH):
        return None  # 模板不存在，调用方回退到手工生成

    buf = io.BytesIO()
    with open(_TEMPLATE_PATH, 'rb') as f:
        buf.write(f.read())
    buf.seek(0)
    doc = DocxDocument(buf)

    proj  = meta.project_name.strip() or '工程名称'
    stage = meta.design_stage.strip()  or '施工详图设计'
    now   = datetime.datetime.now()
    paras = doc.paragraphs  # 总共37段（0-36）

    # ------------------------------------------------------------------
    # Para[1] — 标题行  "工程   阶段 计算 产品运行卡"
    # run[3]='工程'  run[4,5,6]=空格  run[7]='阶段'
    # ------------------------------------------------------------------
    p1 = paras[1]
    p1.runs[3].text = proj          # 工程名称（无下划线）
    p1.runs[4].text = ' '           # 带下划线的间距
    p1.runs[5].text = stage         # 阶段文字放入带下划线的 run，呈现下划线效果
    p1.runs[6].text = ' '           # 带下划线的间距
    p1.runs[7].text = '  阶段'      # 静态标签"阶段"（无下划线）

    # ------------------------------------------------------------------
    # Para[2] — "产品级别：   级  ...  记录编号："
    # run[1,2,3]=产品级别前空格(带下划线)  run[4]='级 '(静态文字)
    # 只写"三"/"二"/"一"，run[4]的"级"保留不动
    # ------------------------------------------------------------------
    p2 = paras[2]
    level_prefix = meta.product_level.rstrip('级') or meta.product_level
    p2.runs[1].text = f' {level_prefix}'
    p2.runs[2].text = ''
    p2.runs[3].text = ''
    p2.runs[7].text = f' 记录编号：{meta.record_number or ""}'

    # Para[6] / Para[7] — 第/共册
    paras[6].runs[0].text = f'第  {meta.volume_current or ""}  册   '
    paras[7].runs[0].text = f'共  {meta.volume_total  or ""}  册   '

    # ------------------------------------------------------------------
    # Para[10] — 封面项目名称 + 设计阶段
    # run[2]='工程 '(bold)  run[6]='设计阶段'(bold)
    # ------------------------------------------------------------------
    paras[10].runs[2].text = f'{proj}  '
    paras[10].runs[6].text = f'{stage}  设计阶段'

    # ------------------------------------------------------------------
    # Para[12] — 计算书标题（红框占位文字），覆盖全部 runs
    # ------------------------------------------------------------------
    for r in paras[12].runs:
        r.text = ''
    paras[12].runs[0].text = calc_title

    # Para[17-20] — 专业/审查/校核/计算人员
    # 用 ljust 补足原始长度，保持下划线宽度与模板一致
    paras[17].runs[1].text = f' {meta.specialty or "水工"} '.ljust(18)
    paras[18].runs[1].text = f'  {meta.reviewer   or ""}'.ljust(20)
    paras[19].runs[1].text = f'  {meta.checker    or ""}'.ljust(20)
    paras[20].runs[1].text = f'  {meta.calculator or ""}'.ljust(20)
    for i in [2, 3, 4]:
        if i < len(paras[20].runs):
            paras[20].runs[i].text = ''

    # Para[29] — 日期  "20  年  月  日"
    # run[0]='20'  run[2]='年'  run[4]='月'  run[5]='  '  run[6]='日'
    paras[29].runs[0].text = str(now.year)
    paras[29].runs[4].text = f'{now.month:02d}月'
    if len(paras[29].runs) > 5: paras[29].runs[5].text = ''
    if len(paras[29].runs) > 6: paras[29].runs[6].text = ''

    # ------------------------------------------------------------------
    # Table[0] — 14行×5列 运行卡大表
    # ------------------------------------------------------------------
    tbl0 = doc.tables[0]
    _fill_template_cell(tbl0.cell(0, 1), calc_title,        10, left=True)
    _fill_template_cell(tbl0.cell(1, 4), meta.calculator or '', 9)
    _fill_template_cell(tbl0.cell(2, 1), calc_content_desc, 10, left=True)
    clause = meta.mandatory_clause.strip() or '无'
    if clause != '无':
        _fill_template_cell(tbl0.cell(3, 1), f'产品采用的强标条款：{clause}', 10, left=True)

    # ------------------------------------------------------------------
    # Table[1] — 强标检查表（6行×6列）
    # ------------------------------------------------------------------
    tbl1 = doc.tables[1]
    _fill_template_cell(tbl1.cell(1, 1), proj,             10, left=True)
    _fill_template_cell(tbl1.cell(1, 3), stage,            10)
    _fill_template_cell(tbl1.cell(2, 1), calc_title,       10, left=True)
    _fill_template_cell(tbl1.cell(2, 3), calc_content_desc,10, left=True)

    # Para[31]: 空 Heading 1 + 分节符
    # 分节符必须保留：Section 1(封面/目录页)无页脚页码，Section 2(正文)页码从1开始
    p31 = doc.paragraphs[31]
    p31.style = doc.styles['Normal']
    # 删除 Para[31] 之后的多余空段（模板末尾可能残留）
    for p in list(doc.paragraphs[32:]):
        p._element.getparent().remove(p._element)

    program_text = calc_program_text or '渠系建筑物水力计算系统 V1.0'

    def _h4(text):
        """章节标题（Heading 1 样式，WPS 手动生成目录时可识别）"""
        p = doc.add_paragraph(style='Heading 1')
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(4)
        _run_song(p, text, 14, bold=True)

    def _body_p(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.line_spacing  = Pt(20)
        _run_song(p, text, 12)

    _h4('1、计算目的')
    if calc_purpose:
        _body_p(calc_purpose)

    _h4('2、计算依据')
    all_refs = list(references or [])
    extra = list(meta.extra_references or [])
    all_refs = all_refs + [r for r in extra if r not in all_refs]
    for i, ref in enumerate(all_refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        _run_song(p, f'{i}、{ref}', 12)

    _h4('3、基本资料')
    basic = meta.basic_info.strip() or '（基本资料未填写，请在"项目设置"中补充）'
    for seg in basic.split('\n'):
        seg = seg.strip()
        if seg:
            _body_p(seg)

    _h4('4、计算程序')
    _body_p(program_text)

    return doc


def create_engineering_report_doc(meta, calc_title, calc_content_desc,
                                   calc_purpose, references,
                                   calc_program_text=None):
    """创建工程产品运行卡格式 Word 文档（前5页）。
    优先使用模板文件；模板不存在时回退到手工生成。
    Returns:
        doc: Document 对象（已含前5页），调用方 page_break 后追加计算内容
    """
    # ---- 优先：模板替换方式 ----
    doc = create_engineering_report_doc_from_template(
        meta, calc_title, calc_content_desc,
        calc_purpose, references, calc_program_text
    )
    if doc is not None:
        return doc

    # ---- 回退：手工生成方式（模板文件缺失时） ----
    doc = DocxDocument()

    style_normal = doc.styles['Normal']
    style_normal.font.name = '宋体'
    style_normal._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style_normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    style_normal.font.size = Pt(12)
    style_normal.paragraph_format.line_spacing = Pt(20)

    program_text = calc_program_text or '渠系建筑物水力计算系统 V1.0'

    _add_page1_running_card(doc, meta, calc_title, calc_content_desc)
    _add_page2_cover(doc, meta, calc_title)
    _add_eng_header_footer(doc, meta, calc_title)
    _add_page3_mandatory_standards(doc, meta, calc_title, calc_content_desc)
    _add_page4_toc(doc)
    _add_page5_calc_intro(doc, meta, calc_purpose, references, program_text)

    return doc


def update_doc_toc_via_com(filepath):
    """保存后通过 COM 自动刷新目录（支持 Microsoft Word 和 WPS）。
    返回 True 表示刷新成功，返回 False 表示 COM 不可用（无副作用）。"""
    try:
        import win32com.client
    except ImportError:
        return False
    abs_path = os.path.abspath(filepath)
    for app_name in ['Word.Application', 'KWPS.Application', 'WPS.Application']:
        app = None
        try:
            app = win32com.client.DispatchEx(app_name)
            app.Visible = False
            doc = app.Documents.Open(abs_path)
            doc.Fields.Update()
            for i in range(1, doc.TablesOfContents.Count + 1):
                doc.TablesOfContents(i).Update()
            doc.Save()
            doc.Close()
            app.Quit()
            return True
        except Exception:
            try:
                if app is not None:
                    app.Quit()
            except Exception:
                pass
            continue
    return False


def doc_add_eng_h(doc, text):
    """工程报告风格正文标题（Heading 4 等效，宋体加粗14pt）"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    _run_song(p, text, 14, bold=True)
    return p


def doc_add_eng_body(doc, text):
    """工程报告风格正文段落（宋体12pt，首行缩进2字，1.5倍行距）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = Pt(20)
    _run_song(p, text, 12)
    return p


def doc_render_calc_text_eng(doc, text, skip_title_keyword=""):
    """将纯文本计算过程渲染为工程报告风格 Word 内容（宋体，紧凑排版）"""
    if not text:
        return
    for line in text.split('\n'):
        stripped = line.strip()
        if not stripped:
            continue
        if set(stripped) <= {'=', '-', ' '} and len(stripped) > 5:
            continue
        if skip_title_keyword and skip_title_keyword in stripped:
            continue
        if stripped.startswith('{{') and stripped.endswith('}}'):
            continue
        if stripped.startswith('【') and '】' in stripped:
            doc_add_eng_h(doc, stripped.lstrip('【').rstrip('】'))
            continue
        step_m = re.match(r'^(\d+)\.\s+(.+)', stripped)
        if step_m:
            doc_add_eng_h(doc, f'{step_m.group(1)}. {step_m.group(2)}')
            continue
        fl = try_convert_formula_line(stripped)
        if fl:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.2)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            omml = latex_to_omml(fl)
            if omml is not None:
                p._element.append(omml)
            else:
                _run_song(p, f' {stripped}', 11)
            continue
        doc_add_eng_body(doc, stripped)
