# -*- coding: utf-8 -*-
"""
明渠水力计算面板 —— QWidget 版本（可嵌入主导航框架）

支持：梯形/复式梯形/矩形/圆形断面
功能：参数输入、计算、结果显示、断面图、导出Word/TXT/图表
"""

import sys
import os
import math
import re
import copy
import json
import html as html_mod
from pathlib import Path

# 将计算模块目录加入搜索路径
_pkg_root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
sys.path.insert(0, os.path.join(_pkg_root, "calc_渠系计算算法内核"))

from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QGroupBox,
    QSplitter, QFrame, QTabWidget, QFileDialog, QScrollArea,
    QPushButton, QApplication, QRadioButton, QButtonGroup,
    QTableWidget,
)
from PySide6.QtCore import Qt, Signal, QTimer, QEvent
from PySide6.QtGui import QFont
from app_渠系计算前端.webview_compat import (
    create_web_view,
    get_web_engine_import_error,
    run_view_javascript,
    scroll_view_to_anchor,
    view_supports_scripted_html,
)

from qfluentwidgets import (
    ComboBox, PushButton, PrimaryPushButton, LineEdit,
    CheckBox, InfoBar, InfoBarPosition
)

from app_渠系计算前端.case_manager import (
    FlowLayout as _FlowLayout,
    CaseTagChip as _CaseTagChip,
    DashedButton as _DashedButton,
    MAX_CASES,
    _SUB, _sub,
    CASE_TAG_ACTIVE_SS as _CASE_TAG_ACTIVE_SS,
    CASE_TAG_INACTIVE_SS as _CASE_TAG_INACTIVE_SS,
    CASE_QUICK_SS as _CASE_QUICK_SS,
    CaseTagNavigator as _CaseTagNavigator,
    CaseWorkbenchStrip as _CaseWorkbenchStrip,
)

import matplotlib
matplotlib.use('QtAgg')
import matplotlib.pyplot as plt
try:
    from matplotlib.backends.backend_qtagg import FigureCanvasQTAgg as FigureCanvas
    from matplotlib.backends.backend_qtagg import NavigationToolbar2QT as NavToolbar
except ImportError:
    from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
    from matplotlib.backends.backend_qt5agg import NavigationToolbar2QT as NavToolbar
from matplotlib.figure import Figure
import numpy as np

plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'SimSun']
plt.rcParams['axes.unicode_minus'] = False

# 计算引擎
from 明渠设计 import (
    quick_calculate as mingqu_calculate,
    quick_calculate_compound_trapezoidal as mingqu_compound_calculate,
    quick_calculate_circular as circular_calculate,
    quick_calculate_u_section as mingqu_u_calculate,
    _u_arc_geometry,
    calculate_area, calculate_wetted_perimeter, calculate_hydraulic_radius,
    calculate_compound_trapezoid_water_width,
    get_flow_increase_percent, MAX_BETA,
    PI, MIN_FREEBOARD, MIN_FREE_AREA_PERCENT, MIN_FLOW_FACTOR
)

# 共享模块
from app_渠系计算前端.styles import P, S, W, E, BG, CARD, BD, T1, T2, INPUT_LABEL_STYLE, INPUT_SECTION_STYLE, INPUT_HINT_STYLE
from app_渠系计算前端.export_utils import (
    WORD_EXPORT_AVAILABLE, add_formula_to_doc, try_convert_formula_line, ask_open_file,
    create_styled_doc, doc_add_h1, doc_add_h2,
    doc_add_formula, doc_add_styled_table, doc_add_table_caption,
    doc_render_calc_text, doc_add_figure, doc_add_result_table,
    create_engineering_report_doc, doc_add_eng_h, doc_add_eng_body,
    doc_render_calc_text_eng, update_doc_toc_via_com, doc_add_table_caption,
)
from app_渠系计算前端.report_meta import (
    ExportConfirmDialog, build_calc_purpose, REFERENCES_BASE, load_meta
)
from app_渠系计算前端.dxf_multi_export import (
    DxfExportCaseEntry,
    choose_scale_denom,
    export_combined_case_dxf,
    export_single_case_dxf,
    format_empty_export_warning,
    format_export_result_message,
    partition_valid_case_entries,
    select_case_entries,
    show_multi_case_dxf_dialog,
)
from app_渠系计算前端.open_channel.dxf_export import (
    export_open_channel_dxf,
    draw_open_channel_dxf_on_msp,
    draw_open_channel_comparison_table,
)
from app_渠系计算前端.open_channel.comparison import (
    OPEN_CHANNEL_COMPARISON_SPEC,
    build_open_channel_comparison_tables,
)
from app_渠系计算前端.section_comparison import (
    add_section_comparison_word_tables,
    build_table_clipboard_text,
    fill_comparison_table,
)
from app_渠系计算前端.section_plotting import draw_section
from app_渠系计算前端.section_plot_layout import (
    clear_section_plot_state,
    configure_section_grid_canvas,
    connect_section_tab_refresh,
    connect_section_plot_double_click,
    create_section_plot_scroll_area,
    register_section_axis_dialog,
    reset_section_axis_dialogs,
    schedule_section_plot_restore_refresh,
)
from app_渠系计算前端.section_shapes import (
    WaterState,
    build_circular_shape,
    build_compound_trapezoid_shape,
    build_open_channel_u_shape,
    build_trapezoid_shape,
)
from app_渠系计算前端.open_channel.appendix_e_table import (
    appendix_e_probe_script,
    appendix_e_shared_head_html,
    appendix_e_tabulator_head_html,
    build_appendix_e_error_body,
    build_appendix_e_qt_compatible_body,
    build_appendix_e_static_body,
    build_appendix_e_tabulator_body,
    make_appendix_e_payload,
)
from app_渠系计算前端.formula_renderer import (
    plain_text_to_formula_html, plain_text_to_formula_body,
    wrap_with_katex, load_formula_page, make_plain_html,
    HelpPageBuilder
)
from app_渠系计算前端.increase_input_helper import (
    INCREASE_MODE_PERCENT,
    INCREASE_MODE_Q_INCREASED,
    build_increase_formula_lines,
    build_increase_summary_lines,
    build_increase_hint_text,
    normalize_increase_mode,
    resolve_increase_input,
)
from app_渠系计算前端.result_navigation import (
    CaseResultNavigationBar,
    apply_case_result_state,
    build_result_nav_bar,
    build_result_navigation_head,
    case_result_jump_hint,
    collect_case_result_state,
    has_fresh_case_results,
    is_case_result_stale,
    make_case_result_anchor,
    sync_case_result_nav_bar,
    wrap_case_result_block,
)
from app_渠系计算前端.result_summary import (
    build_result_summary_word_items,
    prepend_result_summary_to_body,
    prepend_result_summary_to_html,
)
if WORD_EXPORT_AVAILABLE:
    from docx import Document as DocxDocument
    from docx.shared import Pt, Cm


def _e(s):
    """HTML转义"""
    return html_mod.escape(str(s))


class OpenChannelPanel(QWidget):
    """明渠水力计算面板"""
    data_changed = Signal()

    def __init__(self, parent=None):
        super().__init__(parent)
        self.input_params = {}
        self.current_result = None
        self._appendix_e_export_text = ""
        self._export_plain_text = ""
        self._scripted_render_token = 0
        self._scripted_load_probe_handler = None
        self._captured_render_body = ""
        self._captured_render_head = ""
        self._cases = [self._default_case()]
        self._current_case_idx = 0
        self._all_results = []          # [(case_idx, input_params, result), ...]
        self._loading_case = False
        self._suppress_result_render = False
        self._panel_key = "open-channel"
        self._results_dirty = False
        self._stale_result_case_indexes = set()
        self._all_results_stale = False
        self._has_rendered_results = False
        self._init_ui()
        self._setup_result_dirty_tracking()
        self._rebuild_case_tags()
        # 首次进入时按当前工况同步界面可见性，避免两套加大流量输入同时显示。
        self._load_case(self._current_case_idx)

    # ================================================================
    # UI 构建
    # ================================================================
    def _init_ui(self):
        main_lay = QHBoxLayout(self)
        main_lay.setContentsMargins(10, 8, 10, 8)

        splitter = QSplitter(Qt.Horizontal)
        splitter.setChildrenCollapsible(False)
        main_lay.addWidget(splitter)

        # 左侧输入
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        scroll.setStyleSheet("QScrollArea{border:none;background:transparent;}")
        inp_w = QWidget()
        self._build_input(inp_w)
        scroll.setWidget(inp_w)
        scroll.setMinimumWidth(280)
        scroll.setMaximumWidth(420)
        splitter.addWidget(scroll)

        # 右侧输出
        out_w = QWidget()
        self._build_output(out_w)
        splitter.addWidget(out_w)
        splitter.setSizes([340, 900])

    # ----------------------------------------------------------------
    # 输入面板
    # ----------------------------------------------------------------
    def _build_input(self, parent):
        lay = QVBoxLayout(parent)
        lay.setContentsMargins(5, 5, 5, 5)
        lay.setSpacing(6)
        self._input_title = QLabel("输入参数")
        self._input_title.setStyleSheet("font-size:18px;font-weight:700;color:#0E5DB8;padding:0 2px 2px 2px;")
        lay.addWidget(self._input_title)
        self._case_strip = _CaseWorkbenchStrip(parent)
        self._case_strip.add_requested.connect(self._add_case)
        self._case_strip.case_switched.connect(self._switch_case)
        self._case_strip.case_renamed.connect(self._on_case_renamed)
        self._case_strip.apply_to_all_requested.connect(self._apply_to_all_cases)
        self._case_strip.copy_from_prev_requested.connect(self._copy_from_prev_case)
        self._case_strip.remove_current_requested.connect(self._remove_current_case)
        self._case_nav = self._case_strip.navigator()
        lay.addWidget(self._case_strip)
        grp = QGroupBox("输入参数")
        self._input_group = grp
        grp.setTitle("")
        fl = QVBoxLayout(grp)
        fl.setSpacing(5)

        # ---- 工况标签导航 ----
        _tag_row = QHBoxLayout()
        _tag_row.setSpacing(4)
        self._case_tag_container = QWidget()
        self._case_tag_flow = _FlowLayout(self._case_tag_container, spacing=5)
        self._case_tag_flow.setContentsMargins(0, 0, 0, 0)
        _tag_row.addWidget(self._case_tag_container, 1)
        self._add_case_btn = _DashedButton("+ 添加")
        self._add_case_btn.setCursor(Qt.PointingHandCursor)
        self._add_case_btn.setFixedHeight(28)
        self._add_case_btn.clicked.connect(self._add_case)
        _tag_row.addWidget(self._add_case_btn)
        fl.addLayout(_tag_row)
        self._case_count_label = QLabel("1 个计算工况")
        self._case_count_label.setStyleSheet("font-size:11px;color:#999;")
        fl.addWidget(self._case_count_label)
        self._case_tag_container.hide()
        self._add_case_btn.hide()
        self._case_count_label.hide()
        self._case_nav = _CaseTagNavigator(parent=grp)
        self._case_nav.add_requested.connect(self._add_case)
        self._case_nav.case_switched.connect(self._switch_case)
        self._case_nav.case_renamed.connect(self._on_case_renamed)
        fl.addWidget(self._case_nav)

        # 工况管理行（与工况标签挂钩）
        _quick_row = QHBoxLayout()
        _quick_row.setSpacing(4)
        _copy_all_btn = QPushButton("复制参数到所有")
        _copy_all_btn.setCursor(Qt.PointingHandCursor)
        _copy_all_btn.setStyleSheet(_CASE_QUICK_SS)
        _copy_all_btn.setToolTip("将当前工况的参数（不含Q）复制到其余所有工况")
        _copy_all_btn.clicked.connect(self._apply_to_all_cases)
        _quick_row.addWidget(_copy_all_btn)
        _copy_prev_btn = QPushButton("从上一个复制")
        _copy_prev_btn.setCursor(Qt.PointingHandCursor)
        _copy_prev_btn.setStyleSheet(_CASE_QUICK_SS)
        _copy_prev_btn.setToolTip("将上一个工况的参数（不含Q）复制到当前工况")
        _copy_prev_btn.clicked.connect(self._copy_from_prev_case)
        _quick_row.addWidget(_copy_prev_btn)
        self._del_case_btn = QPushButton("删除当前")
        self._del_case_btn.setCursor(Qt.PointingHandCursor)
        self._del_case_btn.setStyleSheet(_CASE_QUICK_SS)
        self._del_case_btn.setToolTip("删除当前选中的工况（至少保留一个）")
        self._del_case_btn.clicked.connect(self._remove_current_case)
        _quick_row.addWidget(self._del_case_btn)
        fl.addLayout(_quick_row)
        self._case_tag_container.hide()
        self._add_case_btn.hide()
        self._case_count_label.hide()
        self._case_nav.hide()
        _copy_all_btn.hide()
        _copy_prev_btn.hide()
        self._del_case_btn.hide()
        self._legacy_case_nav = self._case_nav
        self._case_nav = self._case_strip.navigator()

        fl.addWidget(self._sep())

        # 断面类型
        r = QHBoxLayout(); r.addWidget(QLabel("断面类型:"))
        self.section_combo = ComboBox()
        self.section_combo.addItems(["梯形", "复式梯形", "矩形", "圆形", "U形"])
        self.section_combo.currentTextChanged.connect(self._on_section_type_changed)
        r.addWidget(self.section_combo, 1); fl.addLayout(r)

        self.Q_edit = self._field(fl, "设计流量 Q (m³/s):", "5.0")
        self.Q_edit.textChanged.connect(self._on_q_text_changed)
        self.m_lbl, self.m_edit = self._field2(fl, "边坡系数 m:", "1.0")
        self.n_edit = self._field(fl, "糙率 n:", "0.014")
        self.slope_edit = self._field(fl, "水力坡降 1/", "3000")

        fl.addWidget(self._slbl("【流速参数】"))
        self.vmin_edit = self._field(fl, "不淤流速 (m/s):", "0.1")
        self.vmax_edit = self._field(fl, "不冲流速 (m/s):", "100.0")
        fl.addWidget(self._hint("(一般情况下保持默认数值即可)"))

        self.inc_cb = CheckBox("考虑加大流量比例系数")
        self.inc_cb.setChecked(True)
        self.inc_cb.stateChanged.connect(self._on_inc_toggle)
        fl.addWidget(self.inc_cb)
        self.inc_mode_row = QWidget()
        self.inc_mode_row_lay = QHBoxLayout(self.inc_mode_row)
        self.inc_mode_row_lay.setContentsMargins(0, 0, 0, 0)
        self.inc_mode_row_lay.setSpacing(10)
        self.inc_mode_row_lay.addWidget(self._hint("输入方式:"))
        self.inc_mode_group = QButtonGroup(self)
        self.inc_mode_percent_rb = QRadioButton("按比例")
        self.inc_mode_q_rb = QRadioButton("按Q加大")
        self.inc_mode_group.addButton(self.inc_mode_percent_rb)
        self.inc_mode_group.addButton(self.inc_mode_q_rb)
        self.inc_mode_percent_rb.setChecked(True)
        self.inc_mode_percent_rb.toggled.connect(self._on_inc_mode_changed)
        self.inc_mode_q_rb.toggled.connect(self._on_inc_mode_changed)
        self.inc_mode_row_lay.addWidget(self.inc_mode_percent_rb)
        self.inc_mode_row_lay.addWidget(self.inc_mode_q_rb)
        self.inc_mode_row_lay.addStretch()
        fl.addWidget(self.inc_mode_row)
        self.inc_lbl, self.inc_edit = self._field2(fl, "流量加大比例 (%):", "")
        self.inc_q_lbl, self.inc_q_edit = self._field2(fl, "加大流量 Q加大 (m³/s):", "")
        self.inc_edit.textChanged.connect(self._refresh_increase_hint)
        self.inc_q_edit.textChanged.connect(self._refresh_increase_hint)
        self.inc_hint = QLabel("(留空则自动计算)")
        self.inc_hint.setStyleSheet(INPUT_HINT_STYLE)
        self.inc_derived_hint = self.inc_hint
        fl.addWidget(self.inc_hint)

        fl.addWidget(self._sep())
        fl.addWidget(self._slbl("【可选参数】"))
        self.beta_lbl, self.beta_edit = self._field2(fl, "指定宽深比 β:", "")
        self.b_lbl, self.b_edit = self._field2(fl, "指定底宽 B (m):", "")
        self.bb_hint = self._hint("(二选一输入，留空则自动计算)")
        fl.addWidget(self.bb_hint)

        self.m1_lbl, self.m1_edit = self._field2(fl, "左上坡 m1:", "")
        self.B1_lbl, self.B1_edit = self._field2(fl, "平台宽 B1 (m):", "")
        self.m2_lbl, self.m2_edit = self._field2(fl, "左下坡 m2:", "")
        self.B2_lbl, self.B2_edit = self._field2(fl, "渠底宽 B2 (m):", "")
        self.m3_lbl, self.m3_edit = self._field2(fl, "右坡 m3:", "")
        self.h1_lbl, self.h1_edit = self._field2(fl, "平台高差 h1 (m):", "")
        self.compound_hint = self._hint("(复式梯形需填写全部 6 个固定几何参数)")
        for w in (
            self.m1_lbl, self.m1_edit, self.B1_lbl, self.B1_edit,
            self.m2_lbl, self.m2_edit, self.B2_lbl, self.B2_edit,
            self.m3_lbl, self.m3_edit, self.h1_lbl, self.h1_edit,
            self.compound_hint,
        ):
            w.hide()
        fl.addWidget(self.compound_hint)

        self.D_lbl, self.D_edit = self._field2(fl, "指定直径 D (m):", "")
        self.D_hint_lbl = self._hint("(留空则自动计算)")
        fl.addWidget(self.D_hint_lbl)
        for w in (self.D_lbl, self.D_edit, self.D_hint_lbl): w.hide()

        # U形专有字段
        self.R_lbl, self.R_edit = self._field2(fl, "圆弧半径 R (m):", "0.8")
        self.alpha_lbl, self.alpha_edit = self._field2(fl, "外倾角 α (°):", "14")
        self.theta_lbl, self.theta_edit = self._field2(fl, "圆心角 θ (°):", "152")
        for w in (self.R_lbl, self.R_edit, self.alpha_lbl, self.alpha_edit,
                  self.theta_lbl, self.theta_edit): w.hide()

        fl.addWidget(self._sep())
        self.detail_cb = CheckBox("输出详细计算过程")
        self.detail_cb.setChecked(True)
        fl.addWidget(self.detail_cb)

        br = QHBoxLayout()
        self._calc_btn = PrimaryPushButton("计算"); self._calc_btn.setCursor(Qt.PointingHandCursor); self._calc_btn.clicked.connect(self._calculate)
        clb = PushButton("清空"); clb.setCursor(Qt.PointingHandCursor); clb.clicked.connect(self._clear)
        br.addWidget(self._calc_btn); br.addWidget(clb); fl.addLayout(br)

        er = QHBoxLayout()
        ec = PushButton("导出DXF"); ec.clicked.connect(self._export_dxf)
        ew = PushButton("导出Word"); ew.clicked.connect(self._export_word)
        er.addWidget(ec); er.addWidget(ew)
        fl.addLayout(er)

        lay.addWidget(grp)
        lay.addStretch()

    def _field(self, lay, label, default=""):
        r = QHBoxLayout(); l = QLabel(label); l.setMinimumWidth(140)
        l.setStyleSheet(INPUT_LABEL_STYLE)
        r.addWidget(l); e = LineEdit(); e.setText(default); r.addWidget(e, 1); lay.addLayout(r)
        return e

    def _field2(self, lay, label, default=""):
        r = QHBoxLayout(); l = QLabel(label); l.setMinimumWidth(140)
        l.setStyleSheet(INPUT_LABEL_STYLE)
        r.addWidget(l); e = LineEdit(); e.setText(default); r.addWidget(e, 1); lay.addLayout(r)
        return l, e

    def _slbl(self, t):
        l = QLabel(t); l.setStyleSheet(INPUT_SECTION_STYLE); return l

    def _hint(self, t):
        l = QLabel(t); l.setStyleSheet(INPUT_HINT_STYLE); return l

    def _sep(self):
        f = QFrame(); f.setFrameShape(QFrame.HLine); f.setStyleSheet(f"color:{BD};"); return f

    def _on_inc_toggle(self, _state):
        enabled = self.inc_cb.isChecked()
        self.inc_mode_row.setVisible(enabled)
        self.inc_lbl.setVisible(enabled and self._current_increase_mode() == INCREASE_MODE_PERCENT)
        self.inc_edit.setVisible(enabled and self._current_increase_mode() == INCREASE_MODE_PERCENT)
        self.inc_q_lbl.setVisible(enabled and self._current_increase_mode() != INCREASE_MODE_PERCENT)
        self.inc_q_edit.setVisible(enabled and self._current_increase_mode() != INCREASE_MODE_PERCENT)
        self.inc_hint.setVisible(enabled)
        self._refresh_increase_hint()

    def _current_increase_mode(self):
        return INCREASE_MODE_Q_INCREASED if self.inc_mode_q_rb.isChecked() else INCREASE_MODE_PERCENT

    def _set_increase_mode(self, mode):
        normalized = normalize_increase_mode(mode)
        if normalized == INCREASE_MODE_Q_INCREASED:
            self.inc_mode_q_rb.setChecked(True)
        else:
            self.inc_mode_percent_rb.setChecked(True)
        self._on_inc_toggle(None)

    def _on_inc_mode_changed(self, _checked):
        self._on_inc_toggle(None)

    def _refresh_increase_hint(self):
        self.inc_hint.setText(build_increase_hint_text(
            use_increase=self.inc_cb.isChecked(),
            mode=self._current_increase_mode(),
            design_q_text=self.Q_edit.text(),
            percent_text=self.inc_edit.text(),
            q_increased_text=self.inc_q_edit.text(),
        ))

    def _get_increase_summary_lines(self, params, result):
        """生成加大流量输入方式说明。"""
        return build_increase_summary_lines(
            use_increase=params.get('use_increase', True),
            mode=params.get('inc_mode', INCREASE_MODE_PERCENT),
            percent_text=params.get('inc_pct_text', ''),
            q_increased_text=params.get('inc_q_text', ''),
            result_increase_percent=result.get('increase_percent', 0.0),
            result_q_increased=result.get('Q_increased', params.get('Q', 0.0)),
        )

    # ----------------------------------------------------------------
    # 输出面板
    # ----------------------------------------------------------------
    def _build_output(self, parent):
        lay = QVBoxLayout(parent)
        lay.setContentsMargins(0, 0, 0, 0)
        self.notebook = QTabWidget()
        lay.addWidget(self.notebook)

        # Tab1: 计算结果
        t1 = QWidget(); t1l = QVBoxLayout(t1); t1l.setContentsMargins(5, 5, 5, 5)
        grp = QGroupBox("计算结果详情"); gl = QVBoxLayout(grp)
        self._result_case_nav = CaseResultNavigationBar(grp)
        self._result_case_nav.case_requested.connect(self._jump_to_case_result)
        gl.addWidget(self._result_case_nav)
        self.result_text = create_web_view()
        gl.addWidget(self.result_text)
        t1l.addWidget(grp)
        self.notebook.addTab(t1, "计算结果")

        # Tab2: 断面图
        t2 = QWidget(); t2l = QVBoxLayout(t2); t2l.setContentsMargins(5, 5, 5, 5)
        self.section_fig = Figure(figsize=(8, 6), dpi=100)
        self.section_canvas = FigureCanvas(self.section_fig)
        self.section_toolbar = NavToolbar(self.section_canvas, t2)
        t2l.addWidget(self.section_toolbar)
        self._section_plot_scroll = create_section_plot_scroll_area(self.section_canvas)
        t2l.addWidget(self._section_plot_scroll)
        connect_section_plot_double_click(self)
        section_tab_index = self.notebook.addTab(t2, "断面图")
        connect_section_tab_refresh(self, section_tab_index)

        # Tab3: 工况对比
        t3 = QWidget(); t3l = QVBoxLayout(t3); t3l.setContentsMargins(5, 5, 5, 5)
        cmp_grp = QGroupBox("工况对比"); cmp_lay = QVBoxLayout(cmp_grp)
        self.comparison_hint = QLabel("请先完成计算，系统会在这里汇总各工况的关键水力结果和结构尺寸。")
        self.comparison_hint.setWordWrap(True)
        self.comparison_hint.setStyleSheet("color:#666; font-size:12px;")
        cmp_lay.addWidget(self.comparison_hint)
        cmp_lay.addWidget(QLabel("水力结果对比表"))
        self.comparison_hydraulic_table = QTableWidget(0, len(OPEN_CHANNEL_COMPARISON_SPEC.hydraulic_columns))
        self._configure_comparison_table(self.comparison_hydraulic_table)
        fill_comparison_table(
            self.comparison_hydraulic_table,
            OPEN_CHANNEL_COMPARISON_SPEC.hydraulic_columns,
            [],
        )
        cmp_lay.addWidget(self.comparison_hydraulic_table)
        cmp_lay.addWidget(QLabel("结构尺寸对比表"))
        self.comparison_dimension_table = QTableWidget(0, len(OPEN_CHANNEL_COMPARISON_SPEC.dimension_columns))
        self._configure_comparison_table(self.comparison_dimension_table)
        fill_comparison_table(
            self.comparison_dimension_table,
            OPEN_CHANNEL_COMPARISON_SPEC.dimension_columns,
            [],
        )
        cmp_lay.addWidget(self.comparison_dimension_table)
        self.comparison_table = self.comparison_hydraulic_table
        t3l.addWidget(cmp_grp)
        self.notebook.addTab(t3, "工况对比")

        self._show_initial_help()

    # ----------------------------------------------------------------
    # 断面类型切换
    # ----------------------------------------------------------------
    def _on_section_type_changed(self, stype):
        _u_widgets = (self.R_lbl, self.R_edit, self.alpha_lbl, self.alpha_edit,
                      self.theta_lbl, self.theta_edit)
        _compound_widgets = (
            self.m1_lbl, self.m1_edit, self.B1_lbl, self.B1_edit,
            self.m2_lbl, self.m2_edit, self.B2_lbl, self.B2_edit,
            self.m3_lbl, self.m3_edit, self.h1_lbl, self.h1_edit,
            self.compound_hint,
        )
        if stype == "矩形":
            self.m_lbl.hide(); self.m_edit.hide()
            self.m_edit.setText("0.0")
            for w in (self.beta_lbl, self.beta_edit, self.b_lbl, self.b_edit, self.bb_hint): w.show()
            for w in _compound_widgets: w.hide()
            for w in (self.D_lbl, self.D_edit, self.D_hint_lbl): w.hide()
            for w in _u_widgets: w.hide()
        elif stype == "梯形":
            self.m_lbl.show(); self.m_edit.show()
            self.m_edit.setText("1.0")
            for w in (self.beta_lbl, self.beta_edit, self.b_lbl, self.b_edit, self.bb_hint): w.show()
            for w in _compound_widgets: w.hide()
            for w in (self.D_lbl, self.D_edit, self.D_hint_lbl): w.hide()
            for w in _u_widgets: w.hide()
        elif stype == "复式梯形":
            self.m_lbl.hide(); self.m_edit.hide()
            for w in (self.beta_lbl, self.beta_edit, self.b_lbl, self.b_edit, self.bb_hint): w.hide()
            for w in _compound_widgets: w.show()
            for w in (self.D_lbl, self.D_edit, self.D_hint_lbl): w.hide()
            for w in _u_widgets: w.hide()
        elif stype == "圆形":
            self.m_lbl.hide(); self.m_edit.hide()
            for w in (self.beta_lbl, self.beta_edit, self.b_lbl, self.b_edit, self.bb_hint): w.hide()
            for w in _compound_widgets: w.hide()
            for w in (self.D_lbl, self.D_edit, self.D_hint_lbl): w.show()
            for w in _u_widgets: w.hide()
        elif stype == "U形":
            self.m_lbl.hide(); self.m_edit.hide()
            for w in (self.beta_lbl, self.beta_edit, self.b_lbl, self.b_edit, self.bb_hint): w.hide()
            for w in _compound_widgets: w.hide()
            for w in (self.D_lbl, self.D_edit, self.D_hint_lbl): w.hide()
            for w in _u_widgets: w.show()

        # 断面类型切换时同步当前工况，确保工况标签实时刷新
        if self._loading_case:
            return
        if not hasattr(self, '_cases'):
            return
        if 0 <= self._current_case_idx < len(self._cases):
            self._cases[self._current_case_idx]['section_type'] = stype
        self._rebuild_case_tags()

    def _setup_result_dirty_tracking(self):
        if not hasattr(self, "_input_group"):
            return
        for widget in self._input_group.findChildren(QWidget):
            if hasattr(widget, "textChanged"):
                try:
                    widget.textChanged.connect(self._on_result_inputs_changed)
                except Exception:
                    pass
            if hasattr(widget, "currentTextChanged"):
                try:
                    widget.currentTextChanged.connect(self._on_result_inputs_changed)
                except Exception:
                    pass
            if hasattr(widget, "stateChanged"):
                try:
                    widget.stateChanged.connect(self._on_result_inputs_changed)
                except Exception:
                    pass

    def _on_result_inputs_changed(self, *_args):
        self._mark_results_dirty()

    def _case_result_state_kwargs(self, case_idx):
        """返回目标工况结果状态判断所需的参数。"""
        return {
            "all_results": self._all_results,
            "has_rendered_results": self._has_rendered_results,
            "results_dirty": self._results_dirty,
            "case_idx": case_idx,
            "stale_case_indexes": getattr(self, "_stale_result_case_indexes", set()),
            "all_results_stale": getattr(self, "_all_results_stale", False),
        }

    def _mark_results_dirty(
        self,
        case_idx=None,
        *,
        all_cases=False,
        mark_case=True,
        case_indexes=None,
    ):
        """标记旧结果过期；默认只标记当前工况。"""
        if self._loading_case:
            return
        if self._has_rendered_results or self._all_results:
            self._results_dirty = True
            if not hasattr(self, "_stale_result_case_indexes"):
                self._stale_result_case_indexes = set()
            if all_cases:
                self._all_results_stale = True
                self._stale_result_case_indexes.clear()
            elif not getattr(self, "_all_results_stale", False):
                targets = case_indexes
                if targets is None and mark_case:
                    targets = [self._current_case_idx if case_idx is None else case_idx]
                for target in targets or []:
                    try:
                        self._stale_result_case_indexes.add(int(target))
                    except (TypeError, ValueError):
                        continue
            self._clear_comparison_tables("参数已变更，请重新计算后查看工况对比。")

    def _mark_results_fresh(self):
        self._results_dirty = False
        self._stale_result_case_indexes = set()
        self._all_results_stale = False
        self._has_rendered_results = bool(self._all_results)

    def _configure_comparison_table(self, table):
        """设置工况对比表的通用交互样式。"""
        table.verticalHeader().setVisible(False)
        table.setAlternatingRowColors(True)
        table.setSelectionBehavior(QTableWidget.SelectRows)
        table.setSelectionMode(QTableWidget.ExtendedSelection)
        table.setEditTriggers(QTableWidget.NoEditTriggers)
        table.horizontalHeader().setStretchLastSection(True)
        table.installEventFilter(self)

    def eventFilter(self, obj, event):
        """处理两张对比表的全选和复制快捷键。"""
        tables = {
            getattr(self, "comparison_hydraulic_table", None),
            getattr(self, "comparison_dimension_table", None),
        }
        if obj in tables and event.type() == QEvent.KeyPress:
            mods = event.modifiers()
            ctrl_only = bool(mods & Qt.ControlModifier) and not (
                mods & (Qt.ShiftModifier | Qt.AltModifier | Qt.MetaModifier)
            )
            if ctrl_only and event.key() == Qt.Key_A:
                obj.selectAll()
                return True
            if ctrl_only and event.key() == Qt.Key_C:
                self._copy_comparison_selection_to_clipboard(obj)
                return True
        return super().eventFilter(obj, event)

    def _copy_comparison_selection_to_clipboard(self, table):
        """复制当前对比表选区到剪贴板。"""
        text, row_count, col_count = build_table_clipboard_text(table)
        if not text:
            return
        QApplication.clipboard().setText(text)
        InfoBar.success(
            "已复制",
            f"已复制 {row_count} 行 × {col_count} 列到剪贴板，可粘贴到 Excel。",
            parent=self._info_parent(),
            duration=2000,
            position=InfoBarPosition.TOP,
        )

    def _set_comparison_hint(self, text):
        """更新工况对比提示。"""
        if hasattr(self, "comparison_hint"):
            self.comparison_hint.setText(text)

    def _clear_comparison_tables(self, hint="请先完成计算，系统会在这里汇总各工况的关键水力结果和结构尺寸。"):
        """清空工况对比两张表。"""
        for table in (
            getattr(self, "comparison_hydraulic_table", None),
            getattr(self, "comparison_dimension_table", None),
        ):
            if table is not None:
                table.setRowCount(0)
        self._set_comparison_hint(hint)

    def _refresh_comparison_tables(self):
        """按成功工况刷新明渠工况对比表。"""
        if not hasattr(self, "comparison_hydraulic_table"):
            return
        tables = build_open_channel_comparison_tables(getattr(self, "_all_results", []))
        fill_comparison_table(
            self.comparison_hydraulic_table,
            OPEN_CHANNEL_COMPARISON_SPEC.hydraulic_columns,
            tables.hydraulic_rows,
        )
        fill_comparison_table(
            self.comparison_dimension_table,
            OPEN_CHANNEL_COMPARISON_SPEC.dimension_columns,
            tables.dimension_rows,
        )
        if tables.hydraulic_rows:
            self._set_comparison_hint("已汇总成功计算的工况；加大流量未启用时对应列留空。")
        else:
            self._set_comparison_hint("当前没有可汇总的成功工况，请检查计算结果。")

    def _show_result_jump_hint(self, stale=False, reason=None):
        title, content = case_result_jump_hint(stale=stale, reason=reason)
        InfoBar.warning(
            title=title,
            content=content,
            parent=self._info_parent(),
            position=InfoBarPosition.TOP,
            duration=2500,
        )

    def _case_result_nav_label(self, case_idx):
        if 0 <= case_idx < len(self._cases):
            return self._case_label(self._cases[case_idx], case_idx)
        return f"工况 {case_idx + 1}"

    def _case_result_nav_summary(self, case_idx, params, result):
        case = self._cases[case_idx] if 0 <= case_idx < len(self._cases) else {}
        stype = params.get("section_type") or case.get("section_type", "梯形")
        q_raw = params.get("Q") if "Q" in params else case.get("Q", "")
        try:
            q_text = f"Q={float(q_raw):.3f}"
        except Exception:
            q_text = f"Q={str(q_raw).strip() or '?'}"
        return f"{stype} · {'计算失败' if not result.get('success') else q_text}"

    def _build_case_nav_items(self):
        items = []
        for case_idx, params, result in self._all_results:
            items.append({
                "case_idx": case_idx,
                "anchor_id": make_case_result_anchor(self._panel_key, case_idx),
                "label": self._case_result_nav_label(case_idx),
                "summary": self._case_result_nav_summary(case_idx, params, result),
                "is_error": not result.get("success"),
            })
        return items

    def _jump_to_case_result(self, case_idx, *, defer_until_load=False):
        if not has_fresh_case_results(**self._case_result_state_kwargs(case_idx)):
            all_results_stale = getattr(self, "_all_results_stale", False)
            self._show_result_jump_hint(
                stale=is_case_result_stale(
                    case_idx=case_idx,
                    results_dirty=self._results_dirty,
                    stale_case_indexes=getattr(self, "_stale_result_case_indexes", set()),
                    all_results_stale=all_results_stale,
                ),
                reason="structure_stale" if all_results_stale else None,
            )
            return False
        self.notebook.setCurrentIndex(0)
        return scroll_view_to_anchor(
            self.result_text,
            make_case_result_anchor(self._panel_key, case_idx),
            highlight=True,
            smooth=True,
            defer_until_load=defer_until_load,
        )

    # ----------------------------------------------------------------
    # 初始帮助
    # ----------------------------------------------------------------
    def _show_initial_help(self):
        sync_case_result_nav_bar(getattr(self, "_result_case_nav", None), [])
        h = HelpPageBuilder("明渠水力计算", '请选择断面类型并输入参数后点击“计算”按钮')
        h.section("支持断面类型")
        h.numbered_list([
            ("矩形断面", "m = 0，附录E自动寻优底宽；可指定宽深比或底宽"),
            ("梯形断面", "用户设定边坡系数 m，附录E自动寻优底宽；可指定宽深比或底宽"),
            ("复式梯形断面", "输入 m1/B1/m2/B2/m3/h1 共 6 个固定几何参数，程序只反算水深与校核结果"),
            ("圆形明渠", "自动搜索最优直径；可指定直径 D"),
            ("U形明渠", "圆弧底+斜直线壁；输入R、外倾角α、圆心角θ；自动反算水深"),
        ])
        h.section("计算模式总览")
        h.table(
            ["断面类型 / 可选参数填写方式", "程序行为"],
            [
                ["矩形/梯形 — 全部留空", "附录E自动寻优最优底宽 B"],
                ["矩形/梯形 — 指定宽深比 β", "以 β=B/h 为约束，自动搜索最优 B"],
                ["矩形/梯形 — 指定底宽 B", "固定 B，反算水深并验算流速"],
                ["复式梯形 — 输入 m1/B1/m2/B2/m3/h1", "固定几何，反算设计水深和加大流量工况，不做附录E寻优"],
                ["圆形 — 留空直径 D", "自动搜索满足约束的最小 D"],
                ["圆形 — 指定直径 D", "固定 D，反算水深并验算流速"],
                ["U形 — 输入R/α/θ", "固定几何，反算水深并验算流速"],
            ]
        )
        h.section("复式梯形几何公式")
        h.formula("A = B_2·h + ((m_2+m_3)/2)·h²", "平台以下面积（h ≤ h_1）")
        h.formula("χ = B_2 + h·√(1+m_2²) + h·√(1+m_3²)", "平台以下湿周（h ≤ h_1）")
        h.formula("A = A_1 + W_1·h_s + ((m_1+m_3)/2)·h_s²", "越过平台面积（h > h_1）")
        h.formula("χ = B_2 + h_1·√(1+m_2²) + B_1 + h_s·√(1+m_1²) + h·√(1+m_3²)", "越过平台湿周（h > h_1）")
        h.hint("复式梯形参数固定为：m1=左上坡，B1=平台宽，m2=左下坡，B2=渠底宽，m3=右坡，h1=平台高差")
        h.section("U形断面几何公式")
        h.formula("h_0 = R·(1 − cos(θ/2))", "弧区高度")
        h.formula("A = R²·arccos((R−h)/R) − (R−h)·√(2Rh−h²)", "纯弧区面积（h ≤ h_0）")
        h.formula("χ = 2R·arccos((R−h)/R)", "纯弧区湿周（h ≤ h_0）")
        h.formula("A = A_{arc} + (b_{arc} + m·h_s)·h_s", "直线段区面积（h > h_0）")
        h.formula("χ = θ/180·π·R + 2·h_s·√(1+m²)", "直线段区湿周（h > h_0）")
        h.hint("矩形/梯形：宽深比 β 与底宽 B 不可同时填写（二选一）")
        h.section("曼宁公式")
        h.text("本程序基于曼宁公式进行计算：")
        h.formula("Q = (1/n) × A × R^(2/3) × i^(1/2)", "流量公式")
        h.section("断面几何公式")
        h.formula("A = (B + m×h) × h", "过水面积")
        h.formula("χ = B + 2×h×√(1+m²)", "湿周")
        h.formula("R = A/χ", "水力半径")
        h.section("宽深比说明")
        h.bullet_list([
            "定义：β = B/h（底宽 / 设计水深）",
            "可选参数中可指定宽深比或底宽",
            "二选一输入，留空则自动寻优计算",
        ])
        h.section("约束条件")
        h.bullet_list(["流速范围：不淤流速 < V < 不冲流速"])
        h.section("加大流量比例规范表")
        h.table(
            ["设计流量 Q (m³/s)", "加大比例"],
            [
                ["Q < 1", "30%"],
                ["1 ≤ Q < 5", "25%"],
                ["5 ≤ Q < 20", "20%"],
                ["20 ≤ Q < 50", "15%"],
                ["50 ≤ Q < 100", "10%"],
                ["Q ≥ 100", "5%"],
            ]
        )
        self.result_text.setHtml(h.build())

    # ----------------------------------------------------------------
    # 辅助：读取输入值
    # ----------------------------------------------------------------
    def _fval(self, edit, default=0.0):
        t = edit.text().strip()
        if not t: return default
        try: return float(t)
        except ValueError: return default

    def _fval_opt(self, edit):
        t = edit.text().strip()
        if not t: return None
        try: return float(t)
        except ValueError: return None

    def _info_parent(self):
        """获取InfoBar宿主，优先当前页面。"""
        return self

    # ================================================================
    # 工况管理
    # ================================================================
    @staticmethod
    def _default_case():
        return {
            'custom_label': None,
            'section_type': '梯形',
            'Q': '5.0', 'm': '1.0', 'n': '0.014', 'slope_inv': '3000',
            'v_min': '0.1', 'v_max': '100.0',
            'inc_checked': True, 'inc_pct': '', 'inc_mode': INCREASE_MODE_PERCENT, 'inc_q_text': '',
            'detail_checked': True,
            'beta': '', 'b': '',
            'm1': '', 'B1': '', 'm2': '', 'B2': '', 'm3': '', 'h1': '',
            'D': '',
            'R': '0.8', 'alpha': '14', 'theta': '152',
        }

    def _save_current_case(self):
        if not (0 <= self._current_case_idx < len(self._cases)):
            return
        c = self._cases[self._current_case_idx]
        c['section_type'] = self.section_combo.currentText()
        c['Q'] = self.Q_edit.text()
        c['m'] = self.m_edit.text()
        c['n'] = self.n_edit.text()
        c['slope_inv'] = self.slope_edit.text()
        c['v_min'] = self.vmin_edit.text()
        c['v_max'] = self.vmax_edit.text()
        c['inc_checked'] = self.inc_cb.isChecked()
        c['inc_pct'] = self.inc_edit.text()
        c['inc_mode'] = self._current_increase_mode()
        c['inc_q_text'] = self.inc_q_edit.text()
        c['detail_checked'] = self.detail_cb.isChecked()
        c['beta'] = self.beta_edit.text()
        c['b'] = self.b_edit.text()
        c['m1'] = self.m1_edit.text()
        c['B1'] = self.B1_edit.text()
        c['m2'] = self.m2_edit.text()
        c['B2'] = self.B2_edit.text()
        c['m3'] = self.m3_edit.text()
        c['h1'] = self.h1_edit.text()
        c['D'] = self.D_edit.text()
        c['R'] = self.R_edit.text()
        c['alpha'] = self.alpha_edit.text()
        c['theta'] = self.theta_edit.text()

    def _load_case(self, idx):
        if not (0 <= idx < len(self._cases)):
            return
        c = self._cases[idx]
        self._loading_case = True
        # Section type first (triggers show/hide)
        self.section_combo.blockSignals(True)
        self.section_combo.setCurrentText(c.get('section_type', '梯形'))
        self.section_combo.blockSignals(False)
        self._on_section_type_changed(c.get('section_type', '梯形'))
        self.Q_edit.blockSignals(True)
        self.Q_edit.setText(c.get('Q', ''))
        self.Q_edit.blockSignals(False)
        self.m_edit.setText(c.get('m', '1.0'))
        self.n_edit.setText(c.get('n', '0.014'))
        self.slope_edit.setText(c.get('slope_inv', '3000'))
        self.vmin_edit.setText(c.get('v_min', '0.1'))
        self.vmax_edit.setText(c.get('v_max', '100.0'))
        self.inc_cb.setChecked(c.get('inc_checked', True))
        self.inc_edit.setText(c.get('inc_pct', ''))
        self.inc_q_edit.setText(c.get('inc_q_text', ''))
        self._set_increase_mode(c.get('inc_mode', INCREASE_MODE_PERCENT))
        self.detail_cb.setChecked(c.get('detail_checked', True))
        self.beta_edit.setText(c.get('beta', ''))
        self.b_edit.setText(c.get('b', ''))
        self.m1_edit.setText(c.get('m1', ''))
        self.B1_edit.setText(c.get('B1', ''))
        self.m2_edit.setText(c.get('m2', ''))
        self.B2_edit.setText(c.get('B2', ''))
        self.m3_edit.setText(c.get('m3', ''))
        self.h1_edit.setText(c.get('h1', ''))
        self.D_edit.setText(c.get('D', ''))
        self.R_edit.setText(c.get('R', '0.8'))
        self.alpha_edit.setText(c.get('alpha', '14'))
        self.theta_edit.setText(c.get('theta', '152'))
        self._on_inc_toggle(None)
        self._loading_case = False

    def _switch_case(self, idx):
        if idx != self._current_case_idx:
            self._save_current_case()
            self._current_case_idx = idx
            self._load_case(idx)
            self._rebuild_case_tags()
        if has_fresh_case_results(
            all_results=self._all_results,
            has_rendered_results=self._has_rendered_results,
            results_dirty=self._results_dirty,
            case_idx=idx,
            stale_case_indexes=getattr(self, "_stale_result_case_indexes", set()),
            all_results_stale=getattr(self, "_all_results_stale", False),
        ):
            self._jump_to_case_result(idx)

    def _add_case(self):
        if len(self._cases) >= MAX_CASES:
            InfoBar.warning(title="提示", content=f"最多支持 {MAX_CASES} 个工况",
                            parent=self._info_parent(), position=InfoBarPosition.TOP, duration=2000)
            return
        self._save_current_case()
        self._mark_results_dirty(mark_case=False)
        new_case = copy.deepcopy(self._cases[self._current_case_idx])
        new_case['Q'] = ''
        new_case['custom_label'] = None
        self._cases.append(new_case)
        self._current_case_idx = len(self._cases) - 1
        self._load_case(self._current_case_idx)
        self._rebuild_case_tags()
        self._update_calc_btn_text()
        self.Q_edit.setFocus()

    def _remove_current_case(self):
        if len(self._cases) <= 1:
            InfoBar.warning(title="提示", content="至少保留一个工况",
                            parent=self._info_parent(), position=InfoBarPosition.TOP, duration=2000)
            return
        self._mark_results_dirty(all_cases=True)
        idx = self._current_case_idx
        self._cases.pop(idx)
        if self._current_case_idx >= len(self._cases):
            self._current_case_idx = len(self._cases) - 1
        self._load_case(self._current_case_idx)
        self._rebuild_case_tags()
        self._update_calc_btn_text()
        InfoBar.success(title="已删除", content=f"工况{idx + 1} 已删除，当前 {len(self._cases)} 个工况",
                        parent=self._info_parent(), position=InfoBarPosition.TOP, duration=2000)

    def _rebuild_case_tags(self):
        if hasattr(self, '_case_strip'):
            self._case_strip.sync_cases(self._cases, self._current_case_idx, self._case_view)
            self._case_strip.set_remove_enabled(len(self._cases) > 1)
            self._case_nav = self._case_strip.navigator()
            return
        if hasattr(self, '_case_nav'):
            self._case_nav.sync_cases(self._cases, self._current_case_idx, self._case_view)
            return
        if not hasattr(self, '_case_tag_flow'):
            return
        layout = self._case_tag_flow
        while layout.count():
            item = layout.takeAt(0)
            w = item.widget()
            if w:
                w.setParent(None)
                w.deleteLater()
        for i, case in enumerate(self._cases):
            label = case.get('custom_label') or self._auto_label(case, i)
            chip = _CaseTagChip(i, label, active=(i == self._current_case_idx))
            chip.switched.connect(self._switch_case)
            chip.renamed.connect(self._on_case_renamed)
            layout.addWidget(chip)
        n = len(self._cases)
        self._case_count_label.setText(f"{n} 个计算工况")
        self._case_tag_container.updateGeometry()
        self._case_tag_container.update()

    def _case_view(self, case, idx):
        stype = case.get('section_type', '梯形')
        q_text = (case.get('Q', '') or '').strip() or '?'
        custom = (case.get('custom_label') or '').strip()
        label = f"{custom or stype} · Q={q_text}"
        return {
            "label": label,
            "tooltip": f"{label}\n断面类型：{stype}\n设计流量 Q={q_text} m³/s",
        }

    def _case_label(self, case, idx):
        return self._case_view(case, idx)["label"]

    def _auto_label(self, case, idx):
        stype = case.get('section_type', '梯形')
        q_text = (case.get('Q', '') or '').strip() or '?'
        return f"{stype}-Q{_sub(idx + 1)}={q_text}"

    def _on_case_renamed(self, idx, new_name):
        if 0 <= idx < len(self._cases):
            self._cases[idx]['custom_label'] = new_name
            self._rebuild_case_tags()

    def _update_calc_btn_text(self):
        n = len(self._cases)
        if n <= 1:
            self._calc_btn.setText("计算")
        else:
            self._calc_btn.setText(f"计算全部 ({n}个工况)")

    def _on_q_text_changed(self, text):
        if self._loading_case:
            return
        if not hasattr(self, '_cases'):
            return
        if 0 <= self._current_case_idx < len(self._cases):
            self._cases[self._current_case_idx]['Q'] = text
        self._refresh_increase_hint()
        self._rebuild_case_tags()

    def _apply_to_all_cases(self):
        self._save_current_case()
        self._mark_results_dirty(
            case_indexes=[i for i in range(len(self._cases)) if i != self._current_case_idx]
        )
        src = self._cases[self._current_case_idx]
        keys = ('section_type', 'm', 'n', 'slope_inv', 'v_min', 'v_max',
                'inc_checked', 'inc_pct', 'inc_mode', 'inc_q_text', 'detail_checked',
                'beta', 'b', 'm1', 'B1', 'm2', 'B2', 'm3', 'h1', 'D', 'R', 'alpha', 'theta')
        for i, case in enumerate(self._cases):
            if i != self._current_case_idx:
                for k in keys:
                    case[k] = src[k]
        n_copied = len(self._cases) - 1
        if n_copied == 0:
            InfoBar.warning(title="提示", content="当前只有一个工况，无需复制",
                            parent=self._info_parent(), position=InfoBarPosition.TOP, duration=2000)
            return
        InfoBar.success(title="已复制", content=f"参数已复制到其余 {n_copied} 个工况",
                        parent=self._info_parent(), position=InfoBarPosition.TOP, duration=2000)

    def _copy_from_prev_case(self):
        if self._current_case_idx == 0:
            InfoBar.warning(title="提示", content="当前已是第一个工况",
                            parent=self._info_parent(), position=InfoBarPosition.TOP, duration=2000)
            return
        self._save_current_case()
        self._mark_results_dirty()
        prev = self._cases[self._current_case_idx - 1]
        curr = self._cases[self._current_case_idx]
        for k in ('section_type', 'm', 'n', 'slope_inv', 'v_min', 'v_max',
                   'inc_checked', 'inc_pct', 'inc_mode', 'inc_q_text', 'detail_checked',
                   'beta', 'b', 'm1', 'B1', 'm2', 'B2', 'm3', 'h1', 'D', 'R', 'alpha', 'theta'):
            curr[k] = prev[k]
        self._load_case(self._current_case_idx)
        InfoBar.success(title="已复制", content=f"已从工况{self._current_case_idx}复制参数",
                        parent=self._info_parent(), position=InfoBarPosition.TOP, duration=2000)

    # ================================================================
    # 计算
    # ================================================================
    def _prepare_calculation_run(self):
        # Flush pending UI updates before reading the active case snapshot.
        try:
            QApplication.sendPostedEvents(None, 0)
            QApplication.processEvents()
        except Exception:
            pass
        self._save_current_case()
        self._rebuild_case_tags()
        self._all_results = []
        self.current_result = None
        self.input_params = {}
        self._export_plain_text = ""

    def _parse_and_calc_case(self, case, case_num):
        """解析单个工况并执行计算，返回 (input_params, result)"""
        def _fv(key, label, must_positive=True):
            t = (case.get(key, '') or '').strip()
            if not t:
                raise ValueError(f"工况{case_num}: 请输入{label}")
            try:
                v = float(t)
            except ValueError:
                raise ValueError(f"工况{case_num}: {label}输入无效")
            if must_positive and v <= 0:
                raise ValueError(f"工况{case_num}: {label}必须大于0")
            return v

        def _fv_opt(key):
            t = (case.get(key, '') or '').strip()
            if not t:
                return None
            try:
                return float(t)
            except ValueError:
                return None

        stype = case.get('section_type', '梯形')
        Q = _fv('Q', '设计流量 Q')
        n = _fv('n', '糙率 n')
        slope_inv = _fv('slope_inv', '水力坡降倒数')
        v_min = _fv('v_min', '不淤流速', must_positive=False)
        v_max = _fv('v_max', '不冲流速', must_positive=False)
        if v_min >= v_max:
            raise ValueError(f"工况{case_num}: 不淤流速必须小于不冲流速")

        use_increase = case.get('inc_checked', True)
        increase_resolution = resolve_increase_input(
            use_increase=use_increase,
            mode=case.get('inc_mode', INCREASE_MODE_PERCENT),
            design_q=Q,
            percent_text=case.get('inc_pct', ''),
            q_increased_text=case.get('inc_q_text', ''),
            disabled_percent=0.0,
        )
        manual_increase = increase_resolution.manual_increase_percent
        inc_mode = increase_resolution.mode

        if stype == "圆形":
            manual_D = _fv_opt('D')
            params = {
                'Q': Q, 'n': n, 'slope_inv': slope_inv,
                'v_min': v_min, 'v_max': v_max,
                'section_type': stype, 'manual_D': manual_D,
                'detail_checked': case.get('detail_checked', True),
                'manual_increase': manual_increase,
                'use_increase': use_increase,
                'inc_mode': inc_mode,
                'inc_pct_text': case.get('inc_pct', ''),
                'inc_q_text': case.get('inc_q_text', ''),
            }
            result = circular_calculate(
                Q=Q, n=n, slope_inv=slope_inv,
                v_min=v_min, v_max=v_max,
                manual_D=manual_D,
                increase_percent=manual_increase
            )
        elif stype == "复式梯形":
            m1 = _fv('m1', '左上坡 m1', must_positive=False)
            B1 = _fv('B1', '平台宽 B1')
            m2 = _fv('m2', '左下坡 m2', must_positive=False)
            B2 = _fv('B2', '渠底宽 B2')
            m3 = _fv('m3', '右坡 m3', must_positive=False)
            h1 = _fv('h1', '平台高差 h1')
            params = {
                'Q': Q, 'n': n, 'slope_inv': slope_inv,
                'v_min': v_min, 'v_max': v_max,
                'section_type': stype,
                'm1': m1, 'B1': B1, 'm2': m2, 'B2': B2, 'm3': m3, 'h1': h1,
                'detail_checked': case.get('detail_checked', True),
                'manual_increase': manual_increase,
                'use_increase': use_increase,
                'inc_mode': inc_mode,
                'inc_pct_text': case.get('inc_pct', ''),
                'inc_q_text': case.get('inc_q_text', ''),
            }
            result = mingqu_compound_calculate(
                Q=Q, m1=m1, B1=B1, m2=m2, B2=B2, m3=m3, h1=h1,
                n=n, slope_inv=slope_inv,
                v_min=v_min, v_max=v_max,
                manual_increase_percent=manual_increase
            )
        elif stype == "U形":
            R_val = _fv('R', '圆弧半径 R')
            alpha_val = _fv('alpha', '外倾角 α', must_positive=False)
            theta_val = _fv('theta', '圆心角 θ')
            if theta_val <= 0 or theta_val > 360:
                raise ValueError(f"工况{case_num}: 圆心角 θ 需在 0°~360° 之间")
            params = {
                'Q': Q, 'n': n, 'slope_inv': slope_inv,
                'v_min': v_min, 'v_max': v_max,
                'section_type': stype,
                'R': R_val, 'alpha_deg': alpha_val, 'theta_deg': theta_val,
                'detail_checked': case.get('detail_checked', True),
                'manual_increase': manual_increase,
                'use_increase': use_increase,
                'inc_mode': inc_mode,
                'inc_pct_text': case.get('inc_pct', ''),
                'inc_q_text': case.get('inc_q_text', ''),
            }
            result = mingqu_u_calculate(
                Q=Q, R=R_val, alpha_deg=alpha_val, theta_deg=theta_val,
                n=n, slope_inv=slope_inv,
                v_min=v_min, v_max=v_max,
                manual_increase_percent=manual_increase
            )
        else:
            m = float((case.get('m', '0') or '0').strip() or '0') if stype == "梯形" else 0.0
            if stype == "梯形" and m < 0:
                raise ValueError(f"工况{case_num}: 边坡系数 m 不能为负")
            manual_beta = _fv_opt('beta')
            manual_b = _fv_opt('b')
            params = {
                'Q': Q, 'm': m, 'n': n, 'slope_inv': slope_inv,
                'v_min': v_min, 'v_max': v_max,
                'section_type': stype,
                'manual_beta': manual_beta, 'manual_b': manual_b,
                'detail_checked': case.get('detail_checked', True),
                'manual_increase': manual_increase,
                'use_increase': use_increase,
                'inc_mode': inc_mode,
                'inc_pct_text': case.get('inc_pct', ''),
                'inc_q_text': case.get('inc_q_text', ''),
            }
            result = mingqu_calculate(
                Q=Q, m=m, n=n, slope_inv=slope_inv,
                v_min=v_min, v_max=v_max,
                manual_beta=manual_beta,
                manual_b=manual_b,
                manual_increase_percent=manual_increase
            )
        return params, result

    def _calculate(self):
        self._prepare_calculation_run()
        errors = []

        for i, case in enumerate(self._cases):
            try:
                params, result = self._parse_and_calc_case(case, i + 1)
                self._all_results.append((i, params, result))
            except (ValueError, TypeError) as ex:
                msg = str(ex)
                errors.append(msg)
                q_text = str(case.get('Q', '') or '').strip()
                try:
                    q_val = float(q_text) if q_text else 0.0
                except Exception:
                    q_val = 0.0
                self._all_results.append((
                    i,
                    {
                        'section_type': case.get('section_type', '梯形'),
                        'Q': q_val,
                    },
                    {'success': False, 'error_message': msg}
                ))
            except Exception as ex:
                msg = f"工况{i+1}: 计算出错 - {str(ex)}"
                errors.append(msg)
                q_text = str(case.get('Q', '') or '').strip()
                try:
                    q_val = float(q_text) if q_text else 0.0
                except Exception:
                    q_val = 0.0
                self._all_results.append((
                    i,
                    {
                        'section_type': case.get('section_type', '梯形'),
                        'Q': q_val,
                    },
                    {'success': False, 'error_message': msg}
                ))

        if errors:
            InfoBar.error(title="输入错误", content="\n".join(errors),
                          parent=self._info_parent(), position=InfoBarPosition.TOP, duration=6000)
        if not self._all_results:
            return

        # 兼容旧属性
        _, first_params, first_result = self._all_results[0]
        self.input_params = first_params
        self.current_result = first_result

        try:
            self._display_all_results()
            self.notebook.setCurrentIndex(0)
            self.data_changed.emit()
        except Exception as ex:
            message = str(ex) or ex.__class__.__name__
            InfoBar.error(
                title="结果显示失败",
                content=f"本次计算已完成，但结果渲染失败：{message}",
                parent=self._info_parent(),
                position=InfoBarPosition.TOP,
                duration=6000,
            )
            self._show_render_failure(message)

    def _build_multi_case_summary_text(self):
        total = len(self._all_results)
        success = sum(1 for _, _, result in self._all_results if result.get('success'))
        failed = max(0, total - success)
        return "\n".join([
            "【本次计算总览】",
            f"共 {total} 个工况，成功 {success} 个，失败 {failed} 个",
            "结果已按工况顺序重新刷新显示。",
        ])

    def _show_render_failure(self, message):
        total = len(self._all_results)
        success = sum(1 for _, _, result in self._all_results if result.get('success'))
        failed = max(0, total - success)
        text = "\n".join([
            "本次计算已完成，但结果渲染失败。",
            "",
            f"本次共计算 {total} 个工况，成功 {success} 个，失败 {failed} 个。",
            f"错误信息：{message}",
            "",
            "请重新计算，或切换工况后重试。",
        ])
        self._export_plain_text = text
        self.notebook.setCurrentIndex(0)
        failure_html = make_plain_html(text)
        try:
            self._render_result_html(failure_html)
        except Exception:
            try:
                load_formula_page(self.result_text, failure_html)
            except Exception:
                try:
                    self.result_text.setHtml(failure_html)
                except Exception:
                    pass

    def _display_all_results(self):
        """多工况结果显示：合并结果并提供快捷导航。"""
        _multi = len(self._all_results) > 1
        all_plain_parts = []
        all_html_parts = []
        nav_items = []
        extra_heads = []
        seen_heads = set()
        panel_key = getattr(self, "_panel_key", "open-channel")
        label_getter = getattr(self, "_case_result_nav_label", None)
        summary_getter = getattr(self, "_case_result_nav_summary", None)
        self._suppress_result_render = True

        try:
            for case_idx, params, result in self._all_results:
                rendered = OpenChannelPanel._render_case_result_content(self, params, result)
                plain = rendered["plain_text"]
                raw_plain = plain
                body_html = rendered["body_html"]
                body_html = prepend_result_summary_to_body(
                    "open_channel", params, result, body_html
                )
                extra_head = (rendered["extra_head"] or "").strip()
                nav_label = (
                    label_getter(case_idx)
                    if callable(label_getter)
                    else f"工况 {case_idx + 1}"
                )
                if callable(summary_getter):
                    nav_summary = summary_getter(case_idx, params, result)
                else:
                    section_type = params.get("section_type", "梯形")
                    q_raw = params.get("Q", 0.0)
                    try:
                        nav_q_text = f"Q={float(q_raw):.3f}"
                    except Exception:
                        nav_q_text = f"Q={str(q_raw).strip() or '?'}"
                    nav_summary = (
                        f"{section_type} · {'计算失败' if not result.get('success') else nav_q_text}"
                    )
                if _multi:
                    section_type = params.get("section_type", "梯形")
                    q_raw = params.get("Q", 0.0)
                    try:
                        q_text = f"{float(q_raw):.3f}"
                    except Exception:
                        q_text = str(q_raw).strip() or "-"
                    plain = "【工况 {}｜{}断面｜Q = {} m³/s】\n\n{}".format(
                        case_idx + 1,
                        section_type,
                        q_text,
                        plain,
                    )
                if _multi:
                    header_text = "【工况 {}｜{}断面｜Q = {} m³/s】".format(
                        case_idx + 1,
                        section_type,
                        q_text,
                    )
                    plain = header_text + "\n\n" + raw_plain
                    body_html = (
                        '<div class="codex-case-block__summary">'
                        f"{html_mod.escape(header_text)}"
                        "</div>"
                        + body_html
                    )
                all_plain_parts.append(plain)
                all_html_parts.append(
                    wrap_case_result_block(
                        panel_key,
                        case_idx,
                        f"工况 {case_idx + 1}",
                        body_html,
                        subtitle=nav_label,
                        is_error=not result.get("success"),
                    )
                )
                if extra_head and extra_head not in seen_heads:
                    seen_heads.add(extra_head)
                    extra_heads.append(extra_head)
        finally:
            self._suppress_result_render = False

        _, first_params, first_result = self._all_results[0]
        self.input_params = first_params
        self.current_result = first_result

        combined_text = "\n\n".join(all_plain_parts)
        summary_text = ""
        if _multi:
            summary_builder = getattr(self, "_build_multi_case_summary_text", None)
            if callable(summary_builder):
                summary_text = summary_builder()
            else:
                success_count = sum(1 for _ci, _params, result in self._all_results if result.get("success"))
                fail_count = len(self._all_results) - success_count
                summary_text = f"共 {len(self._all_results)} 个工况，成功 {success_count} 个，失败 {fail_count} 个。"
            combined_text = summary_text + "\n\n" + combined_text
        self._export_plain_text = combined_text

        nav_builder = getattr(self, "_build_case_nav_items", None)
        nav_items = nav_builder() if callable(nav_builder) else []
        nav_html = build_result_nav_bar(nav_items, hidden=True)
        summary_body = ""
        if summary_text:
            summary_body = (
                '<div style="margin-bottom:24px;">'
                f"{plain_text_to_formula_body(summary_text)}"
                "</div>"
            )
        combined_body = nav_html + summary_body + "\n".join(all_html_parts)
        combined_head = build_result_navigation_head()
        if extra_heads:
            combined_head += "\n" + "\n".join(extra_heads)
        self._render_result_html(wrap_with_katex(combined_body, extra_head=combined_head))
        sync_case_result_nav_bar(getattr(self, "_result_case_nav", None), nav_items)
        if not getattr(self, "_suppress_project_restore_side_effects", False):
            mark_fresh = getattr(self, "_mark_results_fresh", None)
            if callable(mark_fresh):
                mark_fresh()
            jump_to_case = getattr(self, "_jump_to_case_result", None)
            if callable(jump_to_case):
                jump_to_case(getattr(self, "_current_case_idx", 0), defer_until_load=True)
            self._update_section_plot_all()
            refresh_comparison = getattr(self, "_refresh_comparison_tables", None)
            if callable(refresh_comparison):
                refresh_comparison()

    def _update_section_plot_all(self):
        """多工况断面图"""
        success_results = [(ci, p, r) for ci, p, r in self._all_results if r.get('success')]
        reset_section_axis_dialogs(self)
        if not success_results:
            self.section_fig.clear()
            configure_section_grid_canvas(self, 1)
            self.section_canvas.draw()
            return
        if len(success_results) == 1:
            # 单工况走原有逻辑
            ci, p, r = success_results[0]
            self.input_params = p
            configure_section_grid_canvas(self, 1)
            self._update_section_plot(r)
            axes = getattr(self.section_fig, "axes", [])
            if axes:
                stype = p.get('section_type', '梯形')
                title = self._multi_case_section_plot_title(ci, stype)
                register_section_axis_dialog(
                    self,
                    axes[0],
                    title,
                    lambda target_ax, p=p, r=r, title=title: self._draw_case_section_plot(target_ax, p, r, title),
                )
            return
        self.section_fig.clear()
        n = len(success_results)
        layout = configure_section_grid_canvas(self, n)
        ncols = layout.columns
        nrows = layout.rows
        axes = self.section_fig.subplots(nrows, ncols, squeeze=False)
        for idx_r, (ci, p, r) in enumerate(success_results):
            row, col = divmod(idx_r, ncols)
            ax = axes[row][col]
            stype = p.get('section_type', '梯形')
            title = self._multi_case_section_plot_title(ci, stype)
            self._draw_case_section_plot(ax, p, r, title)
            register_section_axis_dialog(
                self,
                ax,
                title,
                lambda target_ax, p=p, r=r, title=title: self._draw_case_section_plot(target_ax, p, r, title),
            )
        for idx_r in range(n, nrows * ncols):
            row, col = divmod(idx_r, ncols)
            axes[row][col].set_visible(False)
        self.section_fig.tight_layout()
        self.section_canvas.draw()

    def _multi_case_section_plot_title(self, case_idx, section_type):
        """生成多工况断面图标题，优先使用用户自定义工况名。"""
        case = {}
        cases = getattr(self, "_cases", [])
        if 0 <= case_idx < len(cases):
            case = cases[case_idx] or {}
        custom_label = str(case.get("custom_label") or "").strip()
        if custom_label:
            return custom_label
        return f"工况{case_idx + 1} {section_type}"

    def _draw_case_section_plot(self, ax, params, result, title):
        """按单个工况类型复用完整断面绘图，确保多工况也保留尺寸标注。"""
        stype = params.get('section_type', '梯形')
        Q = float(params.get('Q', 0.0) or 0.0)
        increase_depth, _, _ = OpenChannelPanel._increase_plot_values(params, result)
        if stype == '圆形':
            self._draw_circular(
                ax,
                result.get('D_design', 0.0),
                result.get('y_d', 0.0),
                result.get('V_d', result.get('V_design', 0.0)),
                Q,
                title,
            )
            OpenChannelPanel._draw_increase_water_level(
                self, ax, stype, params, result, increase_depth
            )
        elif stype == '复式梯形':
            h_w = result.get('h_design', 0.0)
            h_ch = result.get('h_prime', 0.0) if result.get('h_prime', 0.0) > 0 else h_w * 1.35
            self._draw_compound_trapezoid(
                ax,
                result.get('b_design', params.get('B2', 0.0)),
                params.get('m1', 0.0),
                params.get('B1', 0.0),
                params.get('m2', 0.0),
                params.get('m3', 0.0),
                params.get('h1', 0.0),
                h_ch,
                result.get('V_design', 0.0),
                Q,
                h_w,
                title,
            )
            OpenChannelPanel._draw_increase_water_level(
                self, ax, stype, params, result, increase_depth
            )
        elif stype == 'U形':
            h_w = result.get('h_design', 0.0)
            h_ch = result.get('h_prime', 0.0) if result.get('h_prime', 0.0) > 0 else h_w * 1.35
            self._draw_u_section(
                ax,
                result.get('R', 0.0),
                result.get('alpha_deg', 0.0),
                result.get('theta_deg', 0.0),
                h_w,
                h_ch,
                result.get('V_design', 0.0),
                Q,
                title,
            )
            OpenChannelPanel._draw_increase_water_level(
                self, ax, stype, params, result, increase_depth
            )
        else:
            b = result.get('b_design', 0.0)
            h = result.get('h_design', 0.0)
            h_ch = result.get('h_prime', 0.0) if increase_depth > 0 and result.get('h_prime', 0.0) > 0 else h
            m = params.get('m', 0.0)
            self._draw_trapezoid(
                ax,
                b,
                h_ch,
                m,
                result.get('V_design', 0.0),
                Q,
                h,
                title,
            )
            OpenChannelPanel._draw_increase_water_level(
                self, ax, stype, params, result, increase_depth
            )

    @staticmethod
    def _positive_plot_number(value, default=0.0):
        """把绘图字段转换为正数，非法值按默认值处理。"""
        try:
            number = float(value)
        except (TypeError, ValueError):
            return default
        if not math.isfinite(number):
            return default
        return number if number > 0 else default

    @staticmethod
    def _increase_plot_values(params, result):
        """统一取得加大工况用于断面图的水深、流量和流速。"""
        if not params.get('use_increase', True):
            return 0.0, 0.0, 0.0
        if params.get('section_type') == '圆形':
            depth = OpenChannelPanel._positive_plot_number(result.get('y_i'))
            q_inc = OpenChannelPanel._positive_plot_number(
                result.get('Q_inc', result.get('Q_increased'))
            )
            v_inc = OpenChannelPanel._positive_plot_number(
                result.get('V_i', result.get('V_increased'))
            )
        else:
            depth = OpenChannelPanel._positive_plot_number(result.get('h_increased'))
            q_inc = OpenChannelPanel._positive_plot_number(result.get('Q_increased'))
            v_inc = OpenChannelPanel._positive_plot_number(result.get('V_increased'))
        return depth, q_inc, v_inc

    def _draw_increase_water_level(self, ax, stype, params, result, h_w):
        """在多工况同图上叠加加大水位线。"""
        if h_w <= 0:
            return
        if stype == '圆形':
            span = OpenChannelPanel._draw_circular_water_line(
                self, ax, result.get('D_design', 0.0), h_w, '--'
            )
        elif stype == '复式梯形':
            h_ch = result.get('h_prime', 0.0) if result.get('h_prime', 0.0) > 0 else h_w * 1.35
            span = OpenChannelPanel._draw_compound_trapezoid_water_line(
                self,
                ax,
                result.get('b_design', params.get('B2', 0.0)),
                params.get('m1', 0.0),
                params.get('B1', 0.0),
                params.get('m2', 0.0),
                params.get('m3', 0.0),
                params.get('h1', 0.0),
                h_ch,
                h_w,
                '--',
            )
        elif stype == 'U形':
            span = OpenChannelPanel._draw_u_section_water_line(
                self,
                ax,
                result.get('R', 0.0),
                result.get('alpha_deg', 0.0),
                result.get('theta_deg', 0.0),
                h_w,
                '--',
            )
        else:
            span = OpenChannelPanel._draw_trapezoid_water_line(
                self,
                ax,
                result.get('b_design', 0.0),
                params.get('m', 0.0),
                h_w,
                '--',
            )
        if span:
            OpenChannelPanel._draw_increase_water_label(self, ax, span[1], h_w)
            OpenChannelPanel._draw_increase_depth_dimension(self, ax, span[1], h_w)

    def _draw_increase_water_label(self, ax, right_x, h_w):
        """在加大水位线旁标注加大水深。"""
        y_min, y_max = ax.get_ylim()
        y_offset = max((y_max - y_min) * 0.015, 0.03)
        label_y = h_w + y_offset
        va = 'bottom'
        if label_y > y_max:
            label_y = h_w - y_offset
            va = 'top'
        ax.text(
            right_x,
            label_y,
            f'加大水位 {h_w:.2f}m',
            ha='right',
            va=va,
            fontsize=8,
            color='blue',
        )

    def _draw_increase_depth_dimension(self, ax, right_x, h_w):
        """绘制加大水深竖向尺寸箭头，并避开既有断面尺寸标注。"""
        if h_w <= 0:
            return
        x_min, x_max = ax.get_xlim()
        y_min, y_max = ax.get_ylim()
        x_span = max(x_max - x_min, 1.0)
        y_span = max(y_max - y_min, 1.0)
        arrow_x = max(right_x, x_max) + x_span * 0.06
        text_x = arrow_x + x_span * 0.035
        ax.annotate(
            '',
            xy=(arrow_x, h_w),
            xytext=(arrow_x, 0),
            arrowprops=dict(arrowstyle='<->', color='blue', lw=1.5),
        )
        ax.text(
            text_x,
            h_w / 2,
            f'h加大={h_w:.2f}m',
            fontsize=8,
            color='blue',
            rotation=90,
            va='center',
            ha='left',
        )
        ax.set_xlim(x_min, max(x_max, text_x + x_span * 0.08))
        ax.set_ylim(min(y_min, -y_span * 0.02), max(y_max, h_w + y_span * 0.04))

    def _draw_trapezoid_water_line(self, ax, b, m, h_w, linestyle='-'):
        """绘制梯形、矩形断面的水位线。"""
        b = OpenChannelPanel._positive_plot_number(b)
        h_w = OpenChannelPanel._positive_plot_number(h_w)
        if b <= 0 or h_w <= 0:
            return
        ww = b + 2 * m * h_w
        ax.plot([-ww / 2, ww / 2], [h_w, h_w], 'b', lw=1.5, linestyle=linestyle)
        return -ww / 2, ww / 2

    def _draw_compound_trapezoid_water_line(self, ax, B2, m1, B1, m2, m3, h1, h_ch, h_w, linestyle='-'):
        """绘制复式梯形断面的水位线。"""
        h_w = OpenChannelPanel._positive_plot_number(h_w)
        if h_w <= 0:
            return
        geometry = self._compound_trapezoid_geometry(B2, m1, B1, m2, m3, h1, h_ch)
        water_points = self._compound_trapezoid_water_points(geometry, B2, m1, m2, m3, h1, h_w)
        if not water_points:
            return
        left_water = water_points[-1][0] if h_w <= h1 else water_points[3][0]
        right_water = water_points[2][0]
        ax.plot([left_water, right_water], [h_w, h_w], 'b', lw=1.5, linestyle=linestyle)
        return left_water, right_water

    def _draw_u_section_water_line(self, ax, R, alpha_deg, theta_deg, h_w, linestyle='-'):
        """绘制 U 形断面的水位线。"""
        R = OpenChannelPanel._positive_plot_number(R)
        h_w = OpenChannelPanel._positive_plot_number(h_w)
        if R <= 0 or h_w <= 0:
            return
        theta_rad = math.radians(theta_deg)
        h0 = R * (1.0 - math.cos(theta_rad / 2.0))
        if h_w <= h0:
            half_bw = math.sqrt(max(0.0, R * R - (R - h_w) ** 2))
        else:
            m = math.tan(math.radians(alpha_deg))
            b_arc = 2.0 * R * math.sin(theta_rad / 2.0)
            half_bw = (b_arc + 2 * m * (h_w - h0)) / 2
        ax.plot([-half_bw, half_bw], [h_w, h_w], 'b', lw=1.5, linestyle=linestyle)
        return -half_bw, half_bw

    def _draw_circular_water_line(self, ax, D, y, linestyle='-'):
        """绘制圆形断面的水位线。"""
        D = OpenChannelPanel._positive_plot_number(D)
        y = OpenChannelPanel._positive_plot_number(y)
        if D <= 0 or y <= 0 or y >= D:
            return
        R = D / 2
        h_off = y - R
        if abs(h_off) > R:
            return
        water_w = math.sqrt(max(0.0, R ** 2 - h_off ** 2))
        ax.plot([-water_w, water_w], [y, y], 'b', lw=1.5, linestyle=linestyle)
        return -water_w, water_w

    def _show_error(self, title, msg):
        sync_case_result_nav_bar(getattr(self, "_result_case_nav", None), [])
        out = []
        out.append("=" * 70)
        out.append(f"  {title}")
        out.append("=" * 70)
        out.append("")
        out.append(msg)
        out.append("")
        out.append("-" * 70)
        out.append("请修正后重新计算。")
        out.append("=" * 70)
        plain_text = "\n".join(out)
        self._export_plain_text = plain_text
        html = make_plain_html(plain_text)
        if self._suppress_result_render:
            OpenChannelPanel._capture_render_output(self, html)
            return
        self.result_text.setHtml(html)

    def _reset_render_capture(self):
        self._captured_render_body = ""
        self._captured_render_head = ""

    @staticmethod
    def _split_rendered_html_document(html):
        if not html:
            return "", ""
        head_match = re.search(r"<head[^>]*>(.*?)</head>", html, flags=re.IGNORECASE | re.DOTALL)
        body_match = re.search(r"<body[^>]*>(.*?)</body>", html, flags=re.IGNORECASE | re.DOTALL)
        head = head_match.group(1) if head_match else ""
        body = body_match.group(1) if body_match else html
        head = re.sub(r"<meta\b[^>]*>\s*", "", head, flags=re.IGNORECASE)
        return head.strip(), body.strip()

    def _capture_render_output(self, html):
        head, body = OpenChannelPanel._split_rendered_html_document(html)
        self._captured_render_head = head
        self._captured_render_body = body

    def _consume_render_capture(self):
        body_html = getattr(self, "_captured_render_body", "")
        extra_head = getattr(self, "_captured_render_head", "")
        OpenChannelPanel._reset_render_capture(self)
        if not body_html and not extra_head:
            return None
        return {
            "body_html": body_html,
            "extra_head": extra_head,
        }

    def _render_case_result_content(self, params, result):
        OpenChannelPanel._reset_render_capture(self)
        self.input_params = params
        self.current_result = result
        self._update_result_display(result)
        plain_text = self._export_plain_text or ""
        captured = OpenChannelPanel._consume_render_capture(self)
        if captured is None:
            captured = {
                "body_html": plain_text_to_formula_body(plain_text),
                "extra_head": "",
            }
        captured["plain_text"] = plain_text
        return captured

    def _disconnect_scripted_probe_handler(self):
        signal = getattr(self.result_text, "loadFinished", None)
        handler = getattr(self, "_scripted_load_probe_handler", None)
        if signal is not None and handler is not None:
            try:
                signal.disconnect(handler)
            except (RuntimeError, TypeError):
                pass
        self._scripted_load_probe_handler = None

    def _appendix_e_runtime_mode_label(self):
        if view_supports_scripted_html(self.result_text):
            return "QWebEngineView"
        err = get_web_engine_import_error()
        if err is not None:
            return f"QTextBrowser 降级视图（QtWebEngine 导入失败: {err.__class__.__name__}）"
        return "QTextBrowser 降级视图"

    def _appendix_e_guidance_lines(self):
        return [
            "请重新计算后再试。",
            "若仍无法显示，请重启程序后再试。",
            "若问题持续存在，请从完整安装目录启动程序后再试。",
        ]

    def _build_appendix_e_failure_html(self, payload, pre1, pre2, reason_text):
        error_body = build_appendix_e_error_body(
            payload,
            summary_text="附录E断面方案对比表暂时无法显示，请重新计算后再试。",
            runtime_mode=self._appendix_e_runtime_mode_label(),
            reason_text=reason_text,
            guidance_lines=self._appendix_e_guidance_lines(),
        )
        full_body = plain_text_to_formula_body(pre1) + error_body + plain_text_to_formula_body(pre2)
        return wrap_with_katex(full_body, extra_head=appendix_e_shared_head_html())

    def _handle_scripted_probe_result(self, token, probe_script, failure_builder, attempt, raw_result):
        if token != self._scripted_render_token:
            return

        probe = {}
        if isinstance(raw_result, str) and raw_result:
            try:
                probe = json.loads(raw_result)
            except json.JSONDecodeError:
                probe = {"state": "invalid-json", "errorText": raw_result}
        elif isinstance(raw_result, dict):
            probe = raw_result

        if probe.get("ready"):
            return

        if attempt < 6:
            QTimer.singleShot(
                180,
                lambda: self._probe_scripted_render_state(
                    token, probe_script, failure_builder, attempt + 1
                ),
            )
            return

        state = probe.get("state") or "unknown"
        error_text = probe.get("errorText") or ""
        if state == "loading":
            reason = "页面已载入，但多次探测后仍停留在初始化阶段，未检测到 Tabulator 表格 DOM。"
        elif state == "error":
            reason = "页面脚本已报告 Tabulator 初始化失败。"
        elif probe.get("hasTable"):
            reason = "检测到了表格容器，但握手标记未进入 ready 状态。"
        else:
            reason = "页面已载入，但未检测到 Tabulator 表格 DOM，说明第三方组件未真正完成渲染。"
        if error_text:
            reason = f"{reason} {error_text}"
        load_formula_page(self.result_text, failure_builder(reason))

    def _probe_scripted_render_state(self, token, probe_script, failure_builder, attempt=1):
        if token != self._scripted_render_token:
            return

        started = run_view_javascript(
            self.result_text,
            probe_script,
            lambda raw_result: self._handle_scripted_probe_result(
                token, probe_script, failure_builder, attempt, raw_result
            ),
        )
        if not started:
            load_formula_page(
                self.result_text,
                failure_builder("当前结果页无法执行 JavaScript 探针，第三方表格运行环境不可用。"),
            )

    def _render_result_html(self, html, base_dir=None, scripted_probe=None, scripted_failure_builder=None):
        """统一结果页渲染入口，支持批量收集时抑制中间渲染。"""
        if self._suppress_result_render:
            OpenChannelPanel._capture_render_output(self, html)
            return

        html = prepend_result_summary_to_html(
            "open_channel",
            getattr(self, "input_params", {}),
            getattr(self, "current_result", {}),
            html,
        )

        self._scripted_render_token += 1
        token = self._scripted_render_token
        self._disconnect_scripted_probe_handler()

        if scripted_probe and not view_supports_scripted_html(self.result_text):
            if scripted_failure_builder is not None:
                reason = "当前结果页未进入 QWebEngineView，桌面端无法执行第三方表格脚本。"
                load_formula_page(self.result_text, scripted_failure_builder(reason))
            else:
                load_formula_page(self.result_text, html, base_path=base_dir)
            return

        if scripted_probe and scripted_failure_builder is not None:
            def _on_load_finished(ok):
                if token != self._scripted_render_token:
                    return
                self._disconnect_scripted_probe_handler()
                if not ok:
                    load_formula_page(
                        self.result_text,
                        scripted_failure_builder("QWebEngine 未能完成附录E结果页载入。"),
                    )
                    return
                self._probe_scripted_render_state(token, scripted_probe, scripted_failure_builder)

            self._scripted_load_probe_handler = _on_load_finished
            self.result_text.loadFinished.connect(_on_load_finished)

        load_formula_page(self.result_text, html, base_path=base_dir)

    @staticmethod
    def _appendix_e_resource_dir():
        return Path(__file__).resolve().parents[1] / "resources"

    def _build_appendix_e_payload(self, schemes, sel_b, sel_h, v_min, v_max):
        return make_appendix_e_payload(schemes, sel_b, sel_h, v_min, v_max)

    def _build_appendix_e_markup(self, payload):
        scripted_view = view_supports_scripted_html(getattr(self, "result_text", None))
        if scripted_view:
            body = build_appendix_e_static_body(payload)
            head = appendix_e_shared_head_html()
            mode = "webengine-static"
        else:
            body = build_appendix_e_qt_compatible_body(payload, runtime_mode="")
            head = ""
            mode = "qtextbrowser-compatible"
        try:
            print(f"[AppendixE] render_mode={mode}")
        except Exception:
            pass
        return {
            "body": body,
            "head": head,
            "mode": mode,
        }

    def _render_appendix_e_result_html(self, pre1, pre2, ae_payload, ae_markup):
        full_body = (
            plain_text_to_formula_body(pre1)
            + ae_markup["body"]
            + plain_text_to_formula_body(pre2)
        )
        full_html = wrap_with_katex(full_body, extra_head=ae_markup["head"])
        self._render_result_html(full_html)

    # ================================================================
    # 结果显示分发
    # ================================================================
    def _update_result_display(self, result):
        if not result['success']:
            self._show_error("计算失败", result.get('error_message', '未知错误'))
            return
        stype = self.input_params.get('section_type', '梯形')
        detail = self.input_params.get('detail_checked', self.detail_cb.isChecked())
        if stype == '圆形':
            if detail: self._show_circular_detail(result)
            else: self._show_circular_brief(result)
        elif stype == '复式梯形':
            if detail: self._show_compound_trapezoid_detail(result)
            else: self._show_compound_trapezoid_brief(result)
        elif stype == 'U形':
            if detail: self._show_u_detail(result)
            else: self._show_u_brief(result)
        else:
            if detail: self._show_trapezoid_detail(result)
            else: self._show_trapezoid_brief(result)

    # ================================================================
    # 梯形/矩形 - 简要结果
    # ================================================================
    def _show_trapezoid_brief(self, result):
        p = self.input_params
        Q, m, n = p['Q'], p['m'], p['n']
        slope_inv = p['slope_inv']
        v_min, v_max = p['v_min'], p['v_max']
        stype = p.get('section_type', '梯形')

        b = result['b_design']; h = result['h_design']
        V = result['V_design']; A = result['A_design']
        R = result['R_design']; beta = result['Beta_design']
        inc_pct = result['increase_percent']
        Q_inc = result['Q_increased']
        h_inc = result['h_increased']; V_inc = result['V_increased']
        Fb = result['Fb']; H = result['h_prime']
        increase_summary_lines = self._get_increase_summary_lines(p, result)

        o = []
        o.append("=" * 70)
        o.append(f"              明渠水力计算结果（{stype}断面）")
        o.append("=" * 70)
        o.append("")
        o.append("【输入参数】")
        o.append("")
        _n = 1
        o.append(f"  {_n}. 断面类型:")
        o.append(f"     {stype}")
        o.append("")
        _n += 1
        o.append(f"  {_n}. 设计流量:")
        o.append(f"     Q = {Q:.3f} m³/s")
        o.append("")
        if stype == "梯形":
            _n += 1
            o.append(f"  {_n}. 边坡系数:")
            o.append(f"     m = {m}")
            o.append("")
        _n += 1
        o.append(f"  {_n}. 糙率:")
        o.append(f"     n = {n}")
        o.append("")
        _n += 1
        o.append(f"  {_n}. 水力坡降:")
        o.append(f"     = 1/{int(slope_inv)}")
        o.append("")
        _n += 1
        o.append(f"  {_n}. 不淤流速:")
        o.append(f"     = {v_min} m/s")
        o.append("")
        _n += 1
        o.append(f"  {_n}. 不冲流速:")
        o.append(f"     = {v_max} m/s")
        o.append("")
        o.append("【设计方法】")
        o.append("")
        o.append(f"  1. 采用方法:")
        o.append(f"     {result['design_method']}")
        o.append("")
        o.append("【设计结果】")
        o.append(f"  底宽 B = {b:.3f} m")
        o.append(f"  水深 h = {h:.3f} m")
        o.append(f"  宽深比 β = {beta:.3f}")
        o.append(f"  过水面积 A = {A:.3f} m²")
        o.append(f"  水力半径 R = {R:.3f} m")
        o.append(f"  设计流速 V = {V:.3f} m/s")
        o.append("")

        # 附录E表格
        schemes = result.get('appendix_e_schemes', [])
        if schemes:
            pre1 = "\n".join(o)
            ae_payload = self._build_appendix_e_payload(schemes, b, h, v_min, v_max)
            ae_markup = self._build_appendix_e_markup(ae_payload)
            self._appendix_e_export_text = self._build_ae_text(schemes, b, h, v_min, v_max)

            o2 = []
            use_increase = p.get('use_increase', True)
            if use_increase:
                o2.append("【加大流量工况】")
                o2.extend([f"  {line}" for line in increase_summary_lines])
                if h_inc > 0:
                    o2.append(f"  加大水深 h加大 = {h_inc:.3f} m")
                    o2.append(f"  加大流速 V加大 = {V_inc:.3f} m/s")
                    o2.append(f"  岕顶超高 Fb = {Fb:.3f} m")  # 岕顶超高
                    o2.append(f"  渠道高度 H = {H:.3f} m")
            else:
                Fb_d = round(0.25 * h + 0.2, 3)
                H_d = round(h + Fb_d, 3)
                o2.append("【渠道尺寸计算】")
                o2.append(f"  (不考虑加大流量，以设计水深计算渠道高度)")
                o2.append(f"  超高 Fb = 0.25 × h + 0.2 = 0.25 × {h:.3f} + 0.2 = {Fb_d:.3f} m")
                o2.append(f"  渠道高度 H = h + Fb = {h:.3f} + {Fb_d:.3f} = {H_d:.3f} m")
            o2.append("")
            o2.append("【验证结果】")
            vel_ok = v_min < V < v_max
            if use_increase:
                fb_req = 0.25 * h_inc + 0.2 if h_inc > 0 else 0
                fb_ok = Fb >= (fb_req - 0.001) if h_inc > 0 else False
                o2.append(f"  流速验证: {'✓ 通过' if vel_ok else '✗ 未通过'}")
                o2.append(f"  超高复核: {'✓ 通过' if fb_ok else '✗ 未通过'} (Fb={Fb:.3f}m, 规范要求≥{fb_req:.3f}m)")
                all_pass = vel_ok and fb_ok
            else:
                o2.append(f"  流速验证: {'✓ 通过' if vel_ok else '✗ 未通过'}")
                all_pass = vel_ok
            o2.append("")
            o2.append("=" * 70)
            o2.append(f"  综合验证结果: {'全部通过 ✓' if all_pass else '未通过 ✗'}")
            o2.append("=" * 70)
            pre2 = "\n".join(o2)

            self._export_plain_text = pre1 + "\n\n" + self._appendix_e_export_text + "\n\n" + pre2
            self._render_appendix_e_result_html(pre1, pre2, ae_payload, ae_markup)
            return

        use_increase = p.get('use_increase', True)
        if use_increase:
            o.append("【加大流量工况】")
            o.extend([f"  {line}" for line in increase_summary_lines])
            if h_inc > 0:
                o.append(f"  加大水深 h加大 = {h_inc:.3f} m")
                o.append(f"  加大流速 V加大 = {V_inc:.3f} m/s")
                o.append(f"  岕顶超高 Fb = {Fb:.3f} m")
                o.append(f"  渠道高度 H = {H:.3f} m")
        else:
            Fb_d = round(0.25 * h + 0.2, 3)
            H_d = round(h + Fb_d, 3)
            o.append("【渠道尺寸计算】")
            o.append(f"  (不考虑加大流量，以设计水深计算渠道高度)")
            o.append(f"  超高 Fb = 0.25 × h + 0.2 = 0.25 × {h:.3f} + 0.2 = {Fb_d:.3f} m")
            o.append(f"  渠道高度 H = h + Fb = {h:.3f} + {Fb_d:.3f} = {H_d:.3f} m")
        o.append("")
        o.append("【验证结果】")
        vel_ok = v_min < V < v_max
        if use_increase:
            fb_req = 0.25 * h_inc + 0.2 if h_inc > 0 else 0
            fb_ok = Fb >= (fb_req - 0.001) if h_inc > 0 else False
            o.append(f"  流速验证: {'✓ 通过' if vel_ok else '✗ 未通过'}")
            o.append(f"  超高复核: {'✓ 通过' if fb_ok else '✗ 未通过'} (Fb={Fb:.3f}m, 规范要求≥{fb_req:.3f}m)")
            all_pass = vel_ok and fb_ok
        else:
            o.append(f"  流速验证: {'✓ 通过' if vel_ok else '✗ 未通过'}")
            all_pass = vel_ok
        o.append("=" * 70)
        o.append(f"  综合验证结果: {'全部通过 ✓' if all_pass else '未通过 ✗'}")
        o.append("=" * 70)
        txt = "\n".join(o)
        self._export_plain_text = txt
        self._render_result_html(plain_text_to_formula_html(txt))

    # ================================================================
    # 梯形/矩形 - 详细结果
    # ================================================================
    def _show_trapezoid_detail(self, result):
        p = self.input_params
        Q, m, n = p['Q'], p['m'], p['n']
        slope_inv = p['slope_inv']; i = 1.0 / slope_inv
        v_min, v_max = p['v_min'], p['v_max']
        stype = p.get('section_type', '梯形')

        b = result['b_design']; h = result['h_design']
        V = result['V_design']; A = result['A_design']
        X = result['X_design']; R = result['R_design']
        beta = result['Beta_design']; Q_calc = result['Q_calc']

        inc_pct = result['increase_percent']
        Q_inc = result['Q_increased']; h_inc = result['h_increased']
        V_inc = result['V_increased']
        A_inc = result.get('A_increased', -1)
        X_inc = result.get('X_increased', -1)
        R_inc = result.get('R_increased', -1)
        Fb = result['Fb']; H = result['h_prime']
        increase_summary_lines = self._get_increase_summary_lines(p, result)

        o = []
        o.append("=" * 70)
        o.append(f"              明渠水力计算结果（{stype}断面）")
        o.append("=" * 70)
        o.append("")
        o.append("【一、输入参数】")
        o.append("")
        _n = 1
        o.append(f"  {_n}. 断面类型:")
        o.append(f"     {stype}")
        o.append("")
        _n += 1
        o.append(f"  {_n}. 设计流量:")
        o.append(f"     Q = {Q:.3f} m³/s")
        o.append("")
        if stype == "梯形":
            _n += 1
            o.append(f"  {_n}. 边坡系数:")
            o.append(f"     m = {m}")
            o.append("")
        _n += 1
        o.append(f"  {_n}. 糙率:")
        o.append(f"     n = {n}")
        o.append("")
        _n += 1
        o.append(f"  {_n}. 水力坡降:")
        o.append(f"     = 1/{int(slope_inv)}")
        o.append("")
        _n += 1
        o.append(f"  {_n}. 不淤流速:")
        o.append(f"     = {v_min} m/s")
        o.append("")
        _n += 1
        o.append(f"  {_n}. 不冲流速:")
        o.append(f"     = {v_max} m/s")
        o.append("")
        if p.get('manual_beta'):
            _n += 1
            o.append(f"  {_n}. 指定宽深比:")
            o.append(f"     β = {p['manual_beta']}")
            o.append("")
        if p.get('manual_b'):
            _n += 1
            o.append(f"  {_n}. 指定底宽:")
            o.append(f"     B = {p['manual_b']} m")
            o.append("")
        if p.get('manual_increase'):
            _n += 1
            o.append(f"  {_n}. 加大流量输入方式:")
            o.append(f"     {increase_summary_lines[0]}")
            o.append("")

        o.append("【二、设计方法】")
        o.append("")
        o.append(f"  1. 采用方法:")
        o.append(f"     {result['design_method']}")
        o.append("")

        schemes = result.get('appendix_e_schemes', [])
        has_ae = bool(schemes)
        if has_ae:
            pre1 = "\n".join(o)
            ae_payload = self._build_appendix_e_payload(schemes, b, h, v_min, v_max)
            ae_markup = self._build_appendix_e_markup(ae_payload)
            self._appendix_e_export_text = self._build_ae_text(schemes, b, h, v_min, v_max)
            o = []

        o.append("【三、设计结果】")
        o.append("")
        o.append("  1. 设计底宽:")
        o.append(f"     B = {b:.3f} m")
        o.append("")
        o.append("  2. 设计水深:")
        o.append(f"     h = {h:.3f} m")
        o.append("")
        o.append("  3. 宽深比:")
        o.append(f"     β = B/h = {b:.3f}/{h:.3f} = {beta:.3f}")
        o.append("")
        o.append("  4. 过水面积计算:")
        if stype == "梯形":
            o.append(f"     A = (B + m×h) × h")
            o.append(f"       = ({b:.3f} + {m}×{h:.3f}) × {h:.3f}")
            o.append(f"       = {b + m * h:.3f} × {h:.3f}")
        else:
            o.append(f"     A = B × h")
            o.append(f"       = {b:.3f} × {h:.3f}")
        o.append(f"       = {A:.3f} m²")
        o.append("")
        o.append("  5. 湿周计算:")
        sq = math.sqrt(1 + m * m)
        if stype == "梯形":
            o.append(f"     χ = B + 2×h×√(1+m²)")
            o.append(f"       = {b:.3f} + 2×{h:.3f}×√(1+{m}²)")
            o.append(f"       = {b:.3f} + 2×{h:.3f}×{sq:.4f}")
            o.append(f"       = {b:.3f} + {2 * h * sq:.3f}")
        else:
            o.append(f"     χ = B + 2×h")
            o.append(f"       = {b:.3f} + 2×{h:.3f}")
            o.append(f"       = {b:.3f} + {2 * h:.3f}")
        o.append(f"       = {X:.3f} m")
        o.append("")
        o.append("  6. 水力半径计算:")
        o.append(f"     R = A/χ = {A:.3f}/{X:.3f} = {R:.3f} m")
        o.append("")
        o.append("  7. 设计流速计算 (曼宁公式):")
        o.append(f"     V = (1/n) × R^(2/3) × i^(1/2)")
        o.append(f"       = (1/{n}) × {R:.3f}^(2/3) × {i:.6f}^(1/2)")
        o.append(f"       = {1/n:.2f} × {R**(2/3):.4f} × {math.sqrt(i):.6f}")
        o.append(f"       = {V:.3f} m/s")
        o.append("")
        o.append("  8. 流量校核:")
        o.append(f"      Q计算 = V × A = {V:.3f} × {A:.3f} = {V * A:.3f} m³/s")
        o.append(f"      误差 = {abs(V * A - Q)/Q*100:.2f}%")
        o.append("")

        use_increase = p.get('use_increase', True)
        if use_increase:
          o.append("【四、加大流量工况计算】")
          o.append("")
          o.append("  1. 输入与换算:")
          o.extend([f"      {line}" for line in increase_summary_lines])
          for formula_line in build_increase_formula_lines(
              design_q=Q,
              increase_percent=inc_pct,
              q_increased=Q_inc,
          ):
              o.append(f"      {formula_line}")
          o.append("")

        if use_increase and h_inc > 0:
            if A_inc <= 0: A_inc = (b + m * h_inc) * h_inc
            if X_inc <= 0: X_inc = b + 2 * h_inc * math.sqrt(1 + m * m)
            if R_inc <= 0 and X_inc > 0: R_inc = A_inc / X_inc

            o.append("  2. 加大水深计算:")
            o.append(f"      根据加大流量 Q加大 = {Q_inc:.3f} m³/s 和设计底宽 B = {b:.3f} m，")
            o.append(f"      利用曼宁公式反算水深:")
            o.append(f"      h加大 = {h_inc:.3f} m")
            o.append("")
            o.append("  3. 加大过水面积计算:")
            if stype == "梯形":
                o.append(f"      A加大 = (B + m×h加大) × h加大")
                o.append(f"           = ({b:.3f} + {m}×{h_inc:.3f}) × {h_inc:.3f}")
                o.append(f"           = {b + m * h_inc:.3f} × {h_inc:.3f}")
            else:
                o.append(f"      A加大 = B × h加大")
                o.append(f"           = {b:.3f} × {h_inc:.3f}")
            o.append(f"           = {A_inc:.3f} m²")
            o.append("")
            o.append("  4. 加大湿周计算:")
            sq2 = math.sqrt(1 + m * m)
            if stype == "梯形":
                o.append(f"      χ加大 = B + 2×h加大×√(1+m²)")
                o.append(f"           = {b:.3f} + 2×{h_inc:.3f}×√(1+{m}²)")
                o.append(f"           = {b:.3f} + 2×{h_inc:.3f}×{sq2:.4f}")
                o.append(f"           = {b:.3f} + {2 * h_inc * sq2:.3f}")
            else:
                o.append(f"      χ加大 = B + 2×h加大")
                o.append(f"           = {b:.3f} + 2×{h_inc:.3f}")
                o.append(f"           = {b:.3f} + {2 * h_inc:.3f}")
            o.append(f"           = {X_inc:.3f} m")
            o.append("")
            o.append("  5. 加大水力半径计算:")
            o.append(f"      R加大 = A加大 / χ加大")
            o.append(f"           = {A_inc:.3f} / {X_inc:.3f}")
            o.append(f"           = {R_inc:.3f} m")
            o.append("")
            o.append("  6. 加大流速计算 (曼宁公式):")
            o.append(f"      V加大 = (1/n) × R加大^(2/3) × i^(1/2)")
            o.append(f"           = (1/{n}) × {R_inc:.3f}^(2/3) × {i:.6f}^(1/2)")
            o.append(f"           = {1/n:.2f} × {R_inc**(2/3):.4f} × {math.sqrt(i):.6f}")
            o.append(f"           = {V_inc:.3f} m/s")
            o.append("")
            Q_chk = V_inc * A_inc
            o.append("  7. 流量校核:")
            o.append(f"      Q校核 = V加大 × A加大 = {V_inc:.3f} × {A_inc:.3f} = {Q_chk:.3f} m³/s")
            o.append(f"      误差 = {abs(Q_chk - Q_inc) / Q_inc * 100:.2f}%")
            o.append("")
            o.append("  8. 渠道岸顶超高计算（规范 6.4.8-2）:")
            o.append(f"      Fb = (1/4) × h加大 + 0.2")
            o.append(f"         = (1/4) × {h_inc:.3f} + 0.2")
            o.append(f"         = {Fb:.3f} m")
            o.append("")
            o.append("  9. 渠道高度计算:")
            o.append(f"      H = h加大 + Fb")
            o.append(f"        = {h_inc:.3f} + {Fb:.3f}")
            o.append(f"        = {H:.3f} m")
        elif use_increase:
            o.append("  加大水深计算失败")
        o.append("")

        if not use_increase:
            Fb_d = round(0.25 * h + 0.2, 3)
            H_d = round(h + Fb_d, 3)
            o.append("【四、渠道尺寸计算】")
            o.append("")
            o.append(f"  (不考虑加大流量，以设计水深计算渠道高度)")
            o.append(f"  1. 超高计算（规范 6.4.8-2）:")
            o.append(f"      Fb = (1/4) × h + 0.2 = (1/4) × {h:.3f} + 0.2 = {Fb_d:.3f} m")
            o.append(f"  2. 渠道高度计算:")
            o.append(f"      H = h + Fb = {h:.3f} + {Fb_d:.3f} = {H_d:.3f} m")
            o.append("")

        o.append("【五、设计验证】")
        o.append("")
        vel_ok = v_min < V < v_max
        o.append(f"  1. 流速验证:")
        o.append(f"      范围要求: {v_min} < V < {v_max} m/s")
        o.append(f"      设计流速: V = {V:.3f} m/s")
        o.append(f"      结果: {'通过 ✓' if vel_ok else '未通过 ✗'}")
        o.append("")
        if use_increase:
            fb_req = 0.25 * h_inc + 0.2 if h_inc > 0 else 0
            fb_ok = Fb >= (fb_req - 0.001) if h_inc > 0 else False
            o.append(f"  2. 超高复核（规范 6.4.8-2）:")
            o.append(f"      规范要求: Fb ≥ (1/4)×h加大 + 0.2 = {fb_req:.3f} m")
            o.append(f"      计算结果: Fb = {Fb:.3f} m")
            o.append(f"      结果: {'通过 ✓' if fb_ok else '未通过 ✗'}")
            o.append("")
            all_pass = vel_ok and fb_ok
        else:
            all_pass = vel_ok
        o.append("=" * 70)
        o.append(f"  综合验证结果: {'全部通过 ✓' if all_pass else '未通过 ✗'}")
        o.append("=" * 70)

        if has_ae:
            pre2 = "\n".join(o)
            self._export_plain_text = pre1 + "\n\n" + self._appendix_e_export_text + "\n\n" + pre2
            self._render_appendix_e_result_html(pre1, pre2, ae_payload, ae_markup)
        else:
            txt = "\n".join(o)
            self._export_plain_text = txt
            self._render_result_html(plain_text_to_formula_html(txt))

    def _build_compound_trapezoid_text(self, result, detail=False):
        """生成复式梯形明渠结果文本。"""
        p = self.input_params
        stype = p.get('section_type', '复式梯形')
        Q = p['Q']; n = p['n']; slope_inv = p['slope_inv']
        v_min = p['v_min']; v_max = p['v_max']
        m1 = p['m1']; B1 = p['B1']; m2 = p['m2']; B2 = p['B2']; m3 = p['m3']; h1 = p['h1']
        b = result['b_design']; h = result['h_design']
        V = result['V_design']; A = result['A_design']
        X = result['X_design']; R = result['R_design']
        Q_calc = result['Q_calc']
        inc_pct = result['increase_percent']
        Q_inc = result['Q_increased']; h_inc = result['h_increased']
        V_inc = result['V_increased']; Fb = result['Fb']; H = result['h_prime']
        increase_summary_lines = self._get_increase_summary_lines(p, result)

        o = []
        o.append("=" * 70)
        o.append(f"              明渠水力计算结果（{stype}断面）")
        o.append("=" * 70)
        o.append("")
        o.append("【一、输入参数】")
        o.append(f"  断面类型 = {stype}")
        o.append(f"  Q = {Q:.3f} m³/s, n = {n}, 水力坡降 = 1/{int(slope_inv)}")
        o.append(f"  不淤流速 = {v_min} m/s, 不冲流速 = {v_max} m/s")
        o.append("")
        o.append("【二、固定几何参数】")
        o.append(f"  左上坡 m1 = {m1}")
        o.append(f"  平台宽 B1 = {B1:.3f} m")
        o.append(f"  左下坡 m2 = {m2}")
        o.append(f"  渠底宽 B2 = {B2:.3f} m")
        o.append(f"  右坡 m3 = {m3}")
        o.append(f"  平台高差 h1 = {h1:.3f} m")
        o.append("")
        o.append("【三、设计结果】")
        o.append(f"  采用方法 = {result['design_method']}")
        o.append(f"  设计水深 h = {h:.3f} m")
        o.append(f"  过水面积 A = {A:.3f} m²")
        o.append(f"  湿周 χ = {X:.3f} m")
        o.append(f"  水力半径 R = {R:.3f} m")
        o.append(f"  设计流速 V = {V:.3f} m/s")
        o.append(f"  反算流量 Q校核 = {Q_calc:.3f} m³/s")
        if detail:
            water_width = calculate_compound_trapezoid_water_width(b, h, m2, m3, h1, B1, m1)
            o.append(f"  设计水面宽 = {water_width:.3f} m")
        o.append("")

        use_increase = p.get('use_increase', True)
        if use_increase:
            o.append("【四、加大流量工况】")
            o.extend([f"  {line}" for line in increase_summary_lines])
            if h_inc > 0:
                o.append(f"  加大水深 h加大 = {h_inc:.3f} m")
                o.append(f"  加大流速 V加大 = {V_inc:.3f} m/s")
                o.append(f"  超高 Fb = {Fb:.3f} m")
                o.append(f"  渠道高度 H = {H:.3f} m")
        else:
            Fb_d = round(0.25 * h + 0.2, 3)
            H_d = round(h + Fb_d, 3)
            o.append("【四、渠道尺寸】")
            o.append(f"  超高 Fb = {Fb_d:.3f} m")
            o.append(f"  渠道高度 H = {H_d:.3f} m")
        o.append("")
        o.append("【五、验证结果】")
        vel_ok = v_min < V < v_max
        o.append(f"  流速验证: {'✓ 通过' if vel_ok else '✗ 未通过'}")
        if use_increase and h_inc > 0:
            fb_req = 0.25 * h_inc + 0.2
            fb_ok = Fb >= (fb_req - 0.001)
            o.append(f"  超高复核: {'✓ 通过' if fb_ok else '✗ 未通过'} (Fb={Fb:.3f}m, 规范要求≥{fb_req:.3f}m)")
        if result.get('constraint_warnings'):
            o.append("  约束提示:")
            for warning in result['constraint_warnings']:
                o.append(f"    - {warning}")
        o.append("=" * 70)
        return "\n".join(o)

    def _show_compound_trapezoid_brief(self, result):
        """显示复式梯形明渠简要结果。"""
        txt = self._build_compound_trapezoid_text(result, detail=False)
        self._export_plain_text = txt
        self._render_result_html(plain_text_to_formula_html(txt))

    def _show_compound_trapezoid_detail(self, result):
        """显示复式梯形明渠详细结果。"""
        txt = self._build_compound_trapezoid_text(result, detail=True)
        self._export_plain_text = txt
        self._render_result_html(plain_text_to_formula_html(txt))

    # ================================================================
    # 圆形 - 简要结果
    # ================================================================
    def _show_circular_brief(self, result):
        p = self.input_params
        Q, n = p['Q'], p['n']; slope_inv = p['slope_inv']
        v_min, v_max = p['v_min'], p['v_max']
        D = result.get('D_design', 0)
        h = result.get('y_d', 0); V = result.get('V_d', 0)
        A_d = result.get('A_d', 0); FB_d = result.get('FB_d', 0)
        PA_d = result.get('PA_d', 0)
        inc_info = result.get('increase_percent', '')
        Q_inc = result.get('Q_inc', 0)
        h_i = result.get('y_i', 0); V_i = result.get('V_i', 0)
        FB_i = result.get('FB_i', 0); PA_i = result.get('PA_i', 0)

        o = []
        o.append("=" * 70)
        o.append("              明渠水力计算结果（圆形断面）")
        o.append("=" * 70)
        o.append("")
        o.append("【输入参数】")
        o.append("")
        o.append(f"  1. 断面类型:")
        o.append(f"     圆形")
        o.append("")
        o.append(f"  2. 设计流量:")
        o.append(f"     Q = {Q:.3f} m³/s")
        o.append("")
        o.append(f"  3. 糙率:")
        o.append(f"     n = {n}")
        o.append("")
        o.append(f"  4. 水力坡降:")
        o.append(f"     = 1/{int(slope_inv)}")
        o.append("")
        o.append(f"  5. 不淤流速:")
        o.append(f"     = {v_min} m/s")
        o.append("")
        o.append(f"  6. 不冲流速:")
        o.append(f"     = {v_max} m/s")
        o.append("")
        o.append("【断面尺寸】")
        o.append("")
        o.append(f"  1. 设计直径:")
        o.append(f"     D = {D:.2f} m")
        o.append("")
        o.append("【设计流量工况】")
        o.append(f"  设计水深 h = {h:.3f} m")
        o.append(f"  设计流速 V = {V:.3f} m/s")
        o.append(f"  过水面积 A = {A_d:.3f} m²")
        o.append(f"  净空高度 Fb = {FB_d:.3f} m")
        o.append(f"  净空比例 = {PA_d:.1f}%")
        o.append("")
        use_increase = p.get('use_increase', True)
        if use_increase:
            o.append("【加大流量工况】")
            o.extend([f"  {line}" for line in self._get_increase_summary_lines(p, result)])
            o.append(f"  加大水深 h加大 = {h_i:.3f} m")
            o.append(f"  加大流速 V加大 = {V_i:.3f} m/s")
            o.append(f"  净空高度 Fb加大 = {FB_i:.3f} m")
            o.append(f"  净空比例 = {PA_i:.1f}%")
            o.append("")
        o.append("【验证结果】")
        vel_ok = V is not None and v_min <= V <= v_max
        o.append(f"  1. 设计流速验证")
        o.append(f"     范围要求: {v_min} ≤ V ≤ {v_max} m/s")
        o.append(f"     计算结果: V = {V:.3f} m/s")
        o.append(f"     验证结果: {'通过 ✓' if vel_ok else '未通过 ✗'}")
        o.append("")
        if use_increase:
            vel_i_ok = V_i is not None and v_min <= V_i <= v_max if V_i else True
            o.append(f"  2. 加大流速验证")
            o.append(f"     范围要求: {v_min} ≤ V ≤ {v_max} m/s")
            if V_i:
                o.append(f"     计算结果: V加大 = {V_i:.3f} m/s")
                o.append(f"     验证结果: {'通过 ✓' if vel_i_ok else '未通过 ✗'}")
            else:
                o.append(f"     计算结果: 无数据")
            o.append("")
            all_ok = vel_ok and vel_i_ok
        else:
            all_ok = vel_ok
        o.append("=" * 70)
        o.append(f"  综合验证结果: {'全部通过 ✓' if all_ok else '未通过 ✗'}")
        o.append("=" * 70)
        txt = "\n".join(o)
        self._export_plain_text = txt
        self._render_result_html(plain_text_to_formula_html(txt))

    # ================================================================
    # 圆形 - 详细结果
    # ================================================================
    def _show_circular_detail(self, result):
        p = self.input_params
        Q, n = p['Q'], p['n']; slope_inv = p['slope_inv']
        v_min, v_max = p['v_min'], p['v_max']
        i = 1.0 / slope_inv

        D_calc = result.get('D_calculated', 0)
        D = result.get('D_design', 0)
        pipe_area = PI * D**2 / 4 if D > 0 else 0

        h_d = result.get('y_d', 0); V_d = result.get('V_d', 0)
        A_d = result.get('A_d', 0); P_d = result.get('P_d', 0)
        R_d = result.get('R_d', 0); PA_d = result.get('PA_d', 0)
        FB_d = result.get('FB_d', 0); Q_chk_d = result.get('Q_check_d', 0)

        inc_info = result.get('increase_percent', '')
        Q_inc = result.get('Q_inc', 0)
        h_i = result.get('y_i', 0); V_i = result.get('V_i', 0)
        A_i = result.get('A_i', 0); P_i = result.get('P_i', 0)
        R_i = result.get('R_i', 0); PA_i = result.get('PA_i', 0)
        FB_i = result.get('FB_i', 0)

        try: inc_pct = float(inc_info.split('%')[0])
        except: inc_pct = 20

        Q_min = result.get('Q_min', 0)
        h_m = result.get('y_m', 0); V_m = result.get('V_m', 0)
        A_m = result.get('A_m', 0); P_m = result.get('P_m', 0)
        R_m = result.get('R_m', 0)

        o = []
        o.append("=" * 70)
        o.append("              明渠水力计算结果（圆形断面）")
        o.append("=" * 70)
        o.append("")
        o.append("【一、输入参数】")
        o.append("")
        _n = 1
        o.append(f"  {_n}. 断面类型:")
        o.append(f"     圆形")
        o.append("")
        _n += 1
        o.append(f"  {_n}. 设计流量:")
        o.append(f"     Q = {Q:.3f} m³/s")
        o.append("")
        _n += 1
        o.append(f"  {_n}. 糙率:")
        o.append(f"     n = {n}")
        o.append("")
        _n += 1
        o.append(f"  {_n}. 水力坡降:")
        o.append(f"     = 1/{int(slope_inv)}")
        o.append("")
        _n += 1
        o.append(f"  {_n}. 不淤流速:")
        o.append(f"     = {v_min} m/s")
        o.append("")
        _n += 1
        o.append(f"  {_n}. 不冲流速:")
        o.append(f"     = {v_max} m/s")
        o.append("")
        manual_d_input = p.get('manual_D', p.get('manual_b'))
        if manual_d_input:
            _n += 1
            o.append(f"  {_n}. 指定直径:")
            o.append(f"     D = {manual_d_input} m")
            o.append("")

        o.append("【二、直径确定】")
        o.append("")
        if D_calc and D_calc > 0:
            o.append(f"  1. 计算直径: D计算 = {D_calc:.3f} m")
        o.append(f"  2. 设计直径: D = {D:.2f} m")
        o.append("")
        o.append("  3. 管道总断面积计算:")
        o.append(f"     A总 = π × D² / 4")
        o.append(f"        = {PI:.4f} × {D:.2f}² / 4")
        o.append(f"        = {PI:.4f} × {D**2:.4f} / 4")
        o.append(f"        = {pipe_area:.3f} m²")
        o.append("")

        o.append("【三、设计流量工况计算】")
        o.append("")
        o.append("  1. 设计水深计算:")
        o.append(f"     根据设计流量 Q = {Q:.3f} m³/s，利用曼宁公式反算水深:")
        o.append(f"     h = {h_d:.3f} m")
        o.append("")

        if h_d > 0 and D > 0 and h_d <= D:
            Rr = D / 2
            theta = 2 * math.acos(max(-1, min(1, (Rr - h_d) / Rr)))
            o.append(f"  2. 圆心角计算:")
            o.append(f"     θ = 2 × arccos((R - h) / R)")
            o.append(f"       = 2 × arccos(({Rr:.3f} - {h_d:.3f}) / {Rr:.3f})")
            o.append(f"       = 2 × arccos({(Rr - h_d)/Rr:.4f})")
            o.append(f"       = {math.degrees(theta):.2f}° ({theta:.4f} rad)")
            o.append("")
            o.append(f"  3. 过水面积计算:")
            o.append(f"     A = (D²/8) × (θ - sinθ)")
            o.append(f"       = ({D:.3f}²/8) × ({theta:.4f} - sin{theta:.4f})")
            o.append(f"       = {D**2/8:.4f} × {theta - math.sin(theta):.4f}")
            o.append(f"       = {A_d:.3f} m²")
            o.append("")
            o.append(f"  4. 湿周计算:")
            o.append(f"      χ = (D/2) × θ")
            o.append(f"        = ({D:.3f}/2) × {theta:.4f}")
            o.append(f"        = {Rr:.3f} × {theta:.4f}")
            o.append(f"        = {P_d:.3f} m")
            o.append("")
        else:
            o.append(f"  3. 过水面积: A = {A_d:.3f} m²")
            o.append("")
            o.append(f"  4. 湿周: χ = {P_d:.3f} m")
            o.append("")

        o.append(f"  5. 水力半径计算:")
        o.append(f"      R = A / χ")
        o.append(f"        = {A_d:.3f} / {P_d:.3f}")
        o.append(f"        = {R_d:.3f} m")
        o.append("")
        o.append(f"  6. 设计流速计算 (曼宁公式):")
        o.append(f"      V = (1/n) × R^(2/3) × i^(1/2)")
        o.append(f"        = (1/{n}) × {R_d:.3f}^(2/3) × {i:.6f}^(1/2)")
        if R_d > 0:
            o.append(f"        = {1/n:.2f} × {R_d**(2/3):.4f} × {math.sqrt(i):.6f}")
        o.append(f"        = {V_d:.3f} m/s")
        o.append("")
        o.append(f"  7. 流量校核:")
        o.append(f"      Q计算 = V × A")
        o.append(f"           = {V_d:.3f} × {A_d:.3f}")
        o.append(f"           = {V_d * A_d:.3f} m³/s")
        if V_d * A_d > 0:
            o.append(f"      误差 = {abs(V_d * A_d - Q)/Q*100:.2f}%")
        o.append("")
        o.append(f"  8. 净空高度:")
        o.append(f"      Fb = D - h = {D:.3f} - {h_d:.3f} = {FB_d:.3f} m")
        o.append("")
        o.append(f"  9. 净空面积:")
        o.append(f"      PA = (A总 - A) / A总 × 100%")
        o.append(f"         = ({pipe_area:.3f} - {A_d:.3f}) / {pipe_area:.3f} × 100%")
        o.append(f"         = {PA_d:.1f}%")
        o.append("")

        use_increase_circ = p.get('use_increase', True)
        if use_increase_circ:
            o.append("【四、加大流量工况计算】")
            o.append("")
            o.append("  1. 输入与换算:")
            o.extend([f"      {line}" for line in self._get_increase_summary_lines(p, result)])
            for formula_line in build_increase_formula_lines(
                design_q=Q,
                increase_percent=inc_pct,
                q_increased=Q_inc,
            ):
                o.append(f"      {formula_line}")
            o.append("")
            if h_i is not None and h_i > 0 and D > 0:
                o.append("  2. 加大水深计算:")
                o.append(f"      根据加大流量 Q加大 = {Q_inc:.3f} m³/s，利用曼宁公式反算水深:")
                o.append(f"      h加大 = {h_i:.3f} m")
                o.append("")
                Rr_i = D / 2
                theta_i = 2 * math.acos(max(-1, min(1, (Rr_i - h_i) / Rr_i)))
                o.append(f"  3. 圆心角计算:")
                o.append(f"      θ加大 = 2 × arccos((R - h加大) / R)")
                o.append(f"           = 2 × arccos(({Rr_i:.3f} - {h_i:.3f}) / {Rr_i:.3f})")
                o.append(f"           = 2 × arccos({(Rr_i - h_i)/Rr_i:.4f})")
                o.append(f"           = {math.degrees(theta_i):.2f}° ({theta_i:.4f} rad)")
                o.append("")
                o.append(f"  4. 过水面积计算:")
                o.append(f"      A加大 = (D²/8) × (θ加大 - sinθ加大)")
                o.append(f"           = ({D:.3f}²/8) × ({theta_i:.4f} - sin{theta_i:.4f})")
                o.append(f"           = {D**2/8:.4f} × {theta_i - math.sin(theta_i):.4f}")
                o.append(f"           = {A_i:.3f} m²")
                o.append("")
                o.append(f"  5. 湿周计算:")
                o.append(f"      χ加大 = (D/2) × θ加大")
                o.append(f"           = ({D:.3f}/2) × {theta_i:.4f}")
                o.append(f"           = {Rr_i:.3f} × {theta_i:.4f}")
                o.append(f"           = {P_i:.3f} m")
                o.append("")
                o.append(f"  6. 水力半径计算:")
                o.append(f"      R加大 = A加大 / χ加大")
                if A_i and P_i:
                    o.append(f"           = {A_i:.3f} / {P_i:.3f}")
                    o.append(f"           = {R_i:.3f} m")
                o.append("")
                o.append(f"  7. 加大流速计算 (曼宁公式):")
                o.append(f"      V加大 = (1/n) × R加大^(2/3) × i^(1/2)")
                if R_i and R_i > 0:
                    o.append(f"           = (1/{n}) × {R_i:.3f}^(2/3) × {i:.6f}^(1/2)")
                    o.append(f"           = {1/n:.2f} × {R_i**(2/3):.4f} × {math.sqrt(i):.6f}")
                o.append(f"           = {V_i:.3f} m/s")
                o.append("")
                o.append(f"  8. 流量校核:")
                if V_i and A_i:
                    o.append(f"      Q计算 = V加大 × A加大")
                    o.append(f"           = {V_i:.3f} × {A_i:.3f}")
                    o.append(f"           = {V_i * A_i:.3f} m³/s")
                    if Q_inc > 0:
                        o.append(f"      误差 = {abs(V_i * A_i - Q_inc) / Q_inc * 100:.2f}%")
                o.append("")
                o.append(f"  9. 净空高度计算:")
                o.append(f"      Fb加大 = D - h加大")
                if h_i:
                    o.append(f"           = {D:.3f} - {h_i:.3f}")
                    o.append(f"           = {FB_i:.3f} m")
                o.append("")
                o.append(f"  10. 净空面积计算:")
                if A_i:
                    o.append(f"      PA加大 = (A总 - A加大) / A总 × 100%")
                    o.append(f"           = ({pipe_area:.3f} - {A_i:.3f}) / {pipe_area:.3f} × 100%")
                    o.append(f"           = {PA_i:.1f}%")
                o.append("")
            else:
                o.append(f"  2. 加大水深: h加大 = N/A")
                o.append("")

        o.append("【五、最小流量工况计算】")
        o.append("")
        o.append("  1. 最小流量计算:")
        o.append(f"      Q最小 = Q × 最小流量系数")
        o.append(f"           = {Q:.3f} × 0.4")
        if Q_min is not None and Q_min > 0:
            o.append(f"           = {Q_min:.3f} m³/s")
        else:
            o.append(f"           = N/A")
        o.append("")

        if Q_min is not None and Q_min > 0 and h_m is not None and h_m > 0 and D > 0:
            o.append("  2. 最小水深计算:")
            o.append(f"      根据最小流量 Q最小 = {Q_min:.3f} m³/s，利用曼宁公式反算水深:")
            o.append(f"      h最小 = {h_m:.3f} m")
            o.append("")

            Rr_m = D / 2
            if h_m <= D:
                theta_m = 2 * math.acos(max(-1, min(1, (Rr_m - h_m) / Rr_m)))
                o.append(f"  3. 圆心角计算:")
                o.append(f"      θ最小 = 2 × arccos((R - h最小) / R)")
                o.append(f"           = 2 × arccos(({Rr_m:.3f} - {h_m:.3f}) / {Rr_m:.3f})")
                o.append(f"           = 2 × arccos({(Rr_m - h_m)/Rr_m:.4f})")
                o.append(f"           = {math.degrees(theta_m):.2f}° ({theta_m:.4f} rad)")
                o.append("")
                o.append(f"  4. 过水面积计算:")
                o.append(f"      A最小 = (D²/8) × (θ最小 - sinθ最小)")
                o.append(f"           = ({D:.3f}²/8) × ({theta_m:.4f} - sin{theta_m:.4f})")
                o.append(f"           = {D**2/8:.4f} × {theta_m - math.sin(theta_m):.4f}")
                if A_m:
                    o.append(f"           = {A_m:.3f} m²")
                o.append("")
                o.append(f"  5. 湿周计算:")
                o.append(f"      χ最小 = (D/2) × θ最小")
                o.append(f"           = ({D:.3f}/2) × {theta_m:.4f}")
                o.append(f"           = {Rr_m:.3f} × {theta_m:.4f}")
                if P_m:
                    o.append(f"           = {P_m:.3f} m")
                o.append("")
            else:
                o.append(f"  3. 过水面积: A最小 = {A_m:.3f} m²" if A_m else "  3. 过水面积: A最小 = N/A")
                o.append("")
                o.append(f"  4. 湿周: χ最小 = {P_m:.3f} m" if P_m else "  4. 湿周: χ最小 = N/A")
                o.append("")

            o.append(f"  6. 水力半径计算:")
            o.append(f"      R最小 = A最小 / χ最小")
            if A_m and P_m:
                o.append(f"           = {A_m:.3f} / {P_m:.3f}")
                o.append(f"           = {R_m:.3f} m")
            o.append("")

            o.append(f"  7. 最小流速计算 (曼宁公式):")
            o.append(f"      V最小 = (1/n) × R最小^(2/3) × i^(1/2)")
            if R_m and R_m > 0:
                o.append(f"           = (1/{n}) × {R_m:.3f}^(2/3) × {i:.6f}^(1/2)")
                o.append(f"           = {1/n:.2f} × {R_m**(2/3):.4f} × {math.sqrt(i):.6f}")
            if V_m is not None:
                o.append(f"           = {V_m:.3f} m/s")
            o.append("")

            o.append(f"  8. 流量校核:")
            if V_m and A_m:
                o.append(f"      Q计算 = V最小 × A最小")
                o.append(f"           = {V_m:.3f} × {A_m:.3f}")
                o.append(f"           = {V_m * A_m:.3f} m³/s")
                if Q_min > 0:
                    o.append(f"      误差 = {abs(V_m * A_m - Q_min) / Q_min * 100:.2f}%")
            o.append("")
        else:
            o.append("  2. 最小水深: h最小 = N/A")
            o.append("  3. 最小流速: V最小 = N/A")
            o.append("")

        o.append("【六、设计验证】")
        o.append("")
        vel_ok = V_d is not None and v_min <= V_d <= v_max
        mv_ok = V_m is not None and V_m >= v_min
        o.append(f"  1. 流速验证:")
        o.append(f"      范围要求: {v_min} ≤ V ≤ {v_max} m/s")
        if V_d is not None:
            o.append(f"      设计流速: V = {V_d:.3f} m/s")
            o.append(f"      结果: {'通过 ✓' if vel_ok else '未通过 ✗'}")
        else:
            o.append(f"      计算失败")
        o.append("")
        if use_increase_circ:
            fb_ok = FB_i is not None and FB_i >= MIN_FREEBOARD
            pa_ok = PA_i is not None and PA_i >= MIN_FREE_AREA_PERCENT
            o.append(f"  2. 净空高度验证:")
            o.append(f"      规范要求: Fb ≥ {MIN_FREEBOARD} m")
            if FB_i is not None:
                o.append(f"      计算结果: Fb = {FB_i:.3f} m")
                o.append(f"      结果: {'通过 ✓' if fb_ok else '未通过 ✗'}")
            else:
                o.append(f"      计算失败")
            o.append("")
            o.append(f"  3. 净空面积验证:")
            o.append(f"      规范要求: PA ≥ {MIN_FREE_AREA_PERCENT}%")
            if PA_i is not None:
                o.append(f"      计算结果: PA = {PA_i:.1f}%")
                o.append(f"      结果: {'通过 ✓' if pa_ok else '未通过 ✗'}")
            else:
                o.append(f"      计算失败")
            o.append("")
            next_idx = 4
        else:
            fb_ok = pa_ok = True
            next_idx = 2
        o.append(f"  {next_idx}. 最小流速验证:")
        o.append(f"      规范要求: V ≥ {v_min} m/s")
        if V_m is not None:
            o.append(f"      计算结果: V = {V_m:.3f} m/s")
            o.append(f"      结果: {'通过 ✓' if mv_ok else '未通过 ✗'}")
        else:
            o.append(f"      计算失败")
        o.append("")
        all_pass = vel_ok and fb_ok and pa_ok and mv_ok
        o.append("=" * 70)
        o.append(f"  综合验证结果: {'全部通过 ✓' if all_pass else '未通过 ✗'}")
        o.append("=" * 70)
        txt = "\n".join(o)
        self._export_plain_text = txt
        self._render_result_html(plain_text_to_formula_html(txt))

    # ================================================================
    # U形 - 简要结果
    # ================================================================
    def _show_u_brief(self, result):
        p = self.input_params
        Q, n, slope_inv = p['Q'], p['n'], p['slope_inv']
        v_min, v_max = p['v_min'], p['v_max']
        R = result['R']; alpha_deg = result['alpha_deg']; theta_deg = result['theta_deg']
        m = result['m']; h0 = result['h0']; b_arc = result['b_arc']
        h = result['h_design']; V = result['V_design']; A = result['A_design']
        X = result['X_design']; Rh = result['R_design']
        inc_pct = result['increase_percent']; Q_inc = result['Q_increased']
        h_inc = result['h_increased']; V_inc = result['V_increased']
        Fb = result['Fb']; H = result['h_prime']
        increase_summary_lines = self._get_increase_summary_lines(p, result)
        o = []
        o.append("=" * 70)
        o.append("              明渠水力计算结果（U形断面）")
        o.append("=" * 70)
        o.append("")
        o.append("【输入参数】")
        o.append(f"  断面类型: U形    R = {R:.3f} m, α = {alpha_deg}°, θ = {theta_deg}°")
        o.append(f"  Q = {Q:.3f} m³/s,  n = {n},  i = 1/{int(slope_inv)}")
        o.append("")
        o.append("【断面几何参数】")
        o.append(f"  m = tan(α) = {m:.4f},  h_0 = {h0:.3f} m,  b_{{arc}} = {b_arc:.3f} m")
        o.append("")
        o.append("【设计流量工况】")
        o.append(f"  设计水深 h = {h:.3f} m")
        o.append(f"  过水面积 A = {A:.3f} m²")
        o.append(f"  湿周 χ = {X:.3f} m")
        o.append(f"  水力半径 R_h = {Rh:.3f} m")
        o.append(f"  设计流速 V = {V:.3f} m/s")
        o.append("")
        use_inc = p.get('use_increase', True)
        if use_inc:
            o.append("【加大流量工况】")
            o.extend([f"  {line}" for line in increase_summary_lines])
            if h_inc > 0:
                o.append(f"  h加大 = {h_inc:.3f} m,  V加大 = {V_inc:.3f} m/s")
                o.append(f"  超高 Fb = {Fb:.3f} m,  渠道高度 H = {H:.3f} m")
        o.append("")
        o.append("【验证结果】")
        vel_ok = v_min < V < v_max
        o.append(f"  流速: {v_min} < V={V:.3f} < {v_max} → {'✓ 通过' if vel_ok else '✗ 未通过'}")
        if use_inc and h_inc > 0:
            fb_req = 0.25 * h_inc + 0.2
            fb_ok = Fb >= (fb_req - 0.001)
            o.append(f"  超高: Fb={Fb:.3f}m ≥ {fb_req:.3f}m → {'✓ 通过' if fb_ok else '✗ 未通过'}")
            all_pass = vel_ok and fb_ok
        else:
            all_pass = vel_ok
        o.append("=" * 70)
        o.append(f"  综合验证: {'全部通过 ✓' if all_pass else '未通过 ✗'}")
        o.append("=" * 70)
        txt = "\n".join(o)
        self._export_plain_text = txt
        self._render_result_html(plain_text_to_formula_html(txt))

    # ================================================================
    # U形 - 详细结果
    # ================================================================
    def _show_u_detail(self, result):
        p = self.input_params
        Q, n, slope_inv = p['Q'], p['n'], p['slope_inv']
        i = 1.0 / slope_inv
        v_min, v_max = p['v_min'], p['v_max']
        R = result['R']; alpha_deg = result['alpha_deg']; theta_deg = result['theta_deg']
        m = result['m']; h0 = result['h0']; b_arc = result['b_arc']
        h = result['h_design']; V = result['V_design']; A = result['A_design']
        X = result['X_design']; Rh = result['R_design']; Q_calc = result['Q_calc']
        inc_pct = result['increase_percent']; Q_inc = result['Q_increased']
        h_inc = result['h_increased']; V_inc = result['V_increased']
        A_inc = result.get('A_increased', -1); X_inc = result.get('X_increased', -1)
        R_inc = result.get('R_increased', -1)
        Fb = result['Fb']; H = result['h_prime']
        increase_summary_lines = self._get_increase_summary_lines(p, result)
        theta_rad = math.radians(theta_deg)
        o = []
        o.append("=" * 70)
        o.append("              明渠水力计算结果（U形断面）")
        o.append("=" * 70)
        o.append("")
        o.append("【一、输入参数】")
        o.append(f"  断面类型: U形明渠")
        o.append(f"  设计流量 Q = {Q:.3f} m³/s")
        o.append(f"  圆弧半径 R = {R:.3f} m")
        o.append(f"  直线段外倾角 α = {alpha_deg}°")
        o.append(f"  圆弧段圆心角 θ = {theta_deg}°")
        o.append(f"  糙率 n = {n}")
        o.append(f"  水力坡降 i = 1/{int(slope_inv)} = {i:.6f}")
        o.append(f"  不淤流速 = {v_min} m/s,  不冲流速 = {v_max} m/s")
        o.append("")
        o.append("【二、断面几何参数】")
        o.append(f"  m = tan(α) = tan({alpha_deg}°) = {m:.6f}")
        o.append(f"  h_0 = R·(1-cos(θ/2)) = {R:.3f}×(1-cos({theta_deg/2:.1f}°)) = {h0:.3f} m")
        o.append(f"  b_{{arc}} = 2·R·sin(θ/2) = 2×{R:.3f}×sin({theta_deg/2:.1f}°) = {b_arc:.3f} m")
        o.append("")
        o.append("【三、设计水深计算】")
        o.append(f"  根据Q={Q:.3f} m³/s，曼宁公式二分法反算水深: h = {h:.3f} m")
        o.append(f"  水深区间: h {'≤' if h <= h0 else '>'} h_0={h0:.3f} m → {'纯弧区' if h <= h0 else '直线段区'}")
        o.append("")
        if h <= h0:
            cos_arg = max(-1.0, min(1.0, (R - h) / R))
            acos_val = math.acos(cos_arg)
            sqrt_val = math.sqrt(max(0.0, R * R - (R - h) ** 2))
            o.append("  【纯弧区公式】")
            o.append(f"  过水面积 A = R²·arccos((R-h)/R) - (R-h)·√(R²-(R-h)²)")
            o.append(f"           = {R:.3f}²×{acos_val:.4f} - {R-h:.3f}×{sqrt_val:.4f}")
            o.append(f"           = {A:.3f} m²")
            o.append(f"  湿周 χ = 2·R·arccos((R-h)/R) = 2×{R:.3f}×{acos_val:.4f} = {X:.3f} m")
        else:
            h_s = h - h0
            A_arc = R * R * (theta_rad / 2.0 - math.sin(theta_rad / 2.0) * math.cos(theta_rad / 2.0))
            o.append("  【直线段区公式】")
            o.append(f"  弧面积 A_{{arc}} = R²·(θ/2-sin(θ/2)·cos(θ/2)) = {A_arc:.4f} m²")
            o.append(f"  h_s = h - h_0 = {h:.3f} - {h0:.3f} = {h_s:.3f} m")
            o.append("  过水面积 A = A_{arc} + (b_{arc} + m·h_s)·h_s")
            o.append(f"           = {A_arc:.4f} + ({b_arc:.3f}+{m:.4f}×{h_s:.3f})×{h_s:.3f}")
            o.append(f"           = {A:.3f} m²")
            chi_arc = theta_rad * R
            o.append(f"  湿周 χ = θ·R + 2·h_s·√(1+m²)")
            o.append(f"       = {theta_rad:.4f}×{R:.3f} + 2×{h_s:.3f}×√(1+{m:.4f}²)")
            o.append(f"       = {X:.3f} m")
        o.append(f"  水力半径 R_h = A/χ = {A:.3f}/{X:.3f} = {Rh:.3f} m")
        o.append(f"  设计流速 V = (1/n)·R_h^(2/3)·i^(1/2)")
        o.append(f"           = (1/{n})×{Rh:.3f}^(2/3)×{i:.6f}^(1/2) = {V:.3f} m/s")
        o.append(f"  流量校核 Q计算 = {V:.3f}×{A:.3f} = {Q_calc:.3f} m³/s (误差{abs(Q_calc-Q)/Q*100:.2f}%)")
        o.append("")
        use_inc = p.get('use_increase', True)
        if use_inc:
            o.append("【四、加大流量工况】")
            o.extend([f"  {line}" for line in increase_summary_lines])
            for formula_line in build_increase_formula_lines(
                design_q=Q,
                increase_percent=inc_pct,
                q_increased=Q_inc,
            ):
                o.append(f"  {formula_line}")
            if h_inc > 0:
                o.append(f"  h加大 = {h_inc:.3f} m")
                if A_inc > 0 and X_inc > 0:
                    o.append(f"  A加大 = {A_inc:.3f} m²,  χ加大 = {X_inc:.3f} m,  R加大 = {R_inc:.3f} m")
                o.append(f"  V加大 = {V_inc:.3f} m/s")
                o.append(f"  超高 Fb = 0.25×{h_inc:.3f}+0.2 = {Fb:.3f} m")
                o.append(f"  渠道高度 H = {h_inc:.3f}+{Fb:.3f} = {H:.3f} m")
            else:
                o.append("  加大水深计算失败")
        o.append("")
        o.append("【五、设计验证】")
        vel_ok = v_min < V < v_max
        o.append(f"  流速: {v_min} < V={V:.3f} < {v_max} → {'通过 ✓' if vel_ok else '未通过 ✗'}")
        if use_inc and h_inc > 0:
            fb_req = 0.25 * h_inc + 0.2
            fb_ok = Fb >= (fb_req - 0.001)
            o.append(f"  超高: Fb={Fb:.3f}m ≥ {fb_req:.3f}m → {'通过 ✓' if fb_ok else '未通过 ✗'}")
            all_pass = vel_ok and fb_ok
        else:
            all_pass = vel_ok
        o.append("=" * 70)
        o.append(f"  综合验证: {'全部通过 ✓' if all_pass else '未通过 ✗'}")
        o.append("=" * 70)
        txt = "\n".join(o)
        self._export_plain_text = txt
        self._render_result_html(plain_text_to_formula_html(txt))

    # ================================================================
    # 附录E HTML表格
    # ================================================================
    def _build_ae_html(self, schemes, sel_b, sel_h, v_min, v_max):
        h = '<table class="ae"><tr>'
        for hdr in ['α值', '方案类型', '底宽B(m)', '水深h(m)', '宽深比β', '流速V(m/s)', '面积增加', '状态']:
            h += f'<th>{hdr}</th>'
        h += '</tr>'
        for idx, s in enumerate(schemes):
            alpha, stype = s['alpha'], s['scheme_type']
            sb, sh, sbeta, sV = s['b'], s['h'], s['beta'], s['V']
            area_inc = s['area_increase']
            is_sel = abs(sb - sel_b) < 0.01 and abs(sh - sel_h) < 0.01
            v_ok = v_min < sV < v_max
            if is_sel: cls = "sel"; status = "★ 选中"
            elif not v_ok: cls = "err"; status = "流速不符"
            else: cls = "even" if idx % 2 == 0 else "odd"; status = ""
            h += f'<tr class="{cls}">'
            h += f'<td>{alpha:.2f}</td><td>{stype}</td><td>{sb:.3f}</td><td>{sh:.3f}</td>'
            h += f'<td>{sbeta:.3f}</td><td>{sV:.3f}</td><td>+{area_inc:.0f}%</td>'
            h += f'<td><b>{status}</b></td></tr>'
        h += '</table>'
        return h

    def _build_ae_text(self, schemes, sel_b, sel_h, v_min, v_max):
        lines = []
        lines.append("【附录E断面方案对比表】")
        lines.append("  说明: α=1.00为水力最佳断面(深窄)，α越大断面越宽浅，面积增加但流速降低")
        lines.append("")
        lines.append("  α值    方案类型        底宽B(m)  水深h(m)  宽深比β   流速V(m/s)  面积+  状态")
        lines.append("  " + "-" * 78)
        for s in schemes:
            alpha, stype = s['alpha'], s['scheme_type']
            sb, sh, sbeta, sV = s['b'], s['h'], s['beta'], s['V']
            area_inc = s['area_increase']
            is_sel = abs(sb - sel_b) < 0.01 and abs(sh - sel_h) < 0.01
            v_ok = v_min < sV < v_max
            status = "★选中" if is_sel else ("流速不符" if not v_ok else "")
            lines.append(f"  {alpha:.2f}   {stype:<12}  {sb:8.3f}  {sh:8.3f}  {sbeta:8.3f}  {sV:10.3f}  +{area_inc:.0f}%   {status}")
        lines.append("")
        lines.append(f"  注: 流速约束范围 {v_min} ~ {v_max} m/s")
        return "\n".join(lines)

    # ================================================================
    # 断面图
    # ================================================================
    def _update_section_plot(self, result):
        self.section_fig.clear()
        if not result['success']:
            self.section_canvas.draw(); return
        stype = self.input_params.get('section_type', '梯形')
        ax = self.section_fig.add_subplot(111)
        if stype == '圆形':
            D = result.get('D_design', 0)
            y_d = result.get('y_d', 0); V_d = result.get('V_d', 0)
            Q = self.input_params['Q']
            y_i, _, _ = OpenChannelPanel._increase_plot_values(self.input_params, result)
            self._draw_circular(ax, D, y_d, V_d, Q, '设计流量')
            OpenChannelPanel._draw_increase_water_level(self, ax, stype, self.input_params, result, y_i)
        elif stype == '复式梯形':
            m1 = self.input_params.get('m1', 0)
            B1 = self.input_params.get('B1', 0)
            m2 = self.input_params.get('m2', 0)
            B2 = result.get('b_design', self.input_params.get('B2', 0))
            m3 = self.input_params.get('m3', 0)
            h1 = self.input_params.get('h1', 0)
            h_w = result['h_design']
            H_ch = result['h_prime'] if result['h_prime'] > 0 else h_w * 1.35
            V = result['V_design']; Q = self.input_params['Q']
            h_inc, _, _ = OpenChannelPanel._increase_plot_values(self.input_params, result)
            self._draw_compound_trapezoid(ax, B2, m1, B1, m2, m3, h1, H_ch, V, Q, h_w, "设计流量")
            OpenChannelPanel._draw_increase_water_level(self, ax, stype, self.input_params, result, h_inc)
        elif stype == 'U形':
            R = result['R']; alpha_deg = result['alpha_deg']; theta_deg = result['theta_deg']
            h_w = result['h_design']
            H_ch = result['h_prime'] if result['h_prime'] > 0 else h_w * 1.35
            V = result['V_design']; Q = self.input_params['Q']
            h_inc, _, _ = OpenChannelPanel._increase_plot_values(self.input_params, result)
            self._draw_u_section(ax, R, alpha_deg, theta_deg, h_w, H_ch, V, Q, '设计流量')
            OpenChannelPanel._draw_increase_water_level(self, ax, stype, self.input_params, result, h_inc)
        else:
            b = result['b_design']; h = result['h_design']
            m = self.input_params.get('m', 0); Q = self.input_params['Q']
            V = result['V_design']; h_inc, _, _ = OpenChannelPanel._increase_plot_values(self.input_params, result)
            h_prime = result['h_prime']
            use_inc = self.input_params.get('use_increase', True)
            if use_inc and h_inc > 0 and h_prime > 0:
                H_d = h_prime
            else:
                Fb_d = 0.25 * h + 0.2
                H_d = h + Fb_d
            self._draw_trapezoid(ax, b, H_d, m, V, Q, h, "设计流量")
            OpenChannelPanel._draw_increase_water_level(self, ax, stype, self.input_params, result, h_inc)
        self.section_fig.tight_layout()
        self.section_canvas.draw()

    def _draw_trapezoid(self, ax, b, h_ch, m, V, Q, h_w, title):
        """用共享绘图器绘制梯形或矩形明渠断面。"""
        shape = build_trapezoid_shape(b, h_ch, m)
        draw_section(ax, shape, WaterState(depth=h_w, flow=Q, velocity=V), title)

    def _compound_trapezoid_geometry(self, B2, m1, B1, m2, m3, h1, h_ch):
        """生成复式梯形断面的关键轮廓点。"""
        left_break_x = -(B2 / 2 + m2 * h1)
        left_platform_x = left_break_x - B1
        left_top_x = left_platform_x - m1 * max(h_ch - h1, 0.0)
        right_top_x = B2 / 2 + m3 * h_ch
        outline_points = [
            (-B2 / 2, 0.0),
            (B2 / 2, 0.0),
            (right_top_x, h_ch),
            (left_top_x, h_ch),
            (left_platform_x, h1),
            (left_break_x, h1),
        ]
        return {
            "bottom_left": outline_points[0],
            "bottom_right": outline_points[1],
            "right_top": outline_points[2],
            "left_top": outline_points[3],
            "left_platform": outline_points[4],
            "left_break": outline_points[5],
            "outline_points": outline_points,
            "width_ref": max(right_top_x - left_top_x, 1.0),
        }

    def _compound_trapezoid_water_points(self, geometry, B2, m1, m2, m3, h1, h_w):
        """按水深生成真实的复式梯形过水轮廓。"""
        if h_w <= 0:
            return []

        bottom_left = geometry["bottom_left"]
        bottom_right = geometry["bottom_right"]
        left_platform = geometry["left_platform"]
        left_break = geometry["left_break"]

        if h_w <= h1:
            left_water = -(B2 / 2 + m2 * h_w)
            right_water = B2 / 2 + m3 * h_w
            return [
                bottom_left,
                bottom_right,
                (right_water, h_w),
                (left_water, h_w),
            ]

        hs = h_w - h1
        left_water = left_platform[0] - m1 * hs
        right_water = B2 / 2 + m3 * h_w
        return [
            bottom_left,
            bottom_right,
            (right_water, h_w),
            (left_water, h_w),
            left_platform,
            left_break,
        ]

    def _compound_trapezoid_view_limits(self, geometry, h_ch, h1, h_w):
        """根据轮廓和尺寸箭头锚点收紧默认视图范围。"""
        width_ref = geometry["width_ref"]
        left_top_x = geometry["left_top"][0]
        right_top_x = geometry["right_top"][0]
        left_platform_x = geometry["left_platform"][0]
        left_break_x = geometry["left_break"][0]
        dim_offset = max(h_ch * 0.12, 0.15)

        x_candidates = [
            point[0] for point in geometry["outline_points"]
        ] + [
            -geometry["bottom_right"][0],
            geometry["bottom_right"][0],
            left_platform_x,
            left_break_x,
            left_top_x - width_ref * 0.08,
            right_top_x + width_ref * 0.08,
        ]
        if h_w > 0:
            x_candidates.append(left_top_x - width_ref * 0.18)

        y_candidates = [
            point[1] for point in geometry["outline_points"]
        ] + [
            -dim_offset,
            h1 + dim_offset * 0.6,
            h_ch,
        ]
        if h_w > 0:
            y_candidates.append(h_w)

        x_padding = max(width_ref * 0.04, 0.2)
        y_padding = max(h_ch * 0.08, 0.18)
        return (
            min(x_candidates) - x_padding,
            max(x_candidates) + x_padding,
            min(y_candidates) - y_padding,
            max(y_candidates) + y_padding,
        )

    def _draw_compound_trapezoid(self, ax, B2, m1, B1, m2, m3, h1, h_ch, V, Q, h_w, title):
        """用共享绘图器绘制复式梯形明渠断面。"""
        shape = build_compound_trapezoid_shape(B2, m1, B1, m2, m3, h1, h_ch)
        draw_section(ax, shape, WaterState(depth=h_w, flow=Q, velocity=V), title)

    def _draw_u_section(self, ax, R, alpha_deg, theta_deg, h_w, H_ch, V, Q, title):
        """用共享绘图器绘制 U 形明渠断面。"""
        shape = build_open_channel_u_shape(R, alpha_deg, theta_deg, H_ch)
        draw_section(ax, shape, WaterState(depth=h_w, flow=Q, velocity=V), title)

    def _draw_circular(self, ax, D, y, V, Q, title):
        """用共享绘图器绘制圆形明渠断面。"""
        shape = build_circular_shape(D, water_label="y")
        draw_section(ax, shape, WaterState(depth=y, flow=Q, velocity=V), title)

    # ================================================================
    # 清空
    # ================================================================
    def _clear(self):
        self._save_current_case()
        self._all_results = []
        self._results_dirty = False
        self._stale_result_case_indexes = set()
        self._all_results_stale = False
        self._has_rendered_results = False
        self._rebuild_case_tags()
        self._update_calc_btn_text()
        self._show_initial_help()
        self.section_fig.clear()
        self.section_canvas.draw()
        self._refresh_increase_hint()
        self.current_result = None
        self._export_plain_text = ""
        self._clear_comparison_tables()

    # ================================================================
    # 导出
    # ================================================================
    def _export_dxf(self):
        case_entries = self._build_dxf_export_case_entries()
        current_entry = self._get_current_dxf_export_entry(case_entries)
        if len(self._cases) <= 1:
            if current_entry is None or not current_entry.is_valid:
                self._warn_single_dxf_entry_unavailable(current_entry)
                return
            try:
                filepath = self._export_single_dxf_entry(current_entry)
                if not filepath:
                    return
                InfoBar.success("导出成功", f"DXF已保存到: {filepath}", parent=self._info_parent(), duration=4000, position=InfoBarPosition.TOP)
                ask_open_file(filepath, self._info_parent())
            except ImportError as e:
                InfoBar.error("缺少依赖", str(e), parent=self._info_parent(), duration=6000, position=InfoBarPosition.TOP)
            except PermissionError:
                InfoBar.error("文件被占用", "无法写入文件，请先关闭已打开的同名DXF文件，然后重新操作。", parent=self._info_parent(), duration=8000, position=InfoBarPosition.TOP)
            except Exception as e:
                InfoBar.error("导出失败", f"DXF导出失败: {str(e)}", parent=self._info_parent(), duration=5000, position=InfoBarPosition.TOP)
            return

        dialog_result = show_multi_case_dxf_dialog(
            self._info_parent(),
            "明渠断面",
            case_entries,
            self._current_case_idx,
        )
        if dialog_result is None:
            return
        selected_entries = select_case_entries(
            case_entries,
            dialog_result.scope,
            self._current_case_idx,
            dialog_result.checked_case_indexes,
        )
        valid_entries, invalid_entries = partition_valid_case_entries(selected_entries)
        if not valid_entries:
            InfoBar.warning("提示", format_empty_export_warning(invalid_entries), parent=self._info_parent(), duration=4000, position=InfoBarPosition.TOP)
            return
        try:
            if len(valid_entries) == 1:
                filepath = self._export_single_dxf_entry(valid_entries[0], scale_denom=dialog_result.scale_denom)
            else:
                filepath = self._export_combined_dxf_entries(valid_entries, dialog_result.scale_denom)
            if not filepath:
                return
            InfoBar.success(
                "导出成功",
                f"{format_export_result_message(len(valid_entries), invalid_entries)}\n文件：{filepath}",
                parent=self._info_parent(),
                duration=5000,
                position=InfoBarPosition.TOP,
            )
            ask_open_file(filepath, self._info_parent())
        except ImportError as e:
            InfoBar.error("缺少依赖", str(e), parent=self._info_parent(), duration=6000, position=InfoBarPosition.TOP)
        except PermissionError:
            InfoBar.error("文件被占用", "无法写入文件，请先关闭已打开的同名DXF文件，然后重新操作。", parent=self._info_parent(), duration=8000, position=InfoBarPosition.TOP)
        except Exception as e:
            InfoBar.error("导出失败", f"DXF导出失败: {str(e)}", parent=self._info_parent(), duration=5000, position=InfoBarPosition.TOP)

    def _build_dxf_export_case_entries(self):
        entries = []
        results_dirty = bool(getattr(self, "_results_dirty", False))
        result_map = {case_idx: (params, result) for case_idx, params, result in self._all_results}
        for case_idx, case in enumerate(self._cases):
            params, result = result_map.get(case_idx, ({}, None))
            invalid_reason = None
            if results_dirty and self._all_results:
                invalid_reason = "结果已失效"
            elif result is None:
                invalid_reason = "无计算结果"
            elif not result.get("success"):
                invalid_reason = "计算失败"
            entries.append(
                DxfExportCaseEntry(
                    case_idx=case_idx,
                    label=self._case_label(case, case_idx),
                    input_params=params or {},
                    result=result,
                    is_valid=invalid_reason is None,
                    invalid_reason=invalid_reason,
                )
            )
        return entries

    def _get_current_dxf_export_entry(self, case_entries):
        for entry in case_entries:
            if entry.case_idx == self._current_case_idx:
                return entry
        return case_entries[0] if case_entries else None

    def _warn_single_dxf_entry_unavailable(self, entry):
        content = "参数已变更，请先重新计算后再导出。" if entry is not None and entry.invalid_reason == "结果已失效" else "请先进行计算后再导出。"
        InfoBar.warning("提示", content, parent=self._info_parent(), duration=3000, position=InfoBarPosition.TOP)

    def _single_dxf_default_name(self, entry):
        params = entry.input_params or {}
        result = entry.result or {}
        stype = params.get('section_type', '梯形')
        if stype == '圆形':
            return f"明渠断面_圆形_D{result.get('D_design', 0.0):.2f}.dxf"
        if stype == 'U形':
            return f"明渠断面_U形_R{result.get('R', 0.0):.2f}.dxf"
        h_val = result.get('h_prime', 0.0)
        if h_val <= 0:
            h_inc = result.get('h_increased', 0.0)
            h_val = (h_inc + result.get('Fb', 0.3)) if h_inc > 0 else result.get('h_design', 0.0) * 1.35
        return f"明渠断面_{stype}_B{result.get('b_design', 0.0):.2f}xH{h_val:.2f}.dxf"

    def _combined_dxf_default_name(self, count):
        return f"明渠断面_{count}个工况_合并.dxf"

    def _choose_dxf_filepath(self, default_name):
        filepath, _ = QFileDialog.getSaveFileName(self, "保存DXF文件", default_name, "DXF文件 (*.dxf);;所有文件 (*.*)")
        if not filepath:
            return None
        return filepath if filepath.lower().endswith(".dxf") else f"{filepath}.dxf"

    def _export_single_dxf_entry(self, entry, scale_denom=None):
        scale = scale_denom if scale_denom is not None else choose_scale_denom(self)
        if scale is None:
            return None
        filepath = self._choose_dxf_filepath(self._single_dxf_default_name(entry))
        if not filepath:
            return None
        export_single_case_dxf(
            filepath,
            entry,
            scale,
            draw_open_channel_dxf_on_msp,
            draw_summary_table=draw_open_channel_comparison_table,
        )
        return filepath

    def _export_combined_dxf_entries(self, entries, scale_denom):
        filepath = self._choose_dxf_filepath(self._combined_dxf_default_name(len(entries)))
        if not filepath:
            return None
        return export_combined_case_dxf(
            filepath,
            entries,
            scale_denom,
            draw_open_channel_dxf_on_msp,
            draw_summary_table=draw_open_channel_comparison_table,
        )

    def _export_report(self):
        if not self.current_result or not self.current_result.get('success'):
            InfoBar.warning("提示", "请先进行计算后再导出。", parent=self._info_parent(), duration=3000, position=InfoBarPosition.TOP)
            return
        filepath, _ = QFileDialog.getSaveFileName(self, "保存报告", "", "文本文件 (*.txt);;所有文件 (*.*)")
        if not filepath: return
        try:
            content = self._export_plain_text if self._export_plain_text else ''
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(content)
            InfoBar.success("导出成功", f"报告已保存到: {filepath}", parent=self._info_parent(), duration=4000, position=InfoBarPosition.TOP)
            ask_open_file(filepath, self._info_parent())
        except PermissionError:
            InfoBar.error("文件被占用", "无法写入文件，请先关闭已打开的同名文件（如记事本等），然后重新操作。", parent=self._info_parent(), duration=8000, position=InfoBarPosition.TOP)
        except Exception as e:
            InfoBar.error("导出失败", f"保存失败: {str(e)}", parent=self._info_parent(), duration=5000, position=InfoBarPosition.TOP)

    def _export_word(self):
        if not WORD_EXPORT_AVAILABLE:
            InfoBar.warning("缺少依赖",
                "Word导出需要安装 python-docx、latex2mathml、lxml。请执行: pip install python-docx latex2mathml lxml",
                parent=self._info_parent(), duration=6000, position=InfoBarPosition.TOP)
            return
        if not self._all_results:
            InfoBar.warning("提示", "请先进行计算后再导出。", parent=self._info_parent(), duration=3000, position=InfoBarPosition.TOP)
            return
        p = self.input_params
        stype = p.get('section_type', '梯形')
        channel_name = p.get('channel_name', '') or getattr(self, '_channel_name_text', '')
        meta = load_meta()
        auto_purpose = build_calc_purpose('open_channel',
            project=meta.project_name, name=channel_name, section_type=stype)
        n_cases = len(self._all_results)
        current_label = self._auto_label(self._cases[self._current_case_idx], self._current_case_idx) if self._cases else '工况1'
        dlg = ExportConfirmDialog('open_channel', '明渠水力计算书', auto_purpose,
                                  parent=self._info_parent(),
                                  n_cases=n_cases, current_case_label=current_label)
        from PySide6.QtWidgets import QDialog
        if dlg.exec() != QDialog.Accepted:
            return
        self._word_export_meta = dlg.get_meta()
        self._word_export_purpose = dlg.get_calc_purpose()
        self._word_export_refs = dlg.get_references()
        self._word_export_scope = dlg.get_export_scope() if n_cases > 1 else 'all'
        filepath, _ = QFileDialog.getSaveFileName(self, "保存Word报告", "", "Word文档 (*.docx);;所有文件 (*.*)")
        if not filepath: return
        try:
            self._build_word_report(filepath)
            InfoBar.success("导出成功", f"Word报告已保存到: {filepath}", parent=self._info_parent(), duration=4000, position=InfoBarPosition.TOP)
            ask_open_file(filepath, self._info_parent())
        except PermissionError:
            InfoBar.error("文件被占用", "无法写入文件，请先关闭已打开的同名Word文档，然后重新操作。", parent=self._info_parent(), duration=8000, position=InfoBarPosition.TOP)
        except Exception as e:
            InfoBar.error("导出失败", f"Word导出失败: {str(e)}", parent=self._info_parent(), duration=5000, position=InfoBarPosition.TOP)

    def _build_word_report(self, filepath):
        """构建Word报告文档（工程产品运行卡格式），支持多工况"""
        meta = getattr(self, '_word_export_meta', load_meta())
        purpose = getattr(self, '_word_export_purpose', '')
        refs = getattr(self, '_word_export_refs', REFERENCES_BASE.get('open_channel', []))
        scope = getattr(self, '_word_export_scope', 'all')

        # 确定要导出的工况
        if scope == 'current':
            export_results = [(ci, p, r) for ci, p, r in self._all_results if ci == self._current_case_idx]
        else:
            export_results = list(self._all_results)

        n_export = len(export_results)
        first_stype = export_results[0][1].get('section_type', '梯形') if export_results else '梯形'
        first_method = export_results[0][2].get('design_method', '') if export_results else ''
        if n_export == 1:
            content_desc = f'明渠水力断面设计计算（{first_stype}断面）'
        else:
            content_desc = f'明渠水力断面设计计算（{n_export}个工况）'

        doc = create_engineering_report_doc(
            meta=meta,
            calc_title='明渠水力计算书',
            calc_content_desc=content_desc,
            calc_purpose=purpose,
            references=refs,
            calc_program_text=f'渠系建筑物水力计算系统 V1.0\n{content_desc}',
        )
        doc.add_page_break()

        # 5. 基础公式
        doc_add_eng_h(doc, '5、基础公式')
        doc_add_formula(doc, r'Q = \frac{1}{n} \cdot A \cdot R^{2/3} \cdot i^{1/2}', '曼宁公式：')
        # 根据各工况断面类型添加公式
        stypes_used = set(p.get('section_type', '梯形') for _, p, _ in export_results)
        if '梯形' in stypes_used:
            doc_add_formula(doc, r'A = (B + m \cdot h) \cdot h', '梯形过水面积：')
            doc_add_formula(doc, r'\chi = B + 2h\sqrt{1+m^2}', '梯形湿周：')
        if '复式梯形' in stypes_used:
            doc_add_formula(doc, r'A = B_2 h + \frac{m_2+m_3}{2} h^2', '复式梯形平台以下面积：')
            doc_add_formula(doc, r'\chi = B_2 + h\sqrt{1+m_2^2} + h\sqrt{1+m_3^2}', '复式梯形平台以下湿周：')
            doc_add_formula(doc, r'A = A_1 + W_1 h_s + \frac{m_1+m_3}{2} h_s^2', '复式梯形越平台面积：')
            doc_add_formula(doc, r'\chi = B_2 + h_1\sqrt{1+m_2^2} + B_1 + h_s\sqrt{1+m_1^2} + h\sqrt{1+m_3^2}', '复式梯形越平台湿周：')
        if '矩形' in stypes_used:
            doc_add_formula(doc, r'A = B \cdot h', '矩形过水面积：')
            doc_add_formula(doc, r'\chi = B + 2h', '矩形湿周：')
        if '圆形' in stypes_used:
            doc_add_formula(doc, r'A = \frac{D^2}{8}(\theta - \sin\theta)', '圆形过水面积：')
            doc_add_formula(doc, r'\chi = \frac{D}{2} \cdot \theta', '圆形湿周：')
        if 'U形' in stypes_used:
            doc_add_formula(doc, r'h_0 = R(1-\cos(\theta/2))', '弧区临界水深：')
        doc_add_formula(doc, r'R = \frac{A}{\chi}', '水力半径：')
        doc_add_formula(doc, r'V = \frac{1}{n} \cdot R^{2/3} \cdot i^{1/2}', '流速公式：')

        # 6. 计算过程
        doc_add_eng_h(doc, '6、计算过程')
        _multi = n_export > 1
        add_section_comparison_word_tables(
            doc,
            export_results,
            OPEN_CHANNEL_COMPARISON_SPEC,
            heading_func=doc_add_eng_h,
            table_func=doc_add_styled_table,
        )

        for ri, (case_idx, params, result) in enumerate(export_results):
            if not result.get('success'):
                doc_add_eng_body(doc, f'工况{case_idx+1}: 计算失败 - {result.get("error_message", "未知错误")}')
                continue

            stype = params.get('section_type', '梯形')

            if _multi:
                doc_add_eng_h(doc, f'6.{ri+1}、工况{case_idx+1} ({stype}断面, Q={params["Q"]:.3f} m³/s)')

            # 临时设置 input_params 以生成文本
            self.input_params = params
            self.current_result = result
            self._update_result_display(result)
            calc_text = self._export_plain_text or ''

            summary_items = build_result_summary_word_items("open_channel", params, result)
            if summary_items:
                doc_add_eng_h(doc, '重点结果汇总')
                doc_add_result_table(doc, summary_items)

            doc_render_calc_text_eng(doc, calc_text, skip_title_keyword='明渠水力计算结果')

            # 附录E断面方案对比
            schemes = result.get('appendix_e_schemes', [])
            if schemes and stype != '圆形':
                ae_section = f'6.{ri+1}.1' if _multi else '7'
                doc_add_eng_h(doc, f'{ae_section}、断面方案对比（附录E）')
                doc_add_table_caption(doc, '表 1  附录E断面方案对比表')
                b_sel = result['b_design']; h_sel = result['h_design']
                v_min, v_max = params['v_min'], params['v_max']
                headers = ['α值', '方案类型', '底宽B(m)', '水深h(m)', '宽深比β', '流速V(m/s)', '面积增加', '状态']
                data = []
                for s in schemes:
                    is_sel = abs(s['b'] - b_sel) < 0.01 and abs(s['h'] - h_sel) < 0.01
                    v_ok = v_min < s['V'] < v_max
                    status = '★选中' if is_sel else ('流速不符' if not v_ok else '')
                    data.append([
                        f"{s['alpha']:.2f}", s['scheme_type'],
                        f"{s['b']:.3f}", f"{s['h']:.3f}", f"{s['beta']:.3f}",
                        f"{s['V']:.3f}", f"+{s['area_increase']:.0f}%", status
                    ])
                doc_add_styled_table(doc, headers, data, highlight_col=7, highlight_val='★选中',
                                     with_full_border=True)

        # 恢复
        _, first_p, first_r = export_results[0] if export_results else self._all_results[0]
        self.input_params = first_p
        self.current_result = first_r

        # 断面图
        try:
            import tempfile
            tmp = os.path.join(tempfile.gettempdir(), '_mingqu_section.png')
            self.section_fig.savefig(tmp, dpi=150, bbox_inches='tight')
            section_no = '7' if not _multi else str(6 + 1)
            doc_add_eng_h(doc, f'{section_no}、断面图')
            doc_add_figure(doc, tmp, width_cm=14)
            os.remove(tmp)
        except Exception:
            pass

        doc.save(filepath)

    # ================================================================
    # 项目序列化
    # ================================================================
    def to_project_dict(self):
        self._save_current_case()
        return {
            'cases': copy.deepcopy(self._cases),
            'current_case_idx': self._current_case_idx,
            'all_results': copy.deepcopy(self._all_results),
            'current_result': copy.deepcopy(self.current_result),
            'input_params': copy.deepcopy(getattr(self, 'input_params', None)),
            'result_state': collect_case_result_state(self),
            'notebook_idx': self.notebook.currentIndex() if hasattr(self, 'notebook') else 0,
        }

    def from_project_dict(self, data):
        cases = data.get('cases')
        if not cases or not isinstance(cases, list):
            return
        self._cases = cases
        self._current_case_idx = min(data.get('current_case_idx', 0), len(self._cases) - 1)
        self._load_case(self._current_case_idx)
        self._rebuild_case_tags()
        self._update_calc_btn_text()
        self._all_results = data.get('all_results', []) or []
        self.current_result = data.get('current_result')
        self.input_params = data.get('input_params') or {}
        result_state = data.get('result_state')
        if self._all_results:
            try:
                self._suppress_project_restore_side_effects = True
                self._display_all_results()
            except Exception:
                self._all_results = []
                self.current_result = None
                self._results_dirty = False
                self._stale_result_case_indexes = set()
                self._all_results_stale = False
                self._has_rendered_results = False
                clear_section_plot_state(self)
                self._clear_comparison_tables()
                self._show_initial_help()
            finally:
                self._suppress_project_restore_side_effects = False
            if self._all_results:
                apply_case_result_state(self, result_state)
                try:
                    self._update_section_plot_all()
                except Exception:
                    clear_section_plot_state(self)
                try:
                    self._refresh_comparison_tables()
                except Exception:
                    self._clear_comparison_tables()
        else:
            self._all_results = []
            self.current_result = None
            self._results_dirty = False
            self._stale_result_case_indexes = set()
            self._all_results_stale = False
            self._has_rendered_results = False
            clear_section_plot_state(self)
            self._clear_comparison_tables()
            self._show_initial_help()
        if self._all_results:
            apply_case_result_state(self, result_state)
        if hasattr(self, 'notebook'):
            idx = data.get('notebook_idx')
            if isinstance(idx, int):
                idx = max(0, min(idx, self.notebook.count() - 1))
                self.notebook.setCurrentIndex(idx)
        if self._all_results and hasattr(self, 'notebook'):
            schedule_section_plot_restore_refresh(self)
