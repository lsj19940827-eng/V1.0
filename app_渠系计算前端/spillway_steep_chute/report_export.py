# -*- coding: utf-8 -*-
"""泄水渠与陡坡文档和表格导出函数。"""

from pathlib import Path
from types import SimpleNamespace
from typing import Any

from docx import Document
from openpyxl import Workbook

from app_渠系计算前端.export_utils import create_engineering_report_doc
from app_渠系计算前端.report_meta import load_meta

from .models import normalize_result
from .principles import build_calculation_principles, display_start_control_source


def _text(value: Any) -> str:
    """把导出值统一为文本。"""
    if value is None:
        return ""
    return str(value)


def _mapping(value: Any) -> dict[str, Any]:
    """把字典或对象整理成字典。"""
    if isinstance(value, dict):
        return value
    if hasattr(value, "__dict__"):
        return dict(value.__dict__)
    return {}


def _list(value: Any) -> list[Any]:
    """把可能为空的结果整理成列表。"""
    if value is None:
        return []
    if isinstance(value, list):
        return value
    if isinstance(value, tuple):
        return list(value)
    return [value]


def _value(*candidates: Any, default: Any = "") -> Any:
    """按顺序读取第一个非空值，保留零值。"""
    for value in candidates:
        if value not in (None, ""):
            return value
    return default


def _word_meta(meta: Any = None) -> Any:
    """整理 Word 固定格式需要的项目元数据。"""
    source = meta if meta is not None else load_meta()
    fields = {
        "project_name": "",
        "design_stage": "施工图",
        "product_level": "三级",
        "record_number": "",
        "specialty": "水工",
        "calculator": "",
        "checker": "",
        "reviewer": "",
        "approver": "",
        "volume_current": "",
        "volume_total": "",
        "basic_info": "",
        "mandatory_clause": "无",
        "extra_references": [],
    }
    for key in fields:
        fields[key] = getattr(source, key, fields[key])
    return SimpleNamespace(**fields)


def _word_content_desc(case_count: int) -> str:
    """生成泄水渠计算书产品内容说明。"""
    if case_count > 1:
        return f"泄水渠与陡坡水力设计计算（{case_count}个工况）"
    return "泄水渠与陡坡水力设计计算"


def _display_control_source(value: Any) -> str:
    """把内部起点水深来源转换为中文展示。"""
    return display_start_control_source(value)


def _result_cases(result: Any) -> list[dict[str, Any]]:
    """把单工况或多工况导出载荷统一成列表。"""
    data = _mapping(result)
    raw_cases = data.get("export_cases") if data else None
    if raw_cases:
        cases = []
        for idx, item in enumerate(_list(raw_cases)):
            case = _mapping(item)
            case_result = case.get("result") or case.get("data") or case
            summary = normalize_result(case_result).summary
            label = case.get("label") or _mapping(case_result).get("case_label") or summary.get("工程名称") or f"工况{idx + 1}"
            cases.append({"label": str(label), "params": _mapping(case.get("params")), "result": case_result})
        return cases
    summary = normalize_result(result).summary
    label = data.get("case_label") or summary.get("工程名称") or "当前工况"
    return [{"label": str(label), "params": {}, "result": result}]


def _append_word_table(document: Document, headers: list[str], rows: list[list[Any]]) -> None:
    """向文档追加带表头的表格。"""
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in rows or [[""] * len(headers)]:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = _text(value)


def _append_word_paragraph(document: Document, text: Any, style: str | None = None) -> None:
    """追加段落，兼容模板缺少内置英文样式名的情况。"""
    try:
        document.add_paragraph(_text(text), style=style)
    except KeyError:
        document.add_paragraph(_text(text))


def _profile_row(point: dict[str, Any]) -> list[Any]:
    """整理沿程点导出行。"""
    return [
        _value(point.get("x"), point.get("distance_m"), point.get("distance"), point.get("L"), default=0),
        _value(point.get("bed_elevation"), point.get("bed_elevation_m"), point.get("渠底高程"), point.get("z_bed")),
        _value(point.get("water_elevation"), point.get("water_elevation_m"), point.get("水面高程"), point.get("z_water")),
        _value(point.get("depth"), point.get("depth_m"), point.get("水深"), point.get("h")),
        _value(point.get("velocity_ms"), point.get("流速")),
        _value(point.get("froude"), point.get("弗劳德数")),
        _value(point.get("aerated_depth_m"), point.get("掺气水深")),
        _value(point.get("sidewall_top_elevation_m"), point.get("侧墙顶高程")),
    ]


def _connection_rows(data: dict[str, Any]) -> list[list[Any]]:
    """整理上下游衔接导出行。"""
    upstream = _mapping(data.get("upstream_connection"))
    start_control = _mapping(data.get("start_control"))
    jump = _mapping(data.get("hydraulic_jump") or data.get("downstream_energy_dissipation"))
    rows = [
        ["陡槽起点水深来源", _display_control_source(_value(start_control.get("source"), upstream.get("start_depth_source")))],
        ["陡槽起点控制水深（米）", _value(start_control.get("depth_m"), upstream.get("control_depth_m"))],
        ["上游正常水深（米）", _value(upstream.get("upstream_normal_depth_m"))],
        ["上游衔接说明", _value(upstream.get("message"))],
    ]
    if jump:
        rows.extend(
            [
                ["跃前水深（米）", _value(jump.get("pre_jump_depth_m"))],
                ["跃前流速（米/秒）", _value(jump.get("pre_jump_velocity_ms"))],
                ["跃前弗劳德数", _value(jump.get("pre_jump_froude"))],
                ["跃后共轭水深（米）", _value(jump.get("conjugate_depth_m"))],
                ["下游控制水深（米）", _value(jump.get("control_depth_m"))],
                ["尾水判断", _value(jump.get("tailwater_judgement"))],
                ["下游衔接说明", _value(jump.get("message"))],
            ]
        )
    return rows


def _energy_rows(data: dict[str, Any]) -> list[list[Any]]:
    """整理消力池和出口整流导出行。"""
    jump = _mapping(data.get("hydraulic_jump") or data.get("downstream_energy_dissipation"))
    aeration = _mapping(data.get("aeration_and_sidewall"))
    rows = [
        ["最大掺气水深（米）", _value(aeration.get("max_aerated_depth_m"))],
        ["建议侧墙高度（米）", _value(aeration.get("recommended_sidewall_height_m"))],
        ["建议消力池长度（米）", _value(jump.get("recommended_pool_length_m"))],
        ["建议消力池深度（米）", _value(jump.get("recommended_pool_depth_m"))],
        ["出口整流段长度（米）", _value(jump.get("recommended_transition_length_m"), jump.get("outlet_rectification_length_m"))],
        ["消能说明", _value(jump.get("message"))],
        ["掺气侧墙说明", _value(aeration.get("message"))],
    ]
    return rows


def _multi_flow_rows(data: dict[str, Any]) -> list[list[Any]]:
    """整理多流量控制工况导出行。"""
    multi = _mapping(data.get("multi_flow_control"))
    rows = []
    for item in _list(multi.get("cases")):
        case = _mapping(item)
        rows.append(
            [
                _value(case.get("name")),
                _value(case.get("Q")),
                _value(case.get("pre_jump_depth_m")),
                _value(case.get("conjugate_depth_m")),
                _value(case.get("control_depth_m")),
                _value(case.get("control_depth_deficit_m")),
                _value(case.get("recommended_pool_length_m")),
                _value(case.get("recommended_pool_depth_m")),
            ]
        )
    return rows


def _table3_rows(data: dict[str, Any]) -> list[list[Any]]:
    """整理表3轻量接口沿程点导出行。"""
    table3 = _mapping(data.get("water_profile_export"))
    rows = []
    for item in _list(table3.get("points")):
        point = _mapping(item)
        rows.append(
            [
                _value(point.get("桩号_m")),
                _value(point.get("渠底高程_m")),
                _value(point.get("水深_m")),
                _value(point.get("水位_m")),
                _value(point.get("流速_m_s")),
                _value(point.get("弗劳德数")),
            ]
        )
    return rows


def _principle_rows(result: Any) -> list[list[Any]]:
    """整理计算原理导出行。"""
    return [
        [
            item.get("step", ""),
            item.get("purpose", ""),
            item.get("formula_text") or item.get("formula", ""),
            item.get("substitution", ""),
            item.get("result", ""),
            item.get("explanation", ""),
            item.get("source", ""),
        ]
        for item in build_calculation_principles(result)
    ]


def _append_report_meta(document: Document, meta: Any = None, calc_purpose: str = "", references: list[str] | None = None) -> None:
    """写入项目设置、计算目的和计算依据。"""
    if meta is not None:
        rows = [
            ["工程名称", getattr(meta, "project_name", "")],
            ["设计阶段", getattr(meta, "design_stage", "")],
            ["成果等级", getattr(meta, "product_level", "")],
            ["成果编号", getattr(meta, "record_number", "")],
            ["专业", getattr(meta, "specialty", "")],
            ["计算", getattr(meta, "calculator", "")],
            ["校核", getattr(meta, "checker", "")],
            ["审核", getattr(meta, "reviewer", "")],
            ["审定", getattr(meta, "approver", "")],
            ["卷册", f"{getattr(meta, 'volume_current', '')}/{getattr(meta, 'volume_total', '')}".strip("/")],
            ["基本资料", getattr(meta, "basic_info", "")],
            ["强制性条文", getattr(meta, "mandatory_clause", "")],
        ]
        document.add_heading("项目资料", level=2)
        _append_word_table(document, ["项目", "内容"], rows)
    if calc_purpose:
        document.add_heading("计算目的", level=2)
        document.add_paragraph(_text(calc_purpose))
    if references:
        document.add_heading("计算依据", level=2)
        for reference in references:
            document.add_paragraph(_text(reference), style="List Number")


def _append_result_sections(document: Document, label: str, result: Any, *, multi_case: bool) -> None:
    """向 Word 文档写入一个工况的完整结果。"""
    data = _mapping(result)
    view_data = normalize_result(result)
    base_level = 3 if multi_case else 2
    if multi_case:
        document.add_heading(f"工况：{label}", level=2)

    document.add_heading("计算原理", level=base_level)
    _append_word_table(
        document,
        ["步骤", "计算目的", "公式", "本次代入", "计算结果", "原理说明", "来源"],
        _principle_rows(result),
    )

    document.add_heading("结果汇总", level=base_level)
    _append_word_table(
        document,
        ["项目", "数值"],
        [[key, value] for key, value in view_data.summary.items()],
    )

    document.add_heading("沿程水面线", level=base_level)
    _append_word_table(
        document,
        ["桩号（米）", "渠底高程（米）", "水面高程（米）", "水深（米）", "流速（米/秒）", "弗劳德数", "掺气水深（米）", "侧墙顶高程（米）"],
        [_profile_row(point) for point in view_data.profile_points],
    )

    document.add_heading("上下游衔接", level=base_level)
    _append_word_table(document, ["项目", "数值"], _connection_rows(data))

    document.add_heading("消能与出口整流", level=base_level)
    _append_word_table(document, ["项目", "数值"], _energy_rows(data))

    document.add_heading("多流量控制工况", level=base_level)
    _append_word_table(
        document,
        ["工况", "流量（立方米/秒）", "跃前水深（米）", "跃后共轭水深（米）", "控制水深（米）", "控制差值（米）", "建议池长（米）", "建议池深（米）"],
        _multi_flow_rows(data),
    )

    document.add_heading("表3轻量接口", level=base_level)
    _append_word_table(
        document,
        ["桩号（米）", "渠底高程（米）", "水深（米）", "水位（米）", "流速（米/秒）", "弗劳德数"],
        _table3_rows(data),
    )
    table3 = _mapping(data.get("water_profile_export"))
    if table3.get("说明"):
        document.add_paragraph(_text(table3.get("说明")))

    document.add_heading("规范校核", level=base_level)
    _append_word_table(
        document,
        ["项目", "结论", "说明"],
        [
            [
                item.get("name") or item.get("项目") or "",
                item.get("result") or item.get("结论") or "",
                item.get("message") or item.get("说明") or "",
            ]
            for item in view_data.checks
        ],
    )

    document.add_heading("风险提示", level=base_level)
    if view_data.risks:
        for risk in view_data.risks:
            _append_word_paragraph(document, risk, style="List Bullet")
    else:
        document.add_paragraph("无")


def export_spillway_steep_chute_excel(path: str | Path, result: Any) -> Path:
    """导出泄水渠与陡坡计算结果到表格文件。"""
    output_path = Path(path)
    cases = _result_cases(result)
    multi_case = len(cases) > 1
    workbook = Workbook()

    principles_ws = workbook.active
    principles_ws.title = "计算原理"
    principle_header = ["步骤", "计算目的", "公式", "本次代入", "计算结果", "原理说明", "来源"]
    principles_ws.append(["工况", *principle_header] if multi_case else principle_header)
    for case in cases:
        for row in _principle_rows(case["result"]):
            principles_ws.append([case["label"], *row] if multi_case else row)

    summary_ws = workbook.create_sheet("结果汇总")
    summary_ws.title = "结果汇总"
    summary_ws.append(["工况", "项目", "数值"] if multi_case else ["项目", "数值"])
    for case in cases:
        view_data = normalize_result(case["result"])
        for key, value in view_data.summary.items():
            summary_ws.append([case["label"], key, value] if multi_case else [key, value])

    profile_ws = workbook.create_sheet("沿程水面线")
    profile_header = ["桩号（米）", "渠底高程（米）", "水面高程（米）", "水深（米）", "流速（米/秒）", "弗劳德数", "掺气水深（米）", "侧墙顶高程（米）"]
    profile_ws.append(["工况", *profile_header] if multi_case else profile_header)
    for case in cases:
        for point in normalize_result(case["result"]).profile_points:
            row = _profile_row(point)
            profile_ws.append([case["label"], *row] if multi_case else row)

    connection_ws = workbook.create_sheet("上下游衔接")
    connection_ws.append(["工况", "项目", "数值"] if multi_case else ["项目", "数值"])
    for case in cases:
        for row in _connection_rows(_mapping(case["result"])):
            connection_ws.append([case["label"], *row] if multi_case else row)

    energy_ws = workbook.create_sheet("消能与出口")
    energy_ws.append(["工况", "项目", "数值"] if multi_case else ["项目", "数值"])
    for case in cases:
        for row in _energy_rows(_mapping(case["result"])):
            energy_ws.append([case["label"], *row] if multi_case else row)

    flow_ws = workbook.create_sheet("多流量控制")
    flow_header = ["工况", "流量（立方米/秒）", "跃前水深（米）", "跃后共轭水深（米）", "控制水深（米）", "控制差值（米）", "建议池长（米）", "建议池深（米）"]
    flow_ws.append(["来源工况", *flow_header] if multi_case else flow_header)
    for case in cases:
        for row in _multi_flow_rows(_mapping(case["result"])):
            flow_ws.append([case["label"], *row] if multi_case else row)

    table3_ws = workbook.create_sheet("表3轻量接口")
    table3_header = ["桩号（米）", "渠底高程（米）", "水深（米）", "水位（米）", "流速（米/秒）", "弗劳德数"]
    table3_ws.append(["工况", *table3_header] if multi_case else table3_header)
    for case in cases:
        data = _mapping(case["result"])
        for row in _table3_rows(data):
            table3_ws.append([case["label"], *row] if multi_case else row)
        table3 = _mapping(data.get("water_profile_export"))
        if table3.get("说明"):
            table3_ws.append([])
            table3_ws.append([case["label"], "说明", table3.get("说明")] if multi_case else ["说明", table3.get("说明")])

    checks_ws = workbook.create_sheet("规范校核")
    checks_ws.append(["工况", "项目", "结论", "说明"] if multi_case else ["项目", "结论", "说明"])
    for case in cases:
        for item in normalize_result(case["result"]).checks:
            row = [
                item.get("name") or item.get("项目") or "",
                item.get("result") or item.get("结论") or "",
                item.get("message") or item.get("说明") or "",
            ]
            checks_ws.append([case["label"], *row] if multi_case else row)

    risks_ws = workbook.create_sheet("风险提示")
    risks_ws.append(["工况", "风险"] if multi_case else ["风险"])
    for case in cases:
        for risk in normalize_result(case["result"]).risks:
            risks_ws.append([case["label"], risk] if multi_case else [risk])

    output_path.parent.mkdir(parents=True, exist_ok=True)
    workbook.save(output_path)
    return output_path


def export_spillway_steep_chute_word(
    path: str | Path,
    result: Any,
    *,
    meta: Any = None,
    calc_purpose: str = "",
    references: list[str] | None = None,
) -> Path:
    """导出泄水渠与陡坡计算结果到文档文件。"""
    output_path = Path(path)
    cases = _result_cases(result)
    multi_case = len(cases) > 1
    content_desc = _word_content_desc(len(cases))
    document = create_engineering_report_doc(
        meta=_word_meta(meta),
        calc_title="泄水渠与陡坡计算书",
        calc_content_desc=content_desc,
        calc_purpose=calc_purpose,
        references=list(references or []),
        calc_program_text=f"渠系建筑物水力计算系统 V1.0\n{content_desc}",
    )
    document.add_page_break()
    for case in cases:
        _append_result_sections(document, case["label"], case["result"], multi_case=multi_case)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path
