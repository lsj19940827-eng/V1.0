"""
文档转 Markdown 工具
支持将 Word (.docx)、PDF (.pdf)、Excel (.xlsx, .xls) 转换为 Markdown 格式
"""

import os
from pathlib import Path
import docx2txt
import pdfplumber
import pandas as pd
from docx import Document


class DocumentToMarkdown:
    """文档转Markdown转换器"""
    
    def __init__(self):
        self.supported_formats = ['.docx', '.pdf', '.xlsx', '.xls']
    
    def convert_file(self, input_path, output_path=None, verbose=True):
        """
        转换文档为Markdown格式
        
        参数:
            input_path: 输入文件路径
            output_path: 输出文件路径（可选，默认为同名.md文件）
            verbose: 是否输出详细过程
        
        返回:
            转换后的Markdown文本
        """
        input_path = Path(input_path)
        
        if not input_path.exists():
            raise FileNotFoundError(f"文件不存在: {input_path}")
        
        # 检查文件格式
        ext = input_path.suffix.lower()
        if ext not in self.supported_formats:
            raise ValueError(f"不支持的文件格式: {ext}。支持的格式: {', '.join(self.supported_formats)}")
        
        if verbose:
            print(f"开始转换: {input_path.name}")
            print(f"文件格式: {ext}")
        
        # 根据文件类型选择转换方法
        if ext == '.docx':
            markdown_text = self._convert_docx(input_path, verbose)
        elif ext == '.pdf':
            markdown_text = self._convert_pdf(input_path, verbose)
        elif ext in ['.xlsx', '.xls']:
            markdown_text = self._convert_excel(input_path, verbose)
        
        # 保存结果
        if output_path is None:
            output_path = input_path.with_suffix('.md')
        else:
            output_path = Path(output_path)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(markdown_text)
        
        if verbose:
            print(f"转换完成！输出文件: {output_path}")
            print(f"文件大小: {output_path.stat().st_size} 字节")
        
        return markdown_text
    
    def _convert_docx(self, file_path, verbose=False):
        """转换Word文档"""
        if verbose:
            print("正在解析Word文档...")
        
        # 提取文本内容
        text = docx2txt.process(str(file_path))
        
        # 使用python-docx获取更详细的结构信息
        doc = Document(file_path)
        markdown_lines = [f"# {file_path.stem}\n"]
        
        if verbose:
            print(f"文档包含 {len(doc.paragraphs)} 个段落")
        
        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            
            # 根据样式判断标题层级
            style_name = para.style.name.lower()
            if 'heading 1' in style_name or 'title' in style_name:
                markdown_lines.append(f"# {para.text}\n")
            elif 'heading 2' in style_name:
                markdown_lines.append(f"## {para.text}\n")
            elif 'heading 3' in style_name:
                markdown_lines.append(f"### {para.text}\n")
            elif 'heading 4' in style_name:
                markdown_lines.append(f"#### {para.text}\n")
            else:
                # 检查是否为粗体
                if para.runs and any(run.bold for run in para.runs):
                    markdown_lines.append(f"**{para.text}**\n")
                else:
                    markdown_lines.append(f"{para.text}\n")
        
        # 处理表格
        if doc.tables and verbose:
            print(f"文档包含 {len(doc.tables)} 个表格")
        
        for i, table in enumerate(doc.tables):
            markdown_lines.append(f"\n### 表格 {i+1}\n")
            
            # 转换为Markdown表格
            if table.rows:
                # 表头
                headers = [cell.text.strip() for cell in table.rows[0].cells]
                markdown_lines.append("| " + " | ".join(headers) + " |")
                markdown_lines.append("| " + " | ".join(['---'] * len(headers)) + " |")
                
                # 表格内容
                for row in table.rows[1:]:
                    cells = [cell.text.strip() for cell in row.cells]
                    markdown_lines.append("| " + " | ".join(cells) + " |")
            
            markdown_lines.append("\n")
        
        return "\n".join(markdown_lines)
    
    def _convert_pdf(self, file_path, verbose=False):
        """转换PDF文档"""
        if verbose:
            print("正在解析PDF文档...")
        
        markdown_lines = [f"# {file_path.stem}\n"]
        
        with pdfplumber.open(file_path) as pdf:
            total_pages = len(pdf.pages)
            if verbose:
                print(f"PDF包含 {total_pages} 页")
            
            for page_num, page in enumerate(pdf.pages, 1):
                if verbose and page_num % 10 == 0:
                    print(f"已处理 {page_num}/{total_pages} 页...")
                
                markdown_lines.append(f"\n## 第 {page_num} 页\n")
                
                # 提取文本
                text = page.extract_text()
                if text:
                    markdown_lines.append(text)
                
                # 提取表格
                tables = page.extract_tables()
                for i, table in enumerate(tables):
                    if table:
                        markdown_lines.append(f"\n### 表格 {i+1}\n")
                        # 转换为Markdown表格
                        for row_idx, row in enumerate(table):
                            cleaned_row = [str(cell).strip() if cell else '' for cell in row]
                            markdown_lines.append("| " + " | ".join(cleaned_row) + " |")
                            if row_idx == 0:
                                markdown_lines.append("| " + " | ".join(['---'] * len(cleaned_row)) + " |")
                        markdown_lines.append("\n")
        
        if verbose:
            print(f"PDF解析完成")
        
        return "\n".join(markdown_lines)
    
    def _convert_excel(self, file_path, verbose=False):
        """转换Excel文档"""
        if verbose:
            print("正在解析Excel文档...")
        
        markdown_lines = [f"# {file_path.stem}\n"]
        
        # 读取所有工作表
        excel_file = pd.ExcelFile(file_path)
        
        if verbose:
            print(f"Excel包含 {len(excel_file.sheet_names)} 个工作表")
        
        for sheet_name in excel_file.sheet_names:
            if verbose:
                print(f"正在处理工作表: {sheet_name}")
            
            markdown_lines.append(f"\n## {sheet_name}\n")
            
            # 读取工作表数据
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            
            if verbose:
                print(f"工作表包含 {len(df)} 行, {len(df.columns)} 列")
            
            # 转换为Markdown表格
            markdown_table = df.to_markdown(index=False)
            if markdown_table:
                markdown_lines.append(markdown_table)
            else:
                # 备用方案：手动构建表格
                headers = df.columns.tolist()
                markdown_lines.append("| " + " | ".join(str(h) for h in headers) + " |")
                markdown_lines.append("| " + " | ".join(['---'] * len(headers)) + " |")
                
                for _, row in df.iterrows():
                    cells = [str(cell) if pd.notna(cell) else '' for cell in row]
                    markdown_lines.append("| " + " | ".join(cells) + " |")
            
            markdown_lines.append("\n")
        
        if verbose:
            print("Excel解析完成")
        
        return "\n".join(markdown_lines)
    
    def batch_convert(self, input_dir, output_dir=None, verbose=True):
        """
        批量转换文件夹中的文档
        
        参数:
            input_dir: 输入文件夹路径
            output_dir: 输出文件夹路径（可选，默认为输入文件夹）
            verbose: 是否输出详细过程
        """
        input_dir = Path(input_dir)
        if output_dir is None:
            output_dir = input_dir
        else:
            output_dir = Path(output_dir)
            output_dir.mkdir(parents=True, exist_ok=True)
        
        # 查找所有支持的文件
        files_to_convert = []
        for ext in self.supported_formats:
            files_to_convert.extend(input_dir.glob(f"*{ext}"))
        
        if verbose:
            print(f"找到 {len(files_to_convert)} 个文件待转换")
        
        results = []
        for file_path in files_to_convert:
            try:
                output_path = output_dir / f"{file_path.stem}.md"
                self.convert_file(file_path, output_path, verbose)
                results.append((file_path.name, "成功"))
            except Exception as e:
                results.append((file_path.name, f"失败: {str(e)}"))
                if verbose:
                    print(f"错误: {file_path.name} - {str(e)}")
        
        if verbose:
            print("\n转换结果汇总:")
            for filename, status in results:
                print(f"  {filename}: {status}")
        
        return results


def main():
    """示例用法"""
    converter = DocumentToMarkdown()
    
    print("=" * 60)
    print("文档转Markdown工具")
    print("=" * 60)
    print("\n支持的格式: Word (.docx), PDF (.pdf), Excel (.xlsx, .xls)")
    print("\n使用方法:")
    print("1. 转换单个文件:")
    print("   converter.convert_file('input.docx', 'output.md', verbose=True)")
    print("\n2. 批量转换:")
    print("   converter.batch_convert('input_folder', 'output_folder', verbose=True)")
    print("\n3. 在代码中使用:")
    print("   from document_to_markdown import DocumentToMarkdown")
    print("   converter = DocumentToMarkdown()")
    print("   markdown_text = converter.convert_file('your_file.pdf')")
    print("\n" + "=" * 60)


if __name__ == "__main__":
    main()
