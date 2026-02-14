# -*- coding: utf-8 -*-
"""
PySide6 完整复刻 —— 明渠水力计算面板

完整复刻 OpenChannelPanel 的全部功能：
1. 断面类型切换（梯形/矩形/圆形）
2. 参数输入面板（含可选参数动态显示/隐藏）
3. 计算结果文本输出（简要模式 + 详细模式）
4. 附录E方案对比表格（内嵌HTML表格）
5. 断面图绘制（matplotlib）
6. 导出图表 / 导出报告
"""

import sys
import os
import math
import re
import html as html_mod

_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, os.path.join(_root, "渠系建筑物断面计算"))

from PySide6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QGroupBox, QSplitter, QFrame, QTabWidget,
    QTextEdit, QFileDialog, QScrollArea, QStatusBar
)
from PySide6.QtCore import Qt
from PySide6.QtGui import QFont, QDoubleValidator

# QFluentWidgets —— Win11 Fluent Design 组件
from qfluentwidgets import (
    ComboBox, PushButton, PrimaryPushButton, LineEdit,
    CheckBox, InfoBar, InfoBarPosition
)

import matplotlib
matplotlib.use('QtAgg')
import matplotlib.pyplot as plt
try:
    from matplotlib.backends.backend_qtagg import FigureCanvasQTAgg as FigureCanvas
    from matplotlib.backends.backend_qtagg import NavigationToolbar2QT as NavToolbar
except ImportError:
    from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
    from matplotlib.backends.backend_qt5agg import NavigationToolbar2QT as NavToolbar
from matplotlib.figure import Figure
import numpy as np

plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'SimSun']
plt.rcParams['axes.unicode_minus'] = False

from 明渠设计 import (
    quick_calculate as mingqu_calculate,
    quick_calculate_circular as circular_calculate,
    calculate_area, calculate_wetted_perimeter, calculate_hydraulic_radius,
    get_flow_increase_percent, MAX_BETA,
    PI, MIN_FREEBOARD, MIN_FREE_AREA_PERCENT, MIN_FLOW_FACTOR
)

# Word导出依赖
try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    import latex2mathml.converter
    from lxml import etree
    WORD_EXPORT_AVAILABLE = True
except ImportError:
    WORD_EXPORT_AVAILABLE = False


def _find_mml2omml_xsl():
    """查找Office的MML2OMML.XSL文件"""
    candidates = [
        r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL",
        r"C:\Program Files (x86)\Microsoft Office\root\Office16\MML2OMML.XSL",
        r"C:\Program Files\Microsoft Office\Office16\MML2OMML.XSL",
        r"C:\Program Files\Microsoft Office\root\Office15\MML2OMML.XSL",
    ]
    for p in candidates:
        if os.path.exists(p):
            return p
    return None


_MML2OMML_XSL_PATH = _find_mml2omml_xsl() if WORD_EXPORT_AVAILABLE else None
_MML2OMML_TRANSFORM = None

def _get_omml_transform():
    """延迟加载XSLT转换器"""
    global _MML2OMML_TRANSFORM
    if _MML2OMML_TRANSFORM is None and _MML2OMML_XSL_PATH:
        xslt_tree = etree.parse(_MML2OMML_XSL_PATH)
        _MML2OMML_TRANSFORM = etree.XSLT(xslt_tree)
    return _MML2OMML_TRANSFORM


def _latex_to_omml(latex_str):
    """将LaTeX公式转换为Word可编辑的OMML元素"""
    transform = _get_omml_transform()
    if transform is None:
        return None
    try:
        mathml_str = latex2mathml.converter.convert(latex_str)
        mathml_tree = etree.fromstring(mathml_str.encode('utf-8'))
        omml_result = transform(mathml_tree)
        return omml_result.getroot()
    except Exception:
        return None


def _add_formula_to_doc(doc, latex_str, prefix=""):
    """向Word文档添加一个可编辑公式段落"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if prefix:
        run = p.add_run(prefix)
        run.font.name = 'Microsoft YaHei'
        run.font.size = Pt(11)
    omml = _latex_to_omml(latex_str)
    if omml is not None:
        p._element.append(omml)
    else:
        run = p.add_run(f" {latex_str}")
        run.font.name = 'Cambria Math'
        run.font.size = Pt(11)
    return p


# ============================================================
# 样式
# ============================================================
P = "#1976D2"; S = "#2E7D32"; W = "#F57C00"; E = "#D32F2F"
BG = "#F5F7FA"; CARD = "#FFFFFF"; BD = "#E0E0E0"; T1 = "#212121"; T2 = "#757575"

STYLE = f"""
QMainWindow {{ background: {BG}; }}
QGroupBox {{
    font-size: 13px; font-weight: bold; color: {P};
    border: 1px solid {BD}; border-radius: 6px;
    margin-top: 12px; padding: 14px 10px 10px 10px; background: {CARD};
}}
QGroupBox::title {{ subcontrol-origin: margin; left: 12px; padding: 0 6px; background: {CARD}; }}
QLabel {{ color: {T1}; font-size: 12px; }}
QTextEdit {{ border: 1px solid {BD}; border-radius: 4px; background: #f5f5f5; }}
QTabWidget::pane {{ border: 1px solid {BD}; border-radius: 4px; background: {CARD}; }}
QTabBar::tab {{ background: #E8EAF6; color: {T1}; padding: 8px 18px; margin-right: 2px; border-top-left-radius: 4px; border-top-right-radius: 4px; font-size: 12px; }}
QTabBar::tab:selected {{ background: {P}; color: white; font-weight: bold; }}
QStatusBar {{ background: {CARD}; border-top: 1px solid {BD}; font-size: 11px; color: {T2}; }}
"""

AE_CSS = """
<style>
table.ae { border-collapse:collapse; width:96%; margin:8px auto; font-family:'Microsoft YaHei',sans-serif; font-size:10pt; }
table.ae th { background:#f0f0f0; color:#1a1a1a; padding:6px 8px; border:1px solid #e0e0e0; text-align:center; font-weight:bold; }
table.ae td { padding:5px 8px; border:1px solid #e0e0e0; text-align:center; color:#333; }
tr.even { background:#fff; } tr.odd { background:#fafafa; }
tr.sel { background:#d6eaf8; color:#0055a3; font-weight:bold; }
tr.err { background:#fdf0e7; color:#a34400; }
pre { margin:0; padding:0; white-space:pre-wrap; font-family:Consolas,'Courier New',monospace; font-size:11pt; }
</style>
"""


def _e(s):
    """HTML转义"""
    return html_mod.escape(str(s))


class OpenChannelPanel(QMainWindow):
    """明渠水力计算面板 —— PySide6 完整复刻"""

    def __init__(self):
        super().__init__()
        self.setWindowTitle("明渠水力计算 —— PySide6 完整复刻")
        self.setMinimumSize(1100, 750)
        self.resize(1300, 850)
        self.input_params = {}
        self.current_result = None
        self._appendix_e_export_text = ""
        self._export_plain_text = ""  # 导出报告时使用的纯文本版本
        self._init_ui()
        self.statusBar().showMessage("就绪 | PySide6 完整复刻版")

    # ================================================================
    # UI 构建
    # ================================================================
    def _init_ui(self):
        central = QWidget()
        self.setCentralWidget(central)
        main_lay = QHBoxLayout(central)
        main_lay.setContentsMargins(10, 8, 10, 8)

        splitter = QSplitter(Qt.Horizontal)
        splitter.setChildrenCollapsible(False)
        main_lay.addWidget(splitter)

        # 左侧输入
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        scroll.setStyleSheet("QScrollArea{border:none;background:transparent;}")
        inp_w = QWidget()
        self._build_input(inp_w)
        scroll.setWidget(inp_w)
        scroll.setMinimumWidth(280)
        scroll.setMaximumWidth(420)
        splitter.addWidget(scroll)

        # 右侧输出
        out_w = QWidget()
        self._build_output(out_w)
        splitter.addWidget(out_w)
        splitter.setSizes([340, 900])

    # ----------------------------------------------------------------
    # 输入面板
    # ----------------------------------------------------------------
    def _build_input(self, parent):
        lay = QVBoxLayout(parent)
        lay.setContentsMargins(5, 5, 5, 5)
        lay.setSpacing(6)
        grp = QGroupBox("输入参数")
        fl = QVBoxLayout(grp)
        fl.setSpacing(5)

        # 断面类型（FluentWidgets ComboBox）
        r = QHBoxLayout(); r.addWidget(QLabel("断面类型:"))
        self.section_combo = ComboBox()
        self.section_combo.addItems(["梯形", "矩形", "圆形"])
        self.section_combo.currentTextChanged.connect(self._on_section_type_changed)
        r.addWidget(self.section_combo, 1); fl.addLayout(r)

        self.Q_edit = self._field(fl, "设计流量 Q (m³/s):", "5.0")
        self.m_lbl, self.m_edit = self._field2(fl, "边坡系数 m:", "1.0")
        self.n_edit = self._field(fl, "糙率 n:", "0.014")
        self.slope_edit = self._field(fl, "水力坡降 1/", "3000")

        fl.addWidget(self._slbl("【流速参数】"))
        self.vmin_edit = self._field(fl, "不淤流速 (m/s):", "0.1")
        self.vmax_edit = self._field(fl, "不冲流速 (m/s):", "100.0")
        fl.addWidget(self._hint("(一般情况下保持默认数值即可)"))

        fl.addWidget(self._slbl("【流量加大】"))
        self.inc_edit = self._field(fl, "流量加大比例 (%):", "")
        self.inc_hint = QLabel("(留空则自动计算)")
        self.inc_hint.setStyleSheet("font-size:10px;color:black;")
        fl.addWidget(self.inc_hint)

        fl.addWidget(self._sep())
        fl.addWidget(self._slbl("【可选参数】"))
        self.beta_lbl, self.beta_edit = self._field2(fl, "手动宽深比 β:", "")
        self.b_lbl, self.b_edit = self._field2(fl, "手动底宽 B (m):", "")
        self.bb_hint = self._hint("(二选一输入，留空则自动计算)")
        fl.addWidget(self.bb_hint)

        self.D_lbl, self.D_edit = self._field2(fl, "手动直径 D (m):", "")
        self.D_hint_lbl = self._hint("(留空则自动计算)")
        fl.addWidget(self.D_hint_lbl)
        for w in (self.D_lbl, self.D_edit, self.D_hint_lbl): w.hide()

        fl.addWidget(self._sep())
        self.detail_cb = CheckBox("输出详细计算过程")
        self.detail_cb.setChecked(True)
        fl.addWidget(self.detail_cb)

        br = QHBoxLayout()
        cb = PrimaryPushButton("计算"); cb.setCursor(Qt.PointingHandCursor); cb.clicked.connect(self._calculate)
        clb = PushButton("清空"); clb.setCursor(Qt.PointingHandCursor); clb.clicked.connect(self._clear)
        br.addWidget(cb); br.addWidget(clb); fl.addLayout(br)

        er = QHBoxLayout()
        ec = PushButton("导出图表"); ec.clicked.connect(self._export_charts)
        ew = PushButton("导出Word"); ew.clicked.connect(self._export_word)
        er.addWidget(ec); er.addWidget(ew)
        fl.addLayout(er)
        er2 = QHBoxLayout()
        et = PushButton("导出TXT"); et.clicked.connect(self._export_report)
        er2.addWidget(et)
        fl.addLayout(er2)

        lay.addWidget(grp)
        lay.addStretch()

    def _field(self, lay, label, default=""):
        r = QHBoxLayout(); l = QLabel(label); l.setMinimumWidth(130)
        r.addWidget(l); e = LineEdit(); e.setText(default); r.addWidget(e, 1); lay.addLayout(r)
        return e

    def _field2(self, lay, label, default=""):
        r = QHBoxLayout(); l = QLabel(label); l.setMinimumWidth(130)
        r.addWidget(l); e = LineEdit(); e.setText(default); r.addWidget(e, 1); lay.addLayout(r)
        return l, e

    def _slbl(self, t):
        l = QLabel(t); l.setStyleSheet(f"font-size:12px;font-weight:bold;color:{T1};"); return l

    def _hint(self, t):
        l = QLabel(t); l.setStyleSheet("font-size:10px;color:black;"); return l

    def _sep(self):
        f = QFrame(); f.setFrameShape(QFrame.HLine); f.setStyleSheet(f"color:{BD};"); return f

    # ----------------------------------------------------------------
    # 输出面板
    # ----------------------------------------------------------------
    def _build_output(self, parent):
        lay = QVBoxLayout(parent)
        lay.setContentsMargins(0, 0, 0, 0)
        self.notebook = QTabWidget()
        lay.addWidget(self.notebook)

        # Tab1: 计算结果
        t1 = QWidget(); t1l = QVBoxLayout(t1); t1l.setContentsMargins(5, 5, 5, 5)
        grp = QGroupBox("计算结果详情"); gl = QVBoxLayout(grp)
        self.result_text = QTextEdit()
        self.result_text.setReadOnly(True)
        self.result_text.setFont(QFont("Consolas", 11))
        gl.addWidget(self.result_text)
        t1l.addWidget(grp)
        self.notebook.addTab(t1, "计算结果")

        # Tab2: 断面图
        t2 = QWidget(); t2l = QVBoxLayout(t2); t2l.setContentsMargins(5, 5, 5, 5)
        self.section_fig = Figure(figsize=(8, 6), dpi=100)
        self.section_canvas = FigureCanvas(self.section_fig)
        self.section_toolbar = NavToolbar(self.section_canvas, t2)
        t2l.addWidget(self.section_toolbar)
        t2l.addWidget(self.section_canvas)
        self.notebook.addTab(t2, "断面图")

        self._show_initial_help()

    # ----------------------------------------------------------------
    # 断面类型切换
    # ----------------------------------------------------------------
    def _on_section_type_changed(self, stype):
        if stype == "矩形":
            self.m_lbl.hide(); self.m_edit.hide()
            self.m_edit.setText("0.0")
            for w in (self.beta_lbl, self.beta_edit, self.b_lbl, self.b_edit, self.bb_hint): w.show()
            for w in (self.D_lbl, self.D_edit, self.D_hint_lbl): w.hide()
        elif stype == "梯形":
            self.m_lbl.show(); self.m_edit.show()
            self.m_edit.setText("1.0")
            for w in (self.beta_lbl, self.beta_edit, self.b_lbl, self.b_edit, self.bb_hint): w.show()
            for w in (self.D_lbl, self.D_edit, self.D_hint_lbl): w.hide()
        elif stype == "圆形":
            self.m_lbl.hide(); self.m_edit.hide()
            for w in (self.beta_lbl, self.beta_edit, self.b_lbl, self.b_edit, self.bb_hint): w.hide()
            for w in (self.D_lbl, self.D_edit, self.D_hint_lbl): w.show()

    # ----------------------------------------------------------------
    # 初始帮助
    # ----------------------------------------------------------------
    def _show_initial_help(self):
        lines = []
        lines.append('请选择断面类型并输入参数后点击"计算"按钮开始计算...\n')
        lines.append("=" * 50)
        lines.append("明渠水力计算说明")
        lines.append("=" * 50 + "\n")
        lines.append("支持断面类型：")
        lines.append("  1. 矩形断面（边坡系数 m=0）")
        lines.append("  2. 梯形断面（用户自定义边坡系数 m）")
        lines.append("  3. 圆形明渠\n")
        lines.append("本程序基于曼宁公式进行计算：")
        lines.append("  - Q = (1/n) × A × R^(2/3) × i^(1/2)\n")
        lines.append("断面几何公式：")
        lines.append("  - 过水面积: A = (B + m×h) × h")
        lines.append("  - 湿周: X = B + 2×h×√(1+m²)")
        lines.append("  - 水力半径: R = A/X\n")
        lines.append("宽深比说明：")
        lines.append("  - 定义: β = B/h (底宽/设计水深)")
        lines.append(f"  - 推荐范围: 0 < β ≤ {MAX_BETA}")
        lines.append("  - 可选参数中可手动指定宽深比或底宽")
        lines.append("  - 二选一输入，留空则自动寻优计算\n")
        lines.append("约束条件：")
        lines.append("  - 流速范围: 不淤流速 < V < 不冲流速")
        self.result_text.setPlainText("\n".join(lines))

    # ----------------------------------------------------------------
    # 辅助：读取输入值
    # ----------------------------------------------------------------
    def _fval(self, edit, default=0.0):
        t = edit.text().strip()
        if not t:
            return default
        try:
            return float(t)
        except ValueError:
            return default

    def _fval_opt(self, edit):
        t = edit.text().strip()
        if not t:
            return None
        try:
            return float(t)
        except ValueError:
            return None

    # ================================================================
    # 计算
    # ================================================================
    def _calculate(self):
        stype = self.section_combo.currentText()
        try:
            Q = self._fval(self.Q_edit)
            n = self._fval(self.n_edit)
            slope_inv = self._fval(self.slope_edit)
            v_min = self._fval(self.vmin_edit)
            v_max = self._fval(self.vmax_edit)

            if Q <= 0:
                self._show_error("参数错误", "请输入有效的设计流量 Q（必须大于0）。"); return
            if n <= 0:
                self._show_error("参数错误", "请输入有效的糙率 n（必须大于0）。"); return
            if slope_inv <= 0:
                self._show_error("参数错误", "请输入有效的水力坡降倒数（必须大于0）。"); return
            if v_min >= v_max:
                self._show_error("参数错误", "不淤流速必须小于不冲流速。"); return

            manual_increase = self._fval_opt(self.inc_edit)

            if stype == "圆形":
                manual_D = self._fval_opt(self.D_edit)
                self.input_params = {
                    'Q': Q, 'n': n, 'slope_inv': slope_inv,
                    'v_min': v_min, 'v_max': v_max,
                    'section_type': stype, 'manual_b': manual_D,
                    'manual_increase': manual_increase
                }
                result = circular_calculate(
                    Q=Q, n=n, slope_inv=slope_inv,
                    v_min=v_min, v_max=v_max,
                    manual_D=manual_D,
                    increase_percent=manual_increase
                )
            else:
                m = self._fval(self.m_edit) if stype == "梯形" else 0.0
                if stype == "梯形" and m < 0:
                    self._show_error("参数错误", "请输入有效的边坡系数 m（不能为负）。"); return
                manual_beta = self._fval_opt(self.beta_edit)
                manual_b = self._fval_opt(self.b_edit)
                self.input_params = {
                    'Q': Q, 'm': m, 'n': n, 'slope_inv': slope_inv,
                    'v_min': v_min, 'v_max': v_max,
                    'section_type': stype,
                    'manual_beta': manual_beta, 'manual_b': manual_b,
                    'manual_increase': manual_increase
                }
                result = mingqu_calculate(
                    Q=Q, m=m, n=n, slope_inv=slope_inv,
                    v_min=v_min, v_max=v_max,
                    manual_beta=manual_beta,
                    manual_b=manual_b,
                    manual_increase_percent=manual_increase
                )

            self.current_result = result

            # 更新加大比例提示
            if result.get('success') and 'increase_percent' in result:
                ap = result['increase_percent']
                src = "手动指定" if self.inc_edit.text().strip() else "自动计算"
                if isinstance(ap, str):
                    self.inc_hint.setText(f"({src}: {ap})")
                else:
                    self.inc_hint.setText(f"({src}: {ap:.1f}%)")

            self._update_result_display(result)
            self._update_section_plot(result)
            if result['success']:
                self.statusBar().showMessage(f"计算完成 | {result.get('design_method', '')}", 10000)
            else:
                self.statusBar().showMessage("计算失败", 5000)

        except Exception as e:
            self._show_error("计算错误", f"计算过程出错: {str(e)}")

    def _show_error(self, title, msg):
        out = []
        out.append("=" * 70)
        out.append(f"  {title}")
        out.append("=" * 70)
        out.append("")
        out.append(msg)
        out.append("")
        out.append("-" * 70)
        out.append("请修正后重新计算。")
        out.append("=" * 70)
        self.result_text.setPlainText("\n".join(out))

    # ================================================================
    # 结果显示分发
    # ================================================================
    def _update_result_display(self, result):
        if not result['success']:
            self._show_error("计算失败", result.get('error_message', '未知错误'))
            return
        stype = self.input_params.get('section_type', '梯形')
        detail = self.detail_cb.isChecked()
        if stype == '圆形':
            if detail:
                self._show_circular_detail(result)
            else:
                self._show_circular_brief(result)
        else:
            if detail:
                self._show_trapezoid_detail(result)
            else:
                self._show_trapezoid_brief(result)

    # ================================================================
    # 梯形/矩形 - 简要结果
    # ================================================================
    def _show_trapezoid_brief(self, result):
        p = self.input_params
        Q, m, n = p['Q'], p['m'], p['n']
        slope_inv = p['slope_inv']
        v_min, v_max = p['v_min'], p['v_max']
        stype = p.get('section_type', '梯形')

        b = result['b_design']; h = result['h_design']
        V = result['V_design']; A = result['A_design']
        R = result['R_design']; beta = result['Beta_design']
        inc_pct = result['increase_percent']
        Q_inc = result['Q_increased']
        h_inc = result['h_increased']; V_inc = result['V_increased']
        Fb = result['Fb']; H = result['h_prime']
        inc_source = "(手动指定)" if p.get('manual_increase') else "(自动计算)"

        o = []
        o.append("=" * 70)
        o.append(f"              明渠水力计算结果（{stype}断面）")
        o.append("=" * 70)
        o.append("")
        o.append("【输入参数】")
        o.append(f"  设计流量 Q = {Q:.3f} m³/s")
        if stype == "梯形":
            o.append(f"  边坡系数 m = {m}")
        o.append(f"  糙率 n = {n}")
        o.append(f"  水力坡降 1/{int(slope_inv)}")
        o.append("")
        o.append("【设计结果】")
        o.append(f"  设计方法: {result['design_method']}")
        o.append(f"  底宽 B = {b:.3f} m")
        o.append(f"  水深 h = {h:.3f} m")
        o.append(f"  宽深比 β = {beta:.3f}")
        o.append(f"  过水面积 A = {A:.3f} m²")
        o.append(f"  水力半径 R = {R:.3f} m")
        o.append(f"  设计流速 V = {V:.3f} m/s")
        o.append("")

        # 附录E表格
        schemes = result.get('appendix_e_schemes', [])
        if schemes:
            pre1 = "\n".join(o)
            ae_html = self._build_ae_html(schemes, b, h, v_min, v_max)
            self._appendix_e_export_text = self._build_ae_text(schemes, b, h, v_min, v_max)

            o2 = []
            o2.append("【加大流量工况】")
            o2.append(f"  流量加大比例 = {inc_pct:.1f}% {inc_source}")
            o2.append(f"  加大流量 Q加大 = {Q_inc:.3f} m³/s")
            if h_inc > 0:
                o2.append(f"  加大水深 h加大 = {h_inc:.3f} m")
                o2.append(f"  加大流速 V加大 = {V_inc:.3f} m/s")
                o2.append(f"  岸顶超高 Fb = {Fb:.3f} m")
                o2.append(f"  渠道高度 H = {H:.3f} m")
            o2.append("")
            o2.append("【验证结果】")
            vel_ok = v_min < V < v_max
            fb_req = 0.25 * h_inc + 0.2 if h_inc > 0 else 0
            fb_ok = Fb >= (fb_req - 0.001) if h_inc > 0 else False
            o2.append(f"  流速验证: {'✓ 通过' if vel_ok else '✗ 未通过'}")
            o2.append(f"  超高复核: {'✓ 通过' if fb_ok else '✗ 未通过'} (Fb={Fb:.3f}m, 规范要求≥{fb_req:.3f}m)")
            o2.append("")
            all_pass = vel_ok and fb_ok
            o2.append("=" * 70)
            o2.append(f"  综合验证结果: {'全部通过 ✓' if all_pass else '未通过 ✗'}")
            o2.append("=" * 70)
            pre2 = "\n".join(o2)

            full_html = f"<html><head>{AE_CSS}</head><body>"
            full_html += f"<pre>{_e(pre1)}</pre>"
            full_html += "\n<b>【附录E断面方案对比表】</b><br>"
            full_html += "  说明: α=1.00为水力最佳断面(深窄)，α越大断面越宽浅，面积增加但流速降低<br><br>"
            full_html += ae_html
            full_html += f"<br>  注: 流速约束范围 {v_min} ~ {v_max} m/s<br><br>"
            full_html += f"<pre>{_e(pre2)}</pre></body></html>"
            self._export_plain_text = pre1 + "\n\n" + self._appendix_e_export_text + "\n\n" + pre2
            self.result_text.setHtml(full_html)
            return

        # 无附录E表格时
        o.append("【加大流量工况】")
        o.append(f"  流量加大比例 = {inc_pct:.1f}% {inc_source}")
        o.append(f"  加大流量 Q加大 = {Q_inc:.3f} m³/s")
        if h_inc > 0:
            o.append(f"  加大水深 h加大 = {h_inc:.3f} m")
            o.append(f"  加大流速 V加大 = {V_inc:.3f} m/s")
            o.append(f"  岸顶超高 Fb = {Fb:.3f} m")
            o.append(f"  渠道高度 H = {H:.3f} m")
        o.append("")
        o.append("【验证结果】")
        vel_ok = v_min < V < v_max
        fb_req = 0.25 * h_inc + 0.2 if h_inc > 0 else 0
        fb_ok = Fb >= (fb_req - 0.001) if h_inc > 0 else False
        o.append(f"  流速验证: {'✓ 通过' if vel_ok else '✗ 未通过'}")
        o.append(f"  超高复核: {'✓ 通过' if fb_ok else '✗ 未通过'} (Fb={Fb:.3f}m, 规范要求≥{fb_req:.3f}m)")
        o.append("")
        all_pass = vel_ok and fb_ok
        o.append("=" * 70)
        o.append(f"  综合验证结果: {'全部通过 ✓' if all_pass else '未通过 ✗'}")
        o.append("=" * 70)
        txt = "\n".join(o)
        self._export_plain_text = txt
        self.result_text.setPlainText(txt)

    # ================================================================
    # 梯形/矩形 - 详细结果
    # ================================================================
    def _show_trapezoid_detail(self, result):
        p = self.input_params
        Q, m, n = p['Q'], p['m'], p['n']
        slope_inv = p['slope_inv']; i = 1.0 / slope_inv
        v_min, v_max = p['v_min'], p['v_max']
        stype = p.get('section_type', '梯形')

        b = result['b_design']; h = result['h_design']
        V = result['V_design']; A = result['A_design']
        X = result['X_design']; R = result['R_design']
        beta = result['Beta_design']; Q_calc = result['Q_calc']

        inc_pct = result['increase_percent']
        Q_inc = result['Q_increased']; h_inc = result['h_increased']
        V_inc = result['V_increased']
        A_inc = result.get('A_increased', -1)
        X_inc = result.get('X_increased', -1)
        R_inc = result.get('R_increased', -1)
        Fb = result['Fb']; H = result['h_prime']
        inc_source = "(手动指定)" if p.get('manual_increase') else "(自动计算)"

        o = []
        o.append("=" * 70)
        o.append(f"              明渠水力计算结果（{stype}断面）")
        o.append("=" * 70)
        o.append("")
        o.append("【一、输入参数】")
        o.append(f"  断面类型 = {stype}")
        o.append(f"  设计流量 Q = {Q:.3f} m³/s")
        if stype == "梯形": o.append(f"  边坡系数 m = {m}")
        o.append(f"  糙率 n = {n}")
        o.append(f"  水力坡降 1/{int(slope_inv)}")
        o.append(f"  不淤流速 = {v_min} m/s")
        o.append(f"  不冲流速 = {v_max} m/s")
        if p.get('manual_beta'): o.append(f"  [手动] 宽深比 β = {p['manual_beta']}")
        if p.get('manual_b'): o.append(f"  [手动] 底宽 B = {p['manual_b']} m")
        if p.get('manual_increase'): o.append(f"  [手动] 加大比例 = {p['manual_increase']}%")
        o.append("")

        o.append("【二、设计方法】")
        o.append(f"  采用方法: {result['design_method']}")
        o.append("")

        # 附录E表格 —— 需要在文本中间插入
        schemes = result.get('appendix_e_schemes', [])
        has_ae = bool(schemes)
        if has_ae:
            pre1 = "\n".join(o)
            ae_html = self._build_ae_html(schemes, b, h, v_min, v_max)
            self._appendix_e_export_text = self._build_ae_text(schemes, b, h, v_min, v_max)
            o = []  # 重置，后续继续追加

        # 设计结果详细过程
        o.append("【三、设计结果】")
        o.append("")
        o.append("  1. 设计底宽:")
        o.append(f"     B = {b:.3f} m")
        o.append("")
        o.append("  2. 设计水深:")
        o.append(f"     h = {h:.3f} m")
        o.append("")
        o.append("  3. 宽深比:")
        o.append(f"     β = B/h = {b:.3f}/{h:.3f} = {beta:.3f}")
        o.append("")
        o.append("  4. 过水面积计算:")
        if stype == "梯形":
            o.append(f"     A = (B + m×h) × h")
            o.append(f"       = ({b:.3f} + {m}×{h:.3f}) × {h:.3f}")
            o.append(f"       = {b + m * h:.3f} × {h:.3f}")
        else:
            o.append(f"     A = B × h")
            o.append(f"       = {b:.3f} × {h:.3f}")
        o.append(f"       = {A:.3f} m²")
        o.append("")
        o.append("  5. 湿周计算:")
        sq = math.sqrt(1 + m * m)
        if stype == "梯形":
            o.append(f"     χ = B + 2×h×√(1+m²)")
            o.append(f"       = {b:.3f} + 2×{h:.3f}×√(1+{m}²)")
            o.append(f"       = {b:.3f} + 2×{h:.3f}×{sq:.4f}")
            o.append(f"       = {b:.3f} + {2 * h * sq:.3f}")
        else:
            o.append(f"     χ = B + 2×h")
            o.append(f"       = {b:.3f} + 2×{h:.3f}")
            o.append(f"       = {b:.3f} + {2 * h:.3f}")
        o.append(f"       = {X:.3f} m")
        o.append("")
        o.append("  6. 水力半径计算:")
        o.append(f"     R = A/χ = {A:.3f}/{X:.3f} = {R:.3f} m")
        o.append("")
        o.append("  7. 设计流速计算 (曼宁公式):")
        o.append(f"     V = (1/n) × R^(2/3) × i^(1/2)")
        o.append(f"       = (1/{n}) × {R:.3f}^(2/3) × {i:.6f}^(1/2)")
        o.append(f"       = {1/n:.2f} × {R**(2/3):.4f} × {math.sqrt(i):.6f}")
        o.append(f"       = {V:.3f} m/s")
        o.append("")
        o.append("  8. 流量校核:")
        o.append(f"      Q计算 = V × A = {V:.3f} × {A:.3f} = {V * A:.3f} m³/s")
        o.append(f"      误差 = {abs(V * A - Q)/Q*100:.2f}%")
        o.append("")

        # 加大流量工况
        o.append("【四、加大流量工况计算】")
        o.append("")
        o.append("  9. 加大流量计算:")
        o.append(f"      流量加大比例 = {inc_pct:.1f}% {inc_source}")
        o.append(f"      Q加大 = Q × (1 + {inc_pct/100:.2f})")
        o.append(f"           = {Q:.3f} × {1+inc_pct/100:.2f}")
        o.append(f"           = {Q_inc:.3f} m³/s")
        o.append("")

        if h_inc > 0:
            if A_inc <= 0: A_inc = (b + m * h_inc) * h_inc
            if X_inc <= 0: X_inc = b + 2 * h_inc * math.sqrt(1 + m * m)
            if R_inc <= 0 and X_inc > 0: R_inc = A_inc / X_inc

            o.append("  10. 加大水深计算:")
            o.append(f"      根据加大流量 Q加大 = {Q_inc:.3f} m³/s 和设计底宽 B = {b:.3f} m，")
            o.append(f"      利用曼宁公式反算水深:")
            o.append(f"      h加大 = {h_inc:.3f} m")
            o.append("")
            o.append("  11. 加大过水面积计算:")
            if stype == "梯形":
                o.append(f"      A加大 = (B + m×h加大) × h加大")
                o.append(f"           = ({b:.3f} + {m}×{h_inc:.3f}) × {h_inc:.3f}")
                o.append(f"           = {b + m * h_inc:.3f} × {h_inc:.3f}")
            else:
                o.append(f"      A加大 = B × h加大")
                o.append(f"           = {b:.3f} × {h_inc:.3f}")
            o.append(f"           = {A_inc:.3f} m²")
            o.append("")
            o.append("  12. 加大湿周计算:")
            sq2 = math.sqrt(1 + m * m)
            if stype == "梯形":
                o.append(f"      χ加大 = B + 2×h加大×√(1+m²)")
                o.append(f"           = {b:.3f} + 2×{h_inc:.3f}×√(1+{m}²)")
                o.append(f"           = {b:.3f} + 2×{h_inc:.3f}×{sq2:.4f}")
                o.append(f"           = {b:.3f} + {2 * h_inc * sq2:.3f}")
            else:
                o.append(f"      χ加大 = B + 2×h加大")
                o.append(f"           = {b:.3f} + 2×{h_inc:.3f}")
                o.append(f"           = {b:.3f} + {2 * h_inc:.3f}")
            o.append(f"           = {X_inc:.3f} m")
            o.append("")
            o.append("  13. 加大水力半径计算:")
            o.append(f"      R加大 = A加大 / χ加大")
            o.append(f"           = {A_inc:.3f} / {X_inc:.3f}")
            o.append(f"           = {R_inc:.3f} m")
            o.append("")
            o.append("  14. 加大流速计算 (曼宁公式):")
            o.append(f"      V加大 = (1/n) × R加大^(2/3) × i^(1/2)")
            o.append(f"           = (1/{n}) × {R_inc:.3f}^(2/3) × {i:.6f}^(1/2)")
            o.append(f"           = {1/n:.2f} × {R_inc**(2/3):.4f} × {math.sqrt(i):.6f}")
            o.append(f"           = {V_inc:.3f} m/s")
            o.append("")
            Q_chk = V_inc * A_inc
            o.append("  15. 流量校核:")
            o.append(f"      Q校核 = V加大 × A加大 = {V_inc:.3f} × {A_inc:.3f} = {Q_chk:.3f} m³/s")
            o.append(f"      误差 = {abs(Q_chk - Q_inc) / Q_inc * 100:.2f}%")
            o.append("")
            o.append("  16. 渠道岸顶超高计算（规范 6.4.8-2）:")
            o.append(f"      Fb = (1/4) × h加大 + 0.2")
            o.append(f"         = (1/4) × {h_inc:.3f} + 0.2")
            o.append(f"         = {Fb:.3f} m")
            o.append("")
            o.append("  17. 渠道高度计算:")
            o.append(f"      H = h加大 + Fb")
            o.append(f"        = {h_inc:.3f} + {Fb:.3f}")
            o.append(f"        = {H:.3f} m")
        else:
            o.append("  加大水深计算失败")
        o.append("")

        # 验证
        o.append("【五、设计验证】")
        o.append("")
        vel_ok = v_min < V < v_max
        o.append(f"  18. 流速验证:")
        o.append(f"      范围要求: {v_min} < V < {v_max} m/s")
        o.append(f"      设计流速: V = {V:.3f} m/s")
        o.append(f"      结果: {'通过 ✓' if vel_ok else '未通过 ✗'}")
        o.append("")
        fb_req = 0.25 * h_inc + 0.2 if h_inc > 0 else 0
        fb_ok = Fb >= (fb_req - 0.001) if h_inc > 0 else False
        o.append(f"  19. 超高复核（规范 6.4.8-2）:")
        o.append(f"      规范要求: Fb ≥ (1/4)×h加大 + 0.2 = {fb_req:.3f} m")
        o.append(f"      计算结果: Fb = {Fb:.3f} m")
        o.append(f"      结果: {'通过 ✓' if fb_ok else '未通过 ✗'}")
        o.append("")
        all_pass = vel_ok and fb_ok
        o.append("=" * 70)
        o.append(f"  综合验证结果: {'全部通过 ✓' if all_pass else '未通过 ✗'}")
        o.append("=" * 70)

        if has_ae:
            pre2 = "\n".join(o)
            full_html = f"<html><head>{AE_CSS}</head><body>"
            full_html += f"<pre>{_e(pre1)}</pre>"
            full_html += "\n<b>【附录E断面方案对比表】</b><br>"
            full_html += "  说明: α=1.00为水力最佳断面(深窄)，α越大断面越宽浅，面积增加但流速降低<br><br>"
            full_html += ae_html
            full_html += f"<br>  注: 流速约束范围 {v_min} ~ {v_max} m/s<br><br>"
            full_html += f"<pre>{_e(pre2)}</pre></body></html>"
            self._export_plain_text = pre1 + "\n\n" + self._appendix_e_export_text + "\n\n" + pre2
            self.result_text.setHtml(full_html)
        else:
            txt = "\n".join(o)
            self._export_plain_text = txt
            self.result_text.setPlainText(txt)

    # ================================================================
    # 圆形 - 简要结果
    # ================================================================
    def _show_circular_brief(self, result):
        p = self.input_params
        Q, n = p['Q'], p['n']; slope_inv = p['slope_inv']
        v_min, v_max = p['v_min'], p['v_max']
        D = result.get('D_design', 0)
        h = result.get('y_d', 0); V = result.get('V_d', 0)
        A_d = result.get('A_d', 0); FB_d = result.get('FB_d', 0)
        PA_d = result.get('PA_d', 0)
        inc_info = result.get('increase_percent', '')
        Q_inc = result.get('Q_inc', 0)
        h_i = result.get('y_i', 0); V_i = result.get('V_i', 0)
        FB_i = result.get('FB_i', 0); PA_i = result.get('PA_i', 0)

        o = []
        o.append("=" * 70)
        o.append("              明渠水力计算结果（圆形断面）")
        o.append("=" * 70)
        o.append("")
        o.append("【输入参数】")
        o.append(f"  设计流量 Q = {Q:.3f} m³/s")
        o.append(f"  糙率 n = {n}")
        o.append(f"  水力坡降 1/{int(slope_inv)}")
        o.append("")
        o.append("【断面尺寸】")
        o.append(f"  设计直径 D = {D:.2f} m")
        o.append("")
        o.append("【设计流量工况】")
        o.append(f"  设计水深 h = {h:.3f} m")
        o.append(f"  设计流速 V = {V:.3f} m/s")
        o.append(f"  过水面积 A = {A_d:.3f} m²")
        o.append(f"  净空高度 Fb = {FB_d:.3f} m")
        o.append(f"  净空比例 = {PA_d:.1f}%")
        o.append("")
        o.append("【加大流量工况】")
        o.append(f"  流量加大比例 = {inc_info}")
        o.append(f"  加大流量 Q加大 = {Q_inc:.3f} m³/s")
        o.append(f"  加大水深 h加大 = {h_i:.3f} m")
        o.append(f"  加大流速 V加大 = {V_i:.3f} m/s")
        o.append(f"  净空高度 Fb加大 = {FB_i:.3f} m")
        o.append(f"  净空比例 = {PA_i:.1f}%")
        o.append("")
        o.append("【验证结果】")
        vel_ok = V is not None and v_min <= V <= v_max
        vel_i_ok = V_i is not None and v_min <= V_i <= v_max if V_i else True
        o.append(f"  1. 设计流速验证")
        o.append(f"     范围要求: {v_min} ≤ V ≤ {v_max} m/s")
        o.append(f"     计算结果: V = {V:.3f} m/s")
        o.append(f"     验证结果: {'通过 ✓' if vel_ok else '未通过 ✗'}")
        o.append("")
        o.append(f"  2. 加大流速验证")
        o.append(f"     范围要求: {v_min} ≤ V ≤ {v_max} m/s")
        if V_i:
            o.append(f"     计算结果: V加大 = {V_i:.3f} m/s")
            o.append(f"     验证结果: {'通过 ✓' if vel_i_ok else '未通过 ✗'}")
        else:
            o.append(f"     计算结果: 无数据")
        o.append("")
        all_ok = vel_ok and vel_i_ok
        o.append("=" * 70)
        o.append(f"  综合验证结果: {'通过 ✓' if all_ok else '未通过 ✗'}")
        o.append("=" * 70)
        txt = "\n".join(o)
        self._export_plain_text = txt
        self.result_text.setPlainText(txt)

    # ================================================================
    # 圆形 - 详细结果
    # ================================================================
    def _show_circular_detail(self, result):
        p = self.input_params
        Q, n = p['Q'], p['n']; slope_inv = p['slope_inv']
        v_min, v_max = p['v_min'], p['v_max']
        i = 1.0 / slope_inv

        D_calc = result.get('D_calculated', 0)
        D = result.get('D_design', 0)
        pipe_area = PI * D**2 / 4 if D > 0 else 0

        h_d = result.get('y_d', 0); V_d = result.get('V_d', 0)
        A_d = result.get('A_d', 0); P_d = result.get('P_d', 0)
        R_d = result.get('R_d', 0); PA_d = result.get('PA_d', 0)
        FB_d = result.get('FB_d', 0); Q_chk_d = result.get('Q_check_d', 0)

        inc_info = result.get('increase_percent', '')
        Q_inc = result.get('Q_inc', 0)
        h_i = result.get('y_i', 0); V_i = result.get('V_i', 0)
        A_i = result.get('A_i', 0); P_i = result.get('P_i', 0)
        R_i = result.get('R_i', 0); PA_i = result.get('PA_i', 0)
        FB_i = result.get('FB_i', 0)

        try: inc_pct = float(inc_info.split('%')[0])
        except: inc_pct = 20

        Q_min = result.get('Q_min', 0)
        h_m = result.get('y_m', 0); V_m = result.get('V_m', 0)

        o = []
        o.append("=" * 70)
        o.append("              明渠水力计算结果（圆形断面）")
        o.append("=" * 70)
        o.append("")
        o.append("【一、输入参数】")
        o.append(f"  断面类型 = 圆形")
        o.append(f"  设计流量 Q = {Q:.3f} m³/s")
        o.append(f"  糙率 n = {n}")
        o.append(f"  水力坡降 1/{int(slope_inv)}")
        o.append(f"  不淤流速 = {v_min} m/s")
        o.append(f"  不冲流速 = {v_max} m/s")
        if p.get('manual_b'):
            o.append(f"  [手动] 直径 D = {p['manual_b']} m")
        o.append("")

        o.append("【二、直径确定】")
        o.append("")
        if D_calc and D_calc > 0:
            o.append(f"  1. 计算直径: D计算 = {D_calc:.3f} m")
        o.append(f"  2. 设计直径: D = {D:.2f} m")
        o.append("")
        o.append("  3. 管道总断面积计算:")
        o.append(f"     A总 = π × D² / 4")
        o.append(f"        = {PI:.4f} × {D:.2f}² / 4")
        o.append(f"        = {PI:.4f} × {D**2:.4f} / 4")
        o.append(f"        = {pipe_area:.3f} m²")
        o.append("")

        o.append("【三、设计流量工况计算】")
        o.append(f"  Q = {Q:.3f} m³/s")
        o.append("")
        o.append("  4. 设计水深计算:")
        o.append(f"     根据设计流量 Q = {Q:.3f} m³/s，利用曼宁公式反算水深:")
        o.append(f"     h = {h_d:.3f} m")
        o.append("")

        if h_d > 0 and D > 0 and h_d <= D:
            Rr = D / 2
            theta = 2 * math.acos(max(-1, min(1, (Rr - h_d) / Rr)))
            o.append(f"  5. 圆心角计算:")
            o.append(f"     θ = 2 × arccos((R - h) / R)")
            o.append(f"       = 2 × arccos(({Rr:.3f} - {h_d:.3f}) / {Rr:.3f})")
            o.append(f"       = 2 × arccos({(Rr - h_d)/Rr:.4f})")
            o.append(f"       = {math.degrees(theta):.2f}° ({theta:.4f} rad)")
            o.append("")
            o.append(f"  6. 过水面积计算:")
            o.append(f"     A = (D²/8) × (θ - sinθ)")
            o.append(f"       = ({D:.3f}²/8) × ({theta:.4f} - sin{theta:.4f})")
            o.append(f"       = {D**2/8:.4f} × {theta - math.sin(theta):.4f}")
            o.append(f"       = {A_d:.3f} m²")
            o.append("")
            o.append(f"  7. 湿周计算:")
            o.append(f"      χ = (D/2) × θ")
            o.append(f"        = ({D:.3f}/2) × {theta:.4f}")
            o.append(f"        = {Rr:.3f} × {theta:.4f}")
            o.append(f"        = {P_d:.3f} m")
            o.append("")
        else:
            o.append(f"  6. 过水面积: A = {A_d:.3f} m²")
            o.append("")
            o.append(f"  7. 湿周: χ = {P_d:.3f} m")
            o.append("")

        o.append(f"  8. 水力半径计算:")
        o.append(f"      R = A / χ")
        o.append(f"        = {A_d:.3f} / {P_d:.3f}")
        o.append(f"        = {R_d:.3f} m")
        o.append("")
        o.append(f"  9. 设计流速计算 (曼宁公式):")
        o.append(f"      V = (1/n) × R^(2/3) × i^(1/2)")
        o.append(f"        = (1/{n}) × {R_d:.3f}^(2/3) × {i:.6f}^(1/2)")
        if R_d > 0:
            o.append(f"        = {1/n:.2f} × {R_d**(2/3):.4f} × {math.sqrt(i):.6f}")
        o.append(f"        = {V_d:.3f} m/s")
        o.append("")
        o.append(f"  10. 流量校核:")
        o.append(f"      Q计算 = V × A")
        o.append(f"           = {V_d:.3f} × {A_d:.3f}")
        o.append(f"           = {V_d * A_d:.3f} m³/s")
        if V_d * A_d > 0:
            o.append(f"      误差 = {abs(V_d * A_d - Q)/Q*100:.2f}%")
        o.append("")
        o.append(f"  11. 净空高度:")
        o.append(f"      Fb = D - h = {D:.3f} - {h_d:.3f} = {FB_d:.3f} m")
        o.append("")
        o.append(f"  12. 净空面积:")
        o.append(f"      PA = (A总 - A) / A总 × 100%")
        o.append(f"         = ({pipe_area:.3f} - {A_d:.3f}) / {pipe_area:.3f} × 100%")
        o.append(f"         = {PA_d:.1f}%")
        o.append("")

        # 加大流量工况
        o.append("【四、加大流量工况计算】")
        o.append("")
        o.append(f"  13. 加大流量计算:")
        o.append(f"      流量加大比例 = {inc_info}")
        o.append(f"      Q加大 = Q × (1 + {inc_pct/100:.2f})")
        o.append(f"           = {Q:.3f} × {1+inc_pct/100:.2f}")
        o.append(f"           = {Q_inc:.3f} m³/s")
        o.append("")

        if h_i is not None and h_i > 0 and D > 0:
            o.append("  14. 加大水深计算:")
            o.append(f"      根据加大流量 Q加大 = {Q_inc:.3f} m³/s，利用曼宁公式反算水深:")
            o.append(f"      h加大 = {h_i:.3f} m")
            o.append("")
            Rr_i = D / 2
            theta_i = 2 * math.acos(max(-1, min(1, (Rr_i - h_i) / Rr_i)))
            o.append(f"  15. 圆心角计算:")
            o.append(f"      θ加大 = 2 × arccos((R - h加大) / R)")
            o.append(f"           = 2 × arccos(({Rr_i:.3f} - {h_i:.3f}) / {Rr_i:.3f})")
            o.append(f"           = 2 × arccos({(Rr_i - h_i)/Rr_i:.4f})")
            o.append(f"           = {math.degrees(theta_i):.2f}° ({theta_i:.4f} rad)")
            o.append("")
            o.append(f"  16. 过水面积计算:")
            o.append(f"      A加大 = (D²/8) × (θ加大 - sinθ加大)")
            o.append(f"           = ({D:.3f}²/8) × ({theta_i:.4f} - sin{theta_i:.4f})")
            o.append(f"           = {D**2/8:.4f} × {theta_i - math.sin(theta_i):.4f}")
            o.append(f"           = {A_i:.3f} m²")
            o.append("")
            o.append(f"  17. 湿周计算:")
            o.append(f"      χ加大 = (D/2) × θ加大")
            o.append(f"           = ({D:.3f}/2) × {theta_i:.4f}")
            o.append(f"           = {Rr_i:.3f} × {theta_i:.4f}")
            o.append(f"           = {P_i:.3f} m")
            o.append("")
        else:
            o.append(f"  14. 加大水深: h加大 = N/A")
            o.append("")

        o.append(f"  18. 水力半径计算:")
        o.append(f"      R加大 = A加大 / χ加大")
        if A_i and P_i:
            o.append(f"           = {A_i:.3f} / {P_i:.3f}")
            o.append(f"           = {R_i:.3f} m")
        o.append("")
        o.append(f"  19. 加大流速计算 (曼宁公式):")
        o.append(f"      V加大 = (1/n) × R加大^(2/3) × i^(1/2)")
        if R_i and R_i > 0:
            o.append(f"           = (1/{n}) × {R_i:.3f}^(2/3) × {i:.6f}^(1/2)")
            o.append(f"           = {1/n:.2f} × {R_i**(2/3):.4f} × {math.sqrt(i):.6f}")
        o.append(f"           = {V_i:.3f} m/s")
        o.append("")
        o.append(f"  20. 流量校核:")
        if V_i and A_i:
            o.append(f"      Q计算 = V加大 × A加大")
            o.append(f"           = {V_i:.3f} × {A_i:.3f}")
            o.append(f"           = {V_i * A_i:.3f} m³/s")
            if Q_inc > 0:
                o.append(f"      误差 = {abs(V_i * A_i - Q_inc) / Q_inc * 100:.2f}%")
        o.append("")
        o.append(f"  21. 净空高度计算:")
        o.append(f"      Fb加大 = D - h加大")
        if h_i:
            o.append(f"           = {D:.3f} - {h_i:.3f}")
            o.append(f"           = {FB_i:.3f} m")
        o.append("")
        o.append(f"  22. 净空面积计算:")
        if A_i:
            o.append(f"      PA加大 = (A总 - A加大) / A总 × 100%")
            o.append(f"           = ({pipe_area:.3f} - {A_i:.3f}) / {pipe_area:.3f} × 100%")
            o.append(f"           = {PA_i:.1f}%")
        o.append("")

        # 最小流量工况
        o.append("【五、最小流量工况计算】")
        o.append(f"  Q最小 = {Q_min:.3f} m³/s" if Q_min else "  Q最小 = N/A")
        o.append(f"  水深 h最小 = {h_m:.3f} m" if h_m else "  水深 h最小 = N/A")
        o.append(f"  流速 V最小 = {V_m:.3f} m/s" if V_m else "  流速 V最小 = N/A")
        o.append("")

        # 验证
        o.append("【六、设计验证】")
        o.append("")
        vel_ok = V_d is not None and v_min <= V_d <= v_max
        fb_ok = FB_i is not None and FB_i >= MIN_FREEBOARD
        pa_ok = PA_i is not None and PA_i >= MIN_FREE_AREA_PERCENT
        mv_ok = V_m is not None and V_m >= v_min

        o.append(f"  23. 流速验证:")
        o.append(f"      不淤流速 ≤ V ≤ 不冲流速")
        if V_d is not None:
            o.append(f"      {v_min} ≤ {V_d:.3f} ≤ {v_max}")
            o.append(f"      结果: {'通过 ✓' if vel_ok else '未通过 ✗'}")
        o.append("")
        o.append(f"  24. 净空高度验证:")
        o.append(f"      Fb加大 ≥ {MIN_FREEBOARD}")
        if FB_i is not None:
            o.append(f"      {FB_i:.3f} ≥ {MIN_FREEBOARD}")
            o.append(f"      结果: {'通过 ✓' if fb_ok else '未通过 ✗'}")
        o.append("")
        o.append(f"  25. 净空面积验证:")
        o.append(f"      PA加大 ≥ {MIN_FREE_AREA_PERCENT}%")
        if PA_i is not None:
            o.append(f"      {PA_i:.1f}% ≥ {MIN_FREE_AREA_PERCENT}%")
            o.append(f"      结果: {'通过 ✓' if pa_ok else '未通过 ✗'}")
        o.append("")
        o.append(f"  26. 最小流速验证:")
        o.append(f"      V最小 ≥ 不淤流速")
        if V_m is not None:
            o.append(f"      {V_m:.3f} ≥ {v_min}")
            o.append(f"      结果: {'通过 ✓' if mv_ok else '未通过 ✗'}")
        o.append("")
        all_pass = vel_ok and fb_ok and pa_ok and mv_ok
        o.append("=" * 70)
        o.append(f"  综合验证结果: {'全部通过 ✓' if all_pass else '未通过 ✗'}")
        o.append("=" * 70)
        txt = "\n".join(o)
        self._export_plain_text = txt
        self.result_text.setPlainText(txt)

    # ================================================================
    # 附录E HTML表格
    # ================================================================
    def _build_ae_html(self, schemes, sel_b, sel_h, v_min, v_max):
        h = '<table class="ae"><tr>'
        for hdr in ['α值', '方案类型', '底宽B(m)', '水深h(m)', '宽深比β', '流速V(m/s)', '面积增加', '状态']:
            h += f'<th>{hdr}</th>'
        h += '</tr>'
        for idx, s in enumerate(schemes):
            alpha, stype = s['alpha'], s['scheme_type']
            sb, sh, sbeta, sV = s['b'], s['h'], s['beta'], s['V']
            area_inc = s['area_increase']
            is_sel = abs(sb - sel_b) < 0.01 and abs(sh - sel_h) < 0.01
            v_ok = v_min < sV < v_max
            if is_sel:
                cls = "sel"; status = "★ 选中"
            elif not v_ok:
                cls = "err"; status = "流速不符"
            else:
                cls = "even" if idx % 2 == 0 else "odd"; status = ""
            h += f'<tr class="{cls}">'
            h += f'<td>{alpha:.2f}</td><td>{stype}</td><td>{sb:.3f}</td><td>{sh:.3f}</td>'
            h += f'<td>{sbeta:.3f}</td><td>{sV:.3f}</td><td>+{area_inc:.0f}%</td>'
            h += f'<td><b>{status}</b></td></tr>'
        h += '</table>'
        return h

    def _build_ae_text(self, schemes, sel_b, sel_h, v_min, v_max):
        lines = []
        lines.append("【附录E断面方案对比表】")
        lines.append("  说明: α=1.00为水力最佳断面(深窄)，α越大断面越宽浅，面积增加但流速降低")
        lines.append("")
        lines.append("  α值    方案类型        底宽B(m)  水深h(m)  宽深比β   流速V(m/s)  面积+  状态")
        lines.append("  " + "-" * 78)
        for s in schemes:
            alpha, stype = s['alpha'], s['scheme_type']
            sb, sh, sbeta, sV = s['b'], s['h'], s['beta'], s['V']
            area_inc = s['area_increase']
            is_sel = abs(sb - sel_b) < 0.01 and abs(sh - sel_h) < 0.01
            v_ok = v_min < sV < v_max
            status = "★选中" if is_sel else ("流速不符" if not v_ok else "")
            lines.append(f"  {alpha:.2f}   {stype:<12}  {sb:8.3f}  {sh:8.3f}  {sbeta:8.3f}  {sV:10.3f}  +{area_inc:.0f}%   {status}")
        lines.append("")
        lines.append(f"  注: 流速约束范围 {v_min} ~ {v_max} m/s")
        return "\n".join(lines)

    # ================================================================
    # 断面图
    # ================================================================
    def _update_section_plot(self, result):
        self.section_fig.clear()
        if not result['success']:
            self.section_canvas.draw()
            return
        stype = self.input_params.get('section_type', '梯形')
        if stype == '圆形':
            D = result.get('D_design', 0)
            y_d = result.get('y_d', 0); V_d = result.get('V_d', 0)
            Q = self.input_params['Q']
            ax = self.section_fig.add_subplot(111)
            self._draw_circular(ax, D, y_d, V_d, Q, '设计流量')
        else:
            b = result['b_design']; h = result['h_design']
            m = self.input_params.get('m', 0); Q = self.input_params['Q']
            V = result['V_design']; h_inc = result['h_increased']
            Q_inc = result['Q_increased']; V_inc = result['V_increased']
            h_prime = result['h_prime']
            axes = self.section_fig.subplots(1, 2)
            self._draw_trapezoid(axes[0], b, h, m, V, Q, h, "设计流量")
            if h_inc > 0:
                self._draw_trapezoid(axes[1], b, h_prime, m, V_inc, Q_inc, h_inc, "加大流量")
            else:
                axes[1].set_title("加大流量\n数据不可用")
        self.section_fig.tight_layout()
        self.section_canvas.draw()

    def _draw_trapezoid(self, ax, b, h_ch, m, V, Q, h_w, title):
        tw = b + 2 * m * h_ch
        ax.plot([-b/2, b/2], [0, 0], 'k-', lw=2)
        ax.plot([-b/2, -tw/2], [0, h_ch], 'k-', lw=2)
        ax.plot([b/2, tw/2], [0, h_ch], 'k-', lw=2)
        ax.plot([-tw/2, tw/2], [h_ch, h_ch], 'k--', lw=1)
        if h_w > 0:
            ww = b + 2 * m * h_w
            wx = [-b/2, -ww/2, ww/2, b/2]
            wy = [0, h_w, h_w, 0]
            ax.fill(wx, wy, color='lightblue', alpha=0.7)
            ax.plot([-ww/2, ww/2], [h_w, h_w], 'b-', lw=1.5)
        # 标注底宽
        ax.annotate('', xy=(b/2, -0.1*h_ch), xytext=(-b/2, -0.1*h_ch),
                     arrowprops=dict(arrowstyle='<->', color='gray', lw=1.5))
        ax.text(0, -0.2*h_ch, f'B={b:.2f}m', ha='center', fontsize=9, color='gray')
        # 标注渠道高度
        ax.annotate('', xy=(tw/2+0.08*tw, h_ch), xytext=(tw/2+0.08*tw, 0),
                     arrowprops=dict(arrowstyle='<->', color='purple', lw=1.5))
        ax.text(tw/2+0.12*tw, h_ch/2, f'H={h_ch:.2f}m', fontsize=9, color='purple', rotation=90, va='center')
        # 标注水深
        if h_w > 0:
            ax.annotate('', xy=(-tw/2-0.08*tw, h_w), xytext=(-tw/2-0.08*tw, 0),
                         arrowprops=dict(arrowstyle='<->', color='blue', lw=1.5))
            ax.text(-tw/2-0.12*tw, h_w/2, f'h={h_w:.2f}m', fontsize=9, color='blue', rotation=90, va='center', ha='right')
        ax.set_xlim(-tw*0.85, tw*0.85)
        ax.set_ylim(-h_ch*0.4, h_ch*1.2)
        ax.set_aspect('equal')
        ax.set_title(f'{title}\nQ={Q:.2f}m$^3$/s, V={V:.2f}m/s', fontsize=10)
        ax.grid(True, alpha=0.3)
        ax.axhline(y=0, color='brown', lw=3)

    def _draw_circular(self, ax, D, y, V, Q, title):
        R = D / 2
        theta = np.linspace(0, 2*np.pi, 100)
        cx = R * np.cos(theta); cy = R + R * np.sin(theta)
        ax.plot(cx, cy, 'k-', lw=2)
        if y > 0 and y < D:
            h_off = y - R
            if abs(h_off) <= R:
                half_a = math.acos(h_off / R)
                water_w = math.sqrt(R**2 - h_off**2)
                wa = np.linspace(np.pi/2 + half_a, np.pi/2 - half_a + 2*np.pi, 50)
                wx = R * np.cos(wa); wy = R + R * np.sin(wa)
                mask = wy <= y + 0.001
                wxf = wx[mask]; wyf = wy[mask]
                if len(wxf) > 0:
                    px = np.concatenate([[water_w], wxf, [-water_w]])
                    py = np.concatenate([[y], wyf, [y]])
                    ax.fill(px, py, color='lightblue', alpha=0.7)
                    ax.plot([-water_w, water_w], [y, y], 'b-', lw=1.5)
        # 标注直径
        ax.annotate('', xy=(R, R), xytext=(-R, R),
                     arrowprops=dict(arrowstyle='<->', color='gray', lw=1.5))
        ax.text(0, R+0.15*R, f'D={D:.2f}m', ha='center', fontsize=9, color='gray')
        # 标注水深
        if y > 0:
            ax.annotate('', xy=(-R-0.12*R, y), xytext=(-R-0.12*R, 0),
                         arrowprops=dict(arrowstyle='<->', color='blue', lw=1.5))
            ax.text(-R-0.2*R, y/2, f'y={y:.2f}m', ha='right', fontsize=8, color='blue', rotation=90, va='center')
        ax.set_xlim(-R*1.7, R*1.7)
        ax.set_ylim(-R*0.4, D*1.2)
        ax.set_aspect('equal')
        ax.set_title(f'{title}\nQ={Q:.2f}m$^3$/s, V={V:.2f}m/s', fontsize=10)
        ax.grid(True, alpha=0.3)
        ax.axhline(y=0, color='brown', lw=3)

    # ================================================================
    # 清空
    # ================================================================
    def _clear(self):
        self._show_initial_help()
        self.section_fig.clear()
        self.section_canvas.draw()
        self.inc_hint.setText("(留空则自动计算)")
        self.current_result = None
        self.statusBar().showMessage("已清空", 3000)

    # ================================================================
    # 导出
    # ================================================================
    def _export_charts(self):
        if not self.current_result or not self.current_result.get('success'):
            InfoBar.warning("提示", "请先进行计算后再导出。", parent=self, duration=3000, position=InfoBarPosition.TOP)
            return
        folder = QFileDialog.getExistingDirectory(self, "选择保存目录")
        if not folder:
            return
        try:
            self.section_fig.savefig(os.path.join(folder, '明渠断面图.png'), dpi=150, bbox_inches='tight')
            InfoBar.success("导出成功", f"图表已保存到: {folder}", parent=self, duration=4000, position=InfoBarPosition.TOP)
        except Exception as e:
            InfoBar.error("导出失败", f"保存失败: {str(e)}", parent=self, duration=5000, position=InfoBarPosition.TOP)

    def _export_report(self):
        if not self.current_result or not self.current_result.get('success'):
            InfoBar.warning("提示", "请先进行计算后再导出。", parent=self, duration=3000, position=InfoBarPosition.TOP)
            return
        filepath, _ = QFileDialog.getSaveFileName(
            self, "保存报告", "", "文本文件 (*.txt);;所有文件 (*.*)")
        if not filepath:
            return
        try:
            content = self._export_plain_text if self._export_plain_text else self.result_text.toPlainText()
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(content)
            InfoBar.success("导出成功", f"报告已保存到: {filepath}", parent=self, duration=4000, position=InfoBarPosition.TOP)
        except Exception as e:
            InfoBar.error("导出失败", f"保存失败: {str(e)}", parent=self, duration=5000, position=InfoBarPosition.TOP)

    # ================================================================
    # 导出Word（含可编辑公式）
    # ================================================================
    def _export_word(self):
        if not WORD_EXPORT_AVAILABLE:
            InfoBar.warning("缺少依赖",
                "Word导出需要安装 python-docx、latex2mathml、lxml。请执行: pip install python-docx latex2mathml lxml",
                parent=self, duration=6000, position=InfoBarPosition.TOP)
            return
        if not self.current_result or not self.current_result.get('success'):
            InfoBar.warning("提示", "请先进行计算后再导出。", parent=self, duration=3000, position=InfoBarPosition.TOP)
            return

        filepath, _ = QFileDialog.getSaveFileName(
            self, "保存Word报告", "", "Word文档 (*.docx);;所有文件 (*.*)")
        if not filepath:
            return

        try:
            self._build_word_report(filepath)
            InfoBar.success("导出成功", f"Word报告已保存到: {filepath}", parent=self, duration=4000, position=InfoBarPosition.TOP)
        except Exception as e:
            InfoBar.error("导出失败", f"Word导出失败: {str(e)}", parent=self, duration=5000, position=InfoBarPosition.TOP)

    def _build_word_report(self, filepath):
        """构建Word报告文档（含可编辑LaTeX公式）"""
        doc = DocxDocument()
        res = self.current_result
        p = self.input_params
        stype = p.get('section_type', '梯形')

        # 设置默认字体
        style = doc.styles['Normal']
        style.font.name = 'Microsoft YaHei'
        style.font.size = Pt(11)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

        # ---- 标题 ----
        title = doc.add_heading('', level=0)
        run = title.add_run(f'明渠水力计算报告（{stype}断面）')

        # ---- 基本信息 ----
        doc.add_paragraph(f'计算方法：{res.get("design_method", "")}')
        doc.add_paragraph('')

        # ---- 基础公式（可编辑） ----
        doc.add_heading('基础公式', level=1)
        _add_formula_to_doc(doc, r'Q = \frac{1}{n} \cdot A \cdot R^{2/3} \cdot i^{1/2}', '曼宁公式：')
        if stype == '圆形':
            _add_formula_to_doc(doc, r'A = \frac{D^2}{8}(\theta - \sin\theta)', '过水面积：')
            _add_formula_to_doc(doc, r'\chi = \frac{D}{2} \cdot \theta', '湿周：')
            _add_formula_to_doc(doc, r'\theta = 2\arccos\frac{R-h}{R}', '圆心角：')
        elif stype == '梯形':
            _add_formula_to_doc(doc, r'A = (B + m \cdot h) \cdot h', '过水面积：')
            _add_formula_to_doc(doc, r'\chi = B + 2h\sqrt{1+m^2}', '湿周：')
        else:
            _add_formula_to_doc(doc, r'A = B \cdot h', '过水面积：')
            _add_formula_to_doc(doc, r'\chi = B + 2h', '湿周：')
        _add_formula_to_doc(doc, r'R = \frac{A}{\chi}', '水力半径：')
        _add_formula_to_doc(doc, r'V = \frac{1}{n} \cdot R^{2/3} \cdot i^{1/2}', '流速公式：')

        # ---- 计算过程（按行解析纯文本结果） ----
        doc.add_heading('计算过程', level=1)
        text = self._export_plain_text
        if not text:
            text = self.result_text.toPlainText()

        lines = text.split('\n')
        for line in lines:
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph('')
                continue
            # 跳过装饰分隔线
            if set(stripped) <= {'=', '-', ' '} and len(stripped) > 5:
                continue
            # 标题行（如"明渠水力计算结果（梯形断面）"）
            if '明渠水力计算结果' in stripped:
                continue
            # 章节标题 【一、输入参数】
            if stripped.startswith('【') and '】' in stripped:
                heading_text = stripped.lstrip('【').rstrip('】')
                doc.add_heading(heading_text, level=2)
                continue
            # 带编号的计算步骤标题（如"  1. 设计底宽:"）
            step_match = re.match(r'^(\d+)\.\s+(.+)', stripped)
            if step_match:
                doc.add_heading(f'{step_match.group(1)}. {step_match.group(2)}', level=3)
                continue
            # 公式计算行 —— 尝试转为可编辑公式
            # 识别模式："V = (1/n) × R^(2/3) × ..."  或  "= 71.43 × 0.543 × 0.018"
            formula_line = self._try_convert_formula_line(stripped)
            if formula_line:
                _add_formula_to_doc(doc, formula_line, '    ')
                continue
            # 普通文本段落
            para = doc.add_paragraph(stripped)
            para.paragraph_format.space_before = Pt(1)
            para.paragraph_format.space_after = Pt(1)
            for run in para.runs:
                run.font.name = 'Microsoft YaHei'
                run.font.size = Pt(10.5)

        # ---- 附录E表格（如果有） ----
        schemes = res.get('appendix_e_schemes', [])
        if schemes and stype != '圆形':
            doc.add_heading('附录E 断面方案对比表', level=2)
            doc.add_paragraph('说明: α=1.00为水力最佳断面(深窄)，α越大断面越宽浅，面积增加但流速降低')
            b_sel = res['b_design']
            h_sel = res['h_design']
            v_min, v_max = p['v_min'], p['v_max']
            headers = ['α值', '方案类型', '底宽B(m)', '水深h(m)', '宽深比β', '流速V(m/s)', '面积增加', '状态']
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Light Shading Accent 1'
            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = h
            for s in schemes:
                is_sel = abs(s['b'] - b_sel) < 0.01 and abs(s['h'] - h_sel) < 0.01
                v_ok = v_min < s['V'] < v_max
                status = "★选中" if is_sel else ("流速不符" if not v_ok else "")
                row = table.add_row()
                vals = [f"{s['alpha']:.2f}", s['scheme_type'], f"{s['b']:.3f}",
                        f"{s['h']:.3f}", f"{s['beta']:.3f}", f"{s['V']:.3f}",
                        f"+{s['area_increase']:.0f}%", status]
                for i, v in enumerate(vals):
                    row.cells[i].text = v

        # ---- 断面图（嵌入） ----
        try:
            import tempfile
            tmp = os.path.join(tempfile.gettempdir(), '_mingqu_section.png')
            self.section_fig.savefig(tmp, dpi=150, bbox_inches='tight')
            doc.add_heading('断面图', level=2)
            doc.add_picture(tmp, width=Cm(15))
            os.remove(tmp)
        except Exception:
            pass

        doc.save(filepath)

    def _try_convert_formula_line(self, line):
        """尝试将计算行转换为LaTeX公式字符串，不适合则返回None"""
        # 跳过纯文字描述
        if not any(c in line for c in ['=', '×', '√', '^']):
            return None
        # 跳过含中文大段描述的行
        cn_count = sum(1 for c in line if '\u4e00' <= c <= '\u9fff')
        if cn_count > 4:
            return None
        # 跳过验证结果行
        if '✓' in line or '✗' in line or '通过' in line or '未通过' in line:
            return None
        # 匹配公式行模式："X = ... = ... = 数值 单位"
        m = re.match(r'^([A-Za-zα-ωΑ-Ωβχ\u4e00-\u9fff]+\s*=\s*.+)', line)
        if not m:
            m = re.match(r'^(=\s*.+)', line)
        if not m:
            return None
        expr = line.strip()
        # 转换为LaTeX
        latex = expr
        latex = latex.replace('×', r' \times ')
        latex = latex.replace('√', r'\sqrt')
        latex = latex.replace('π', r'\pi')
        latex = latex.replace('χ', r'\chi')
        latex = latex.replace('β', r'\beta')
        latex = latex.replace('θ', r'\theta')
        # 处理上标  R^(2/3) → R^{2/3}
        latex = re.sub(r'\^[\(（]([^)）]+)[\)）]', r'^{\1}', latex)
        # 处理 m² → m^{2}
        latex = latex.replace('²', '^{2}')
        latex = latex.replace('³', '^{3}')
        # 处理 m³/s → \text{m}^3\text{/s}
        latex = re.sub(r'm\^?\{?3\}?/s', r'\\text{ m}^3\\text{/s}', latex)
        # 处理单位 m, m² 后缀
        latex = re.sub(r'\s+m$', r' \\text{ m}', latex)
        latex = re.sub(r'\s+m²$', r' \\text{ m}^2', latex)
        return latex


# ============================================================
# 入口
# ============================================================
def main():
    app = QApplication(sys.argv)
    app.setStyle("Fusion")
    app.setFont(QFont("Microsoft YaHei", 10))
    app.setStyleSheet(STYLE)
    window = OpenChannelPanel()
    window.show()
    sys.exit(app.exec())


if __name__ == "__main__":
    main()
