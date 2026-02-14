# -*- coding: utf-8 -*-
"""
导出工具 —— Word（含可编辑LaTeX公式）/ TXT 导出辅助函数
"""

import os
import re
import subprocess

# Word导出依赖（可选）
try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    import latex2mathml.converter
    from lxml import etree
    WORD_EXPORT_AVAILABLE = True
except ImportError:
    WORD_EXPORT_AVAILABLE = False


# ============================================================
# OMML转换（LaTeX → Word可编辑公式）
# ============================================================

def _find_mml2omml_xsl():
    """查找Office的MML2OMML.XSL文件"""
    candidates = [
        r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL",
        r"C:\Program Files (x86)\Microsoft Office\root\Office16\MML2OMML.XSL",
        r"C:\Program Files\Microsoft Office\Office16\MML2OMML.XSL",
        r"C:\Program Files\Microsoft Office\root\Office15\MML2OMML.XSL",
    ]
    for p in candidates:
        if os.path.exists(p):
            return p
    return None


_MML2OMML_XSL_PATH = _find_mml2omml_xsl() if WORD_EXPORT_AVAILABLE else None
_MML2OMML_TRANSFORM = None


def _get_omml_transform():
    """延迟加载XSLT转换器"""
    global _MML2OMML_TRANSFORM
    if _MML2OMML_TRANSFORM is None and _MML2OMML_XSL_PATH:
        xslt_tree = etree.parse(_MML2OMML_XSL_PATH)
        _MML2OMML_TRANSFORM = etree.XSLT(xslt_tree)
    return _MML2OMML_TRANSFORM


def latex_to_omml(latex_str):
    """将LaTeX公式转换为Word可编辑的OMML元素"""
    transform = _get_omml_transform()
    if transform is None:
        return None
    try:
        mathml_str = latex2mathml.converter.convert(latex_str)
        mathml_tree = etree.fromstring(mathml_str.encode('utf-8'))
        omml_result = transform(mathml_tree)
        return omml_result.getroot()
    except Exception:
        return None


def add_formula_to_doc(doc, latex_str, prefix=""):
    """向Word文档添加一个可编辑公式段落"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if prefix:
        run = p.add_run(prefix)
        run.font.name = 'Microsoft YaHei'
        run.font.size = Pt(11)
    omml = latex_to_omml(latex_str)
    if omml is not None:
        p._element.append(omml)
    else:
        run = p.add_run(f" {latex_str}")
        run.font.name = 'Cambria Math'
        run.font.size = Pt(11)
    return p


# ============================================================
# 公式行转换辅助
# ============================================================

def try_convert_formula_line(line):
    """尝试将计算行转换为LaTeX公式字符串，不适合则返回None"""
    if not any(c in line for c in ['=', '×', '√', '^']):
        return None
    cn_count = sum(1 for c in line if '\u4e00' <= c <= '\u9fff')
    if cn_count > 4:
        return None
    if '✓' in line or '✗' in line or '通过' in line or '未通过' in line:
        return None
    m = re.match(r'^([A-Za-zα-ωΑ-Ωβχ\u4e00-\u9fff]+\s*=\s*.+)', line)
    if not m:
        m = re.match(r'^(=\s*.+)', line)
    if not m:
        return None
    expr = line.strip()
    latex = expr
    latex = latex.replace('×', r' \times ')
    latex = latex.replace('√', r'\sqrt')
    latex = latex.replace('π', r'\pi')
    latex = latex.replace('χ', r'\chi')
    latex = latex.replace('β', r'\beta')
    latex = latex.replace('θ', r'\theta')
    latex = re.sub(r'\^[\(（]([^)）]+)[\)）]', r'^{\1}', latex)
    latex = latex.replace('²', '^{2}')
    latex = latex.replace('³', '^{3}')
    latex = re.sub(r'm\^?\{?3\}?/s', r'\\text{ m}^3\\text{/s}', latex)
    latex = re.sub(r'\s+m$', r' \\text{ m}', latex)
    latex = re.sub(r'\s+m²$', r' \\text{ m}^2', latex)
    return latex


# ============================================================
# 导出后询问并打开文件
# ============================================================

def ask_open_file(filepath, parent=None):
    """导出成功后询问用户是否用默认程序打开文件"""
    from 渠系断面设计.styles import fluent_question
    if fluent_question(
        parent, "导出成功",
        f"文件已保存到:\n{filepath}\n\n是否立即打开该文件？",
        yes_text="打开", no_text="关闭"
    ):
        try:
            os.startfile(filepath)
        except Exception:
            try:
                subprocess.Popen(['start', '', filepath], shell=True)
            except Exception:
                pass


# ============================================================
# 高端咨询报告风格排版工具（方案3）
# ============================================================

# --- 配色常量 ---
WORD_NAVY = "1B2A4A"
WORD_GOLD = "C9A96E"
WORD_WARM_GRAY = "F8F6F3"
WORD_TEXT_DARK = "2C3E50"
WORD_TEXT_MID = "5D6D7E"

_NAVY_RGB = RGBColor(0x1B, 0x2A, 0x4A) if WORD_EXPORT_AVAILABLE else None
_GOLD_RGB = RGBColor(0xC9, 0xA9, 0x6E) if WORD_EXPORT_AVAILABLE else None
_TEXT_DARK_RGB = RGBColor(0x2C, 0x3E, 0x50) if WORD_EXPORT_AVAILABLE else None
_TEXT_MID_RGB = RGBColor(0x5D, 0x6D, 0x7E) if WORD_EXPORT_AVAILABLE else None
_WHITE_RGB = RGBColor(0xFF, 0xFF, 0xFF) if WORD_EXPORT_AVAILABLE else None


def _set_cell_shading(cell, color_hex):
    """设置单元格底色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_border(cell, **kwargs):
    """设置单元格边框 kwargs: top/bottom/left/right = (size, color, style)"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
    for edge, (sz, color, style) in kwargs.items():
        el = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{style}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        )
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _set_paragraph_border_bottom(paragraph, color="000000", size=6):
    """段落底部边框线"""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:space="1" w:color="{color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def _set_table_borders(table, top_color, bottom_color, inside_h_color,
                       inside_v_color=None, left_color=None, right_color=None,
                       top_sz="8", bottom_sz="8", inside_h_sz="2"):
    """设置表格边框样式"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    iv = f'<w:insideV w:val="single" w:sz="2" w:space="0" w:color="{inside_v_color}"/>' if inside_v_color else \
         f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="FFFFFF"/>'
    lv = f'<w:left w:val="single" w:sz="2" w:space="0" w:color="{left_color}"/>' if left_color else \
         f'<w:left w:val="none" w:sz="0" w:space="0" w:color="FFFFFF"/>'
    rv = f'<w:right w:val="single" w:sz="2" w:space="0" w:color="{right_color}"/>' if right_color else \
         f'<w:right w:val="none" w:sz="0" w:space="0" w:color="FFFFFF"/>'
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{top_sz}" w:space="0" w:color="{top_color}"/>'
        f'  <w:bottom w:val="single" w:sz="{bottom_sz}" w:space="0" w:color="{bottom_color}"/>'
        f'  <w:insideH w:val="single" w:sz="{inside_h_sz}" w:space="0" w:color="{inside_h_color}"/>'
        f'  {iv}{lv}{rv}'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def _add_page_number(section, font_name='微软雅黑', font_size=9):
    """添加页脚页码"""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p.add_run("— ")
    run1.font.name = font_name
    run1.font.size = Pt(font_size)
    fld_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run_f1 = p.add_run()
    run_f1._r.append(fld_begin)
    instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run_f2 = p.add_run()
    run_f2._r.append(instr)
    fld_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run_f3 = p.add_run()
    run_f3._r.append(fld_end)
    run2 = p.add_run(" —")
    run2.font.name = font_name
    run2.font.size = Pt(font_size)


def _run_styled(paragraph, text, font_name='微软雅黑', font_size=11,
                color=None, bold=False, italic=False):
    """向段落添加一个带样式的run"""
    r = paragraph.add_run(text)
    r.font.name = font_name
    r._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    r.font.size = Pt(font_size)
    if color:
        r.font.color.rgb = color
    if bold:
        r.bold = True
    if italic:
        r.italic = True
    return r


_LOGO_PNG_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'resources', 'logo.png')
_APP_NAME = "渠系建筑物水力计算系统"
_APP_VERSION = "V1.0"
_DEFAULT_DESIGN_UNIT = "四川水发勘测设计研究有限公司"


def create_styled_doc(title, subtitle="", header_text="",
                      channel_name="", channel_level="",
                      start_station="", end_station=""):
    """创建一个方案3风格的Word文档（含封面、页眉页脚），返回doc对象
    
    Args:
        title: 封面大标题，如 "明渠水力计算书"
        subtitle: 封面副标题，如 "梯形断面 · 经济断面法"
        header_text: 页眉文字
        channel_name: 渠道名称（从用户输入读取）
        channel_level: 渠道级别（从用户输入读取）
        start_station: 起始桩号
        end_station: 终止桩号
    """
    import datetime
    doc = DocxDocument()
    
    # --- 页面设置 ---
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)
    
    # --- 默认样式 ---
    style_normal = doc.styles['Normal']
    style_normal.font.name = '微软雅黑'
    style_normal.font.size = Pt(11)
    style_normal._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    # 必须使用显式比例行距（非固定磅值），否则大字号标题/Logo/图片会被裁切
    style_normal.paragraph_format.line_spacing = 1.15
    
    # ==================== 封面 ====================
    # 上方留白
    for _ in range(2):
        doc.add_paragraph('')
    
    # 软件Logo
    if os.path.exists(_LOGO_PNG_PATH):
        logo_p = doc.add_paragraph()
        logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_p.paragraph_format.line_spacing = 1.0
        logo_p.add_run().add_picture(_LOGO_PNG_PATH, width=Cm(2.2))
    
    # 软件名称
    app_p = doc.add_paragraph()
    app_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    app_p.paragraph_format.space_before = Pt(6)
    app_p.paragraph_format.space_after = Pt(2)
    _run_styled(app_p, _APP_NAME, '微软雅黑', 11, _TEXT_MID_RGB)
    
    # 版本号
    ver_p = doc.add_paragraph()
    ver_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ver_p.paragraph_format.space_before = Pt(0)
    ver_p.paragraph_format.space_after = Pt(16)
    _run_styled(ver_p, _APP_VERSION, '微软雅黑', 9, _GOLD_RGB)
    
    # 金色分割线（上）
    line1 = doc.add_paragraph()
    line1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line1.paragraph_format.space_after = Pt(16)
    _set_paragraph_border_bottom(line1, color=WORD_GOLD, size=4)
    
    # 大标题
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(8)
    title_p.paragraph_format.space_after = Pt(8)
    title_p.paragraph_format.line_spacing = 1.0
    _run_styled(title_p, title, '黑体', 32, _NAVY_RGB, bold=True)
    
    # 副标题
    if subtitle:
        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_p.paragraph_format.space_before = Pt(4)
        sub_p.paragraph_format.space_after = Pt(4)
        _run_styled(sub_p, subtitle, '微软雅黑', 14, _TEXT_MID_RGB)
    
    # 金色分割线（下）
    line2 = doc.add_paragraph()
    line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line2.paragraph_format.space_before = Pt(16)
    _set_paragraph_border_bottom(line2, color=WORD_GOLD, size=4)
    
    # 留白
    for _ in range(3):
        doc.add_paragraph('')
    
    # 封面信息表
    cover_items = []
    cover_items.append(("设计单位", _DEFAULT_DESIGN_UNIT))
    if channel_name:
        cover_items.append(("渠道名称", channel_name))
    if channel_level:
        cover_items.append(("渠道级别", channel_level))
    if start_station or end_station:
        station_text = ""
        if start_station and end_station:
            station_text = f"{start_station} ~ {end_station}"
        elif start_station:
            station_text = f"{start_station} ~"
        else:
            station_text = f"~ {end_station}"
        cover_items.append(("桩  号", station_text))
    cover_items.append(("计算日期", datetime.datetime.now().strftime("%Y年%m月")))
    
    _add_cover_info_table(doc, cover_items)
    
    # ==================== 页眉页脚 ====================
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if header_text:
        _run_styled(hp, header_text, '微软雅黑', 8, _GOLD_RGB)
    _set_paragraph_border_bottom(hp, color=WORD_GOLD, size=3)
    
    _add_page_number(section)
    
    return doc


def _add_cover_info_table(doc, items):
    """在封面添加信息表（内部方法）"""
    info_tbl = doc.add_table(rows=len(items), cols=2)
    info_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(items):
        c0 = info_tbl.cell(i, 0)
        c1 = info_tbl.cell(i, 1)
        c0.text = label
        c1.text = value
        c0.width = Cm(3)
        c1.width = Cm(8.5)
        for p in c0.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for r in p.runs:
                r.font.name = '微软雅黑'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                r.font.size = Pt(12)
                r.font.color.rgb = _GOLD_RGB
        for p in c1.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent = Cm(0.5)
            for r in p.runs:
                r.font.name = '微软雅黑'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                r.font.size = Pt(12)
                r.font.color.rgb = _TEXT_DARK_RGB
        for c in (c0, c1):
            _set_cell_border(c,
                top=("0","FFFFFF","none"), left=("0","FFFFFF","none"),
                right=("0","FFFFFF","none"), bottom=("2", WORD_GOLD, "single"))


def doc_add_h1(doc, text):
    """添加一级标题（黑体 + 金色底线）"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.0
    _run_styled(p, text, '黑体', 18, _NAVY_RGB, bold=True)
    _set_paragraph_border_bottom(p, color=WORD_GOLD, size=6)
    return p


def doc_add_h2(doc, text):
    """添加二级标题（微软雅黑加粗 深藏青）"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    _run_styled(p, text, '微软雅黑', 13, _NAVY_RGB, bold=True)
    return p


def doc_add_h3(doc, text):
    """添加三级标题/步骤标题（微软雅黑加粗 12pt）"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    _run_styled(p, text, '微软雅黑', 12, _NAVY_RGB, bold=True)
    return p


def doc_add_body(doc, text, indent=False):
    """添加正文段落"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(1.2)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    _run_styled(p, text, '微软雅黑', 11, _TEXT_DARK_RGB)
    return p


def doc_add_formula(doc, latex_str, prefix=""):
    """添加公式段落（方案3风格缩进 + OMML可编辑公式）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Cm(1.2)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    if prefix:
        _run_styled(p, prefix, '微软雅黑', 11, _TEXT_DARK_RGB)
    omml = latex_to_omml(latex_str)
    if omml is not None:
        p._element.append(omml)
    else:
        _run_styled(p, f" {latex_str}", 'Cambria Math', 11, _TEXT_DARK_RGB)
    return p


def doc_add_styled_table(doc, headers, data, highlight_col=None, highlight_val=None,
                         with_full_border=False):
    """添加方案3风格的数据表格
    
    Args:
        headers: 列标题列表
        data: 二维数据列表
        highlight_col: 高亮匹配的列索引
        highlight_val: 高亮匹配的值
        with_full_border: 是否显示完整网格线
    """
    table = doc.add_table(rows=len(data)+1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 表头
    for j, hdr in enumerate(headers):
        c = table.cell(0, j)
        c.text = hdr
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = '微软雅黑'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                r.font.size = Pt(9)
                r.bold = True
                r.font.color.rgb = _WHITE_RGB
        _set_cell_shading(c, WORD_NAVY)
    
    # 数据行
    for i, row_data in enumerate(data):
        for j, val in enumerate(row_data):
            c = table.cell(i+1, j)
            c.text = str(val)
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.name = '微软雅黑'
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    r.font.size = Pt(9)
                    r.font.color.rgb = _TEXT_DARK_RGB
                    if highlight_col is not None and j == highlight_col and highlight_val and val == highlight_val:
                        r.font.color.rgb = _GOLD_RGB
                        r.bold = True
        if i % 2 == 0:
            for j in range(len(headers)):
                _set_cell_shading(table.cell(i+1, j), WORD_WARM_GRAY)
    
    # 边框
    iv_color = "E8E4DF" if with_full_border else None
    lv_color = "E8E4DF" if with_full_border else None
    rv_color = "E8E4DF" if with_full_border else None
    _set_table_borders(table, WORD_NAVY, WORD_NAVY, "E8E4DF",
                       inside_v_color=iv_color, left_color=lv_color, right_color=rv_color)
    return table


def doc_add_result_table(doc, items):
    """添加计算结果汇总表（项目+结果 两列，金色表头）
    
    Args:
        items: list of (name, value) 元组
    """
    table = doc.add_table(rows=len(items)+1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, hdr in enumerate(["项  目", "计算结果"]):
        c = table.cell(0, j)
        c.text = hdr
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = '微软雅黑'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                r.font.size = Pt(10)
                r.bold = True
                r.font.color.rgb = _GOLD_RGB
        _set_cell_shading(c, WORD_NAVY)
    
    for i, (name, val) in enumerate(items):
        c0 = table.cell(i+1, 0)
        c1 = table.cell(i+1, 1)
        c0.text = name
        c1.text = val
        for p in c0.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent = Cm(0.3)
            for r in p.runs:
                r.font.name = '微软雅黑'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                r.font.size = Pt(11)
                r.font.color.rgb = _TEXT_MID_RGB
        for p in c1.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = '微软雅黑'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                r.font.size = Pt(11)
                r.font.color.rgb = _NAVY_RGB
                r.bold = True
        if i % 2 == 0:
            for j in range(2):
                _set_cell_shading(table.cell(i+1, j), WORD_WARM_GRAY)
    
    _set_table_borders(table, WORD_NAVY, WORD_NAVY, "E8E4DF")
    return table


def doc_add_param_table(doc, params):
    """添加基本参数表（参数+取值 两列）
    
    Args:
        params: list of (name, value) 元组
    """
    table = doc.add_table(rows=len(params)+1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, hdr in enumerate(["参  数", "取  值"]):
        c = table.cell(0, j)
        c.text = hdr
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = '微软雅黑'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                r.font.size = Pt(10)
                r.bold = True
                r.font.color.rgb = _WHITE_RGB
        _set_cell_shading(c, WORD_NAVY)
    
    for i, (name, val) in enumerate(params):
        for j, txt in enumerate([name, val]):
            c = table.cell(i+1, j)
            c.text = txt
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.name = '微软雅黑'
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    r.font.size = Pt(10)
                    r.font.color.rgb = _TEXT_DARK_RGB
        if i % 2 == 0:
            for j in range(2):
                _set_cell_shading(table.cell(i+1, j), WORD_WARM_GRAY)
    
    _set_table_borders(table, WORD_NAVY, WORD_NAVY, "E8E4DF")
    return table


def doc_add_table_caption(doc, text):
    """添加表标题（居中、斜体、灰色）"""
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(6)
    _run_styled(cap, text, '微软雅黑', 10, _TEXT_MID_RGB, italic=True)
    return cap


def doc_add_figure(doc, image_path, width_cm=14):
    """添加居中图片（独立段落，显式比例行距避免裁切）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run()
    run.add_picture(image_path, width=Cm(width_cm))
    return p


def doc_render_calc_text(doc, text, skip_title_keyword=""):
    """将纯文本计算过程渲染为方案3风格的Word内容
    
    自动识别：
    - 【xxx】 → 二级标题
    - 数字. xxx → 三级标题
    - 公式行 → OMML公式
    - 普通行 → 正文段落
    
    Args:
        text: 纯文本计算过程
        skip_title_keyword: 跳过包含此关键词的行（如"水力计算结果"）
    """
    if not text:
        return
    for line in text.split('\n'):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph('')
            continue
        if set(stripped) <= {'=', '-', ' '} and len(stripped) > 5:
            continue
        if skip_title_keyword and skip_title_keyword in stripped:
            continue
        if stripped.startswith('【') and '】' in stripped:
            doc_add_h2(doc, stripped.lstrip('【').rstrip('】'))
            continue
        step_m = re.match(r'^(\d+)\.\s+(.+)', stripped)
        if step_m:
            doc_add_h3(doc, f'{step_m.group(1)}. {step_m.group(2)}')
            continue
        fl = try_convert_formula_line(stripped)
        if fl:
            doc_add_formula(doc, fl, '    ')
            continue
        doc_add_body(doc, stripped)
