# -*- coding: utf-8 -*-
"""
明渠水力计算面板 —— QWidget 版本（可嵌入主导航框架）

支持：梯形/矩形/圆形断面
功能：参数输入、计算、结果显示、断面图、导出Word/TXT/图表
"""

import sys
import os
import math
import re
import html as html_mod

# 将计算模块目录加入搜索路径
_pkg_root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
sys.path.insert(0, os.path.join(_pkg_root, "渠系建筑物断面计算"))

from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QGroupBox,
    QSplitter, QFrame, QTabWidget, QFileDialog, QScrollArea
)
from PySide6.QtCore import Qt
from PySide6.QtGui import QFont
from PySide6.QtWebEngineWidgets import QWebEngineView

from qfluentwidgets import (
    ComboBox, PushButton, PrimaryPushButton, LineEdit,
    CheckBox, InfoBar, InfoBarPosition
)

import matplotlib
matplotlib.use('QtAgg')
import matplotlib.pyplot as plt
try:
    from matplotlib.backends.backend_qtagg import FigureCanvasQTAgg as FigureCanvas
    from matplotlib.backends.backend_qtagg import NavigationToolbar2QT as NavToolbar
except ImportError:
    from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
    from matplotlib.backends.backend_qt5agg import NavigationToolbar2QT as NavToolbar
from matplotlib.figure import Figure
import numpy as np

plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'SimSun']
plt.rcParams['axes.unicode_minus'] = False

# 计算引擎
from 明渠设计 import (
    quick_calculate as mingqu_calculate,
    quick_calculate_circular as circular_calculate,
    quick_calculate_u_section as mingqu_u_calculate,
    _u_arc_geometry,
    calculate_area, calculate_wetted_perimeter, calculate_hydraulic_radius,
    get_flow_increase_percent, MAX_BETA,
    PI, MIN_FREEBOARD, MIN_FREE_AREA_PERCENT, MIN_FLOW_FACTOR
)

# 共享模块
from 渠系断面设计.styles import P, S, W, E, BG, CARD, BD, T1, T2, AE_CSS, INPUT_LABEL_STYLE, INPUT_SECTION_STYLE, INPUT_HINT_STYLE
from 渠系断面设计.export_utils import (
    WORD_EXPORT_AVAILABLE, add_formula_to_doc, try_convert_formula_line, ask_open_file,
    create_styled_doc, doc_add_h1, doc_add_h2,
    doc_add_formula, doc_add_styled_table, doc_add_table_caption,
    doc_render_calc_text, doc_add_figure,
    create_engineering_report_doc, doc_add_eng_h, doc_add_eng_body,
    doc_render_calc_text_eng, update_doc_toc_via_com, doc_add_table_caption,
)
from 渠系断面设计.report_meta import (
    ExportConfirmDialog, build_calc_purpose, REFERENCES_BASE, load_meta
)
from 渠系断面设计.open_channel.dxf_export import export_open_channel_dxf
from 渠系断面设计.formula_renderer import (
    plain_text_to_formula_html, plain_text_to_formula_body,
    wrap_with_katex, load_formula_page, make_plain_html,
    HelpPageBuilder
)
if WORD_EXPORT_AVAILABLE:
    from docx import Document as DocxDocument
    from docx.shared import Pt, Cm


def _e(s):
    """HTML转义"""
    return html_mod.escape(str(s))


class OpenChannelPanel(QWidget):
    """明渠水力计算面板"""

    def __init__(self, parent=None):
        super().__init__(parent)
        self.input_params = {}
        self.current_result = None
        self._appendix_e_export_text = ""
        self._export_plain_text = ""
        self._init_ui()

    # ================================================================
    # UI 构建
    # ================================================================
    def _init_ui(self):
        main_lay = QHBoxLayout(self)
        main_lay.setContentsMargins(10, 8, 10, 8)

        splitter = QSplitter(Qt.Horizontal)
        splitter.setChildrenCollapsible(False)
        main_lay.addWidget(splitter)

        # 左侧输入
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        scroll.setStyleSheet("QScrollArea{border:none;background:transparent;}")
        inp_w = QWidget()
        self._build_input(inp_w)
        scroll.setWidget(inp_w)
        scroll.setMinimumWidth(280)
        scroll.setMaximumWidth(420)
        splitter.addWidget(scroll)

        # 右侧输出
        out_w = QWidget()
        self._build_output(out_w)
        splitter.addWidget(out_w)
        splitter.setSizes([340, 900])

    # ----------------------------------------------------------------
    # 输入面板
    # ----------------------------------------------------------------
    def _build_input(self, parent):
        lay = QVBoxLayout(parent)
        lay.setContentsMargins(5, 5, 5, 5)
        lay.setSpacing(6)
        grp = QGroupBox("输入参数")
        fl = QVBoxLayout(grp)
        fl.setSpacing(5)

        # 断面类型
        r = QHBoxLayout(); r.addWidget(QLabel("断面类型:"))
        self.section_combo = ComboBox()
        self.section_combo.addItems(["梯形", "矩形", "圆形", "U形"])
        self.section_combo.currentTextChanged.connect(self._on_section_type_changed)
        r.addWidget(self.section_combo, 1); fl.addLayout(r)

        self.Q_edit = self._field(fl, "设计流量 Q (m³/s):", "5.0")
        self.m_lbl, self.m_edit = self._field2(fl, "边坡系数 m:", "1.0")
        self.n_edit = self._field(fl, "糙率 n:", "0.014")
        self.slope_edit = self._field(fl, "水力坡降 1/", "3000")

        fl.addWidget(self._slbl("【流速参数】"))
        self.vmin_edit = self._field(fl, "不淤流速 (m/s):", "0.1")
        self.vmax_edit = self._field(fl, "不冲流速 (m/s):", "100.0")
        fl.addWidget(self._hint("(一般情况下保持默认数值即可)"))

        self.inc_cb = CheckBox("考虑加大流量比例系数")
        self.inc_cb.setChecked(True)
        self.inc_cb.stateChanged.connect(self._on_inc_toggle)
        fl.addWidget(self.inc_cb)
        self.inc_edit = self._field(fl, "流量加大比例 (%):", "")
        self.inc_hint = QLabel("(留空则自动计算)")
        self.inc_hint.setStyleSheet(INPUT_HINT_STYLE)
        fl.addWidget(self.inc_hint)

        fl.addWidget(self._sep())
        fl.addWidget(self._slbl("【可选参数】"))
        self.beta_lbl, self.beta_edit = self._field2(fl, "指定宽深比 β:", "")
        self.b_lbl, self.b_edit = self._field2(fl, "指定底宽 B (m):", "")
        self.bb_hint = self._hint("(二选一输入，留空则自动计算)")
        fl.addWidget(self.bb_hint)

        self.D_lbl, self.D_edit = self._field2(fl, "指定直径 D (m):", "")
        self.D_hint_lbl = self._hint("(留空则自动计算)")
        fl.addWidget(self.D_hint_lbl)
        for w in (self.D_lbl, self.D_edit, self.D_hint_lbl): w.hide()

        # U形专有字段
        self.R_lbl, self.R_edit = self._field2(fl, "圆弧半径 R (m):", "0.8")
        self.alpha_lbl, self.alpha_edit = self._field2(fl, "外倾角 α (°):", "14")
        self.theta_lbl, self.theta_edit = self._field2(fl, "圆心角 θ (°):", "152")
        for w in (self.R_lbl, self.R_edit, self.alpha_lbl, self.alpha_edit,
                  self.theta_lbl, self.theta_edit): w.hide()

        fl.addWidget(self._sep())
        self.detail_cb = CheckBox("输出详细计算过程")
        self.detail_cb.setChecked(True)
        fl.addWidget(self.detail_cb)

        br = QHBoxLayout()
        cb = PrimaryPushButton("计算"); cb.setCursor(Qt.PointingHandCursor); cb.clicked.connect(self._calculate)
        clb = PushButton("清空"); clb.setCursor(Qt.PointingHandCursor); clb.clicked.connect(self._clear)
        br.addWidget(cb); br.addWidget(clb); fl.addLayout(br)

        er = QHBoxLayout()
        ec = PushButton("导出DXF"); ec.clicked.connect(self._export_dxf)
        ew = PushButton("导出Word"); ew.clicked.connect(self._export_word)
        er.addWidget(ec); er.addWidget(ew)
        fl.addLayout(er)

        lay.addWidget(grp)
        lay.addStretch()

    def _field(self, lay, label, default=""):
        r = QHBoxLayout(); l = QLabel(label); l.setMinimumWidth(140)
        l.setStyleSheet(INPUT_LABEL_STYLE)
        r.addWidget(l); e = LineEdit(); e.setText(default); r.addWidget(e, 1); lay.addLayout(r)
        return e

    def _field2(self, lay, label, default=""):
        r = QHBoxLayout(); l = QLabel(label); l.setMinimumWidth(140)
        l.setStyleSheet(INPUT_LABEL_STYLE)
        r.addWidget(l); e = LineEdit(); e.setText(default); r.addWidget(e, 1); lay.addLayout(r)
        return l, e

    def _slbl(self, t):
        l = QLabel(t); l.setStyleSheet(INPUT_SECTION_STYLE); return l

    def _hint(self, t):
        l = QLabel(t); l.setStyleSheet(INPUT_HINT_STYLE); return l

    def _sep(self):
        f = QFrame(); f.setFrameShape(QFrame.HLine); f.setStyleSheet(f"color:{BD};"); return f

    def _on_inc_toggle(self, _state):
        enabled = self.inc_cb.isChecked()
        self.inc_edit.setVisible(enabled)
        self.inc_hint.setVisible(enabled)

    # ----------------------------------------------------------------
    # 输出面板
    # ----------------------------------------------------------------
    def _build_output(self, parent):
        lay = QVBoxLayout(parent)
        lay.setContentsMargins(0, 0, 0, 0)
        self.notebook = QTabWidget()
        lay.addWidget(self.notebook)

        # Tab1: 计算结果
        t1 = QWidget(); t1l = QVBoxLayout(t1); t1l.setContentsMargins(5, 5, 5, 5)
        grp = QGroupBox("计算结果详情"); gl = QVBoxLayout(grp)
        self.result_text = QWebEngineView()
        gl.addWidget(self.result_text)
        t1l.addWidget(grp)
        self.notebook.addTab(t1, "计算结果")

        # Tab2: 断面图
        t2 = QWidget(); t2l = QVBoxLayout(t2); t2l.setContentsMargins(5, 5, 5, 5)
        self.section_fig = Figure(figsize=(8, 6), dpi=100)
        self.section_canvas = FigureCanvas(self.section_fig)
        self.section_toolbar = NavToolbar(self.section_canvas, t2)
        t2l.addWidget(self.section_toolbar)
        t2l.addWidget(self.section_canvas)
        self.notebook.addTab(t2, "断面图")

        self._show_initial_help()

    # ----------------------------------------------------------------
    # 断面类型切换
    # ----------------------------------------------------------------
    def _on_section_type_changed(self, stype):
        _u_widgets = (self.R_lbl, self.R_edit, self.alpha_lbl, self.alpha_edit,
                      self.theta_lbl, self.theta_edit)
        if stype == "矩形":
            self.m_lbl.hide(); self.m_edit.hide()
            self.m_edit.setText("0.0")
            for w in (self.beta_lbl, self.beta_edit, self.b_lbl, self.b_edit, self.bb_hint): w.show()
            for w in (self.D_lbl, self.D_edit, self.D_hint_lbl): w.hide()
            for w in _u_widgets: w.hide()
        elif stype == "梯形":
            self.m_lbl.show(); self.m_edit.show()
            self.m_edit.setText("1.0")
            for w in (self.beta_lbl, self.beta_edit, self.b_lbl, self.b_edit, self.bb_hint): w.show()
            for w in (self.D_lbl, self.D_edit, self.D_hint_lbl): w.hide()
            for w in _u_widgets: w.hide()
        elif stype == "圆形":
            self.m_lbl.hide(); self.m_edit.hide()
            for w in (self.beta_lbl, self.beta_edit, self.b_lbl, self.b_edit, self.bb_hint): w.hide()
            for w in (self.D_lbl, self.D_edit, self.D_hint_lbl): w.show()
            for w in _u_widgets: w.hide()
        elif stype == "U形":
            self.m_lbl.hide(); self.m_edit.hide()
            for w in (self.beta_lbl, self.beta_edit, self.b_lbl, self.b_edit, self.bb_hint): w.hide()
            for w in (self.D_lbl, self.D_edit, self.D_hint_lbl): w.hide()
            for w in _u_widgets: w.show()

    # ----------------------------------------------------------------
    # 初始帮助
    # ----------------------------------------------------------------
    def _show_initial_help(self):
        h = HelpPageBuilder("明渠水力计算", '请选择断面类型并输入参数后点击“计算”按钮')
        h.section("支持断面类型")
        h.numbered_list([
            ("矩形断面", "m = 0，附录E自动寻优底宽；可指定宽深比或底宽"),
            ("梯形断面", "用户设定边坡系数 m，附录E自动寻优底宽；可指定宽深比或底宽"),
            ("圆形明渠", "自动搜索最优直径；可指定直径 D"),
            ("U形明渠", "圆弧底+斜直线壁；输入R、外倾角α、圆心角θ；自动反算水深"),
        ])
        h.section("计算模式总览")
        h.table(
            ["断面类型 / 可选参数填写方式", "程序行为"],
            [
                ["矩形/梯形 — 全部留空", "附录E自动寻优最优底宽 B"],
                ["矩形/梯形 — 指定宽深比 β", "以 β=B/h 为约束，自动搜索最优 B"],
                ["矩形/梯形 — 指定底宽 B", "固定 B，反算水深并验算流速"],
                ["圆形 — 留空直径 D", "自动搜索满足约束的最小 D"],
                ["圆形 — 指定直径 D", "固定 D，反算水深并验算流速"],
                ["U形 — 输入R/α/θ", "固定几何，反算水深并验算流速"],
            ]
        )
        h.section("U形断面几何公式")
        h.formula("h₀ = R·(1 − cos(θ/2))", "弧区高度")
        h.formula("当 h ≤ h₀: A = R²·arccos((R−h)/R) − (R−h)·√(2Rh−h²)", "纯弧区面积")
        h.formula("当 h ≤ h₀: χ = 2R·arccos((R−h)/R)", "纯弧区湿周")
        h.formula("当 h > h₀: A = A_arc + (b_arc + m·h_s)·h_s", "直线段区面积")
        h.formula("当 h > h₀: χ = θ/180·π·R + 2·h_s·√(1+m²)", "直线段区湿周")
        h.hint("矩形/梯形：宽深比 β 与底宽 B 不可同时填写（二选一）")
        h.section("曼宁公式")
        h.text("本程序基于曼宁公式进行计算：")
        h.formula("Q = (1/n) × A × R^(2/3) × i^(1/2)", "流量公式")
        h.section("断面几何公式")
        h.formula("A = (B + m×h) × h", "过水面积")
        h.formula("χ = B + 2×h×√(1+m²)", "湿周")
        h.formula("R = A/χ", "水力半径")
        h.section("宽深比说明")
        h.bullet_list([
            "定义：β = B/h（底宽 / 设计水深）",
            "可选参数中可指定宽深比或底宽",
            "二选一输入，留空则自动寻优计算",
        ])
        h.section("约束条件")
        h.bullet_list(["流速范围：不淤流速 < V < 不冲流速"])
        h.section("加大流量比例规范表")
        h.table(
            ["设计流量 Q (m³/s)", "加大比例"],
            [
                ["Q < 1", "30%"],
                ["1 ≤ Q < 5", "25%"],
                ["5 ≤ Q < 20", "20%"],
                ["20 ≤ Q < 50", "15%"],
                ["50 ≤ Q < 100", "10%"],
                ["Q ≥ 100", "5%"],
            ]
        )
        self.result_text.setHtml(h.build())

    # ----------------------------------------------------------------
    # 辅助：读取输入值
    # ----------------------------------------------------------------
    def _fval(self, edit, default=0.0):
        t = edit.text().strip()
        if not t: return default
        try: return float(t)
        except ValueError: return default

    def _fval_opt(self, edit):
        t = edit.text().strip()
        if not t: return None
        try: return float(t)
        except ValueError: return None

    def _info_parent(self):
        """获取InfoBar的父窗口（向上查找主窗口）"""
        w = self.window()
        return w if w else self

    # ================================================================
    # 计算
    # ================================================================
    def _calculate(self):
        stype = self.section_combo.currentText()
        try:
            Q = self._fval(self.Q_edit)
            n = self._fval(self.n_edit)
            slope_inv = self._fval(self.slope_edit)
            v_min = self._fval(self.vmin_edit)
            v_max = self._fval(self.vmax_edit)

            if Q <= 0:
                self._show_error("参数错误", "请输入有效的设计流量 Q（必须大于0）。"); return
            if n <= 0:
                self._show_error("参数错误", "请输入有效的糙率 n（必须大于0）。"); return
            if slope_inv <= 0:
                self._show_error("参数错误", "请输入有效的水力坡降倒数（必须大于0）。"); return
            if v_min >= v_max:
                self._show_error("参数错误", "不淤流速必须小于不冲流速。"); return

            use_increase = self.inc_cb.isChecked()
            manual_increase = self._fval_opt(self.inc_edit) if use_increase else 0

            if stype == "圆形":
                manual_D = self._fval_opt(self.D_edit)
                self.input_params = {
                    'Q': Q, 'n': n, 'slope_inv': slope_inv,
                    'v_min': v_min, 'v_max': v_max,
                    'section_type': stype, 'manual_b': manual_D,
                    'manual_increase': manual_increase,
                    'use_increase': use_increase
                }
                result = circular_calculate(
                    Q=Q, n=n, slope_inv=slope_inv,
                    v_min=v_min, v_max=v_max,
                    manual_D=manual_D,
                    increase_percent=manual_increase
                )
            elif stype == "U形":
                R_val = self._fval(self.R_edit)
                alpha_val = self._fval(self.alpha_edit)
                theta_val = self._fval(self.theta_edit)
                if R_val <= 0:
                    self._show_error("参数错误", "请输入有效的圆弧半径 R（必须大于0）。"); return
                if theta_val <= 0 or theta_val > 360:
                    self._show_error("参数错误", "请输入有效的圆心角 θ（0°~360°）。"); return
                self.input_params = {
                    'Q': Q, 'n': n, 'slope_inv': slope_inv,
                    'v_min': v_min, 'v_max': v_max,
                    'section_type': stype,
                    'R': R_val, 'alpha_deg': alpha_val, 'theta_deg': theta_val,
                    'manual_increase': manual_increase,
                    'use_increase': use_increase
                }
                result = mingqu_u_calculate(
                    Q=Q, R=R_val, alpha_deg=alpha_val, theta_deg=theta_val,
                    n=n, slope_inv=slope_inv,
                    v_min=v_min, v_max=v_max,
                    manual_increase_percent=manual_increase
                )
            else:
                m = self._fval(self.m_edit) if stype == "梯形" else 0.0
                if stype == "梯形" and m < 0:
                    self._show_error("参数错误", "请输入有效的边坡系数 m（不能为负）。"); return
                manual_beta = self._fval_opt(self.beta_edit)
                manual_b = self._fval_opt(self.b_edit)
                self.input_params = {
                    'Q': Q, 'm': m, 'n': n, 'slope_inv': slope_inv,
                    'v_min': v_min, 'v_max': v_max,
                    'section_type': stype,
                    'manual_beta': manual_beta, 'manual_b': manual_b,
                    'manual_increase': manual_increase,
                    'use_increase': use_increase
                }
                result = mingqu_calculate(
                    Q=Q, m=m, n=n, slope_inv=slope_inv,
                    v_min=v_min, v_max=v_max,
                    manual_beta=manual_beta,
                    manual_b=manual_b,
                    manual_increase_percent=manual_increase
                )

            self.current_result = result

            # 更新加大比例提示
            if use_increase and result.get('success') and 'increase_percent' in result:
                ap = result['increase_percent']
                src = "指定" if self.inc_edit.text().strip() else "自动计算"
                if isinstance(ap, str):
                    self.inc_hint.setText(f"({src}: {ap})")
                else:
                    self.inc_hint.setText(f"({src}: {ap:.1f}%)")

            self._update_result_display(result)
            self._update_section_plot(result)

        except Exception as e:
            self._show_error("计算错误", f"计算过程出错: {str(e)}")

    def _show_error(self, title, msg):
        out = []
        out.append("=" * 70)
        out.append(f"  {title}")
        out.append("=" * 70)
        out.append("")
        out.append(msg)
        out.append("")
        out.append("-" * 70)
        out.append("请修正后重新计算。")
        out.append("=" * 70)
        self.result_text.setHtml(make_plain_html("\n".join(out)))

    # ================================================================
    # 结果显示分发
    # ================================================================
    def _update_result_display(self, result):
        if not result['success']:
            self._show_error("计算失败", result.get('error_message', '未知错误'))
            return
        stype = self.input_params.get('section_type', '梯形')
        detail = self.detail_cb.isChecked()
        if stype == '圆形':
            if detail: self._show_circular_detail(result)
            else: self._show_circular_brief(result)
        elif stype == 'U形':
            if detail: self._show_u_detail(result)
            else: self._show_u_brief(result)
        else:
            if detail: self._show_trapezoid_detail(result)
            else: self._show_trapezoid_brief(result)

    # ================================================================
    # 梯形/矩形 - 简要结果
    # ================================================================
    def _show_trapezoid_brief(self, result):
        p = self.input_params
        Q, m, n = p['Q'], p['m'], p['n']
        slope_inv = p['slope_inv']
        v_min, v_max = p['v_min'], p['v_max']
        stype = p.get('section_type', '梯形')

        b = result['b_design']; h = result['h_design']
        V = result['V_design']; A = result['A_design']
        R = result['R_design']; beta = result['Beta_design']
        inc_pct = result['increase_percent']
        Q_inc = result['Q_increased']
        h_inc = result['h_increased']; V_inc = result['V_increased']
        Fb = result['Fb']; H = result['h_prime']
        inc_source = "(指定)" if p.get('manual_increase') else "(自动计算)"

        o = []
        o.append("=" * 70)
        o.append(f"              明渠水力计算结果（{stype}断面）")
        o.append("=" * 70)
        o.append("")
        o.append("【输入参数】")
        o.append("")
        _n = 1
        o.append(f"  {_n}. 断面类型:")
        o.append(f"     {stype}")
        o.append("")
        _n += 1
        o.append(f"  {_n}. 设计流量:")
        o.append(f"     Q = {Q:.3f} m³/s")
        o.append("")
        if stype == "梯形":
            _n += 1
            o.append(f"  {_n}. 边坡系数:")
            o.append(f"     m = {m}")
            o.append("")
        _n += 1
        o.append(f"  {_n}. 糙率:")
        o.append(f"     n = {n}")
        o.append("")
        _n += 1
        o.append(f"  {_n}. 水力坡降:")
        o.append(f"     = 1/{int(slope_inv)}")
        o.append("")
        _n += 1
        o.append(f"  {_n}. 不淤流速:")
        o.append(f"     = {v_min} m/s")
        o.append("")
        _n += 1
        o.append(f"  {_n}. 不冲流速:")
        o.append(f"     = {v_max} m/s")
        o.append("")
        o.append("【设计方法】")
        o.append("")
        o.append(f"  1. 采用方法:")
        o.append(f"     {result['design_method']}")
        o.append("")
        o.append("【设计结果】")
        o.append(f"  底宽 B = {b:.3f} m")
        o.append(f"  水深 h = {h:.3f} m")
        o.append(f"  宽深比 β = {beta:.3f}")
        o.append(f"  过水面积 A = {A:.3f} m²")
        o.append(f"  水力半径 R = {R:.3f} m")
        o.append(f"  设计流速 V = {V:.3f} m/s")
        o.append("")

        # 附录E表格
        schemes = result.get('appendix_e_schemes', [])
        if schemes:
            pre1 = "\n".join(o)
            ae_html = self._build_ae_html(schemes, b, h, v_min, v_max)
            self._appendix_e_export_text = self._build_ae_text(schemes, b, h, v_min, v_max)

            o2 = []
            use_increase = p.get('use_increase', True)
            if use_increase:
                o2.append("【加大流量工况】")
                o2.append(f"  流量加大比例 = {inc_pct:.1f}% {inc_source}")
                o2.append(f"  加大流量 Q加大 = {Q_inc:.3f} m³/s")
                if h_inc > 0:
                    o2.append(f"  加大水深 h加大 = {h_inc:.3f} m")
                    o2.append(f"  加大流速 V加大 = {V_inc:.3f} m/s")
                    o2.append(f"  岕顶超高 Fb = {Fb:.3f} m")  # 岕顶超高
                    o2.append(f"  渠道高度 H = {H:.3f} m")
            else:
                Fb_d = round(0.25 * h + 0.2, 3)
                H_d = round(h + Fb_d, 3)
                o2.append("【渠道尺寸计算】")
                o2.append(f"  (不考虑加大流量，以设计水深计算渠道高度)")
                o2.append(f"  超高 Fb = 0.25 × h + 0.2 = 0.25 × {h:.3f} + 0.2 = {Fb_d:.3f} m")
                o2.append(f"  渠道高度 H = h + Fb = {h:.3f} + {Fb_d:.3f} = {H_d:.3f} m")
            o2.append("")
            o2.append("【验证结果】")
            vel_ok = v_min < V < v_max
            if use_increase:
                fb_req = 0.25 * h_inc + 0.2 if h_inc > 0 else 0
                fb_ok = Fb >= (fb_req - 0.001) if h_inc > 0 else False
                o2.append(f"  流速验证: {'✓ 通过' if vel_ok else '✗ 未通过'}")
                o2.append(f"  超高复核: {'✓ 通过' if fb_ok else '✗ 未通过'} (Fb={Fb:.3f}m, 规范要求≥{fb_req:.3f}m)")
                all_pass = vel_ok and fb_ok
            else:
                o2.append(f"  流速验证: {'✓ 通过' if vel_ok else '✗ 未通过'}")
                all_pass = vel_ok
            o2.append("")
            o2.append("=" * 70)
            o2.append(f"  综合验证结果: {'全部通过 ✓' if all_pass else '未通过 ✗'}")
            o2.append("=" * 70)
            pre2 = "\n".join(o2)

            body1 = plain_text_to_formula_body(pre1)
            ae_body = "\n<b>【附录E断面方案对比表】</b><br>"
            ae_body += "  说明: α=1.00为水力最佳断面(深窄)，α越大断面越宽浅，面积增加但流速降低<br><br>"
            ae_body += ae_html
            ae_body += f"<br>  注: 流速约束范围 {v_min} ~ {v_max} m/s<br><br>"
            full_body = body1 + ae_body + plain_text_to_formula_body("\n".join(o2))
            full_html = wrap_with_katex(full_body, extra_head=AE_CSS)
            self._export_plain_text = pre1 + "\n\n" + self._appendix_e_export_text + "\n\n" + pre2
            load_formula_page(self.result_text, full_html)
            return

        use_increase = p.get('use_increase', True)
        if use_increase:
            o.append("【加大流量工况】")
            o.append(f"  流量加大比例 = {inc_pct:.1f}% {inc_source}")
            o.append(f"  加大流量 Q加大 = {Q_inc:.3f} m³/s")
            if h_inc > 0:
                o.append(f"  加大水深 h加大 = {h_inc:.3f} m")
                o.append(f"  加大流速 V加大 = {V_inc:.3f} m/s")
                o.append(f"  岕顶超高 Fb = {Fb:.3f} m")
                o.append(f"  渠道高度 H = {H:.3f} m")
        else:
            Fb_d = round(0.25 * h + 0.2, 3)
            H_d = round(h + Fb_d, 3)
            o.append("【渠道尺寸计算】")
            o.append(f"  (不考虑加大流量，以设计水深计算渠道高度)")
            o.append(f"  超高 Fb = 0.25 × h + 0.2 = 0.25 × {h:.3f} + 0.2 = {Fb_d:.3f} m")
            o.append(f"  渠道高度 H = h + Fb = {h:.3f} + {Fb_d:.3f} = {H_d:.3f} m")
        o.append("")
        o.append("【验证结果】")
        vel_ok = v_min < V < v_max
        if use_increase:
            fb_req = 0.25 * h_inc + 0.2 if h_inc > 0 else 0
            fb_ok = Fb >= (fb_req - 0.001) if h_inc > 0 else False
            o.append(f"  流速验证: {'✓ 通过' if vel_ok else '✗ 未通过'}")
            o.append(f"  超高复核: {'✓ 通过' if fb_ok else '✗ 未通过'} (Fb={Fb:.3f}m, 规范要求≥{fb_req:.3f}m)")
            all_pass = vel_ok and fb_ok
        else:
            o.append(f"  流速验证: {'✓ 通过' if vel_ok else '✗ 未通过'}")
            all_pass = vel_ok
        o.append("=" * 70)
        o.append(f"  综合验证结果: {'全部通过 ✓' if all_pass else '未通过 ✗'}")
        o.append("=" * 70)
        txt = "\n".join(o)
        self._export_plain_text = txt
        load_formula_page(self.result_text, plain_text_to_formula_html(txt))

    # ================================================================
    # 梯形/矩形 - 详细结果
    # ================================================================
    def _show_trapezoid_detail(self, result):
        p = self.input_params
        Q, m, n = p['Q'], p['m'], p['n']
        slope_inv = p['slope_inv']; i = 1.0 / slope_inv
        v_min, v_max = p['v_min'], p['v_max']
        stype = p.get('section_type', '梯形')

        b = result['b_design']; h = result['h_design']
        V = result['V_design']; A = result['A_design']
        X = result['X_design']; R = result['R_design']
        beta = result['Beta_design']; Q_calc = result['Q_calc']

        inc_pct = result['increase_percent']
        Q_inc = result['Q_increased']; h_inc = result['h_increased']
        V_inc = result['V_increased']
        A_inc = result.get('A_increased', -1)
        X_inc = result.get('X_increased', -1)
        R_inc = result.get('R_increased', -1)
        Fb = result['Fb']; H = result['h_prime']
        inc_source = "(指定)" if p.get('manual_increase') else "(自动计算)"

        o = []
        o.append("=" * 70)
        o.append(f"              明渠水力计算结果（{stype}断面）")
        o.append("=" * 70)
        o.append("")
        o.append("【一、输入参数】")
        o.append("")
        _n = 1
        o.append(f"  {_n}. 断面类型:")
        o.append(f"     {stype}")
        o.append("")
        _n += 1
        o.append(f"  {_n}. 设计流量:")
        o.append(f"     Q = {Q:.3f} m³/s")
        o.append("")
        if stype == "梯形":
            _n += 1
            o.append(f"  {_n}. 边坡系数:")
            o.append(f"     m = {m}")
            o.append("")
        _n += 1
        o.append(f"  {_n}. 糙率:")
        o.append(f"     n = {n}")
        o.append("")
        _n += 1
        o.append(f"  {_n}. 水力坡降:")
        o.append(f"     = 1/{int(slope_inv)}")
        o.append("")
        _n += 1
        o.append(f"  {_n}. 不淤流速:")
        o.append(f"     = {v_min} m/s")
        o.append("")
        _n += 1
        o.append(f"  {_n}. 不冲流速:")
        o.append(f"     = {v_max} m/s")
        o.append("")
        if p.get('manual_beta'):
            _n += 1
            o.append(f"  {_n}. 指定宽深比:")
            o.append(f"     β = {p['manual_beta']}")
            o.append("")
        if p.get('manual_b'):
            _n += 1
            o.append(f"  {_n}. 指定底宽:")
            o.append(f"     B = {p['manual_b']} m")
            o.append("")
        if p.get('manual_increase'):
            _n += 1
            o.append(f"  {_n}. 指定加大比例:")
            o.append(f"     = {p['manual_increase']}%")
            o.append("")

        o.append("【二、设计方法】")
        o.append("")
        o.append(f"  1. 采用方法:")
        o.append(f"     {result['design_method']}")
        o.append("")

        schemes = result.get('appendix_e_schemes', [])
        has_ae = bool(schemes)
        if has_ae:
            pre1 = "\n".join(o)
            ae_html = self._build_ae_html(schemes, b, h, v_min, v_max)
            self._appendix_e_export_text = self._build_ae_text(schemes, b, h, v_min, v_max)
            o = []

        o.append("【三、设计结果】")
        o.append("")
        o.append("  1. 设计底宽:")
        o.append(f"     B = {b:.3f} m")
        o.append("")
        o.append("  2. 设计水深:")
        o.append(f"     h = {h:.3f} m")
        o.append("")
        o.append("  3. 宽深比:")
        o.append(f"     β = B/h = {b:.3f}/{h:.3f} = {beta:.3f}")
        o.append("")
        o.append("  4. 过水面积计算:")
        if stype == "梯形":
            o.append(f"     A = (B + m×h) × h")
            o.append(f"       = ({b:.3f} + {m}×{h:.3f}) × {h:.3f}")
            o.append(f"       = {b + m * h:.3f} × {h:.3f}")
        else:
            o.append(f"     A = B × h")
            o.append(f"       = {b:.3f} × {h:.3f}")
        o.append(f"       = {A:.3f} m²")
        o.append("")
        o.append("  5. 湿周计算:")
        sq = math.sqrt(1 + m * m)
        if stype == "梯形":
            o.append(f"     χ = B + 2×h×√(1+m²)")
            o.append(f"       = {b:.3f} + 2×{h:.3f}×√(1+{m}²)")
            o.append(f"       = {b:.3f} + 2×{h:.3f}×{sq:.4f}")
            o.append(f"       = {b:.3f} + {2 * h * sq:.3f}")
        else:
            o.append(f"     χ = B + 2×h")
            o.append(f"       = {b:.3f} + 2×{h:.3f}")
            o.append(f"       = {b:.3f} + {2 * h:.3f}")
        o.append(f"       = {X:.3f} m")
        o.append("")
        o.append("  6. 水力半径计算:")
        o.append(f"     R = A/χ = {A:.3f}/{X:.3f} = {R:.3f} m")
        o.append("")
        o.append("  7. 设计流速计算 (曼宁公式):")
        o.append(f"     V = (1/n) × R^(2/3) × i^(1/2)")
        o.append(f"       = (1/{n}) × {R:.3f}^(2/3) × {i:.6f}^(1/2)")
        o.append(f"       = {1/n:.2f} × {R**(2/3):.4f} × {math.sqrt(i):.6f}")
        o.append(f"       = {V:.3f} m/s")
        o.append("")
        o.append("  8. 流量校核:")
        o.append(f"      Q计算 = V × A = {V:.3f} × {A:.3f} = {V * A:.3f} m³/s")
        o.append(f"      误差 = {abs(V * A - Q)/Q*100:.2f}%")
        o.append("")

        use_increase = p.get('use_increase', True)
        if use_increase:
          o.append("【四、加大流量工况计算】")
          o.append("")
          o.append("  1. 加大流量计算:")
          o.append(f"      流量加大比例 = {inc_pct:.1f}% {inc_source}")
          o.append(f"      Q加大 = Q × (1 + {inc_pct/100:.2f})")
          o.append(f"           = {Q:.3f} × {1+inc_pct/100:.2f}")
          o.append(f"           = {Q_inc:.3f} m³/s")
          o.append("")

        if use_increase and h_inc > 0:
            if A_inc <= 0: A_inc = (b + m * h_inc) * h_inc
            if X_inc <= 0: X_inc = b + 2 * h_inc * math.sqrt(1 + m * m)
            if R_inc <= 0 and X_inc > 0: R_inc = A_inc / X_inc

            o.append("  2. 加大水深计算:")
            o.append(f"      根据加大流量 Q加大 = {Q_inc:.3f} m³/s 和设计底宽 B = {b:.3f} m，")
            o.append(f"      利用曼宁公式反算水深:")
            o.append(f"      h加大 = {h_inc:.3f} m")
            o.append("")
            o.append("  3. 加大过水面积计算:")
            if stype == "梯形":
                o.append(f"      A加大 = (B + m×h加大) × h加大")
                o.append(f"           = ({b:.3f} + {m}×{h_inc:.3f}) × {h_inc:.3f}")
                o.append(f"           = {b + m * h_inc:.3f} × {h_inc:.3f}")
            else:
                o.append(f"      A加大 = B × h加大")
                o.append(f"           = {b:.3f} × {h_inc:.3f}")
            o.append(f"           = {A_inc:.3f} m²")
            o.append("")
            o.append("  4. 加大湿周计算:")
            sq2 = math.sqrt(1 + m * m)
            if stype == "梯形":
                o.append(f"      χ加大 = B + 2×h加大×√(1+m²)")
                o.append(f"           = {b:.3f} + 2×{h_inc:.3f}×√(1+{m}²)")
                o.append(f"           = {b:.3f} + 2×{h_inc:.3f}×{sq2:.4f}")
                o.append(f"           = {b:.3f} + {2 * h_inc * sq2:.3f}")
            else:
                o.append(f"      χ加大 = B + 2×h加大")
                o.append(f"           = {b:.3f} + 2×{h_inc:.3f}")
                o.append(f"           = {b:.3f} + {2 * h_inc:.3f}")
            o.append(f"           = {X_inc:.3f} m")
            o.append("")
            o.append("  5. 加大水力半径计算:")
            o.append(f"      R加大 = A加大 / χ加大")
            o.append(f"           = {A_inc:.3f} / {X_inc:.3f}")
            o.append(f"           = {R_inc:.3f} m")
            o.append("")
            o.append("  6. 加大流速计算 (曼宁公式):")
            o.append(f"      V加大 = (1/n) × R加大^(2/3) × i^(1/2)")
            o.append(f"           = (1/{n}) × {R_inc:.3f}^(2/3) × {i:.6f}^(1/2)")
            o.append(f"           = {1/n:.2f} × {R_inc**(2/3):.4f} × {math.sqrt(i):.6f}")
            o.append(f"           = {V_inc:.3f} m/s")
            o.append("")
            Q_chk = V_inc * A_inc
            o.append("  7. 流量校核:")
            o.append(f"      Q校核 = V加大 × A加大 = {V_inc:.3f} × {A_inc:.3f} = {Q_chk:.3f} m³/s")
            o.append(f"      误差 = {abs(Q_chk - Q_inc) / Q_inc * 100:.2f}%")
            o.append("")
            o.append("  8. 渠道岸顶超高计算（规范 6.4.8-2）:")
            o.append(f"      Fb = (1/4) × h加大 + 0.2")
            o.append(f"         = (1/4) × {h_inc:.3f} + 0.2")
            o.append(f"         = {Fb:.3f} m")
            o.append("")
            o.append("  9. 渠道高度计算:")
            o.append(f"      H = h加大 + Fb")
            o.append(f"        = {h_inc:.3f} + {Fb:.3f}")
            o.append(f"        = {H:.3f} m")
        elif use_increase:
            o.append("  加大水深计算失败")
        o.append("")

        if not use_increase:
            Fb_d = round(0.25 * h + 0.2, 3)
            H_d = round(h + Fb_d, 3)
            o.append("【四、渠道尺寸计算】")
            o.append("")
            o.append(f"  (不考虑加大流量，以设计水深计算渠道高度)")
            o.append(f"  1. 超高计算（规范 6.4.8-2）:")
            o.append(f"      Fb = (1/4) × h + 0.2 = (1/4) × {h:.3f} + 0.2 = {Fb_d:.3f} m")
            o.append(f"  2. 渠道高度计算:")
            o.append(f"      H = h + Fb = {h:.3f} + {Fb_d:.3f} = {H_d:.3f} m")
            o.append("")

        o.append("【五、设计验证】")
        o.append("")
        vel_ok = v_min < V < v_max
        o.append(f"  1. 流速验证:")
        o.append(f"      范围要求: {v_min} < V < {v_max} m/s")
        o.append(f"      设计流速: V = {V:.3f} m/s")
        o.append(f"      结果: {'通过 ✓' if vel_ok else '未通过 ✗'}")
        o.append("")
        if use_increase:
            fb_req = 0.25 * h_inc + 0.2 if h_inc > 0 else 0
            fb_ok = Fb >= (fb_req - 0.001) if h_inc > 0 else False
            o.append(f"  2. 超高复核（规范 6.4.8-2）:")
            o.append(f"      规范要求: Fb ≥ (1/4)×h加大 + 0.2 = {fb_req:.3f} m")
            o.append(f"      计算结果: Fb = {Fb:.3f} m")
            o.append(f"      结果: {'通过 ✓' if fb_ok else '未通过 ✗'}")
            o.append("")
            all_pass = vel_ok and fb_ok
        else:
            all_pass = vel_ok
        o.append("=" * 70)
        o.append(f"  综合验证结果: {'全部通过 ✓' if all_pass else '未通过 ✗'}")
        o.append("=" * 70)

        if has_ae:
            pre2 = "\n".join(o)
            ae_body = "\n<b>【附录E断面方案对比表】</b><br>"
            ae_body += "  说明: α=1.00为水力最佳断面(深窄)，α越大断面越宽浅，面积增加但流速降低<br><br>"
            ae_body += ae_html
            ae_body += f"<br>  注: 流速约束范围 {v_min} ~ {v_max} m/s<br><br>"
            full_body = plain_text_to_formula_body(pre1) + ae_body + plain_text_to_formula_body(pre2)
            full_html = wrap_with_katex(full_body, extra_head=AE_CSS)
            self._export_plain_text = pre1 + "\n\n" + self._appendix_e_export_text + "\n\n" + pre2
            load_formula_page(self.result_text, full_html)
        else:
            txt = "\n".join(o)
            self._export_plain_text = txt
            load_formula_page(self.result_text, plain_text_to_formula_html(txt))

    # ================================================================
    # 圆形 - 简要结果
    # ================================================================
    def _show_circular_brief(self, result):
        p = self.input_params
        Q, n = p['Q'], p['n']; slope_inv = p['slope_inv']
        v_min, v_max = p['v_min'], p['v_max']
        D = result.get('D_design', 0)
        h = result.get('y_d', 0); V = result.get('V_d', 0)
        A_d = result.get('A_d', 0); FB_d = result.get('FB_d', 0)
        PA_d = result.get('PA_d', 0)
        inc_info = result.get('increase_percent', '')
        Q_inc = result.get('Q_inc', 0)
        h_i = result.get('y_i', 0); V_i = result.get('V_i', 0)
        FB_i = result.get('FB_i', 0); PA_i = result.get('PA_i', 0)

        o = []
        o.append("=" * 70)
        o.append("              明渠水力计算结果（圆形断面）")
        o.append("=" * 70)
        o.append("")
        o.append("【输入参数】")
        o.append("")
        o.append(f"  1. 断面类型:")
        o.append(f"     圆形")
        o.append("")
        o.append(f"  2. 设计流量:")
        o.append(f"     Q = {Q:.3f} m³/s")
        o.append("")
        o.append(f"  3. 糙率:")
        o.append(f"     n = {n}")
        o.append("")
        o.append(f"  4. 水力坡降:")
        o.append(f"     = 1/{int(slope_inv)}")
        o.append("")
        o.append(f"  5. 不淤流速:")
        o.append(f"     = {v_min} m/s")
        o.append("")
        o.append(f"  6. 不冲流速:")
        o.append(f"     = {v_max} m/s")
        o.append("")
        o.append("【断面尺寸】")
        o.append("")
        o.append(f"  1. 设计直径:")
        o.append(f"     D = {D:.2f} m")
        o.append("")
        o.append("【设计流量工况】")
        o.append(f"  设计水深 h = {h:.3f} m")
        o.append(f"  设计流速 V = {V:.3f} m/s")
        o.append(f"  过水面积 A = {A_d:.3f} m²")
        o.append(f"  净空高度 Fb = {FB_d:.3f} m")
        o.append(f"  净空比例 = {PA_d:.1f}%")
        o.append("")
        use_increase = p.get('use_increase', True)
        if use_increase:
            o.append("【加大流量工况】")
            o.append(f"  流量加大比例 = {inc_info}")
            o.append(f"  加大流量 Q加大 = {Q_inc:.3f} m³/s")
            o.append(f"  加大水深 h加大 = {h_i:.3f} m")
            o.append(f"  加大流速 V加大 = {V_i:.3f} m/s")
            o.append(f"  净空高度 Fb加大 = {FB_i:.3f} m")
            o.append(f"  净空比例 = {PA_i:.1f}%")
            o.append("")
        o.append("【验证结果】")
        vel_ok = V is not None and v_min <= V <= v_max
        o.append(f"  1. 设计流速验证")
        o.append(f"     范围要求: {v_min} ≤ V ≤ {v_max} m/s")
        o.append(f"     计算结果: V = {V:.3f} m/s")
        o.append(f"     验证结果: {'通过 ✓' if vel_ok else '未通过 ✗'}")
        o.append("")
        if use_increase:
            vel_i_ok = V_i is not None and v_min <= V_i <= v_max if V_i else True
            o.append(f"  2. 加大流速验证")
            o.append(f"     范围要求: {v_min} ≤ V ≤ {v_max} m/s")
            if V_i:
                o.append(f"     计算结果: V加大 = {V_i:.3f} m/s")
                o.append(f"     验证结果: {'通过 ✓' if vel_i_ok else '未通过 ✗'}")
            else:
                o.append(f"     计算结果: 无数据")
            o.append("")
            all_ok = vel_ok and vel_i_ok
        else:
            all_ok = vel_ok
        o.append("=" * 70)
        o.append(f"  综合验证结果: {'全部通过 ✓' if all_ok else '未通过 ✗'}")
        o.append("=" * 70)
        txt = "\n".join(o)
        self._export_plain_text = txt
        load_formula_page(self.result_text, plain_text_to_formula_html(txt))

    # ================================================================
    # 圆形 - 详细结果
    # ================================================================
    def _show_circular_detail(self, result):
        p = self.input_params
        Q, n = p['Q'], p['n']; slope_inv = p['slope_inv']
        v_min, v_max = p['v_min'], p['v_max']
        i = 1.0 / slope_inv

        D_calc = result.get('D_calculated', 0)
        D = result.get('D_design', 0)
        pipe_area = PI * D**2 / 4 if D > 0 else 0

        h_d = result.get('y_d', 0); V_d = result.get('V_d', 0)
        A_d = result.get('A_d', 0); P_d = result.get('P_d', 0)
        R_d = result.get('R_d', 0); PA_d = result.get('PA_d', 0)
        FB_d = result.get('FB_d', 0); Q_chk_d = result.get('Q_check_d', 0)

        inc_info = result.get('increase_percent', '')
        Q_inc = result.get('Q_inc', 0)
        h_i = result.get('y_i', 0); V_i = result.get('V_i', 0)
        A_i = result.get('A_i', 0); P_i = result.get('P_i', 0)
        R_i = result.get('R_i', 0); PA_i = result.get('PA_i', 0)
        FB_i = result.get('FB_i', 0)

        try: inc_pct = float(inc_info.split('%')[0])
        except: inc_pct = 20

        Q_min = result.get('Q_min', 0)
        h_m = result.get('y_m', 0); V_m = result.get('V_m', 0)
        A_m = result.get('A_m', 0); P_m = result.get('P_m', 0)
        R_m = result.get('R_m', 0)

        o = []
        o.append("=" * 70)
        o.append("              明渠水力计算结果（圆形断面）")
        o.append("=" * 70)
        o.append("")
        o.append("【一、输入参数】")
        o.append("")
        _n = 1
        o.append(f"  {_n}. 断面类型:")
        o.append(f"     圆形")
        o.append("")
        _n += 1
        o.append(f"  {_n}. 设计流量:")
        o.append(f"     Q = {Q:.3f} m³/s")
        o.append("")
        _n += 1
        o.append(f"  {_n}. 糙率:")
        o.append(f"     n = {n}")
        o.append("")
        _n += 1
        o.append(f"  {_n}. 水力坡降:")
        o.append(f"     = 1/{int(slope_inv)}")
        o.append("")
        _n += 1
        o.append(f"  {_n}. 不淤流速:")
        o.append(f"     = {v_min} m/s")
        o.append("")
        _n += 1
        o.append(f"  {_n}. 不冲流速:")
        o.append(f"     = {v_max} m/s")
        o.append("")
        if p.get('manual_b'):
            _n += 1
            o.append(f"  {_n}. 指定直径:")
            o.append(f"     D = {p['manual_b']} m")
            o.append("")

        o.append("【二、直径确定】")
        o.append("")
        if D_calc and D_calc > 0:
            o.append(f"  1. 计算直径: D计算 = {D_calc:.3f} m")
        o.append(f"  2. 设计直径: D = {D:.2f} m")
        o.append("")
        o.append("  3. 管道总断面积计算:")
        o.append(f"     A总 = π × D² / 4")
        o.append(f"        = {PI:.4f} × {D:.2f}² / 4")
        o.append(f"        = {PI:.4f} × {D**2:.4f} / 4")
        o.append(f"        = {pipe_area:.3f} m²")
        o.append("")

        o.append("【三、设计流量工况计算】")
        o.append("")
        o.append("  1. 设计水深计算:")
        o.append(f"     根据设计流量 Q = {Q:.3f} m³/s，利用曼宁公式反算水深:")
        o.append(f"     h = {h_d:.3f} m")
        o.append("")

        if h_d > 0 and D > 0 and h_d <= D:
            Rr = D / 2
            theta = 2 * math.acos(max(-1, min(1, (Rr - h_d) / Rr)))
            o.append(f"  2. 圆心角计算:")
            o.append(f"     θ = 2 × arccos((R - h) / R)")
            o.append(f"       = 2 × arccos(({Rr:.3f} - {h_d:.3f}) / {Rr:.3f})")
            o.append(f"       = 2 × arccos({(Rr - h_d)/Rr:.4f})")
            o.append(f"       = {math.degrees(theta):.2f}° ({theta:.4f} rad)")
            o.append("")
            o.append(f"  3. 过水面积计算:")
            o.append(f"     A = (D²/8) × (θ - sinθ)")
            o.append(f"       = ({D:.3f}²/8) × ({theta:.4f} - sin{theta:.4f})")
            o.append(f"       = {D**2/8:.4f} × {theta - math.sin(theta):.4f}")
            o.append(f"       = {A_d:.3f} m²")
            o.append("")
            o.append(f"  4. 湿周计算:")
            o.append(f"      χ = (D/2) × θ")
            o.append(f"        = ({D:.3f}/2) × {theta:.4f}")
            o.append(f"        = {Rr:.3f} × {theta:.4f}")
            o.append(f"        = {P_d:.3f} m")
            o.append("")
        else:
            o.append(f"  3. 过水面积: A = {A_d:.3f} m²")
            o.append("")
            o.append(f"  4. 湿周: χ = {P_d:.3f} m")
            o.append("")

        o.append(f"  5. 水力半径计算:")
        o.append(f"      R = A / χ")
        o.append(f"        = {A_d:.3f} / {P_d:.3f}")
        o.append(f"        = {R_d:.3f} m")
        o.append("")
        o.append(f"  6. 设计流速计算 (曼宁公式):")
        o.append(f"      V = (1/n) × R^(2/3) × i^(1/2)")
        o.append(f"        = (1/{n}) × {R_d:.3f}^(2/3) × {i:.6f}^(1/2)")
        if R_d > 0:
            o.append(f"        = {1/n:.2f} × {R_d**(2/3):.4f} × {math.sqrt(i):.6f}")
        o.append(f"        = {V_d:.3f} m/s")
        o.append("")
        o.append(f"  7. 流量校核:")
        o.append(f"      Q计算 = V × A")
        o.append(f"           = {V_d:.3f} × {A_d:.3f}")
        o.append(f"           = {V_d * A_d:.3f} m³/s")
        if V_d * A_d > 0:
            o.append(f"      误差 = {abs(V_d * A_d - Q)/Q*100:.2f}%")
        o.append("")
        o.append(f"  8. 净空高度:")
        o.append(f"      Fb = D - h = {D:.3f} - {h_d:.3f} = {FB_d:.3f} m")
        o.append("")
        o.append(f"  9. 净空面积:")
        o.append(f"      PA = (A总 - A) / A总 × 100%")
        o.append(f"         = ({pipe_area:.3f} - {A_d:.3f}) / {pipe_area:.3f} × 100%")
        o.append(f"         = {PA_d:.1f}%")
        o.append("")

        use_increase_circ = p.get('use_increase', True)
        if use_increase_circ:
            o.append("【四、加大流量工况计算】")
            o.append("")
            o.append("  1. 加大流量计算:")
            o.append(f"      流量加大比例 = {inc_info}")
            o.append(f"      Q加大 = Q × (1 + {inc_pct/100:.2f})")
            o.append(f"           = {Q:.3f} × {1+inc_pct/100:.2f}")
            o.append(f"           = {Q_inc:.3f} m³/s")
            o.append("")
            if h_i is not None and h_i > 0 and D > 0:
                o.append("  2. 加大水深计算:")
                o.append(f"      根据加大流量 Q加大 = {Q_inc:.3f} m³/s，利用曼宁公式反算水深:")
                o.append(f"      h加大 = {h_i:.3f} m")
                o.append("")
                Rr_i = D / 2
                theta_i = 2 * math.acos(max(-1, min(1, (Rr_i - h_i) / Rr_i)))
                o.append(f"  3. 圆心角计算:")
                o.append(f"      θ加大 = 2 × arccos((R - h加大) / R)")
                o.append(f"           = 2 × arccos(({Rr_i:.3f} - {h_i:.3f}) / {Rr_i:.3f})")
                o.append(f"           = 2 × arccos({(Rr_i - h_i)/Rr_i:.4f})")
                o.append(f"           = {math.degrees(theta_i):.2f}° ({theta_i:.4f} rad)")
                o.append("")
                o.append(f"  4. 过水面积计算:")
                o.append(f"      A加大 = (D²/8) × (θ加大 - sinθ加大)")
                o.append(f"           = ({D:.3f}²/8) × ({theta_i:.4f} - sin{theta_i:.4f})")
                o.append(f"           = {D**2/8:.4f} × {theta_i - math.sin(theta_i):.4f}")
                o.append(f"           = {A_i:.3f} m²")
                o.append("")
                o.append(f"  5. 湿周计算:")
                o.append(f"      χ加大 = (D/2) × θ加大")
                o.append(f"           = ({D:.3f}/2) × {theta_i:.4f}")
                o.append(f"           = {Rr_i:.3f} × {theta_i:.4f}")
                o.append(f"           = {P_i:.3f} m")
                o.append("")
                o.append(f"  6. 水力半径计算:")
                o.append(f"      R加大 = A加大 / χ加大")
                if A_i and P_i:
                    o.append(f"           = {A_i:.3f} / {P_i:.3f}")
                    o.append(f"           = {R_i:.3f} m")
                o.append("")
                o.append(f"  7. 加大流速计算 (曼宁公式):")
                o.append(f"      V加大 = (1/n) × R加大^(2/3) × i^(1/2)")
                if R_i and R_i > 0:
                    o.append(f"           = (1/{n}) × {R_i:.3f}^(2/3) × {i:.6f}^(1/2)")
                    o.append(f"           = {1/n:.2f} × {R_i**(2/3):.4f} × {math.sqrt(i):.6f}")
                o.append(f"           = {V_i:.3f} m/s")
                o.append("")
                o.append(f"  8. 流量校核:")
                if V_i and A_i:
                    o.append(f"      Q计算 = V加大 × A加大")
                    o.append(f"           = {V_i:.3f} × {A_i:.3f}")
                    o.append(f"           = {V_i * A_i:.3f} m³/s")
                    if Q_inc > 0:
                        o.append(f"      误差 = {abs(V_i * A_i - Q_inc) / Q_inc * 100:.2f}%")
                o.append("")
                o.append(f"  9. 净空高度计算:")
                o.append(f"      Fb加大 = D - h加大")
                if h_i:
                    o.append(f"           = {D:.3f} - {h_i:.3f}")
                    o.append(f"           = {FB_i:.3f} m")
                o.append("")
                o.append(f"  10. 净空面积计算:")
                if A_i:
                    o.append(f"      PA加大 = (A总 - A加大) / A总 × 100%")
                    o.append(f"           = ({pipe_area:.3f} - {A_i:.3f}) / {pipe_area:.3f} × 100%")
                    o.append(f"           = {PA_i:.1f}%")
                o.append("")
            else:
                o.append(f"  2. 加大水深: h加大 = N/A")
                o.append("")

        o.append("【五、最小流量工况计算】")
        o.append("")
        o.append("  1. 最小流量计算:")
        o.append(f"      Q最小 = Q × 最小流量系数")
        o.append(f"           = {Q:.3f} × 0.4")
        if Q_min is not None and Q_min > 0:
            o.append(f"           = {Q_min:.3f} m³/s")
        else:
            o.append(f"           = N/A")
        o.append("")

        if Q_min is not None and Q_min > 0 and h_m is not None and h_m > 0 and D > 0:
            o.append("  2. 最小水深计算:")
            o.append(f"      根据最小流量 Q最小 = {Q_min:.3f} m³/s，利用曼宁公式反算水深:")
            o.append(f"      h最小 = {h_m:.3f} m")
            o.append("")

            Rr_m = D / 2
            if h_m <= D:
                theta_m = 2 * math.acos(max(-1, min(1, (Rr_m - h_m) / Rr_m)))
                o.append(f"  3. 圆心角计算:")
                o.append(f"      θ最小 = 2 × arccos((R - h最小) / R)")
                o.append(f"           = 2 × arccos(({Rr_m:.3f} - {h_m:.3f}) / {Rr_m:.3f})")
                o.append(f"           = 2 × arccos({(Rr_m - h_m)/Rr_m:.4f})")
                o.append(f"           = {math.degrees(theta_m):.2f}° ({theta_m:.4f} rad)")
                o.append("")
                o.append(f"  4. 过水面积计算:")
                o.append(f"      A最小 = (D²/8) × (θ最小 - sinθ最小)")
                o.append(f"           = ({D:.3f}²/8) × ({theta_m:.4f} - sin{theta_m:.4f})")
                o.append(f"           = {D**2/8:.4f} × {theta_m - math.sin(theta_m):.4f}")
                if A_m:
                    o.append(f"           = {A_m:.3f} m²")
                o.append("")
                o.append(f"  5. 湿周计算:")
                o.append(f"      χ最小 = (D/2) × θ最小")
                o.append(f"           = ({D:.3f}/2) × {theta_m:.4f}")
                o.append(f"           = {Rr_m:.3f} × {theta_m:.4f}")
                if P_m:
                    o.append(f"           = {P_m:.3f} m")
                o.append("")
            else:
                o.append(f"  3. 过水面积: A最小 = {A_m:.3f} m²" if A_m else "  3. 过水面积: A最小 = N/A")
                o.append("")
                o.append(f"  4. 湿周: χ最小 = {P_m:.3f} m" if P_m else "  4. 湿周: χ最小 = N/A")
                o.append("")

            o.append(f"  6. 水力半径计算:")
            o.append(f"      R最小 = A最小 / χ最小")
            if A_m and P_m:
                o.append(f"           = {A_m:.3f} / {P_m:.3f}")
                o.append(f"           = {R_m:.3f} m")
            o.append("")

            o.append(f"  7. 最小流速计算 (曼宁公式):")
            o.append(f"      V最小 = (1/n) × R最小^(2/3) × i^(1/2)")
            if R_m and R_m > 0:
                o.append(f"           = (1/{n}) × {R_m:.3f}^(2/3) × {i:.6f}^(1/2)")
                o.append(f"           = {1/n:.2f} × {R_m**(2/3):.4f} × {math.sqrt(i):.6f}")
            if V_m is not None:
                o.append(f"           = {V_m:.3f} m/s")
            o.append("")

            o.append(f"  8. 流量校核:")
            if V_m and A_m:
                o.append(f"      Q计算 = V最小 × A最小")
                o.append(f"           = {V_m:.3f} × {A_m:.3f}")
                o.append(f"           = {V_m * A_m:.3f} m³/s")
                if Q_min > 0:
                    o.append(f"      误差 = {abs(V_m * A_m - Q_min) / Q_min * 100:.2f}%")
            o.append("")
        else:
            o.append("  2. 最小水深: h最小 = N/A")
            o.append("  3. 最小流速: V最小 = N/A")
            o.append("")

        o.append("【六、设计验证】")
        o.append("")
        vel_ok = V_d is not None and v_min <= V_d <= v_max
        mv_ok = V_m is not None and V_m >= v_min
        o.append(f"  1. 流速验证:")
        o.append(f"      范围要求: {v_min} ≤ V ≤ {v_max} m/s")
        if V_d is not None:
            o.append(f"      设计流速: V = {V_d:.3f} m/s")
            o.append(f"      结果: {'通过 ✓' if vel_ok else '未通过 ✗'}")
        else:
            o.append(f"      计算失败")
        o.append("")
        if use_increase_circ:
            fb_ok = FB_i is not None and FB_i >= MIN_FREEBOARD
            pa_ok = PA_i is not None and PA_i >= MIN_FREE_AREA_PERCENT
            o.append(f"  2. 净空高度验证:")
            o.append(f"      规范要求: Fb ≥ {MIN_FREEBOARD} m")
            if FB_i is not None:
                o.append(f"      计算结果: Fb = {FB_i:.3f} m")
                o.append(f"      结果: {'通过 ✓' if fb_ok else '未通过 ✗'}")
            else:
                o.append(f"      计算失败")
            o.append("")
            o.append(f"  3. 净空面积验证:")
            o.append(f"      规范要求: PA ≥ {MIN_FREE_AREA_PERCENT}%")
            if PA_i is not None:
                o.append(f"      计算结果: PA = {PA_i:.1f}%")
                o.append(f"      结果: {'通过 ✓' if pa_ok else '未通过 ✗'}")
            else:
                o.append(f"      计算失败")
            o.append("")
            next_idx = 4
        else:
            fb_ok = pa_ok = True
            next_idx = 2
        o.append(f"  {next_idx}. 最小流速验证:")
        o.append(f"      规范要求: V ≥ {v_min} m/s")
        if V_m is not None:
            o.append(f"      计算结果: V = {V_m:.3f} m/s")
            o.append(f"      结果: {'通过 ✓' if mv_ok else '未通过 ✗'}")
        else:
            o.append(f"      计算失败")
        o.append("")
        all_pass = vel_ok and fb_ok and pa_ok and mv_ok
        o.append("=" * 70)
        o.append(f"  综合验证结果: {'全部通过 ✓' if all_pass else '未通过 ✗'}")
        o.append("=" * 70)
        txt = "\n".join(o)
        self._export_plain_text = txt
        load_formula_page(self.result_text, plain_text_to_formula_html(txt))

    # ================================================================
    # U形 - 简要结果
    # ================================================================
    def _show_u_brief(self, result):
        p = self.input_params
        Q, n, slope_inv = p['Q'], p['n'], p['slope_inv']
        v_min, v_max = p['v_min'], p['v_max']
        R = result['R']; alpha_deg = result['alpha_deg']; theta_deg = result['theta_deg']
        m = result['m']; h0 = result['h0']; b_arc = result['b_arc']
        h = result['h_design']; V = result['V_design']; A = result['A_design']
        X = result['X_design']; Rh = result['R_design']
        inc_pct = result['increase_percent']; Q_inc = result['Q_increased']
        h_inc = result['h_increased']; V_inc = result['V_increased']
        Fb = result['Fb']; H = result['h_prime']
        inc_src = "(指定)" if p.get('manual_increase') else "(自动计算)"
        o = []
        o.append("=" * 70)
        o.append("              明渠水力计算结果（U形断面）")
        o.append("=" * 70)
        o.append("")
        o.append("【输入参数】")
        o.append(f"  断面类型: U形    R = {R:.3f} m, α = {alpha_deg}°, θ = {theta_deg}°")
        o.append(f"  Q = {Q:.3f} m³/s,  n = {n},  i = 1/{int(slope_inv)}")
        o.append("")
        o.append("【断面几何参数】")
        o.append(f"  m = tan(α) = {m:.4f},  h₀ = {h0:.3f} m,  b_arc = {b_arc:.3f} m")
        o.append("")
        o.append("【设计流量工况】")
        o.append(f"  设计水深 h = {h:.3f} m")
        o.append(f"  过水面积 A = {A:.3f} m²")
        o.append(f"  湿周 χ = {X:.3f} m")
        o.append(f"  水力半径 R_h = {Rh:.3f} m")
        o.append(f"  设计流速 V = {V:.3f} m/s")
        o.append("")
        use_inc = p.get('use_increase', True)
        if use_inc:
            o.append("【加大流量工况】")
            o.append(f"  加大比例 = {inc_pct:.1f}% {inc_src},  Q加大 = {Q_inc:.3f} m³/s")
            if h_inc > 0:
                o.append(f"  h加大 = {h_inc:.3f} m,  V加大 = {V_inc:.3f} m/s")
                o.append(f"  超高 Fb = {Fb:.3f} m,  渠道高度 H = {H:.3f} m")
        o.append("")
        o.append("【验证结果】")
        vel_ok = v_min < V < v_max
        o.append(f"  流速: {v_min} < V={V:.3f} < {v_max} → {'✓ 通过' if vel_ok else '✗ 未通过'}")
        if use_inc and h_inc > 0:
            fb_req = 0.25 * h_inc + 0.2
            fb_ok = Fb >= (fb_req - 0.001)
            o.append(f"  超高: Fb={Fb:.3f}m ≥ {fb_req:.3f}m → {'✓ 通过' if fb_ok else '✗ 未通过'}")
            all_pass = vel_ok and fb_ok
        else:
            all_pass = vel_ok
        o.append("=" * 70)
        o.append(f"  综合验证: {'全部通过 ✓' if all_pass else '未通过 ✗'}")
        o.append("=" * 70)
        txt = "\n".join(o)
        self._export_plain_text = txt
        load_formula_page(self.result_text, plain_text_to_formula_html(txt))

    # ================================================================
    # U形 - 详细结果
    # ================================================================
    def _show_u_detail(self, result):
        p = self.input_params
        Q, n, slope_inv = p['Q'], p['n'], p['slope_inv']
        i = 1.0 / slope_inv
        v_min, v_max = p['v_min'], p['v_max']
        R = result['R']; alpha_deg = result['alpha_deg']; theta_deg = result['theta_deg']
        m = result['m']; h0 = result['h0']; b_arc = result['b_arc']
        h = result['h_design']; V = result['V_design']; A = result['A_design']
        X = result['X_design']; Rh = result['R_design']; Q_calc = result['Q_calc']
        inc_pct = result['increase_percent']; Q_inc = result['Q_increased']
        h_inc = result['h_increased']; V_inc = result['V_increased']
        A_inc = result.get('A_increased', -1); X_inc = result.get('X_increased', -1)
        R_inc = result.get('R_increased', -1)
        Fb = result['Fb']; H = result['h_prime']
        inc_src = "(指定)" if p.get('manual_increase') else "(自动计算)"
        theta_rad = math.radians(theta_deg)
        o = []
        o.append("=" * 70)
        o.append("              明渠水力计算结果（U形断面）")
        o.append("=" * 70)
        o.append("")
        o.append("【一、输入参数】")
        o.append(f"  断面类型: U形明渠")
        o.append(f"  设计流量 Q = {Q:.3f} m³/s")
        o.append(f"  圆弧半径 R = {R:.3f} m")
        o.append(f"  直线段外倾角 α = {alpha_deg}°")
        o.append(f"  圆弧段圆心角 θ = {theta_deg}°")
        o.append(f"  糙率 n = {n}")
        o.append(f"  水力坡降 i = 1/{int(slope_inv)} = {i:.6f}")
        o.append(f"  不淤流速 = {v_min} m/s,  不冲流速 = {v_max} m/s")
        o.append("")
        o.append("【二、断面几何参数】")
        o.append(f"  m = tan(α) = tan({alpha_deg}°) = {m:.6f}")
        o.append(f"  h₀ = R·(1-cos(θ/2)) = {R:.3f}×(1-cos({theta_deg/2:.1f}°)) = {h0:.3f} m")
        o.append(f"  b_arc = 2·R·sin(θ/2) = 2×{R:.3f}×sin({theta_deg/2:.1f}°) = {b_arc:.3f} m")
        o.append("")
        o.append("【三、设计水深计算】")
        o.append(f"  根据Q={Q:.3f} m³/s，曼宁公式二分法反算水深: h = {h:.3f} m")
        o.append(f"  水深区间: h {'≤' if h <= h0 else '>'} h₀={h0:.3f} m → {'纯弧区' if h <= h0 else '直线段区'}")
        o.append("")
        if h <= h0:
            cos_arg = max(-1.0, min(1.0, (R - h) / R))
            acos_val = math.acos(cos_arg)
            sqrt_val = math.sqrt(max(0.0, R * R - (R - h) ** 2))
            o.append("  【纯弧区公式】")
            o.append(f"  过水面积 A = R²·arccos((R-h)/R) - (R-h)·√(R²-(R-h)²)")
            o.append(f"           = {R:.3f}²×{acos_val:.4f} - {R-h:.3f}×{sqrt_val:.4f}")
            o.append(f"           = {A:.3f} m²")
            o.append(f"  湿周 χ = 2·R·arccos((R-h)/R) = 2×{R:.3f}×{acos_val:.4f} = {X:.3f} m")
        else:
            h_s = h - h0
            A_arc = R * R * (theta_rad / 2.0 - math.sin(theta_rad / 2.0) * math.cos(theta_rad / 2.0))
            o.append("  【直线段区公式】")
            o.append(f"  弧面积 A_arc = R²·(θ/2-sin(θ/2)·cos(θ/2)) = {A_arc:.4f} m²")
            o.append(f"  h_s = h - h₀ = {h:.3f} - {h0:.3f} = {h_s:.3f} m")
            o.append(f"  过水面积 A = A_arc + (b_arc + m·h_s)·h_s")
            o.append(f"           = {A_arc:.4f} + ({b_arc:.3f}+{m:.4f}×{h_s:.3f})×{h_s:.3f}")
            o.append(f"           = {A:.3f} m²")
            chi_arc = theta_rad * R
            o.append(f"  湿周 χ = θ·R + 2·h_s·√(1+m²)")
            o.append(f"       = {theta_rad:.4f}×{R:.3f} + 2×{h_s:.3f}×√(1+{m:.4f}²)")
            o.append(f"       = {X:.3f} m")
        o.append(f"  水力半径 R_h = A/χ = {A:.3f}/{X:.3f} = {Rh:.3f} m")
        o.append(f"  设计流速 V = (1/n)·R_h^(2/3)·i^(1/2)")
        o.append(f"           = (1/{n})×{Rh:.3f}^(2/3)×{i:.6f}^(1/2) = {V:.3f} m/s")
        o.append(f"  流量校核 Q计算 = {V:.3f}×{A:.3f} = {Q_calc:.3f} m³/s (误差{abs(Q_calc-Q)/Q*100:.2f}%)")
        o.append("")
        use_inc = p.get('use_increase', True)
        if use_inc:
            o.append("【四、加大流量工况】")
            o.append(f"  加大比例 = {inc_pct:.1f}% {inc_src}")
            o.append(f"  Q加大 = {Q:.3f}×(1+{inc_pct/100:.2f}) = {Q_inc:.3f} m³/s")
            if h_inc > 0:
                o.append(f"  h加大 = {h_inc:.3f} m")
                if A_inc > 0 and X_inc > 0:
                    o.append(f"  A加大 = {A_inc:.3f} m²,  χ加大 = {X_inc:.3f} m,  R加大 = {R_inc:.3f} m")
                o.append(f"  V加大 = {V_inc:.3f} m/s")
                o.append(f"  超高 Fb = 0.25×{h_inc:.3f}+0.2 = {Fb:.3f} m")
                o.append(f"  渠道高度 H = {h_inc:.3f}+{Fb:.3f} = {H:.3f} m")
            else:
                o.append("  加大水深计算失败")
        o.append("")
        o.append("【五、设计验证】")
        vel_ok = v_min < V < v_max
        o.append(f"  流速: {v_min} < V={V:.3f} < {v_max} → {'通过 ✓' if vel_ok else '未通过 ✗'}")
        if use_inc and h_inc > 0:
            fb_req = 0.25 * h_inc + 0.2
            fb_ok = Fb >= (fb_req - 0.001)
            o.append(f"  超高: Fb={Fb:.3f}m ≥ {fb_req:.3f}m → {'通过 ✓' if fb_ok else '未通过 ✗'}")
            all_pass = vel_ok and fb_ok
        else:
            all_pass = vel_ok
        o.append("=" * 70)
        o.append(f"  综合验证: {'全部通过 ✓' if all_pass else '未通过 ✗'}")
        o.append("=" * 70)
        txt = "\n".join(o)
        self._export_plain_text = txt
        load_formula_page(self.result_text, plain_text_to_formula_html(txt))

    # ================================================================
    # 附录E HTML表格
    # ================================================================
    def _build_ae_html(self, schemes, sel_b, sel_h, v_min, v_max):
        h = '<table class="ae"><tr>'
        for hdr in ['α值', '方案类型', '底宽B(m)', '水深h(m)', '宽深比β', '流速V(m/s)', '面积增加', '状态']:
            h += f'<th>{hdr}</th>'
        h += '</tr>'
        for idx, s in enumerate(schemes):
            alpha, stype = s['alpha'], s['scheme_type']
            sb, sh, sbeta, sV = s['b'], s['h'], s['beta'], s['V']
            area_inc = s['area_increase']
            is_sel = abs(sb - sel_b) < 0.01 and abs(sh - sel_h) < 0.01
            v_ok = v_min < sV < v_max
            if is_sel: cls = "sel"; status = "★ 选中"
            elif not v_ok: cls = "err"; status = "流速不符"
            else: cls = "even" if idx % 2 == 0 else "odd"; status = ""
            h += f'<tr class="{cls}">'
            h += f'<td>{alpha:.2f}</td><td>{stype}</td><td>{sb:.3f}</td><td>{sh:.3f}</td>'
            h += f'<td>{sbeta:.3f}</td><td>{sV:.3f}</td><td>+{area_inc:.0f}%</td>'
            h += f'<td><b>{status}</b></td></tr>'
        h += '</table>'
        return h

    def _build_ae_text(self, schemes, sel_b, sel_h, v_min, v_max):
        lines = []
        lines.append("【附录E断面方案对比表】")
        lines.append("  说明: α=1.00为水力最佳断面(深窄)，α越大断面越宽浅，面积增加但流速降低")
        lines.append("")
        lines.append("  α值    方案类型        底宽B(m)  水深h(m)  宽深比β   流速V(m/s)  面积+  状态")
        lines.append("  " + "-" * 78)
        for s in schemes:
            alpha, stype = s['alpha'], s['scheme_type']
            sb, sh, sbeta, sV = s['b'], s['h'], s['beta'], s['V']
            area_inc = s['area_increase']
            is_sel = abs(sb - sel_b) < 0.01 and abs(sh - sel_h) < 0.01
            v_ok = v_min < sV < v_max
            status = "★选中" if is_sel else ("流速不符" if not v_ok else "")
            lines.append(f"  {alpha:.2f}   {stype:<12}  {sb:8.3f}  {sh:8.3f}  {sbeta:8.3f}  {sV:10.3f}  +{area_inc:.0f}%   {status}")
        lines.append("")
        lines.append(f"  注: 流速约束范围 {v_min} ~ {v_max} m/s")
        return "\n".join(lines)

    # ================================================================
    # 断面图
    # ================================================================
    def _update_section_plot(self, result):
        self.section_fig.clear()
        if not result['success']:
            self.section_canvas.draw(); return
        stype = self.input_params.get('section_type', '梯形')
        if stype == '圆形':
            D = result.get('D_design', 0)
            y_d = result.get('y_d', 0); V_d = result.get('V_d', 0)
            Q = self.input_params['Q']
            ax = self.section_fig.add_subplot(111)
            self._draw_circular(ax, D, y_d, V_d, Q, '设计流量')
        elif stype == 'U形':
            R = result['R']; alpha_deg = result['alpha_deg']; theta_deg = result['theta_deg']
            h_w = result['h_design']
            H_ch = result['h_prime'] if result['h_prime'] > 0 else h_w * 1.35
            V = result['V_design']; Q = self.input_params['Q']
            h_inc = result['h_increased']; Q_inc = result['Q_increased']; V_inc = result['V_increased']
            use_inc = self.input_params.get('use_increase', True)
            if use_inc and h_inc > 0:
                axes = self.section_fig.subplots(1, 2)
                self._draw_u_section(axes[0], R, alpha_deg, theta_deg, h_w, H_ch, V, Q, '设计流量')
                H_ch2 = result['h_prime'] if result['h_prime'] > 0 else h_inc * 1.35
                self._draw_u_section(axes[1], R, alpha_deg, theta_deg, h_inc, H_ch2, V_inc, Q_inc, '加大流量')
            else:
                ax = self.section_fig.add_subplot(111)
                self._draw_u_section(ax, R, alpha_deg, theta_deg, h_w, H_ch, V, Q, '设计流量')
        else:
            b = result['b_design']; h = result['h_design']
            m = self.input_params.get('m', 0); Q = self.input_params['Q']
            V = result['V_design']; h_inc = result['h_increased']
            Q_inc = result['Q_increased']; V_inc = result['V_increased']
            h_prime = result['h_prime']
            use_inc = self.input_params.get('use_increase', True)
            if use_inc:
                axes = self.section_fig.subplots(1, 2)
                self._draw_trapezoid(axes[0], b, h, m, V, Q, h, "设计流量")
                if h_inc > 0:
                    self._draw_trapezoid(axes[1], b, h_prime, m, V_inc, Q_inc, h_inc, "加大流量")
                else:
                    axes[1].set_title("加大流量\n数据不可用")
            else:
                Fb_d = 0.25 * h + 0.2
                H_d = h + Fb_d
                ax = self.section_fig.add_subplot(111)
                self._draw_trapezoid(ax, b, H_d, m, V, Q, h, "设计流量")
        self.section_fig.tight_layout()
        self.section_canvas.draw()

    def _draw_trapezoid(self, ax, b, h_ch, m, V, Q, h_w, title):
        tw = b + 2 * m * h_ch
        ax.plot([-b/2, b/2], [0, 0], 'k-', lw=2)
        ax.plot([-b/2, -tw/2], [0, h_ch], 'k-', lw=2)
        ax.plot([b/2, tw/2], [0, h_ch], 'k-', lw=2)
        ax.plot([-tw/2, tw/2], [h_ch, h_ch], 'k--', lw=1)
        if h_w > 0:
            ww = b + 2 * m * h_w
            wx = [-b/2, -ww/2, ww/2, b/2]
            wy = [0, h_w, h_w, 0]
            ax.fill(wx, wy, color='lightblue', alpha=0.7)
            ax.plot([-ww/2, ww/2], [h_w, h_w], 'b-', lw=1.5)
        ax.annotate('', xy=(b/2, -0.1*h_ch), xytext=(-b/2, -0.1*h_ch),
                     arrowprops=dict(arrowstyle='<->', color='gray', lw=1.5))
        ax.text(0, -0.2*h_ch, f'B={b:.2f}m', ha='center', fontsize=9, color='gray')
        ax.annotate('', xy=(tw/2+0.08*tw, h_ch), xytext=(tw/2+0.08*tw, 0),
                     arrowprops=dict(arrowstyle='<->', color='purple', lw=1.5))
        ax.text(tw/2+0.12*tw, h_ch/2, f'H={h_ch:.2f}m', fontsize=9, color='purple', rotation=90, va='center')
        if h_w > 0:
            ax.annotate('', xy=(-tw/2-0.08*tw, h_w), xytext=(-tw/2-0.08*tw, 0),
                         arrowprops=dict(arrowstyle='<->', color='blue', lw=1.5))
            ax.text(-tw/2-0.12*tw, h_w/2, f'h={h_w:.2f}m', fontsize=9, color='blue', rotation=90, va='center', ha='right')
        ax.set_xlim(-tw*0.85, tw*0.85)
        ax.set_ylim(-h_ch*0.4, h_ch*1.2)
        ax.set_aspect('equal')
        ax.set_title(f'{title}\nQ={Q:.2f}m$^3$/s, V={V:.2f}m/s', fontsize=10)
        ax.grid(True, alpha=0.3)
        ax.axhline(y=0, color='brown', lw=3)

    def _draw_u_section(self, ax, R, alpha_deg, theta_deg, h_w, H_ch, V, Q, title):
        """U形断面维图：圆弧底 + 斜壁 + 水面填充"""
        theta_rad = math.radians(theta_deg)
        m = math.tan(math.radians(alpha_deg))
        h0 = R * (1.0 - math.cos(theta_rad / 2.0))
        b_arc = 2.0 * R * math.sin(theta_rad / 2.0)

        # 弧底轮廓
        half_theta = theta_rad / 2.0
        arc_angles = np.linspace(math.pi * 3 / 2 - half_theta, math.pi * 3 / 2 + half_theta, 60)
        arc_x = R * np.cos(arc_angles)
        arc_y = R + R * np.sin(arc_angles)  # 圆心在 (0, R)

        # 直线段上端
        x_top_r = b_arc / 2.0 + m * H_ch
        x_top_l = -x_top_r
        x_arc_r = b_arc / 2.0
        x_arc_l = -x_arc_r

        outline_x = list(arc_x) + [x_arc_r, x_top_r, x_top_l, x_arc_l] + list(arc_x[:1])
        outline_y = list(arc_y) + [h0, H_ch, H_ch, h0] + list(arc_y[:1])
        ax.plot(outline_x[:len(arc_x)], outline_y[:len(arc_y)], 'k-', lw=2)
        ax.plot([x_arc_r, x_top_r], [h0, H_ch], 'k-', lw=2)
        ax.plot([x_arc_l, x_top_l], [h0, H_ch], 'k-', lw=2)
        ax.plot([x_top_l, x_top_r], [H_ch, H_ch], 'k--', lw=1)

        # 水面填充
        if h_w > 0:
            from matplotlib.patches import Polygon
            from matplotlib.collections import PatchCollection
            if h_w <= h0:
                # 纯弧区
                ang_h = math.acos(max(-1.0, min(1.0, (R - h_w) / R)))
                water_angles = np.linspace(math.pi * 3 / 2 - ang_h, math.pi * 3 / 2 + ang_h, 40)
                wx = list(R * np.cos(water_angles))
                wy = list(R + R * np.sin(water_angles))
                water_pts = list(zip(wx, wy))
                half_bw = math.sqrt(max(0.0, R * R - (R - h_w) ** 2))
                water_pts = [(-half_bw, h_w)] + water_pts + [(half_bw, h_w)]
            else:
                h_s = h_w - h0
                bw = b_arc + 2 * m * h_s
                water_pts = (
                    list(zip(arc_x, arc_y)) +
                    [(x_arc_r, h0), (bw / 2, h_w), (-bw / 2, h_w), (x_arc_l, h0)]
                )
            poly = Polygon(water_pts, closed=True)
            pc = PatchCollection([poly], facecolor='lightblue', alpha=0.7, edgecolor='none')
            ax.add_collection(pc)
            # 水面线
            if h_w <= h0:
                half_bw = math.sqrt(max(0.0, R * R - (R - h_w) ** 2))
                ax.plot([-half_bw, half_bw], [h_w, h_w], 'b-', lw=1.5)
            else:
                h_s = h_w - h0
                bw = b_arc + 2 * m * h_s
                ax.plot([-bw / 2, bw / 2], [h_w, h_w], 'b-', lw=1.5)

        max_x = max(abs(x_top_r), R) * 1.3
        ax.set_xlim(-max_x, max_x)
        ax.set_ylim(-H_ch * 0.3, H_ch * 1.3)
        ax.set_aspect('equal')
        ax.set_title(f'{title}\nQ={Q:.2f}m$^3$/s, V={V:.2f}m/s', fontsize=10)
        ax.grid(True, alpha=0.3)
        ax.axhline(y=0, color='brown', lw=3)
        # 标注
        ax.text(0, -H_ch * 0.15, f'R={R:.2f}m, θ={theta_deg:.0f}°', ha='center', fontsize=8, color='gray')
        if h_w > 0:
            ax.annotate('', xy=(-x_top_r - 0.08 * x_top_r, h_w), xytext=(-x_top_r - 0.08 * x_top_r, 0),
                        arrowprops=dict(arrowstyle='<->', color='blue', lw=1.5))
            ax.text(-x_top_r - 0.15 * x_top_r, h_w / 2, f'h={h_w:.2f}m',
                    ha='right', fontsize=8, color='blue', rotation=90, va='center')

    def _draw_circular(self, ax, D, y, V, Q, title):
        R = D / 2
        theta = np.linspace(0, 2*np.pi, 100)
        cx = R * np.cos(theta); cy = R + R * np.sin(theta)
        ax.plot(cx, cy, 'k-', lw=2)
        if y > 0 and y < D:
            h_off = y - R
            if abs(h_off) <= R:
                half_a = math.acos(h_off / R)
                water_w = math.sqrt(R**2 - h_off**2)
                wa = np.linspace(np.pi/2 + half_a, np.pi/2 - half_a + 2*np.pi, 50)
                wx = R * np.cos(wa); wy = R + R * np.sin(wa)
                mask = wy <= y + 0.001
                wxf = wx[mask]; wyf = wy[mask]
                if len(wxf) > 0:
                    px = np.concatenate([[water_w], wxf, [-water_w]])
                    py = np.concatenate([[y], wyf, [y]])
                    ax.fill(px, py, color='lightblue', alpha=0.7)
                    ax.plot([-water_w, water_w], [y, y], 'b-', lw=1.5)
        ax.annotate('', xy=(R, R), xytext=(-R, R),
                     arrowprops=dict(arrowstyle='<->', color='gray', lw=1.5))
        ax.text(0, R+0.15*R, f'D={D:.2f}m', ha='center', fontsize=9, color='gray')
        if y > 0:
            ax.annotate('', xy=(-R-0.12*R, y), xytext=(-R-0.12*R, 0),
                         arrowprops=dict(arrowstyle='<->', color='blue', lw=1.5))
            ax.text(-R-0.2*R, y/2, f'y={y:.2f}m', ha='right', fontsize=8, color='blue', rotation=90, va='center')
        ax.set_xlim(-R*1.7, R*1.7)
        ax.set_ylim(-R*0.4, D*1.2)
        ax.set_aspect('equal')
        ax.set_title(f'{title}\nQ={Q:.2f}m$^3$/s, V={V:.2f}m/s', fontsize=10)
        ax.grid(True, alpha=0.3)
        ax.axhline(y=0, color='brown', lw=3)

    # ================================================================
    # 清空
    # ================================================================
    def _clear(self):
        self._show_initial_help()
        self.section_fig.clear()
        self.section_canvas.draw()
        self.inc_hint.setText("(留空则自动计算)")
        self.current_result = None

    # ================================================================
    # 导出
    # ================================================================
    def _export_dxf(self):
        if not self.current_result or not self.current_result.get('success'):
            InfoBar.warning("提示", "请先进行计算后再导出。", parent=self._info_parent(), duration=3000, position=InfoBarPosition.TOP)
            return
        p = self.input_params
        res = self.current_result
        stype = p.get('section_type', '梯形')
        if stype == '圆形':
            D = res.get('D_design', 0.0)
            default_name = f'明渠断面_圆形_D{D:.2f}.dxf'
        elif stype == 'U形':
            R_val = res.get('R', 0.0)
            default_name = f'明渠断面_U形_R{R_val:.2f}.dxf'
        else:
            b = res.get('b_design', 0.0)
            H = res.get('h_prime', 0.0)
            if H <= 0:
                h_inc = res.get('h_increased', 0.0)
                H = (h_inc + res.get('Fb', 0.3)) if h_inc > 0 else res.get('h_design', 0.0) * 1.35
            default_name = f'明渠断面_{stype}_B{b:.2f}xH{H:.2f}.dxf'
        scales = ['1:20', '1:50', '1:100', '1:200', '1:500']
        from 渠系断面设计.styles import fluent_select
        scale_str, ok = fluent_select(self, '选择比例尺', '输出比例尺 (图纸单位: mm):', scales, 2)
        if not ok: return
        scale_denom = int(scale_str.split(':')[1])
        filepath, _ = QFileDialog.getSaveFileName(
            self, "保存DXF文件", default_name, "DXF文件 (*.dxf);;所有文件 (*.*)"
        )
        if not filepath:
            return
        try:
            export_open_channel_dxf(filepath, res, p, scale_denom)
            InfoBar.success("导出成功", f"DXF已保存到: {filepath}", parent=self._info_parent(), duration=4000, position=InfoBarPosition.TOP)
            ask_open_file(filepath, self._info_parent())
        except ImportError as e:
            InfoBar.error("缺少依赖", str(e), parent=self._info_parent(), duration=6000, position=InfoBarPosition.TOP)
        except PermissionError:
            InfoBar.error("文件被占用", "无法写入文件，请先关闭已打开的同名DXF文件，然后重新操作。", parent=self._info_parent(), duration=8000, position=InfoBarPosition.TOP)
        except Exception as e:
            InfoBar.error("导出失败", f"DXF导出失败: {str(e)}", parent=self._info_parent(), duration=5000, position=InfoBarPosition.TOP)

    def _export_report(self):
        if not self.current_result or not self.current_result.get('success'):
            InfoBar.warning("提示", "请先进行计算后再导出。", parent=self._info_parent(), duration=3000, position=InfoBarPosition.TOP)
            return
        filepath, _ = QFileDialog.getSaveFileName(self, "保存报告", "", "文本文件 (*.txt);;所有文件 (*.*)")
        if not filepath: return
        try:
            content = self._export_plain_text if self._export_plain_text else ''
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(content)
            InfoBar.success("导出成功", f"报告已保存到: {filepath}", parent=self._info_parent(), duration=4000, position=InfoBarPosition.TOP)
            ask_open_file(filepath, self._info_parent())
        except PermissionError:
            InfoBar.error("文件被占用", "无法写入文件，请先关闭已打开的同名文件（如记事本等），然后重新操作。", parent=self._info_parent(), duration=8000, position=InfoBarPosition.TOP)
        except Exception as e:
            InfoBar.error("导出失败", f"保存失败: {str(e)}", parent=self._info_parent(), duration=5000, position=InfoBarPosition.TOP)

    def _export_word(self):
        if not WORD_EXPORT_AVAILABLE:
            InfoBar.warning("缺少依赖",
                "Word导出需要安装 python-docx、latex2mathml、lxml。请执行: pip install python-docx latex2mathml lxml",
                parent=self._info_parent(), duration=6000, position=InfoBarPosition.TOP)
            return
        if not self.current_result or not self.current_result.get('success'):
            InfoBar.warning("提示", "请先进行计算后再导出。", parent=self._info_parent(), duration=3000, position=InfoBarPosition.TOP)
            return
        p = self.input_params
        stype = p.get('section_type', '梯形')
        channel_name = p.get('channel_name', '') or getattr(self, '_channel_name_text', '')
        meta = load_meta()
        auto_purpose = build_calc_purpose('open_channel',
            project=meta.project_name, name=channel_name, section_type=stype)
        dlg = ExportConfirmDialog('open_channel', '明渠水力计算书', auto_purpose, parent=self._info_parent())
        from PySide6.QtWidgets import QDialog
        if dlg.exec() != QDialog.Accepted:
            return
        self._word_export_meta = dlg.get_meta()
        self._word_export_purpose = dlg.get_calc_purpose()
        self._word_export_refs = dlg.get_references()
        filepath, _ = QFileDialog.getSaveFileName(self, "保存Word报告", "", "Word文档 (*.docx);;所有文件 (*.*)")
        if not filepath: return
        try:
            self._build_word_report(filepath)
            InfoBar.success("导出成功", f"Word报告已保存到: {filepath}", parent=self._info_parent(), duration=4000, position=InfoBarPosition.TOP)
            ask_open_file(filepath, self._info_parent())
        except PermissionError:
            InfoBar.error("文件被占用", "无法写入文件，请先关闭已打开的同名Word文档，然后重新操作。", parent=self._info_parent(), duration=8000, position=InfoBarPosition.TOP)
        except Exception as e:
            InfoBar.error("导出失败", f"Word导出失败: {str(e)}", parent=self._info_parent(), duration=5000, position=InfoBarPosition.TOP)

    def _build_word_report(self, filepath):
        """构建Word报告文档（工程产品运行卡格式）"""
        res = self.current_result
        p = self.input_params
        stype = p.get('section_type', '梯形')
        method = res.get('design_method', '')
        meta = getattr(self, '_word_export_meta', load_meta())
        purpose = getattr(self, '_word_export_purpose', '')
        refs = getattr(self, '_word_export_refs', REFERENCES_BASE.get('open_channel', []))
        content_desc = f'明渠水力断面设计计算（{stype}断面）'

        doc = create_engineering_report_doc(
            meta=meta,
            calc_title='明渠水力计算书',
            calc_content_desc=content_desc,
            calc_purpose=purpose,
            references=refs,
            calc_program_text=f'渠系建筑物水力计算系统 V1.0\n明渠水力计算（{stype}断面 · {method}）',
        )
        doc.add_page_break()

        # 5. 基础公式
        doc_add_eng_h(doc, '5、基础公式')
        doc_add_formula(doc, r'Q = \frac{1}{n} \cdot A \cdot R^{2/3} \cdot i^{1/2}', '曼宁公式：')
        if stype == '圆形':
            doc_add_formula(doc, r'A = \frac{D^2}{8}(\theta - \sin\theta)', '过水面积：')
            doc_add_formula(doc, r'\chi = \frac{D}{2} \cdot \theta', '湿周：')
            doc_add_formula(doc, r'\theta = 2\arccos\frac{R-h}{R}', '圆心角：')
        elif stype == 'U形':
            doc_add_formula(doc, r'h_0 = R(1-\cos(\theta/2))', '弧区临界水深：')
            doc_add_formula(doc, r'A = R^2\arccos\frac{R-h}{R}-(R-h)\sqrt{2Rh-h^2}', '纯弧区面积(h\leq h_0)：')
            doc_add_formula(doc, r'A = A_{arc}+(b_{arc}+m\cdot h_s)\cdot h_s', '直线段区面积(h>h_0)：')
        elif stype == '梯形':
            doc_add_formula(doc, r'A = (B + m \cdot h) \cdot h', '过水面积：')
            doc_add_formula(doc, r'\chi = B + 2h\sqrt{1+m^2}', '湿周：')
        else:
            doc_add_formula(doc, r'A = B \cdot h', '过水面积：')
            doc_add_formula(doc, r'\chi = B + 2h', '湿周：')
        doc_add_formula(doc, r'R = \frac{A}{\chi}', '水力半径：')
        doc_add_formula(doc, r'V = \frac{1}{n} \cdot R^{2/3} \cdot i^{1/2}', '流速公式：')

        # 6. 计算过程
        doc_add_eng_h(doc, '6、计算过程')
        doc_render_calc_text_eng(doc, self._export_plain_text or '', skip_title_keyword='明渠水力计算结果')

        # 7. 断面方案对比（附录E）
        schemes = res.get('appendix_e_schemes', [])
        if schemes and stype != '圆形':
            doc_add_eng_h(doc, '7、断面方案对比（附录E）')
            doc_add_table_caption(doc, '表 1  附录E断面方案对比表')
            b_sel = res['b_design']; h_sel = res['h_design']
            v_min, v_max = p['v_min'], p['v_max']
            headers = ['α值', '方案类型', '底宽B(m)', '水深h(m)', '宽深比β', '流速V(m/s)', '面积增加', '状态']
            data = []
            for s in schemes:
                is_sel = abs(s['b'] - b_sel) < 0.01 and abs(s['h'] - h_sel) < 0.01
                v_ok = v_min < s['V'] < v_max
                status = '★选中' if is_sel else ('流速不符' if not v_ok else '')
                data.append([
                    f"{s['alpha']:.2f}", s['scheme_type'],
                    f"{s['b']:.3f}", f"{s['h']:.3f}", f"{s['beta']:.3f}",
                    f"{s['V']:.3f}", f"+{s['area_increase']:.0f}%", status
                ])
            doc_add_styled_table(doc, headers, data, highlight_col=7, highlight_val='★选中',
                                 with_full_border=True)

        # 断面图
        try:
            import tempfile
            tmp = os.path.join(tempfile.gettempdir(), '_mingqu_section.png')
            self.section_fig.savefig(tmp, dpi=150, bbox_inches='tight')
            section_no = '8' if schemes and stype != '圆形' else '7'
            doc_add_eng_h(doc, f'{section_no}、断面图')
            doc_add_figure(doc, tmp, width_cm=14)
            os.remove(tmp)
        except Exception:
            pass

        doc.save(filepath)
