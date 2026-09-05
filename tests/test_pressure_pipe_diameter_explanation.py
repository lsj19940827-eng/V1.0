# -*- coding: utf-8 -*-
"""验证结果尺寸解释、逐候选数值代入和Word公式，关联有压管道展示与快照。"""

from dataclasses import replace
from pathlib import Path
from types import SimpleNamespace
from uuid import uuid4
import os

os.environ.setdefault('QT_QPA_PLATFORM', 'offscreen')

import pytest
from docx import Document
from PySide6.QtWidgets import QApplication

from app_渠系计算前端.pressure_pipe import diameter_explanation as explanations
from app_渠系计算前端.pressure_pipe.panel import PressurePipePanel
from calc_渠系计算算法内核.有压管道设计 import PressurePipeInput, recommend_diameter


@pytest.fixture(scope='module')
def qapp():
    """提供离屏结果面板所需的Qt应用。"""
    return QApplication.instance() or QApplication([])


def _result(material='HDPE管', **kwargs):
    """按材料生成真实内核候选，避免用展示代码反向制造期望值。"""
    inp = PressurePipeInput(Q=0.5, material_key=material, **kwargs)
    return inp, recommend_diameter(inp)


def test_pe_summary_and_every_candidate_explain_actual_wall_thickness(qapp):
    """用户截图工况必须明确705.2的由来，并逐档显示各自壁厚，不能全复用推荐值。"""
    inp, result = _result()
    panel = PressurePipePanel()
    try:
        panel._all_results = [(0, inp, result)]
        html = panel._build_result_card_html(0, inp, result)
        assert ('单侧壁厚', '47.4 mm') in explanations.explain_diameter(result.recommended).dimensions
        assert '800 − 2 × 47.4 = 705.2 mm' in html
        assert '公称外径 × 壁厚' in html
        assert '公称直径、壁厚与水力内径' in html
        assert html.count('class="candidate-diameter-explanation"') == len(result.top_candidates) - 1
        for candidate in result.top_candidates:
            info = explanations.explain_diameter(candidate)
            assert info.substitution_text in html
            assert f'单侧壁厚 {candidate.nominal_wall_thickness_mm:g} mm' in html
        assert html.count('<svg') >= len(result.top_candidates) + 1
        assert '\\times' not in html and '\\mathrm' not in html
    finally:
        panel.close()
        panel.deleteLater()
        qapp.processEvents()


def test_pe_uses_snapshot_table_thickness_instead_of_outer_divided_by_sdr():
    """锁定表列特殊值90.2，解释不得为凑SDR而改为90.9。"""
    _, result = _result(pe_nominal_pressure_mpa=1.6, manual_nominal_diameter_mm=1000)
    info = explanations.explain_diameter(result.recommended)
    assert info.substitution_text == '1000 − 2 × 90.2 = 819.6 mm'
    assert '表3' in info.source
    assert '不' in info.source and 'SDR' in info.source


def test_di_explanation_deducts_both_wall_and_lining_and_keeps_saved_standard():
    """球墨铸铁从插口外径而非DN扣壁厚和内衬，历史标准仍来自结果快照。"""
    _, result = _result('球墨铸铁管', manual_product_diameter_mm=300)
    info = explanations.explain_diameter(result.recommended)
    assert info.substitution_text == '326 − 2 × (6.2 + 4) = 305.6 mm'
    assert ('插口外径 DE', '326 mm') in info.dimensions
    assert ('单侧内衬厚', '4 mm') in info.dimensions
    saved = replace(result.recommended, product_standard='GB/T 13295—2019',
                    product_source_locator='GB/T 13295—2019 旧结果表列尺寸')
    assert '2019' in explanations.explain_diameter(saved).source
    assert '2026' not in explanations.explain_diameter(saved).source


@pytest.mark.parametrize('material,dn', [('预应力钢筒混凝土管', 1600), ('玻璃钢夹砂管', 2800)])
def test_inner_series_explains_direct_use_without_inventing_wall(material, dn):
    """内径系列直接采用公称内径，不虚构壁厚为零或套用PE公式。"""
    _, result = _result(material, manual_product_diameter_mm=dn)
    info = explanations.explain_diameter(result.recommended)
    assert info.wall_text == '壁厚未由本尺寸目录确定'
    assert '不再扣两次壁厚' in info.meaning
    assert '- 2' not in info.substitution
    assert info.substitution_text == f'直接采用公称内径 {dn} mm，水力内径为 {dn} mm'


def test_legacy_and_incomplete_snapshots_do_not_infer_new_catalog_dimensions():
    """历史任意内径不虚构DN；残缺产品快照明确保留未知壁厚。"""
    legacy = SimpleNamespace(D=0.7, nominal_outer_diameter_mm=None)
    assert explanations.diameter_summary_html(legacy) == ''
    _, result = _result()
    partial = replace(result.recommended, nominal_wall_thickness_mm=None)
    info = explanations.explain_diameter(partial)
    assert info.wall_text == '壁厚信息未完整保存'
    assert info.formula == ''
    assert '无法展示完整换算' in info.method


def test_formula_render_failure_uses_readable_text(monkeypatch):
    """渲染失败时保留可读中文和代入值，不能向用户暴露LaTeX源码。"""
    monkeypatch.setattr(explanations, 'render_latex_svg', lambda *args, **kwargs: None)
    _, result = _result()
    html = explanations.diameter_summary_html(result.recommended)
    assert '800 − 2 × 47.4 = 705.2 mm' in html
    assert '\\times' not in html


@pytest.mark.parametrize('detail', [False, True])
def test_actual_result_page_keeps_one_selected_conversion_and_table_first(qapp, monkeypatch, detail):
    """真实显示入口在详细开关两种状态下均只展示一次推荐换算，候选表仍优先。"""
    from app_渠系计算前端.pressure_pipe import panel as panel_module

    captured = []
    monkeypatch.setattr(panel_module, 'load_formula_page', lambda view, html: captured.append(html))
    monkeypatch.setattr(PressurePipePanel, '_show_initial_help', lambda self: None)
    monkeypatch.setattr(PressurePipePanel, '_jump_to_case_result', lambda *args, **kwargs: None)
    inp, result = _result('球墨铸铁管')
    panel = PressurePipePanel()
    try:
        panel._all_results = [(0, inp, result)]
        panel.detail_cb.setChecked(detail)
        panel._display_all_results()
        html = captured[-1]
        info = explanations.explain_diameter(result.recommended)
        assert html.count(info.substitution_text) == 1
        assert '推荐管径 DN 800' in html
        assert html.index('候选管径对比') < html.index('class="diameter-explanation"')
        assert '五、推荐管径结果' not in html
        assert ('流速与水头损失计算' in html) == detail
        if detail:
            assert 'f 上限' in html and 'f 下限' in html
        assert panel._all_results[0][2].calc_steps == result.calc_steps
    finally:
        panel.close()
        panel.deleteLater()
        qapp.processEvents()


def test_word_summary_and_candidate_table_have_editable_formulas():
    """Word与界面共用相同快照，推荐和每档候选公式均写为OMML。"""
    _, result = _result()
    doc = Document()
    explanations.add_diameter_summary_to_word(doc, result.recommended)
    explanations.add_candidate_diameters_to_word(doc, result.top_candidates)
    assert len(doc._element.xpath('.//m:oMath')) == len(result.top_candidates) + 2
    table = doc.tables[-1]
    assert table.cell(1, 0).text == 'DN800'
    assert table.cell(1, 1).text == '单侧壁厚 47.4 mm'
    assert '705.2' in table.cell(1, 2)._tc.xml
    assert '800' in table.cell(1, 2)._tc.xml


@pytest.mark.parametrize('material', ['HDPE管', '球墨铸铁管'])
def test_actual_word_export_includes_new_diameter_explanations(qapp, material):
    """调用真实报告出口，确保尺寸说明不是只在孤立的辅助函数里存在。"""
    inp, result = _result(material)
    panel = PressurePipePanel()
    destination = Path(__file__).resolve().parents[1] / 'tmp' / f'diameter_report_{uuid4().hex}.docx'
    destination.parent.mkdir(exist_ok=True)
    try:
        panel._all_results = [(0, inp, result)]
        panel._export_plain_text = result.calc_steps
        panel._build_word_report(str(destination))
        doc = Document(destination)
        paragraphs = '\n'.join(p.text for p in doc.paragraphs)
        assert '公称直径、壁厚与水力内径' in paragraphs
        assert '候选规格的壁厚与内径换算' in paragraphs
        assert any(
            row.cells[0].text == '推荐管径' and row.cells[1].text == 'DN 800'
            for table in doc.tables for row in table.rows if len(row.cells) >= 2
        )
        assert any(t.cell(0, 2).text == '水力内径计算（mm）' for t in doc.tables if len(t.columns) == 3)
    finally:
        panel.close()
        panel.deleteLater()
        qapp.processEvents()
        destination.unlink(missing_ok=True)


@pytest.mark.parametrize('manual', [False, True])
@pytest.mark.parametrize('material,manual_field,options', [
    ('球墨铸铁管', 'manual_product_diameter_mm', {}),
    ('HDPE管', 'manual_nominal_diameter_mm', {}),
    ('预应力钢筒混凝土管', 'manual_product_diameter_mm', {}),
    ('玻璃钢夹砂管', 'manual_product_diameter_mm', {}),
    ('钢管', 'manual_steel_diameter_mm', {'steel_dimensions_enabled': True}),
])
def test_final_result_prioritizes_nominal_size_and_preserves_inner_precision(material, manual_field, options, manual):
    """自动和指定结果都须先给DN规格，再给准确水力内径，防止把810.8取整成811当选管结果。"""
    kwargs = dict(options)
    if manual:
        kwargs[manual_field] = 1000 if material == '预应力钢筒混凝土管' else 800
    _, result = _result(material, **kwargs)
    rec = result.recommended
    nominal = rec.nominal_outer_diameter_mm or rec.nominal_diameter_mm
    label = '指定管径' if manual else '推荐管径'
    summary = result.calc_steps.split(f'【五、{label}结果】', 1)[1].split('【六、', 1)[0]
    assert summary.strip().startswith(f'1. {label} DN {nominal:g}')
    assert f'd_i = {rec.hydraulic_inner_diameter_mm:g} mm = {rec.D:g} m' in summary
    if material == '球墨铸铁管':
        assert 'DN 800' in summary
        assert 'DE = 842 mm' in summary
        assert '810.8 mm' in summary
        assert '811 mm' not in summary
    from app_渠系计算前端.formula_renderer import plain_text_to_formula_body
    html = plain_text_to_formula_body(summary)
    assert f'DN {nominal:g}' in html
    assert f'class="step-title">{label} DN {nominal:g}' in html
    assert '810.8:</div>' not in html
    assert '<svg' in html
