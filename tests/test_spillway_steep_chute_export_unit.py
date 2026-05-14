# -*- coding: utf-8 -*-
"""泄水渠与陡坡 Word / Excel 导出的单元测试。"""

import sys
from pathlib import Path
from types import SimpleNamespace

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from docx import Document
from openpyxl import load_workbook

from app_渠系计算前端.spillway_steep_chute.report_export import (
    export_spillway_steep_chute_excel,
    export_spillway_steep_chute_word,
)

FORBIDDEN_USER_VISIBLE_RESIDUE = [
    "trapezoidal",
    "rectangular",
    "manual",
    "backwater",
    "control",
    "wall",
    "cap",
    r"begin{cases}",
    r"end{cases}",
    r"\frac",
    r"\quad",
    r"\text",
    r"\sqrt",
    r"\left",
    r"\right",
    r"\geq",
    r"\max",
    r"\min",
]


def _assert_no_user_visible_residue(text: str) -> None:
    """断言导出成果不暴露源码或英文内部值。"""
    for forbidden in FORBIDDEN_USER_VISIBLE_RESIDUE:
        assert forbidden not in text


def _result_payload():
    """构造导出覆盖所需的结果数据。"""
    return {
        "input_params": {"Q": 12.5, "n": 0.014, "i": 0.02},
        "summary": {
            "工程名称": "熊启钧教学算例",
            "设计流量": "12.50 立方米/秒",
            "最大掺气水深": "1.360 米",
            "正常水深": "1.084 米",
            "临界水深": "1.749 米",
            "坡型": "陡坡",
        },
        "hydraulic": {
            "section_type": "trapezoidal",
            "slope_type": "steep",
            "normal": {"area_m2": 5.2, "hydraulic_radius_m": 0.9, "velocity_ms": 3.8},
            "critical": {"area_m2": 6.3, "hydraulic_radius_m": 0.86, "water_top_width_m": 6.2, "velocity_ms": 3.1},
            "start": {"area_m2": 6.3, "hydraulic_radius_m": 0.86, "water_top_width_m": 6.2, "depth_m": 1.749},
        },
        "profile": {"end_depth_m": 0.9, "water_profile_name": "陡坡降水曲线"},
        "profile_points": [
            {
                "x": 0.0,
                "bed_elevation": 100.0,
                "water_elevation": 101.2,
                "depth": 1.2,
                "velocity_ms": 5.2,
                "froude": 1.6,
                "aerated_depth_m": 1.36,
                "sidewall_top_elevation_m": 101.9,
            },
            {
                "x": 50.0,
                "bed_elevation": 95.0,
                "water_elevation": 95.9,
                "depth": 0.9,
                "velocity_ms": 6.1,
                "froude": 2.0,
                "aerated_depth_m": 1.02,
                "sidewall_top_elevation_m": 96.6,
            },
        ],
        "start_control": {"source": "manual", "depth_m": 1.1},
        "upstream_connection": {"upstream_normal_depth_m": 1.8, "message": "上游缓坡自由接陡坡。"},
        "hydraulic_jump": {
            "pre_jump_depth_m": 0.72,
            "pre_jump_velocity_ms": 8.4,
            "pre_jump_froude": 3.2,
            "conjugate_depth_m": 2.2,
            "control_depth_m": 1.5,
            "tailwater_judgement": "尾水不足",
            "recommended_pool_length_m": 9.9,
            "recommended_pool_depth_m": 0.7,
            "recommended_transition_length_m": 18.0,
            "message": "按跃后共轭水深初拟消力池。",
        },
        "aeration_and_sidewall": {
            "max_aerated_depth_m": 1.36,
            "recommended_sidewall_height_m": 1.8,
            "message": "逐点计算掺气水深和侧墙顶线。",
        },
        "multi_flow_control": {
            "cases": [
                {
                    "name": "设计流量",
                    "Q": 12.5,
                    "pre_jump_depth_m": 0.72,
                    "conjugate_depth_m": 2.2,
                    "control_depth_m": 1.5,
                    "control_depth_deficit_m": 0.7,
                    "recommended_pool_length_m": 9.9,
                    "recommended_pool_depth_m": 0.7,
                }
            ]
        },
        "water_profile_export": {
            "points": [
                {"桩号_m": 0.0, "渠底高程_m": 100.0, "水深_m": 1.2, "水位_m": 101.2, "流速_m_s": 5.2, "弗劳德数": 1.6}
            ],
            "说明": "表3轻量接口仅提供入口、出口和沿程采样点。",
        },
        "checks": [
            {"name": "流速校核", "result": "通过", "message": "设计流速在建议范围内"},
        ],
        "risks": ["出口需复核消能防冲。"],
        "formulas": [
            {"name": "曼宁公式", "latex": r"Q=\frac{1}{n}AR^{2/3}i^{1/2}", "source": "GB 50288-2018"},
        ],
    }


def test_export_excel_writes_principles_before_summary_and_result_sheets(tmp_path):
    """Excel 应先写出计算原理，再写出第二版结果、校核和风险。"""
    output_path = tmp_path / "spillway.xlsx"

    export_spillway_steep_chute_excel(output_path, _result_payload())

    workbook = load_workbook(output_path)
    assert workbook.sheetnames == [
        "计算原理",
        "结果汇总",
        "沿程水面线",
        "上下游衔接",
        "消能与出口",
        "多流量控制",
        "表3轻量接口",
        "规范校核",
        "风险提示",
    ]
    assert workbook["计算原理"]["A1"].value == "步骤"
    assert workbook["计算原理"]["B1"].value == "计算目的"
    assert workbook["计算原理"]["C1"].value == "公式"
    assert workbook["计算原理"]["D1"].value == "本次代入"
    assert workbook["计算原理"]["F1"].value == "原理说明"
    principle_blob = "\n".join(str(cell.value or "") for row in workbook["计算原理"].iter_rows() for cell in row)
    assert "正常水深" in principle_blob
    assert "临界水深" in principle_blob
    assert "水面线逐段计算" in principle_blob
    assert "掺气水深" in principle_blob
    assert "χ=b+2h√(1+m²)" in principle_blob
    assert "R=A/χ" in principle_blob
    assert "P=b" not in principle_blob
    _assert_no_user_visible_residue("\n".join(str(cell.value or "") for sheet in workbook.worksheets for row in sheet.iter_rows() for cell in row))
    assert workbook["结果汇总"]["A2"].value == "工程名称"
    assert workbook["结果汇总"]["B2"].value == "熊启钧教学算例"
    assert workbook["沿程水面线"]["A1"].value == "桩号（米）"
    assert workbook["沿程水面线"]["G1"].value == "掺气水深（米）"
    assert workbook["沿程水面线"]["A2"].value == 0.0
    assert workbook["上下游衔接"]["A2"].value == "陡槽起点水深来源"
    assert workbook["消能与出口"]["A4"].value == "建议消力池长度（米）"
    assert workbook["多流量控制"]["A2"].value == "设计流量"
    assert workbook["表3轻量接口"]["A2"].value == 0.0
    assert workbook["规范校核"]["A2"].value == "流速校核"
    assert "消能防冲" in workbook["风险提示"]["A2"].value


def test_export_excel_preserves_zero_values(tmp_path):
    """导出时合法 0 值应保留为 0，不应写成空白。"""
    output_path = tmp_path / "spillway_zero.xlsx"
    payload = _result_payload()
    payload["start_control"]["depth_m"] = 0.0
    payload["hydraulic_jump"]["recommended_pool_depth_m"] = 0.0
    payload["multi_flow_control"]["cases"][0]["control_depth_deficit_m"] = 0.0

    export_spillway_steep_chute_excel(output_path, payload)

    workbook = load_workbook(output_path)
    assert workbook["上下游衔接"]["B3"].value == 0.0
    assert workbook["消能与出口"]["B5"].value == 0.0
    assert workbook["多流量控制"]["F2"].value == 0.0


def test_export_word_writes_principles_before_summary_profile_checks_and_risks(tmp_path):
    """Word 应先写出计算原理，再写出汇总、沿程点、校核和风险。"""
    output_path = tmp_path / "spillway.docx"

    export_spillway_steep_chute_word(output_path, _result_payload())

    document = Document(output_path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    table_text = "\n".join(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    combined = f"{text}\n{table_text}"

    assert "泄水渠与陡坡计算书" in combined
    assert "工程阶段产品运行卡" in combined
    assert "强制性标准条文执行情况校审检查表" in combined
    assert "计算目的" in combined
    assert "计算依据" in combined
    assert "基本资料" in combined
    assert "计算程序" in combined
    assert "熊启钧教学算例" in combined
    assert combined.index("计算原理") < combined.index("结果汇总")
    assert "水面线逐段计算" in combined
    assert "本次代入" in combined
    assert "沿程水面线" in combined
    assert "上下游衔接" in combined
    assert "消能与出口整流" in combined
    assert "多流量控制工况" in combined
    assert "表3轻量接口" in combined
    assert "掺气水深" in combined
    assert "建议消力池长度" in combined
    assert "流速校核" in combined
    assert "消能防冲" in combined
    assert "GB 50288-2018" in combined
    assert "χ=b+2h√(1+m²)" in combined
    assert "R=A/χ" in combined
    assert "P=b" not in combined
    _assert_no_user_visible_residue(combined)


def test_export_excel_sanitizes_internal_prd_principle_sources(tmp_path):
    """Excel 计算原理不应暴露内部 PRD 表述。"""
    output_path = tmp_path / "spillway_formula_source.xlsx"
    payload = _result_payload()
    payload["formulas"] = [
        {
            "name": "出口整流段",
            "latex": r"L_r=\max(L_{\Delta b},\eta h_c'',L_{\min})",
            "source": "PRD 第二版出口整流段设计口径",
        }
    ]

    export_spillway_steep_chute_excel(output_path, payload)

    workbook = load_workbook(output_path)
    principle_blob = "\n".join(str(cell.value or "") for row in workbook["计算原理"].iter_rows() for cell in row)
    assert "PRD" not in principle_blob
    assert "出口连接段整流布置校核口径" in principle_blob


def test_export_word_sanitizes_internal_prd_principle_sources(tmp_path):
    """Word 计算原理不应暴露内部 PRD 表述。"""
    output_path = tmp_path / "spillway_formula_source.docx"
    payload = _result_payload()
    payload["formulas"] = [
        {
            "name": "消力池初拟尺寸",
            "latex": r"L_d=4.5h_c''",
            "source": "GB 50288-2018 附录 N 与 PRD 第二版口径",
        }
    ]

    export_spillway_steep_chute_word(output_path, payload)

    document = Document(output_path)
    combined = "\n".join(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    assert "PRD" not in combined
    assert "GB 50288-2018 附录 N 与消力池初拟经验口径" in combined


def test_export_word_writes_report_meta_purpose_and_references(tmp_path):
    """Word 导出应写入项目设置、计算目的和计算依据。"""
    output_path = tmp_path / "spillway_meta.docx"
    meta = SimpleNamespace(
        project_name="东干渠工程",
        design_stage="施工图",
        product_level="二级",
        record_number="SF-001",
        specialty="水工",
        calculator="张三",
        checker="李四",
        reviewer="王五",
        approver="赵六",
        volume_current="1",
        volume_total="2",
        basic_info="项目基本资料",
        mandatory_clause="无",
    )

    export_spillway_steep_chute_word(
        output_path,
        _result_payload(),
        meta=meta,
        calc_purpose="用于复核泄水渠与陡坡设计成果。",
        references=["《测试规范》"],
    )

    document = Document(output_path)
    combined = "\n".join(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    combined += "\n" + "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "东干渠工程" in combined
    assert "用于复核泄水渠与陡坡设计成果。" in combined
    assert "《测试规范》" in combined
