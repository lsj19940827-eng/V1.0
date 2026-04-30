# -*- coding: utf-8 -*-
"""共享工况对比表输出工具测试。"""

import os
import sys
from pathlib import Path

import ezdxf
import pytest

os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")

from PySide6.QtWidgets import QApplication, QTableWidget


ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from app_渠系计算前端.dxf_multi_export import DxfExportCaseEntry
from app_渠系计算前端.open_channel.comparison import OPEN_CHANNEL_COMPARISON_SPEC
from app_渠系计算前端.section_comparison import (
    add_section_comparison_word_tables,
    build_table_clipboard_text,
    draw_section_comparison_tables,
    fill_comparison_table,
)


def _qapp():
    """获取测试用 Qt 应用。"""
    return QApplication.instance() or QApplication([])


def _entry():
    """构造一个可成功对比的明渠工况。"""
    return DxfExportCaseEntry(
        case_idx=0,
        label="梯形",
        input_params={"section_type": "梯形", "Q": 5.0, "use_increase": True, "m": 1.5},
        result={
            "success": True,
            "h_design": 1.2,
            "V_design": 1.1,
            "Q_increased": 6.0,
            "h_increased": 1.35,
            "V_increased": 1.2,
            "Fb": 0.4,
            "b_design": 3.2,
            "h_prime": 1.75,
        },
        is_valid=True,
    )


def test_fill_comparison_table_and_copy_text_to_excel():
    """Qt 表格应按列配置填充，并能复制为 Excel 友好的制表符文本。"""
    _qapp()
    table = QTableWidget()
    columns = OPEN_CHANNEL_COMPARISON_SPEC.hydraulic_columns[:4]
    rows = [
        {
            "case_name": "工况 1｜梯形",
            "section_type": "梯形",
            "Q_design": 5.0,
            "Q_increased": 6.0,
        }
    ]

    fill_comparison_table(table, columns, rows)
    table.selectAll()
    text, row_count, col_count = build_table_clipboard_text(table)

    assert row_count == 1
    assert col_count == 4
    assert text == (
        "工况\t断面类型\t设计流量 Q(m³/s)\t加大流量 Q加大(m³/s)\n"
        "工况 1｜梯形\t梯形\t5.000\t6.000"
    )


def test_draw_section_comparison_tables_outputs_two_dxf_titles(tmp_path):
    """DXF 输出应包含水力和结构两张对比表。"""
    doc = ezdxf.new("R2010")
    msp = doc.modelspace()

    height = draw_section_comparison_tables(doc, msp, [_entry()], OPEN_CHANNEL_COMPARISON_SPEC, 0.0, 0.0)

    assert height > 0
    texts = [entity.dxf.text for entity in msp if entity.dxftype() == "TEXT"]
    assert "明渠水力结果对比表" in texts
    assert "明渠结构尺寸对比表" in texts
    assert any(text.startswith("设计流量 Q") for text in texts)
    assert any(text.startswith("渠道高度 H") for text in texts)


def test_add_section_comparison_word_tables_creates_two_tables():
    """Word 输出应在详细工况前追加两张对比表。"""
    from docx import Document

    doc = Document()
    add_section_comparison_word_tables(doc, [_entry()], OPEN_CHANNEL_COMPARISON_SPEC)

    paragraph_text = [p.text for p in doc.paragraphs]
    assert "明渠水力结果对比表" in paragraph_text
    assert "明渠结构尺寸对比表" in paragraph_text
    assert len(doc.tables) == 2
    assert doc.tables[0].cell(0, 0).text == "工况"
    assert doc.tables[1].cell(0, 0).text == "工况"
